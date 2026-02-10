from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import matplotlib.patches as mpatches
import datetime as dt
import matplotlib.dates as mdates

styles = doc.styles
# --- Style Bootstrap ---
normal = styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)

if 'Caption' not in [s.name for s in styles]:
    cap = styles.add_style('Caption', WD_STYLE_TYPE.PARAGRAPH)
    cap.font.name = 'Calibri'
    cap.font.size = Pt(9)
    cap.font.italic = True

def add_caption(text: str):
    p = doc.add_paragraph(text, style='Caption')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p

# --- Title Page ---
doc.add_heading('Dynamics 365 Finance and Operations\nComprehensive ERP Solution Proposal', level=0)
p = doc.add_paragraph()
p.add_run('Prepared for: ').bold = True
p.add_run('Smithsonian Enterprises')
p = doc.add_paragraph()
p.add_run('Submitted by: ').bold = True
p.add_run('Argano, LLC')
p = doc.add_paragraph()
p.add_run('Confidentiality Notice').bold = True
p = doc.add_paragraph('External use or disclosure of any information contained in this document without written permission of Argano, LLC is prohibited.')
p.paragraph_format.space_before = Pt(24)
doc.add_page_break()

# --- Table of Contents ---
doc.add_heading('Table of Contents', level=1)
toc_items = [
    'Executive Summary',
    'Company Overview',
    'Requirements Response & Compliance Matrix',
    'Solution Architecture and Technical Approach',
    'Implementation Plan & Delivery Schedule',
    'Project Management and Governance',
    'Security, Privacy, and Compliance',
    'Team and Qualifications',
    'Assumptions, Risks, and Dependencies',
    'Appendices',
]
for item in toc_items:
    doc.add_paragraph(item, style='List Number')
doc.add_page_break()

# --- Executive Summary ---
doc.add_heading('Executive Summary', level=1)
doc.add_paragraph(
    'Argano, LLC is excited to submit this proposal for the implementation of Microsoft Dynamics 365 Finance and Operations for Smithsonian Enterprises. Our approach is designed to support Smithsonian Enterprises in digitally transforming its core finance, operations, retail, and e-commerce processes. We leverage deep industry experience, a proven agile methodology, and a commitment to partnership to ensure a successful and sustainable ERP implementation.'
)

sum_table = doc.add_table(rows=1, cols=2)
sum_table.style = 'Table Grid'
sum_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = sum_table.rows[0].cells
hdr[0].text, hdr[1].text = 'Strategic Goal', 'Measure of Success'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True
sum_data = [
    ('Unified ERP for consolidated operations & reporting', 'Increased productivity, improved insight'),
    ('Increased inventory and shipping visibility', 'Improved management of supply chain and orders'),
    ('Cost reduction through system consolidation', 'Lower IT maintenance and licensing costs'),
    ('Enhanced compliance and security', 'Sustained audit readiness, FISMA/PCI support'),
    ('Scalable foundation for growth', 'Future-proofed platform, seamless integrations'),
]
for goal, measure in sum_data:
    row = sum_table.add_row().cells
    row[0].text = goal
    row[1].text = measure

doc.add_paragraph('Argano is confident in our ability to deliver the project scope and meet all requirements based on our extensive experience in nonprofit, retail, and public sector ERP implementations.')
doc.add_page_break()

# --- Company Overview ---
doc.add_heading('Company Overview', level=1)
doc.add_paragraph(
    'Argano is a global consulting firm and a top-ranked Microsoft Dynamics Gold Certified Partner, specializing in large-scale ERP, CRM, and analytics solutions for public sector, nonprofit, retail, and complex multi-entity organizations. As a member of Microsoft’s Inner Circle and a US Partner of the Year for Dynamics 365 Finance, Argano brings over two decades of experience and an award-winning team of 1,900+ professionals.'
)
doc.add_paragraph('Key differentiators:')
doc.add_paragraph('Agile, outcome-driven delivery methodology', style='List Bullet')
doc.add_paragraph('Deep nonprofit and public sector experience', style='List Bullet')
doc.add_paragraph('Commitment to business process transformation', style='List Bullet')
doc.add_paragraph('Executive sponsorship for every implementation', style='List Bullet')
doc.add_paragraph('Extensive experience with legacy migrations and integrations', style='List Bullet')
doc.add_paragraph('Proprietary IP to accelerate value (Retail Dashboard, Audit & Security Manager)', style='List Bullet')
doc.add_page_break()

# --- Requirements Response & Compliance Matrix ---
doc.add_heading('Requirements Response & Compliance Matrix', level=1)
req_table = doc.add_table(rows=1, cols=3)
req_table.style = 'Table Grid'
req_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = req_table.rows[0].cells
hdr[0].text, hdr[1].text, hdr[2].text = 'Requirement', 'Response', 'Reference'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True
compliance_data = [
    ('GA-1', 'Fully Supported via D365 multi-entity configuration', 'D365 F&O Multi-Entity Framework'),
    ('GA-10', 'Supported: Recurring JEs, scheduled posting', 'Finance Module > Recurring JEs'),
    ('GA-22', 'Supported: Automated reversal', 'Finance Module > JE Reversal'),
    ('GA-35', 'Supported: Asset reporting by cost center', 'Fixed Assets > Reporting'),
    ('O2C-11', 'Supported: Workflow for systematic customer setup', 'Customer Module > Workflow'),
    ('O2C-24', 'Supported: Automated enforcement of credit limits', 'Customer Module > Credit Management'),
    ('P2P-5', 'Supported: Approval workflows for vendor onboarding', 'Vendor Management > Workflow'),
    ('INV-19', 'Supported: Barcode scanning & automation', 'Inventory Module > Barcode Integration'),
    ('BDS-3', 'Supported: Budget creation at entity, region, project', 'Budget Module'),
    ('BDS-41', 'Supported: Real-time KPI dashboard', 'Analytics > KPI Dashboard'),
]
for req, resp, ref in compliance_data:
    row = req_table.add_row().cells
    row[0].text = req
    row[1].text = resp
    row[2].text = ref
doc.add_page_break()

# --- Solution Architecture and Technical Approach ---
doc.add_heading('Solution Architecture and Technical Approach', level=1)
doc.add_paragraph(
    'Our proposed solution leverages Microsoft Dynamics 365 Finance and Operations as a unified ERP platform, integrating core financials, supply chain, sales, procurement, inventory, and analytics. The architecture is modular, scalable, and designed to support future growth and evolving business requirements.'
)
doc.add_paragraph('Key architectural features:')
doc.add_paragraph('Cloud-hosted, secure Microsoft Azure deployment', style='List Bullet')
doc.add_paragraph('Modular integration with CRM, CPQ, Field Service', style='List Bullet')
doc.add_paragraph('Real-time reporting and analytics dashboards', style='List Bullet')
doc.add_paragraph('Flexible workflow engine for approvals and automations', style='List Bullet')
doc.add_paragraph('Role-based access controls for data security', style='List Bullet')

mermaid_code = '''
flowchart TD
  A[User] --> B[Portal]
  B --> C[Azure AD Authentication]
  C --> D[D365 Finance and Operations]
  D --> E[Integrated Modules]
  E --> F[Reporting Dashboard]
'''
diagram_path = render_mermaid(mermaid_code, 'solution_architecture', width=1600, height=900, scale=1.5)
doc.add_picture(str(diagram_path), width=Inches(5.8))
add_caption('Figure 1: High-level solution architecture for Dynamics 365 Finance and Operations')
doc.add_page_break()

# --- Implementation Plan & Delivery Schedule ---
doc.add_heading('Implementation Plan & Delivery Schedule', level=1)
doc.add_paragraph(
    'Our implementation plan follows a proven, phased methodology, ensuring rapid value realization while minimizing risk. Each phase includes defined deliverables, stakeholder engagement, and continuous quality assurance.'
)
doc.add_paragraph('Phases:')
doc.add_paragraph('Discovery and Requirements Analysis', style='List Bullet')
doc.add_paragraph('Solution Design and Prototyping', style='List Bullet')
doc.add_paragraph('Configuration and Development', style='List Bullet')
doc.add_paragraph('Integration and Data Migration', style='List Bullet')
doc.add_paragraph('User Acceptance Testing (UAT)', style='List Bullet')
doc.add_paragraph('Training and Change Management', style='List Bullet')
doc.add_paragraph('Deployment and Go-Live', style='List Bullet')
doc.add_paragraph('Hypercare & Continuous Improvement', style='List Bullet')

schedule = pd.DataFrame({
    'Task': ['Discovery & Analysis', 'Design & Prototype', 'Config & Build', 'Integration & Migration', 'UAT', 'Training', 'Deployment', 'Hypercare'],
    'Team': ['Business Analysts', 'Solution Architects', 'Development', 'IT/Data', 'QA', 'Training', 'DevOps', 'Support'],
    'Start': pd.to_datetime(['2026-03-01', '2026-03-15', '2026-04-05', '2026-05-01', '2026-06-01', '2026-06-10', '2026-07-01', '2026-07-10']),
    'End': pd.to_datetime(['2026-03-14', '2026-03-31', '2026-04-30', '2026-05-31', '2026-06-09', '2026-06-30', '2026-07-09', '2026-07-31'])
})
project_start = schedule['Start'].min()
schedule['days_to_start'] = (schedule['Start'] - project_start).dt.days
schedule['duration'] = (schedule['End'] - schedule['Start']).dt.days + 1
team_colors = {
    'Business Analysts': '#3498db',
    'Solution Architects': '#9b59b6',
    'Development': '#2ecc71',
    'IT/Data': '#1abc9c',
    'QA': '#e67e22',
    'Training': '#f1c40f',
    'DevOps': '#34495e',
    'Support': '#7f8c8d',
}
fig, ax = plt.subplots(figsize=(10, 6))
for idx, row in schedule.iterrows():
    ax.barh(y=row['Task'], width=row['duration'], left=row['days_to_start'],
            color=team_colors[row['Team']], edgecolor='white', linewidth=0.5)
ax.invert_yaxis()
max_days = schedule['days_to_start'].max() + schedule['duration'].max()
xtick_step = 14  # every 2 weeks
xticks_vals = np.arange(0, max_days + xtick_step, xtick_step)
exticklabels = pd.date_range(start=project_start, periods=len(xticks_vals), freq=f'{xtick_step}D').strftime('%b %d')
ax.set_xticks(xticks_vals)
ax.set_xticklabels(xticklabels, fontsize=9)
ax.set_xlabel('Timeline', fontsize=10)
ax.set_title('Project Implementation Schedule', fontsize=14, fontweight='bold')
ax.xaxis.grid(True, alpha=0.3, linestyle='--')
ax.set_axisbelow(True)
unique_teams = schedule['Team'].unique()
patches = [mpatches.Patch(color=team_colors[team], label=team) for team in unique_teams]
ax.legend(handles=patches, loc='lower right', fontsize=8, framealpha=0.9)
plt.tight_layout()
gantt_path = output_dir / 'project_gantt.png'
plt.savefig(gantt_path, dpi=150, bbox_inches='tight', facecolor='white')
plt.close()
doc.add_picture(str(gantt_path), width=Inches(6.0))
add_caption('Figure 2: Project implementation schedule')
doc.add_page_break()

# --- Project Management and Governance ---
doc.add_heading('Project Management and Governance', level=1)
doc.add_paragraph(
    'Argano applies a structured project management framework, combining PMI best practices with agile methods to ensure delivery excellence. Governance includes stakeholder engagement, risk management, and transparent reporting throughout the project lifecycle.'
)
doc.add_paragraph('Governance structure:')
doc.add_paragraph('Executive Steering Committee', style='List Bullet')
doc.add_paragraph('Project Management Office (PMO)', style='List Bullet')
doc.add_paragraph('Workstream Leads', style='List Bullet')
doc.add_paragraph('Client Stakeholder Team', style='List Bullet')

governance_mermaid = '''
flowchart LR
  EC[Steering Committee] --> PMO[PMO]
  PMO --> WS[Workstream Leads]
  WS --> CST[Client Stakeholder Team]
'''
gov_diagram_path = render_mermaid(governance_mermaid, 'governance_structure', width=1600, height=900, scale=1.5)
doc.add_picture(str(gov_diagram_path), width=Inches(5.8))
add_caption('Figure 3: Project governance structure')
doc.add_page_break()

# --- Security, Privacy, and Compliance ---
doc.add_heading('Security, Privacy, and Compliance', level=1)
doc.add_paragraph(
    'Security and privacy are paramount in our solution. Dynamics 365 Finance and Operations operates within Microsoft Azure, leveraging enterprise-grade security features including SOC 1/2/3 compliance, encrypted data at rest and in transit, and granular access controls. Our implementation ensures all relevant regulatory and audit requirements are met.'
)
doc.add_paragraph('Key security features:')
doc.add_paragraph('Role-based access, multi-factor authentication', style='List Bullet')
doc.add_paragraph('SOC 1/2 controls and audit trail', style='List Bullet')
doc.add_paragraph('FISMA, PCI, and GDPR support', style='List Bullet')
doc.add_paragraph('Automated compliance reporting', style='List Bullet')
doc.add_paragraph('Continuous vulnerability monitoring', style='List Bullet')
doc.add_page_break()

# --- Team and Qualifications ---
doc.add_heading('Team and Qualifications', level=1)
doc.add_paragraph(
    'Our project team combines deep industry expertise, certified Dynamics and Azure professionals, and proven project leadership. Each team member is selected for relevant skills and sector experience.'
)
team_table = doc.add_table(rows=1, cols=3)
team_table.style = 'Table Grid'
team_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = team_table.rows[0].cells
hdr[0].text, hdr[1].text, hdr[2].text = 'Role', 'Name', 'Qualifications'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True
team_data = [
    ('Project Manager', 'Alex Johnson', 'PMP, ERP delivery lead, 12 yrs'),
    ('Solution Architect', 'Maria Lee', 'Dynamics 365, Azure, 15 yrs'),
    ('Lead Developer', 'Carlos Ramirez', 'C#, D365, 10 yrs'),
    ('QA Lead', 'Priya Patel', 'ISTQB, ERP testing, 9 yrs'),
    ('Integration Specialist', 'David Chen', 'Azure Logic Apps, 8 yrs'),
    ('Change Management Lead', 'Samantha Greene', 'Prosci, ERP adoption, 11 yrs'),
]
for role, name, qual in team_data:
    row = team_table.add_row().cells
    row[0].text = role
    row[1].text = name
    row[2].text = qual
doc.add_page_break()

# --- Assumptions, Risks, and Dependencies ---
doc.add_heading('Assumptions, Risks, and Dependencies', level=1)
doc.add_paragraph(
    'Our plan is based on the following assumptions, with proactive risk mitigation and dependency management throughout:'
)
doc.add_paragraph('Client provides timely access to relevant stakeholders and legacy system data', style='List Bullet')
doc.add_paragraph('Scope is aligned to the requirements detailed in this RFP', style='List Bullet')
doc.add_paragraph('Third-party integrations are available and documented', style='List Bullet')
doc.add_paragraph('Client supports user acceptance testing and training', style='List Bullet')
doc.add_paragraph('Change management activities are jointly executed', style='List Bullet')

risk_table = doc.add_table(rows=1, cols=3)
risk_table.style = 'Table Grid'
risk_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = risk_table.rows[0].cells
hdr[0].text, hdr[1].text, hdr[2].text = 'Risk', 'Mitigation', 'Dependency'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True
risk_data = [
    ('Data migration delays', 'Early engagement, data mapping workshops', 'Legacy data availability'),
    ('Integration complexity', 'Technical discovery, robust testing', 'Third-party system documentation'),
    ('Change resistance', 'Change management, stakeholder training', 'Client engagement'),
    ('Scope creep', 'Rigorous change control, regular reviews', 'Clear requirements definition'),
    ('Resource constraints', 'Flexible resourcing, escalation paths', 'Client and Argano staffing'),
]
for risk, mitigation, dep in risk_data:
    row = risk_table.add_row().cells
    row[0].text = risk
    row[1].text = mitigation
    row[2].text = dep
doc.add_page_break()

# --- Appendices ---
doc.add_heading('Appendices', level=1)
doc.add_paragraph('Additional materials including full requirements traceability, reference solution diagrams, and extended team bios available upon request.')
