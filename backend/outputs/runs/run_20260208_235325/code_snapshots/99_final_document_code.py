from docx.shared import Pt, Inches
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import matplotlib.patches as mpatches
import datetime as dt
import matplotlib.dates as mdates

styles = doc.styles

# Normal style
normal = styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)

# Headings
styles['Heading 1'].font.name = 'Calibri'
styles['Heading 1'].font.size = Pt(18)
styles['Heading 1'].font.bold = True
styles['Heading 2'].font.name = 'Calibri'
styles['Heading 2'].font.size = Pt(14)
styles['Heading 2'].font.bold = True
styles['Heading 3'].font.name = 'Calibri'
styles['Heading 3'].font.size = Pt(12)
styles['Heading 3'].font.bold = True
styles['Title'].font.name = 'Calibri'
styles['Title'].font.size = Pt(26)
styles['Title'].font.bold = True

# Caption style
if 'Caption' not in [s.name for s in styles]:
    cap = styles.add_style('Caption', WD_STYLE_TYPE.PARAGRAPH)
    cap.font.name = 'Calibri'
    cap.font.size = Pt(9)
    cap.font.italic = True

def add_caption(text):
    p = doc.add_paragraph(text, style='Caption')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p

# Title Page

doc.add_heading('Kastle Systems ERP Modernization Proposal', level=0)
para = doc.add_paragraph('Prepared for: Kastle Systems')
para.paragraph_format.space_after = Pt(18)
para = doc.add_paragraph('Submitted by: [Vendor Company Name]', style='Normal')
para.paragraph_format.space_after = Pt(18)
doc.add_paragraph('Confidential: External use or disclosure of any information contained in this document without written permission is prohibited.', style='Normal')
doc.add_page_break()

# Letter of Transmittal
from datetime import date
today = date.today().strftime('%B %d, %Y')
doc.add_paragraph(f'{today}', style='Normal')
doc.add_paragraph('Letter of Transmittal to Kastle Systems', style='Normal')
doc.add_paragraph('Re: Kastle Systems ERP Modernization RFP', style='Normal')
doc.add_paragraph('To whom it may concern, this response to the Kastle Systems RFP is hereby submitted by [Vendor Company Name]:', style='Normal')
doc.add_paragraph('Submitted by:', style='Normal')
doc.add_paragraph('[Vendor Company Name]', style='Normal')
doc.add_paragraph('6100 West Plano Parkway', style='Normal')
doc.add_paragraph('Suite 1800, Plano, Texas 75093', style='Normal')
doc.add_paragraph('Attn: Legal: [Authorized Officer Name]', style='Normal')
doc.add_paragraph('Vendor Contact Information:', style='Normal')
doc.add_paragraph('SAM/CSM Name, Position, Phone No., Email', style='Normal')
doc.add_paragraph('DSE/DCE Name, Position, Phone No., Email', style='Normal')
doc.add_paragraph('Kastle Contact Information:', style='Normal')
doc.add_paragraph('Name:', style='Normal')
doc.add_paragraph('Title:', style='Normal')
doc.add_paragraph('Phone:', style='Normal')
doc.add_paragraph('Email:', style='Normal')
doc.add_page_break()

# Table of Contents
contents = [
    '1 Executive Summary',
    '2 Company Overview',
    '3 Service Capabilities',
    '4 Implementation Plan and Details',
    '5 Proposed Timeline',
    '6 Project Management and Governance',
    '7 Security, Privacy, and Compliance',
    '8 Team and Qualifications',
    '9 Assumptions, Risks, and Dependencies',
    '10 Pricing Proposal',
    '11 Client References',
    '12 Appendices',
]
doc.add_heading('TABLE OF CONTENTS', level=1)
for item in contents:
    doc.add_paragraph(item, style='Normal')
doc.add_page_break()

# Executive Summary

doc.add_heading('Executive Summary', level=1)
para = doc.add_paragraph(
    'Kastle Systems seeks an experienced implementation partner to modernize its ERP and Field Service systems. '
    'Our proposal addresses Kastle’s requirements for end-to-end implementation, integration across business platforms, rigorous training, data migration, and ongoing managed services. '
    'We deliver cost-effective, scalable, and secure solutions leveraging deep Microsoft Dynamics expertise and proven delivery methodology.'
)
para.paragraph_format.space_after = Pt(12)

# Strategic Goals Table
str_goals = [
    ('Single ERP application for consolidated operations and reporting', 'Internal productivity increase'),
    ('Increased inventory and shipping visibility', 'Improved management of supply chain and orders'),
    ('Less IT maintenance and licensing costs', 'IT can free up bandwidth to work on other projects'),
]
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'Strategic Goals'
hdr[1].text = 'Measures of Success'
for c in hdr:
    for run in c.paragraphs[0].runs:
        run.bold = True
for goal, measure in str_goals:
    row = table.add_row().cells
    row[0].text = goal
    row[1].text = measure

para = doc.add_paragraph('Our methodology brings strong discipline to ensure a successful implementation and eventual client ownership of the system. The proposed implementation timeline is outlined below.')
para.paragraph_format.space_after = Pt(12)

# Gantt Chart for Implementation Plan
schedule = pd.DataFrame({
    'Task': [
        'Requirements Analysis', 'Solution Design', 'Development Phase 1', 
        'Development Phase 2', 'Integration Testing', 'User Acceptance Testing',
        'Training & Documentation', 'Deployment & Go-Live', 'Hypercare Support'
    ],
    'Team': [
        'Business Analysts', 'Architects', 'Development', 'Development', 
        'QA Team', 'QA Team', 'Training', 'DevOps', 'Support'
    ],
    'Start': pd.to_datetime([
        '2026-03-01', '2026-03-15', '2026-04-01', '2026-05-01',
        '2026-06-01', '2026-06-15', '2026-06-20', '2026-07-01', '2026-07-08'
    ]),
    'End': pd.to_datetime([
        '2026-03-14', '2026-03-31', '2026-04-30', '2026-05-31',
        '2026-06-14', '2026-06-30', '2026-07-05', '2026-07-07', '2026-07-31'
    ])
})

project_start = schedule['Start'].min()
schedule['days_to_start'] = (schedule['Start'] - project_start).dt.days
schedule['duration'] = (schedule['End'] - schedule['Start']).dt.days + 1

team_colors = {
    'Business Analysts': '#3498db',
    'Architects': '#9b59b6',
    'Development': '#2ecc71',
    'QA Team': '#e74c3c',
    'Training': '#f39c12',
    'DevOps': '#1abc9c',
    'Support': '#34495e'
}

fig, ax = plt.subplots(figsize=(10, 6))
for idx, row in schedule.iterrows():
    ax.barh(y=row['Task'], width=row['duration'], left=row['days_to_start'],
            color=team_colors[row['Team']], edgecolor='white', linewidth=0.5)
ax.invert_yaxis()

max_days = schedule['days_to_start'].max() + schedule['duration'].max()
xtick_vals = np.arange(0, max_days + 7, 14)
xticklabels = pd.date_range(start=project_start, periods=len(xtick_vals), freq='14D').strftime('%b %d')
ax.set_xticks(xtick_vals[:len(xticklabels)])
ax.set_xticklabels(xticklabels, fontsize=9)

ax.set_xlabel('Timeline', fontsize=10)
ax.set_title('Project Implementation Schedule', fontsize=14, fontweight='bold')
ax.xaxis.grid(True, alpha=0.3, linestyle='--')
ax.set_axisbelow(True)

unique_teams = schedule['Team'].unique()
patches = [mpatches.Patch(color=team_colors[team], label=team) for team in unique_teams]
ax.legend(handles=patches, loc='lower right', fontsize=8, framealpha=0.9)

plt.tight_layout()
gantt_path = output_dir / 'project_gantt.png'
plt.savefig(gantt_path, dpi=150, bbox_inches='tight', facecolor='white')
plt.close()
doc.add_picture(str(gantt_path), width=Inches(6.0))
add_caption('Figure 1: Project Implementation Schedule')

doc.add_heading('Company Overview', level=1)
doc.add_paragraph(
    'Our company is a leading provider of ERP implementation services, specializing in Microsoft Dynamics platforms. With over 15 years of experience delivering complex modernization projects, we have a proven track record in optimizing operations for organizations across industries, including healthcare, manufacturing, and logistics.'
)

doc.add_heading('Service Capabilities', level=1)
doc.add_paragraph('We offer full lifecycle ERP services, including:')
doc.add_paragraph('ERP and Field Service system implementation', style='List Bullet')
doc.add_paragraph('Platform integration across business applications', style='List Bullet')
doc.add_paragraph('Data migration and conversion', style='List Bullet')
doc.add_paragraph('Managed services and ongoing support', style='List Bullet')
doc.add_paragraph('User training and adoption', style='List Bullet')
doc.add_paragraph('System maintenance and upgrades', style='List Bullet')
doc.add_paragraph('Reporting and analytics', style='List Bullet')
doc.add_paragraph('Security, compliance, and privacy management', style='List Bullet')

# Requirements Traceability Table
reqs = [
    ('REQ-001', 'Proposal follows RFP structure and requirements', 'Section: All'),
    ('REQ-002', 'Acknowledged receipt and intent to respond', 'Section: Executive Summary'),
    ('REQ-003', 'Participation in discovery sessions', 'Section: Implementation Plan'),
    ('REQ-004', 'Methodology covers all requested areas', 'Section: Service Capabilities'),
    ('REQ-005', 'Support for all major ERP functional areas', 'Section: Service Capabilities'),
    ('REQ-006', 'Transparent cost breakdowns', 'Section: Pricing Proposal'),
    ('REQ-007', 'Relevant customer references provided', 'Section: Client References'),
    ('REQ-009', 'Detailed implementation, testing, training, migration plans', 'Section: Implementation Plan'),
    ('REQ-010', 'Support model, help desk, SLA, scalability', 'Section: Project Management'),
    ('REQ-011', 'Confidentiality and non-disclosure compliance', 'Section: Security & Compliance'),
    ('REQ-014', 'Disclosure of conflicts, legal actions', 'Section: Appendices'),
    ('REQ-016', 'Proposal submission format and deadline', 'Section: Appendices'),
]
doc.add_heading('Requirements Traceability Matrix', level=2)
table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'Requirement ID'
hdr[1].text = 'Description'
hdr[2].text = 'Reference Section'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True
for rid, desc, ref in reqs:
    row = table.add_row().cells
    row[0].text = rid
    row[1].text = desc
    row[2].text = ref

# Mermaid Diagram - Project Workflow
mermaid_code = '''
flowchart TD
  A[Discovery] --> B[Design]
  B --> C[Build]
  C --> D[Test]
  D --> E[Deploy]
  E --> F[Hypercare]
'''
diagram_path = render_mermaid(mermaid_code, 'workflow_diagram')
doc.add_picture(str(diagram_path), width=Inches(5.8))
add_caption('Figure 2: Project workflow from discovery to hypercare')

doc.add_heading('Implementation Plan and Details', level=1)
doc.add_paragraph(
    'Our implementation plan is designed to accelerate delivery while minimizing disruption to business operations. The plan includes phased requirements gathering, solution design, iterative build, comprehensive testing, and robust deployment. Each phase is supported by dedicated teams and clear milestones.'
)
doc.add_paragraph('Key milestones include:')
doc.add_paragraph('Requirements Analysis', style='List Bullet')
doc.add_paragraph('Solution Design', style='List Bullet')
doc.add_paragraph('Development Phase 1', style='List Bullet')
doc.add_paragraph('Development Phase 2', style='List Bullet')
doc.add_paragraph('Integration Testing', style='List Bullet')
doc.add_paragraph('User Acceptance Testing', style='List Bullet')
doc.add_paragraph('Training & Documentation', style='List Bullet')
doc.add_paragraph('Deployment & Go-Live', style='List Bullet')
doc.add_paragraph('Hypercare Support', style='List Bullet')

# Milestone Timeline Chart
milestones = pd.DataFrame({
    'Milestone': ['Kickoff', 'Design Signoff', 'Build Complete', 'UAT Complete', 'Go-Live'],
    'Date': pd.to_datetime(['2026-03-01', '2026-03-20', '2026-05-30', '2026-06-25', '2026-07-10']),
})
fig, ax = plt.subplots(figsize=(9, 3.8))
ax.plot(milestones['Date'], [1] * len(milestones), marker='o', linewidth=1.8)
for idx, row in milestones.iterrows():
    ax.text(row['Date'], 1.03, row['Milestone'], rotation=35, ha='left', va='bottom', fontsize=8)
ax.yaxis.set_visible(False)
ax.set_title('Program Milestone Timeline')
ax.xaxis.set_major_locator(mdates.WeekdayLocator(interval=2))
ax.xaxis.set_major_formatter(mdates.DateFormatter('%b %d'))
ax.grid(axis='x', alpha=0.25, linestyle='--')
plt.tight_layout()
milestone_path = output_dir / 'milestone_timeline.png'
plt.savefig(milestone_path, dpi=150, bbox_inches='tight', facecolor='white')
plt.close()
doc.add_picture(str(milestone_path), width=Inches(6.0))
add_caption('Figure 3: Program milestone timeline')

doc.add_heading('Project Management and Governance', level=1)
doc.add_paragraph(
    'Our project management approach is based on PMI and Agile principles, ensuring clear communication, transparency, and rigorous reporting. Governance is structured with defined roles, escalation paths, and regular steering committee meetings. Risk management and change control are embedded throughout.'
)

doc.add_heading('Security, Privacy, and Compliance', level=1)
doc.add_paragraph(
    'Security is integrated at each stage, leveraging industry standards such as ISO 27001 and SOC 2. Confidentiality, privacy, and regulatory compliance are maintained via technical safeguards, access controls, and clear data management policies. All proposal materials are marked for confidentiality as required.'
)

doc.add_heading('Team and Qualifications', level=1)
doc.add_paragraph(
    'Our team consists of seasoned professionals with deep Microsoft Dynamics expertise, project management certifications, and industry experience. Key members include Solution Architects, Business Analysts, Senior Developers, QA leads, and Training Specialists.'
)
doc.add_heading('Assumptions, Risks, and Dependencies', level=1)
doc.add_paragraph(
    'Key assumptions include timely access to Kastle Systems stakeholders, legacy system documentation, and test environments. Risks are managed through early identification, mitigation planning, and escalation procedures. Dependencies include third-party integrations, data migration readiness, and resource availability.'
)
doc.add_heading('Pricing Proposal', level=1)
doc.add_paragraph(
    'Our pricing is transparent and itemized by implementation, integration, training, data migration, managed services, and ongoing costs. Pricing is subject to final scope and agreed deliverables. Detailed cost breakdowns are provided in the appendix.'
)
doc.add_heading('Client References', level=1)
doc.add_paragraph('References available upon request. Two references in healthcare and three additional relevant projects have been included in the appendix.')
doc.add_heading('Appendices', level=1)
doc.add_paragraph('Appendix A: Master Service Agreement and Service Level Agreement drafts')
doc.add_paragraph('Appendix B: Client References and Contact Information')
doc.add_paragraph('Appendix C: Proposal Submission and Confidentiality Details')
doc.add_paragraph('Appendix D: Conflicts of Interest and Legal Action Disclosures')
