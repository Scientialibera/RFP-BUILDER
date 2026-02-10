from docx.enum.style import WD_STYLE_TYPE
from docx.shared import RGBColor
import matplotlib.patches as mpatches
import datetime as dt

styles = doc.styles

# Configure base styles
normal = styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)

# Caption style helper
if 'Caption' not in [s.name for s in styles]:
    cap = styles.add_style('Caption', WD_STYLE_TYPE.PARAGRAPH)
    cap.font.name = 'Calibri'
    cap.font.size = Pt(9)
    cap.font.italic = True


def add_caption(text: str):
    p = doc.add_paragraph(text, style='Caption')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


# Title Page
title = doc.add_heading('Aurelia Digital Systems\nEnterprise Digital Delivery Capabilities Proposal', level=0)
para = doc.add_paragraph('Internal | Confidential')
para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
para.paragraph_format.space_after = Pt(18)

meta = doc.add_paragraph('Date: January 31, 2026')
meta.paragraph_format.space_after = Pt(6)
meta2 = doc.add_paragraph('Prepared by: Aurelia Digital Systems')
meta2.paragraph_format.space_after = Pt(24)

# Executive Summary and Value Proposition
h1 = doc.add_heading('1 Executive Summary and Value Proposition', level=1)
h1.paragraph_format.space_before = Pt(12)

p = doc.add_paragraph(
    'Aurelia Digital Systems is a specialized partner for secure, Azure-centric digital delivery across '
    'public sector, healthcare, utilities, financial services, and logistics. We design, build, and operate '
    'enterprise software, cloud platforms, and data & AI solutions with security-by-design, formal '
    'governance, and measurable business outcomes at the core.'
)

p = doc.add_paragraph(
    'Our delivery model combines agile iteration with stage gates, design reviews, and security sign-offs '
    'to meet demanding regulatory and operational requirements. We employ an Azure-native technology '
    'stack centered on Python (FastAPI), React/TypeScript, and managed Azure services, supported by '
    'DevSecOps automation, infrastructure as code, and policy-as-code. This ensures that every solution '
    'is auditable, resilient, and ready for continuous improvement.'
)

p = doc.add_paragraph(
    'Across public sector portals, healthcare data platforms, and utilities operations tooling, Aurelia has '
    'demonstrated the ability to improve release cadence from quarterly to bi-weekly, reduce manual '
    'reconciliation effort by approximately 40%, and shorten incident triage times by roughly 30%. '
    'These outcomes are achieved through multidisciplinary teams that include dedicated security and '
    'platform engineers, UX designers, and data specialists working in close partnership with client '
    'stakeholders.'
)

# Table: client strategic objectives to capabilities and outcomes
table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'Client Strategic Objective'
hdr[1].text = 'Aurelia Capability'
hdr[2].text = 'Expected Outcome'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

rows = [
    (
        'Increase release cadence without compromising security',
        'Agile delivery with security-by-design, automated testing, and gated releases',
        'Move from quarterly to bi-weekly or monthly releases with maintained compliance'
    ),
    (
        'Modernize legacy platforms to Azure cloud',
        'Cloud & platform engineering, application modernization, and enterprise integration',
        'Reduced infrastructure risk, improved scalability, and lower total cost of ownership'
    ),
    (
        'Improve data-driven decision making',
        'Data integration, governed data platforms, and AI-assisted workflows with human oversight',
        'Fresher, trusted data and faster insight generation for operations and leadership'
    ),
    (
        'Strengthen operational resilience and incident response',
        'Observability, reliability engineering, and incident response playbooks',
        'Reduced mean time to detect and resolve incidents, with clear accountability'
    ),
]

for obj, capab, outcome in rows:
    row_cells = table.add_row().cells
    row_cells[0].text = obj
    row_cells[1].text = capab
    row_cells[2].text = outcome

# Chart: before/after outcomes from representative case studies
sns.set_style('whitegrid')
case_data = pd.DataFrame({
    'Metric': [
        'Release cadence (weeks)',
        'Manual reconciliation effort (hours per week)',
        'Incident triage time (minutes)'
    ],
    'Before': [12, 50, 60],
    'After': [2, 30, 42]
})

fig, ax = plt.subplots(figsize=(8, 4.5))
index = np.arange(len(case_data))
width = 0.35
bars_before = ax.bar(index - width/2, case_data['Before'], width, label='Before')
bars_after = ax.bar(index + width/2, case_data['After'], width, label='After')
ax.set_xticks(index)
ax.set_xticklabels(case_data['Metric'], rotation=20, ha='right')
ax.set_ylabel('Value')
ax.set_title('Representative Outcome Improvements')
ax.legend()
ax.bar_label(bars_before, padding=3, fontsize=8)
ax.bar_label(bars_after, padding=3, fontsize=8)
plt.tight_layout()
chart_path = output_dir / 'exec_summary_outcomes.png'
plt.savefig(chart_path, dpi=150, bbox_inches='tight', facecolor='white')
plt.close(fig)

doc.add_picture(str(chart_path), width=Inches(5.8))
add_caption('Figure 1: Illustrative before-and-after improvements from representative Aurelia engagements')

# Mermaid diagram: engagement lifecycle from Initiate through Operate & Improve
mermaid_code_lifecycle = '''
flowchart LR
  A[Initiate] --> B[Discover]
  B --> C[Design]
  C --> D[Build]
  D --> E[Deploy]
  E --> F[Operate and Improve]
'''

diagram_path = render_mermaid(mermaid_code_lifecycle, 'engagement_lifecycle')
doc.add_picture(str(diagram_path), width=Inches(5.8))
add_caption('Figure 2: High-level Aurelia engagement lifecycle')

# Company Background and Industry Focus
h1 = doc.add_heading('2 Company Background and Industry Focus', level=1)
h1.paragraph_format.space_before = Pt(12)

p = doc.add_paragraph(
    'Founded in 2016 and headquartered in Toronto, Ontario, Aurelia Digital Systems is a '
    '120+ person digital engineering firm focused on enterprise software delivery, cloud '
    'modernization, data & AI, and managed services. Our delivery teams are distributed '
    'across Canada and the United States, enabling follow-the-sun collaboration while '
    'maintaining strong local presence for key clients.'
)

p = doc.add_paragraph(
    'Aurelia operates primarily in regulated and mission-critical environments, including '
    'provincial and municipal government, healthcare providers, utilities, financial services '
    'institutions, and logistics operators. This focus has shaped our security practices, '
    'governance model, and technology choices to align with stringent compliance and '
    'availability requirements.'
)

# Company snapshot table
snapshot = doc.add_table(rows=5, cols=2)
snapshot.style = 'Table Grid'
snapshot.alignment = WD_TABLE_ALIGNMENT.CENTER
rows_snap = snapshot.rows
rows_snap[0].cells[0].text = 'Founded'
rows_snap[0].cells[1].text = '2016'
rows_snap[1].cells[0].text = 'Headquarters'
rows_snap[1].cells[1].text = 'Toronto, Ontario (delivery teams across Canada and the U.S.)'
rows_snap[2].cells[0].text = 'Employees'
rows_snap[2].cells[1].text = '120+ (engineering, delivery, security, UX, and program management)'
rows_snap[3].cells[0].text = 'Primary Focus'
rows_snap[3].cells[1].text = 'Enterprise software delivery, cloud modernization, data & AI, managed services'
rows_snap[4].cells[0].text = 'Industries'
rows_snap[4].cells[1].text = 'Public sector, healthcare, utilities, financial services, logistics'
for row in rows_snap:
    for run in row.cells[0].paragraphs[0].runs:
        run.bold = True

# Industry experience table
industry_table = doc.add_table(rows=1, cols=3)
industry_table.style = 'Table Grid'
industry_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = industry_table.rows[0].cells
hdr[0].text = 'Industry'
hdr[1].text = 'Representative Solutions'
hdr[2].text = 'Example Outcomes'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

industry_rows = [
    ('Public sector', 'Citizen service portals, licensing systems, case management', 'Bi-weekly releases, improved digital adoption'),
    ('Healthcare', 'Data integration hubs, interoperability APIs, clinician portals', '40% reduction in manual data reconciliation'),
    ('Utilities', 'Operations dashboards, field workforce tools, event-driven alerts', '30% faster incident triage and resolution'),
    ('Financial services', 'Customer onboarding, risk reporting, secure document exchange', 'Improved compliance reporting and cycle time'),
    ('Logistics', 'Shipment tracking, capacity planning, partner integration', 'Higher on-time performance and visibility'),
]

for ind, sol, outc in industry_rows:
    row_cells = industry_table.add_row().cells
    row_cells[0].text = ind
    row_cells[1].text = sol
    row_cells[2].text = outc

# Core Service Offerings and Scope Alignment
h1 = doc.add_heading('3 Core Service Offerings and Scope Alignment', level=1)
h1.paragraph_format.space_before = Pt(12)

p = doc.add_paragraph(
    'Aurelia offers a set of integrated service lines that align directly to typical enterprise and '
    'public-sector scopes of work: cloud and platform engineering, application modernization, '
    'data, analytics and AI, security and compliance engineering, enterprise integration, and '
    'UX and service design. These capabilities can be engaged individually or as a cohesive '
    'program to deliver end-to-end solutions.'
)

# Capability-to-requirement mapping table
cap_table = doc.add_table(rows=1, cols=3)
cap_table.style = 'Table Grid'
cap_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = cap_table.rows[0].cells
hdr[0].text = 'Service Offering'
hdr[1].text = 'Description'
hdr[2].text = 'Key Requirements Addressed'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

cap_rows = [
    ('Cloud and Platform Engineering',
     'Design and implementation of Azure landing zones, infrastructure as code, container platforms, and observability.',
     'REQ-001, REQ-003, REQ-011, REQ-018, REQ-019'),
    ('Application Modernization',
     'Domain-driven design, API-first microservices, legacy system decomposition, and performance tuning.',
     'REQ-001, REQ-004, REQ-007, REQ-017, REQ-018'),
    ('Data, Analytics and AI',
     'Data integration, governed storage, semantic modeling, and AI-assisted workflows with human oversight.',
     'REQ-001, REQ-005, REQ-019, REQ-020'),
    ('Security and Compliance Engineering',
     'Threat modeling, secure SDLC, identity and PKI integration, continuous security validation.',
     'REQ-002, REQ-006, REQ-011, REQ-012, REQ-013, REQ-014, REQ-015'),
    ('Enterprise Integration',
     'Event-driven architecture, middleware integration, and API management patterns.',
     'REQ-001, REQ-007, REQ-017, REQ-018'),
    ('UX and Service Design',
     'User research, prototyping, accessibility, and enterprise-grade design systems.',
     'REQ-001, REQ-008, REQ-020'),
]

for name, desc, reqs in cap_rows:
    row_cells = cap_table.add_row().cells
    row_cells[0].text = name
    row_cells[1].text = desc
    row_cells[2].text = reqs

# Solution Delivery Model and Governance
h1 = doc.add_heading('4 Solution Delivery Model and Governance', level=1)
h1.paragraph_format.space_before = Pt(12)

p = doc.add_paragraph(
    'Aurelia follows an agile delivery model with formal governance. Work is organized into '
    'sprints, with clear stage gates at the end of each major phase: Initiate, Discover, Design, '
    'Build, Deploy, and Operate and Improve. Each phase has defined activities, deliverables, '
    'and acceptance criteria, with security, privacy, and compliance embedded from the outset.'
)

p = doc.add_paragraph(
    'Governance is implemented through recurring ceremonies (standups, backlog refinement, '
    'demos, retrospectives), steering committees, design reviews, security checkpoints, and '
    'release readiness assessments. Progress is tracked using burn-up and burn-down charts, '
    'risk registers, and decision logs, providing transparent reporting to stakeholders.'
)

# Phase-by-phase activities table
phase_table = doc.add_table(rows=1, cols=3)
phase_table.style = 'Table Grid'
phase_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = phase_table.rows[0].cells
hdr[0].text = 'Phase'
hdr[1].text = 'Key Activities'
hdr[2].text = 'Primary Outputs'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

phase_rows = [
    ('Initiate', 'Confirm scope, stakeholders, operating model, and success metrics.', 'Charter, RACI, initial roadmap'),
    ('Discover', 'Requirements elaboration, current-state assessment, risk and threat modeling.', 'Requirements backlog, risk register, threat model'),
    ('Design', 'Target architecture, data flows, integration patterns, UX concepts.', 'Architecture diagrams, design decisions, prioritized backlog'),
    ('Build', 'Incremental delivery with automated testing and security checks.', 'Working software, test results, updated documentation'),
    ('Deploy', 'Release readiness, cutover planning, operational handoff.', 'Deployment plan, runbooks, go-live report'),
    ('Operate and Improve', 'Monitoring, incident response, continuous improvement.', 'Operational dashboards, improvement backlog'),
]

for ph, acts, outs in phase_rows:
    row_cells = phase_table.add_row().cells
    row_cells[0].text = ph
    row_cells[1].text = acts
    row_cells[2].text = outs

# Gantt-style schedule for phases
schedule = pd.DataFrame({
    'Task': ['Initiate', 'Discover', 'Design', 'Build', 'Deploy', 'Operate and Improve'],
    'Team': ['Delivery', 'Delivery', 'Architecture', 'Engineering', 'Engineering', 'Operations'],
    'Start': pd.to_datetime(['2026-03-01', '2026-03-08', '2026-03-22', '2026-04-05', '2026-06-01', '2026-06-15']),
    'End': pd.to_datetime(['2026-03-07', '2026-03-21', '2026-04-04', '2026-05-31', '2026-06-14', '2026-12-31'])
})
project_start = schedule['Start'].min()
schedule['days_to_start'] = (schedule['Start'] - project_start).dt.days
schedule['duration'] = (schedule['End'] - schedule['Start']).dt.days + 1
team_colors = {
    'Delivery': '#1f77b4',
    'Architecture': '#9467bd',
    'Engineering': '#2ca02c',
    'Operations': '#7f7f7f'
}
fig, ax = plt.subplots(figsize=(9, 4.8))
for idx, row in schedule.iterrows():
    ax.barh(y=row['Task'], width=row['duration'], left=row['days_to_start'],
            color=team_colors[row['Team']], edgecolor='white', linewidth=0.5)
ax.invert_yaxis()
max_days = schedule['days_to_start'].max() + schedule['duration'].max()
xticks = np.arange(0, max_days + 7, 14)
xticklabels = pd.date_range(start=project_start, periods=len(xticks), freq='14D').strftime('%b %d')
ax.set_xticks(xticks[:len(xticklabels)])
ax.set_xticklabels(xticklabels, fontsize=8)
ax.set_xlabel('Timeline')
ax.set_title('Illustrative Project Implementation Schedule')
ax.xaxis.grid(True, alpha=0.3, linestyle='--')
ax.set_axisbelow(True)
unique_teams = schedule['Team'].unique()
patches = [mpatches.Patch(color=team_colors[t], label=t) for t in unique_teams]
ax.legend(handles=patches, loc='lower right', fontsize=8, framealpha=0.9)
plt.tight_layout()
gantt_path = output_dir / 'implementation_schedule.png'
plt.savefig(gantt_path, dpi=150, bbox_inches='tight', facecolor='white')
plt.close(fig)

doc.add_picture(str(gantt_path), width=Inches(6.0))
add_caption('Figure 3: Illustrative implementation schedule aligned to Aurelia delivery phases')

# Technical Architecture and Azure-Centric Technology Stack
h1 = doc.add_heading('5 Technical Architecture and Azure-Centric Technology Stack', level=1)
h1.paragraph_format.space_before = Pt(12)

p = doc.add_paragraph(
    'Aurelia solutions are built on an Azure-native reference architecture that prioritizes '
    'security, scalability, and maintainability. Frontend applications are implemented with '
    'React and TypeScript, communicating through secure APIs to Python FastAPI services '
    'running on Azure Kubernetes Service or Azure App Service. Data is persisted in '
    'PostgreSQL and Azure storage, with messaging and integration handled through Azure '
    'service bus and event services.'
)

# Technology stack table
stack_table = doc.add_table(rows=1, cols=3)
stack_table.style = 'Table Grid'
stack_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = stack_table.rows[0].cells
hdr[0].text = 'Layer'
hdr[1].text = 'Technologies'
hdr[2].text = 'Role'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

stack_rows = [
    ('Frontend', 'React, TypeScript, component libraries', 'Responsive user interfaces and accessible client experiences.'),
    ('API and services', 'Python, FastAPI, OpenAPI, REST', 'Business logic, orchestration, and integration endpoints.'),
    ('Cloud platform', 'Azure AKS, App Service, Functions', 'Scalable, managed compute for stateless and stateful workloads.'),
    ('Data', 'PostgreSQL, Azure storage, ETL and ELT tooling', 'Transactional data, analytical stores, and governed data pipelines.'),
    ('Integration and messaging', 'Azure Service Bus, Event Grid', 'Reliable asynchronous messaging and event-driven workflows.'),
    ('Identity and access', 'Microsoft Entra ID, Azure AD B2C', 'Single sign-on, conditional access, and role-based access control.'),
    ('Observability and DevSecOps', 'GitHub Actions or Azure DevOps, Terraform or Bicep, monitoring tooling', 'Continuous delivery, policy enforcement, telemetry, and alerting.'),
]

for layer, tech, role in stack_rows:
    row_cells = stack_table.add_row().cells
    row_cells[0].text = layer
    row_cells[1].text = tech
    row_cells[2].text = role

# Mermaid architecture diagram
mermaid_arch = '''
flowchart LR
  U[User] --> B[Browser]
  B --> F[React app]
  F --> G[API gateway]
  G --> S[FastAPI services]
  S --> D[PostgreSQL]
  S --> O[Object storage]
  S --> M[Messaging]
  M --> X[Integration services]
  S --> L[Logging and metrics]
'''
arch_path = render_mermaid(mermaid_arch, 'azure_reference_architecture')
doc.add_picture(str(arch_path), width=Inches(5.8))
add_caption('Figure 4: Typical Aurelia Azure-native reference architecture')

# Cloud and Platform Engineering Approach
h1 = doc.add_heading('6 Cloud and Platform Engineering Approach', level=1)
h1.paragraph_format.space_before = Pt(12)

p = doc.add_paragraph(
    'Aurelia designs and operates Azure landing zones that standardize identity, networking, '
    'governance, and observability across environments. All platform components are defined '
    'using infrastructure as code, typically Terraform or Bicep, and validated through '
    'policy-as-code checks before deployment. Container platforms such as Azure Kubernetes '
    'Service are configured with hardened baselines, workload identity, and integrated logging.'
)

# Platform components and practices table
plat_table = doc.add_table(rows=1, cols=3)
plat_table.style = 'Table Grid'
plat_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = plat_table.rows[0].cells
hdr[0].text = 'Platform Component'
hdr[1].text = 'Engineering Practices'
hdr[2].text = 'Outcomes'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

plat_rows = [
    ('Network and connectivity', 'Hub and spoke topology, private endpoints, firewall rules as code.', 'Segmentation, reduced attack surface, predictable connectivity.'),
    ('Compute and containers', 'AKS clusters, node pools, autoscaling, hardened base images.', 'Scalable and resilient workloads with consistent configuration.'),
    ('Identity and access', 'Entra ID integration, role-based access, just-in-time elevation.', 'Least privilege access and traceable administrative actions.'),
    ('Storage and databases', 'Encrypted storage, backup and restore policies, performance tuning.', 'Durable, performant data platforms aligned to recovery objectives.'),
    ('Monitoring and logging', 'Centralized logs, metrics, dashboards, and alert rules.', 'Faster detection of issues and evidence for audits and post-incident reviews.'),
]

for comp, prac, outc in plat_rows:
    row_cells = plat_table.add_row().cells
    row_cells[0].text = comp
    row_cells[1].text = prac
    row_cells[2].text = outc

# Application Modernization and Integration Strategy
h1 = doc.add_heading('7 Application Modernization and Integration Strategy', level=1)
h1.paragraph_format.space_before = Pt(12)

p = doc.add_paragraph(
    'Aurelia modernizes legacy applications using domain-driven design and API-first '
    'microservices. We identify bounded contexts, extract services incrementally, and '
    'establish clear contracts through REST and event-driven interfaces. Legacy systems are '
    'wrapped with stable APIs and integrated via synchronous and asynchronous patterns '
    'to minimize disruption.'
)

# Legacy vs modern architecture comparison table
mod_table = doc.add_table(rows=1, cols=4)
mod_table.style = 'Table Grid'
mod_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = mod_table.rows[0].cells
hdr[0].text = 'Characteristic'
hdr[1].text = 'Legacy Monolith'
hdr[2].text = 'Modern Microservices'
hdr[3].text = 'Benefit'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

mod_rows = [
    ('Scalability', 'Vertical scaling, limited elasticity.', 'Horizontal scaling by service.', 'Scale only the components that require additional capacity.'),
    ('Resilience', 'Single failure can impact entire system.', 'Isolated failures within services.', 'Improved fault isolation and graceful degradation.'),
    ('Deployment', 'Infrequent, high-risk releases.', 'Frequent, incremental deployments.', 'Reduced change risk and faster time-to-market.'),
    ('Integration', 'Point-to-point, tightly coupled.', 'API-first and event-driven.', 'Simpler integration and extensibility.'),
]

for charac, legacy, modern, benefit in mod_rows:
    row_cells = mod_table.add_row().cells
    row_cells[0].text = charac
    row_cells[1].text = legacy
    row_cells[2].text = modern
    row_cells[3].text = benefit

# Data, Analytics, and AI Capabilities
h1 = doc.add_heading('8 Data, Analytics, and AI Capabilities', level=1)
h1.paragraph_format.space_before = Pt(12)

p = doc.add_paragraph(
    'Aurelia designs governed data platforms that integrate data from operational systems '
    'into PostgreSQL and object storage, using robust ETL and ELT patterns. Semantic '
    'models expose trusted data to analytics tools, while access controls, lineage, and data '
    'quality checks maintain trustworthiness. AI-assisted workflows are implemented with '
    'strong human oversight, ensuring that recommendations are reviewed and approved '
    'before actions are taken in regulated contexts.'
)

# Data platform components table
data_table = doc.add_table(rows=1, cols=3)
data_table.style = 'Table Grid'
data_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = data_table.rows[0].cells
hdr[0].text = 'Component'
hdr[1].text = 'Responsibilities'
hdr[2].text = 'Technologies'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

data_rows = [
    ('Ingestion', 'Capture data from source systems via APIs, events, and batch loads.', 'Integration services, messaging, data movement tooling.'),
    ('Storage', 'Persist raw and curated data with encryption and retention policies.', 'PostgreSQL, object storage, Azure databases.'),
    ('Governance', 'Define ownership, quality rules, and lineage; manage access.', 'Catalogs, policy engines, role-based access control.'),
    ('Semantic layer', 'Provide business-friendly views and metrics.', 'Modeling tools, views, and APIs.'),
    ('Analytics and AI', 'Deliver dashboards, reports, and AI-assisted workflows.', 'BI platforms, machine learning services, orchestration.'),
]

for comp, resp, tech in data_rows:
    row_cells = data_table.add_row().cells
    row_cells[0].text = comp
    row_cells[1].text = resp
    row_cells[2].text = tech

# Security, Privacy, and Compliance-by-Design
h1 = doc.add_heading('9 Security, Privacy, and Compliance-by-Design', level=1)
h1.paragraph_format.space_before = Pt(12)

p = doc.add_paragraph(
    'Security, privacy, and compliance are embedded throughout Aurelia engagements. '
    'We apply defense-in-depth, strong identity controls, encryption by default, and secure '
    'SDLC practices from the earliest phases. Controls are implemented as code wherever '
    'possible, and CI or CD pipelines generate audit-ready evidence of security checks, '
    'test results, and approvals.'
)

p = doc.add_paragraph(
    'Threat modeling is performed during Discover and revisited as the design evolves. '
    'Identity integration with Microsoft Entra ID enables single sign-on, conditional access, '
    'and privileged access workflows. All network traffic is protected with TLS, and data at '
    'rest uses managed keys stored in Azure Key Vault or hardware-backed modules where '
    'required.'
)

# Security control matrix table
sec_table = doc.add_table(rows=1, cols=4)
sec_table.style = 'Table Grid'
sec_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = sec_table.rows[0].cells
hdr[0].text = 'Control Area'
hdr[1].text = 'Practices'
hdr[2].text = 'Tooling and Techniques'
hdr[3].text = 'Requirements Addressed'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

sec_rows = [
    ('Identity and access', 'SSO, conditional access, role-based access control, privileged access workflows.', 'Entra ID, privileged identity, access reviews.', 'REQ-006, REQ-011, REQ-013'),
    ('Network security', 'Segmentation, private endpoints, firewall policies, zero trust principles.', 'Virtual networks, firewalls, network security groups.', 'REQ-006, REQ-012'),
    ('Encryption', 'TLS for data in transit, managed keys for data at rest.', 'Certificates, Azure Key Vault, HSM-backed key storage.', 'REQ-006, REQ-014'),
    ('Secure SDLC', 'Dependency scanning, SAST, DAST, secrets detection, peer review.', 'Security scanners, code review workflows, pipeline gates.', 'REQ-006, REQ-015, REQ-018'),
    ('Monitoring and response', 'Centralized logging, alerting, vulnerability management, incident playbooks.', 'Monitoring platforms, SIEM, runbooks.', 'REQ-011, REQ-016'),
]

for area, prac, tool, reqs in sec_rows:
    row_cells = sec_table.add_row().cells
    row_cells[0].text = area
    row_cells[1].text = prac
    row_cells[2].text = tool
    row_cells[3].text = reqs

# Mermaid lifecycle security diagram
mermaid_sec = '''
flowchart LR
  I[Initiate] --> D[Discover]
  D --> DS[Design]
  DS --> B[Build]
  B --> DP[Deploy]
  DP --> O[Operate and Improve]
  I --> SI[Security objectives]
  D --> TM[Threat modeling]
  B --> SC[Security checks]
  O --> IR[Incident response]
'''
sec_path = render_mermaid(mermaid_sec, 'security_lifecycle')
doc.add_picture(str(sec_path), width=Inches(5.8))
add_caption('Figure 5: Security activities mapped to Aurelia delivery phases')

# DevSecOps, Automation, and Quality Engineering
h1 = doc.add_heading('10 DevSecOps, Automation, and Quality Engineering', level=1)
h1.paragraph_format.space_before = Pt(12)

p = doc.add_paragraph(
    'Aurelia employs DevSecOps practices to automate build, test, security, and deployment '
    'activities. Pipelines in GitHub Actions or Azure DevOps orchestrate unit, integration, and '
    'end-to-end tests; static and dynamic security scans; dependency checks; and deployment '
    'to Azure environments defined by infrastructure as code. Policy-as-code ensures that '
    'only compliant configurations progress to higher environments.'
)

# Test strategy table
test_table = doc.add_table(rows=1, cols=3)
test_table.style = 'Table Grid'
test_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = test_table.rows[0].cells
hdr[0].text = 'Test Type'
hdr[1].text = 'Purpose'
hdr[2].text = 'Typical Tooling'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

test_rows = [
    ('Unit tests', 'Validate individual components and functions.', 'Language-specific test frameworks, mocking tools.'),
    ('Integration tests', 'Verify interactions between services and external systems.', 'API test frameworks, integration test harnesses.'),
    ('End-to-end tests', 'Exercise user journeys across the full stack.', 'Browser automation, synthetic monitoring.'),
    ('Performance tests', 'Assess throughput and latency under load.', 'Load testing tools, profiling.'),
    ('Security tests', 'Identify vulnerabilities in code and configuration.', 'SAST, DAST, dependency scanning, configuration checks.'),
]

for ttype, purpose, tools in test_rows:
    row_cells = test_table.add_row().cells
    row_cells[0].text = ttype
    row_cells[1].text = purpose
    row_cells[2].text = tools

# User Experience and Service Design
h1 = doc.add_heading('11 User Experience and Service Design', level=1)
h1.paragraph_format.space_before = Pt(12)

p = doc.add_paragraph(
    'User experience and service design are core capabilities at Aurelia. UX practitioners '
    'conduct user and stakeholder research, develop personas and journey maps, and '
    'prototype key flows early in the engagement. Design systems aligned to accessibility '
    'standards ensure consistency across products, while usability testing provides feedback '
    'that is integrated into each sprint.'
)

# UX deliverables table
ux_table = doc.add_table(rows=1, cols=3)
ux_table.style = 'Table Grid'
ux_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = ux_table.rows[0].cells
hdr[0].text = 'Deliverable'
hdr[1].text = 'Description'
hdr[2].text = 'Usage in Delivery'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

ux_rows = [
    ('Personas', 'Evidence-based profiles representing key user groups.', 'Guide prioritization and design decisions.'),
    ('Journey maps', 'Visualizations of end-to-end user journeys and pain points.', 'Identify opportunities for improvement and automation.'),
    ('Wireframes and prototypes', 'Low and high fidelity representations of screens and flows.', 'Enable rapid validation with users and stakeholders.'),
    ('Design system', 'Reusable components, patterns, and style guidelines.', 'Ensure consistency, accessibility, and faster UI development.'),
    ('Accessibility audits', 'Assess compliance with standards and assistive technologies.', 'Drive remediation and inclusive design practices.'),
]

for deliv, desc, use in ux_rows:
    row_cells = ux_table.add_row().cells
    row_cells[0].text = deliv
    row_cells[1].text = desc
    row_cells[2].text = use

# Operational Readiness, Monitoring, and Support
h1 = doc.add_heading('12 Operational Readiness, Monitoring, and Support', level=1)
h1.paragraph_format.space_before = Pt(12)

p = doc.add_paragraph(
    'Operational readiness is planned from the earliest phases. Aurelia defines monitoring '
    'and alerting requirements alongside functional requirements, ensuring that systems are '
    'observable on day one of production. Runbooks and incident response playbooks are '
    'co-developed with client teams, and service level objectives are defined for availability, '
    'performance, and response times.'
)

# Operations runbook summary table
ops_table = doc.add_table(rows=1, cols=3)
ops_table.style = 'Table Grid'
ops_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = ops_table.rows[0].cells
hdr[0].text = 'Playbook'
hdr[1].text = 'Trigger'
hdr[2].text = 'Owner'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

ops_rows = [
    ('Incident response', 'High-severity alert or service outage.', 'Operations lead and on-call engineers.'),
    ('Security incident', 'Confirmed or suspected security event.', 'Security engineer and incident commander.'),
    ('Performance tuning', 'Sustained performance degradation or capacity limits.', 'Platform and application engineers.'),
    ('Change management', 'Planned release or configuration change.', 'Release manager and service owners.'),
]

for play, trig, owner in ops_rows:
    row_cells = ops_table.add_row().cells
    row_cells[0].text = play
    row_cells[1].text = trig
    row_cells[2].text = owner

# Representative Experience and Case Studies
h1 = doc.add_heading('13 Representative Experience and Case Studies', level=1)
h1.paragraph_format.space_before = Pt(12)

p = doc.add_paragraph(
    'Aurelia has delivered measurable improvements for clients in public sector, healthcare, '
    'and utilities. The following examples summarize representative engagements and the '
    'outcomes achieved.'
)

# Case study summary table
cs_table = doc.add_table(rows=1, cols=5)
cs_table.style = 'Table Grid'
cs_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = cs_table.rows[0].cells
hdr[0].text = 'Industry'
hdr[1].text = 'Scope'
hdr[2].text = 'Key Technologies'
hdr[3].text = 'Security Posture'
hdr[4].text = 'Quantified Outcomes'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

cs_rows = [
    (
        'Public sector',
        'Service portal modernization and migration to Azure.',
        'React, FastAPI, Azure AKS, Entra ID.',
        'Defense-in-depth, SSO, encrypted data, secure SDLC.',
        'Release cadence improved from quarterly to bi-weekly.'
    ),
    (
        'Healthcare',
        'Data integration program with governed pipelines and APIs.',
        'ETL and ELT tooling, PostgreSQL, object storage, APIs.',
        'Strict access controls, audit logging, encryption, data governance.',
        'Approximately 40 percent reduction in manual reconciliation effort.'
    ),
    (
        'Utilities',
        'Operations dashboards and alerting with audited access.',
        'Event-driven integration, dashboards, observability.',
        'Role-based access, logging, incident playbooks.',
        'Incident triage time reduced by roughly 30 percent.'
    ),
]

for ind, scope, tech, sec, outc in cs_rows:
    row_cells = cs_table.add_row().cells
    row_cells[0].text = ind
    row_cells[1].text = scope
    row_cells[2].text = tech
    row_cells[3].text = sec
    row_cells[4].text = outc

# Delivery Team Structure and Roles
h1 = doc.add_heading('14 Delivery Team Structure and Roles', level=1)
h1.paragraph_format.space_before = Pt(12)

p = doc.add_paragraph(
    'Aurelia fields multidisciplinary teams tailored to the scope and phase of each '
    'engagement. A typical team includes an engagement lead or program manager, solution '
    'architect, security engineer, backend and frontend engineers, a QA or automation '
    'engineer, and optional UX and DevOps or platform engineers. Teams can scale up or '
    'down based on workload and timelines while maintaining clear accountability.'
)

# Team roles and responsibilities table
team_table = doc.add_table(rows=1, cols=3)
team_table.style = 'Table Grid'
team_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = team_table.rows[0].cells
hdr[0].text = 'Role'
hdr[1].text = 'Primary Responsibilities'
hdr[2].text = 'Key Interactions'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

team_rows = [
    ('Engagement lead or program manager', 'Overall delivery governance, stakeholder management, reporting.', 'Client sponsors, product owners, Aurelia leadership.'),
    ('Solution architect', 'End-to-end architecture, technology decisions, non-functional requirements.', 'Engineers, security, platform, client architects.'),
    ('Security engineer', 'Threat modeling, security controls, secure SDLC, incident response planning.', 'Developers, platform engineers, client security teams.'),
    ('Backend engineers', 'Implement APIs, services, and integration logic.', 'Frontend engineers, QA, architects.'),
    ('Frontend engineers', 'Build user interfaces and client-side logic.', 'UX designers, backend engineers, QA.'),
    ('QA or automation engineer', 'Test strategy, automated test suites, quality reporting.', 'Developers, product owners, operations.'),
    ('UX designer', 'Research, design, and usability testing.', 'Product owners, frontend engineers, end users.'),
    ('DevOps or platform engineer', 'CI or CD pipelines, infrastructure as code, observability.', 'Developers, security, operations.'),
]

for role, resp, inter in team_rows:
    row_cells = team_table.add_row().cells
    row_cells[0].text = role
    row_cells[1].text = resp
    row_cells[2].text = inter

# Reporting, Metrics, and Stakeholder Communication
h1 = doc.add_heading('15 Reporting, Metrics, and Stakeholder Communication', level=1)
h1.paragraph_format.space_before = Pt(12)

p = doc.add_paragraph(
    'Aurelia provides transparent reporting aligned to client governance expectations. '
    'Standard reporting includes sprint and release status, burn-up and burn-down charts, '
    'risk and issue registers, decision logs, and metrics on quality, security, and operational '
    'health. Reporting cadences are agreed during Initiate and reviewed regularly.'
)

# Sample status report table
status_table = doc.add_table(rows=1, cols=4)
status_table.style = 'Table Grid'
status_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = status_table.rows[0].cells
hdr[0].text = 'Area'
hdr[1].text = 'Status'
hdr[2].text = 'Summary'
hdr[3].text = 'Key Actions or Decisions'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

status_rows = [
    ('Scope', 'On track', 'Scope remains aligned to agreed backlog.', 'Review any new requests in change control.'),
    ('Schedule', 'At risk', 'One integration dependency is delayed.', 'Escalate to steering committee and adjust plan.'),
    ('Risks', 'Managed', 'Top risks have mitigation owners and dates.', 'Update risk register and communicate changes.'),
    ('Issues', 'Open', 'Two production defects under investigation.', 'Track through incident management and post-incident review.'),
]

for area, stat, summ, act in status_rows:
    row_cells = status_table.add_row().cells
    row_cells[0].text = area
    row_cells[1].text = stat
    row_cells[2].text = summ
    row_cells[3].text = act

# Business Development and Commercial Contact Information
h1 = doc.add_heading('16 Business Development and Commercial Contact Information', level=1)
h1.paragraph_format.space_before = Pt(12)

p = doc.add_paragraph(
    'For business development inquiries, proposal discussions, or commercial matters, '
    'please contact the Aurelia Digital Systems business development office. We provide '
    'timely responses to new requests and governance escalations, typically acknowledging '
    'within one business day and providing a substantive response or plan within three '
    'to five business days depending on complexity.'
)

# Contact information table
contact_table = doc.add_table(rows=1, cols=3)
contact_table.style = 'Table Grid'
contact_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = contact_table.rows[0].cells
hdr[0].text = 'Contact'
hdr[1].text = 'Details'
hdr[2].text = 'Role'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

contact_rows = [
    ('Business Development Office', 'Email: proposals@aurelia-digital.example\nPhone: +1 (416) 555-0199', 'Primary point of contact for proposals and commercial discussions.'),
    ('Engagement Lead (assigned per project)', 'Provided upon project initiation.', 'Day-to-day contact for delivery and governance matters.'),
]

for name, details, role in contact_rows:
    row_cells = contact_table.add_row().cells
    row_cells[0].text = name
    row_cells[1].text = details
    row_cells[2].text = role

# Mermaid escalation path diagram
mermaid_escalation = '''
flowchart TD
  A[Project team] --> B[Engagement lead]
  B --> C[Business development]
  C --> D[Executive sponsor]
'''
esc_path = render_mermaid(mermaid_escalation, 'escalation_path')
doc.add_picture(str(esc_path), width=Inches(5.0))
add_caption('Figure 6: Escalation path from project team to executive sponsors')

# Appendices: Requirements Traceability Matrix
h1 = doc.add_heading('Appendix A: Requirements Traceability', level=1)
h1.paragraph_format.space_before = Pt(12)

rtm = doc.add_table(rows=1, cols=3)
rtm.style = 'Table Grid'
rtm.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = rtm.rows[0].cells
hdr[0].text = 'Requirement ID'
hdr[1].text = 'Description'
hdr[2].text = 'Section(s) Addressed'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

requirements = [
    ('REQ-001', 'Enterprise software, cloud modernization, data and AI, and managed services across target industries.', '1, 2, 3, 5, 6, 7, 8, 11, 12, 13, 14'),
    ('REQ-002', 'Agile delivery with formal governance, security-by-design, and measurable outcomes.', '1, 4, 9, 10, 15'),
    ('REQ-003', 'Cloud and platform engineering including Azure landing zones, IaC, AKS, observability, and reliability.', '3, 5, 6'),
    ('REQ-004', 'Application modernization using domain-driven design and API-first microservices.', '3, 7'),
    ('REQ-005', 'Data, analytics, and AI capabilities with governance and human oversight.', '3, 8'),
    ('REQ-006', 'Security and compliance engineering including threat modeling and continuous validation.', '3, 4, 6, 9, 10, 12'),
    ('REQ-007', 'Enterprise integration with event-driven architecture and API management.', '3, 5, 7'),
    ('REQ-008', 'UX and service design including research, prototyping, and accessibility.', '3, 11'),
    ('REQ-009', 'Solution delivery approach with stage gates, design reviews, and security sign-offs.', '1, 4, 9, 15'),
    ('REQ-010', 'Defined delivery phases from Initiate through Operate and Improve.', '1, 4, 9, 12'),
    ('REQ-011', 'Embedded security, privacy, and compliance with controls as code and audit evidence.', '3, 4, 6, 9, 10, 12'),
    ('REQ-012', 'Defense-in-depth security practices.', '3, 6, 9'),
    ('REQ-013', 'Strong identity controls including SSO and conditional access.', '3, 5, 6, 9'),
    ('REQ-014', 'Encryption by default with TLS and managed keys.', '3, 5, 6, 9'),
    ('REQ-015', 'Secure SDLC practices including scanning and peer review.', '3, 4, 6, 9, 10'),
    ('REQ-016', 'Operational readiness with monitoring, alerting, and incident response.', '1, 4, 6, 9, 10, 12, 15'),
    ('REQ-017', 'Python FastAPI backend, React and TypeScript frontend, Azure services.', '1, 3, 5, 7'),
    ('REQ-018', 'DevSecOps with CI or CD, IaC, and policy-as-code.', '1, 3, 5, 6, 9, 10'),
    ('REQ-019', 'Data technologies including PostgreSQL, object storage, ETL and ELT.', '3, 5, 6, 8, 12'),
    ('REQ-020', 'Version-controlled architecture and process diagrams, including Mermaid.', '3, 4, 5, 7, 8, 9, 10, 11'),
    ('REQ-021', 'Public-sector service portal modernization experience.', '2, 3, 13'),
    ('REQ-022', 'Healthcare data integration experience.', '2, 3, 8, 13'),
    ('REQ-023', 'Utilities operations enablement experience.', '2, 3, 12, 13'),
    ('REQ-024', 'Clear acceptance criteria and definition of done.', '1, 4, 10, 15'),
    ('REQ-025', 'Automated testing strategy across layers.', '3, 4, 5, 6, 8, 9, 10'),
    ('REQ-026', 'Formal design reviews, security checkpoints, release readiness.', '1, 4, 9, 10, 15'),
    ('REQ-027', 'Transparent reporting including burn-up and burn-down, risks, decisions.', '1, 4, 10, 13, 15'),
    ('REQ-028', 'Typical delivery team composition.', '2, 3, 14'),
    ('REQ-029', 'Clear point of contact for business development and proposals.', '1, 16'),
]

for req_id, desc, secs in requirements:
    row_cells = rtm.add_row().cells
    row_cells[0].text = req_id
    row_cells[1].text = desc
    row_cells[2].text = secs
