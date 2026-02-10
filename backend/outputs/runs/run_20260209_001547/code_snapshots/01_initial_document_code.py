from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import matplotlib.patches as mpatches
import datetime as dt
import matplotlib.dates as mdates

styles = doc.styles
# --- Style Bootstrap ---
normal = styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)

if 'Caption' not in [s.name for s in styles]:
    cap = styles.add_style('Caption', WD_STYLE_TYPE.PARAGRAPH)
    cap.font.name = 'Calibri'
    cap.font.size = Pt(9)
    cap.font.italic = True

def add_caption(text: str):
    p = doc.add_paragraph(text, style='Caption')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p

# --- Title Page ---
doc.add_heading('Dynamics 365 Finance and Operations\nComprehensive ERP Solution Proposal', level=0)
p = doc.add_paragraph()
p.add_run('Prepared for: ').bold = True
p.add_run('Smithsonian Enterprises')
p = doc.add_paragraph()
p.add_run('Submitted by: ').bold = True
p.add_run('Argano, LLC')
p = doc.add_paragraph()
p.add_run('Confidentiality Notice').bold = True
p = doc.add_paragraph('External use or disclosure of any information contained in this document without written permission of Argano, LLC is prohibited.')
p.paragraph_format.space_before = Pt(24)
doc.add_page_break()

# --- Table of Contents ---
doc.add_heading('Table of Contents', level=1)
toc_items = [
    'Executive Summary',
    'Company Overview',
    'Requirements Response & Compliance Matrix',
    'Solution Architecture and Technical Approach',
    'Implementation Plan & Delivery Schedule',
    'Project Management and Governance',
    'Security, Privacy, and Compliance',
    'Team and Qualifications',
    'Assumptions, Risks, and Dependencies',
    'Appendices',
]
for item in toc_items:
    doc.add_paragraph(item, style='List Number')
doc.add_page_break()

# --- Executive Summary ---
doc.add_heading('Executive Summary', level=1)
doc.add_paragraph(
    'Argano, LLC is excited to submit this proposal for the implementation of Microsoft Dynamics 365 Finance and Operations for Smithsonian Enterprises. Our approach is designed to support Smithsonian Enterprises in digitally transforming its core finance, operations, retail, and e-commerce processes. We leverage deep industry experience, a proven agile methodology, and a commitment to partnership to ensure a successful and sustainable ERP implementation.'
)

sum_table = doc.add_table(rows=1, cols=2)
sum_table.style = 'Table Grid'
sum_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = sum_table.rows[0].cells
hdr[0].text, hdr[1].text = 'Strategic Goal', 'Measure of Success'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True
sum_data = [
    ('Unified ERP for consolidated operations & reporting', 'Increased productivity, improved insight'),
    ('Increased inventory and shipping visibility', 'Improved management of supply chain and orders'),
    ('Cost reduction through system consolidation', 'Lower IT maintenance and licensing costs'),
    ('Enhanced compliance and security', 'Sustained audit readiness, FISMA/PCI support'),
    ('Scalable foundation for growth', 'Future-proofed platform, seamless integrations'),
]
for goal, measure in sum_data:
    row = sum_table.add_row().cells
    row[0].text = goal
    row[1].text = measure

doc.add_paragraph('Argano is confident in our ability to deliver the project scope and meet all requirements based on our extensive experience in nonprofit, retail, and public sector ERP implementations.')
doc.add_page_break()

# --- Company Overview ---
doc.add_heading('Company Overview', level=1)
doc.add_paragraph(
    'Argano is a global consulting firm and a top-ranked Microsoft Dynamics Gold Certified Partner, specializing in large-scale ERP, CRM, and analytics solutions for public sector, nonprofit, retail, and complex multi-entity organizations. As a member of Microsoft’s Inner Circle and a US Partner of the Year for Dynamics 365 Finance, Argano brings over two decades of experience and an award-winning team of 1,900+ professionals.'
)
doc.add_paragraph('Key differentiators:')
doc.add_paragraph('Agile, outcome-driven delivery methodology', style='List Bullet')
doc.add_paragraph('Deep nonprofit and public sector experience', style='List Bullet')
doc.add_paragraph('Commitment to business process transformation', style='List Bullet')
doc.add_paragraph('Executive sponsorship for every implementation', style='List Bullet')
doc.add_paragraph('Extensive experience with legacy migrations and integrations', style='List Bullet')
doc.add_paragraph('Proprietary IP to accelerate value (Retail Dashboard, Audit & Security Manager)', style='List Bullet')
doc.add_page_break()

# --- Requirements Response & Compliance Matrix ---
doc.add_heading('Requirements Response & Compliance Matrix', level=1)
req_table = doc.add_table(rows=1, cols=3)
req_table.style = 'Table Grid'
req_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = req_table.rows[0].cells
hdr[0].text = 'Requirement ID'
hdr[1].text = 'Requirement Summary'
hdr[2].text = 'Solution Response'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True
# Sample: Show 7 representative requirements per domain
top_reqs = [
    ('GA-1', 'Multi-entity structure with configurable financial rollups', 'Fully supported via D365 organization hierarchy and reporting entities.'),
    ('GA-7', 'Control user access to posting attributes within chart of accounts', 'Role-based security with granular permission assignment.'),
    ('GA-10', 'Configure and automate recurring journal entries', 'Recurring journals and batch posting workflows available.'),
    ('GA-33', 'Run fixed asset depreciation by company or class', 'Automated depreciation by company, class, or asset. Separate books supported.'),
    ('GA-61', 'Automatically eliminate intercompany balances during consolidation', 'Native intercompany elimination during consolidation.'),
    ('O2C-2', 'Integrate customer account creation/updates from CRM', 'Bi-directional integration with D365 CRM; duplicate prevention logic.'),
    ('O2C-34', 'Multiple billing methods per contract, including recurring, milestone, usage-based', 'Flexible billing schedules, methods, and contract templates.'),
    ('O2C-44', 'Integrate CRM quote/project data into ERP', 'End-to-end process from CRM opportunity to ERP project.'),
    ('O2C-64', 'Automate revenue recognition and invoicing based on % complete', 'Rules-based revenue recognition by project progress.'),
    ('P2P-5', 'Approval workflows for new vendor onboarding', 'Configurable approval matrix and full audit trail.'),
    ('P2P-51', '3-way match between invoice, PO, and receipt', 'Automated 3-way match with exception routing.'),
    ('INV-17', 'Track inventory across multiple warehouses and regions', 'Multi-warehouse/location inventory with transfer visibility.'),
    ('BDS-3', 'Create budgets at Entity, Region, and Project levels', 'Budgeting module supports granular creation and consolidated rollup.'),
    ('BDS-12', 'Actual to budget comparisons on prior periods', 'Standard reporting and analytics for actuals vs. budget/forecast.'),
]
for rid, summary, resp in top_reqs:
    row = req_table.add_row().cells
    row[0].text = rid
    row[1].text = summary
    row[2].text = resp
doc.add_paragraph('See Appendix A for a detailed requirements traceability matrix. All requirements in the RFP are supported natively or via standard configuration, unless otherwise noted.')

# --- Solution Architecture and Technical Approach ---
doc.add_page_break()
doc.add_heading('Solution Architecture and Technical Approach', level=1)
doc.add_paragraph('Argano proposes a fully integrated Microsoft Dynamics 365 Finance and Operations solution, leveraging the following technical features:')
doc.add_paragraph('Role-based access control and workflow automation', style='List Bullet')
doc.add_paragraph('Multi-entity, multi-currency, and multi-book support', style='List Bullet')
doc.add_paragraph('Automated approval, segregation of duties, and audit trail', style='List Bullet')
doc.add_paragraph('Robust integration with CRM, Field Service, and external platforms', style='List Bullet')
doc.add_paragraph('Powerful reporting and analytics with drilldown and ad-hoc capabilities', style='List Bullet')
doc.add_paragraph('Support for FISMA, PCI, SOC 1/2, and Smithsonian privacy requirements', style='List Bullet')

mermaid_code = '''
flowchart TD
  A[Smithsonian Users] -->|Web/App| B[D365 Finance & Operations]
  B --> C[GL, AP, AR, FA, Projects]
  B --> D[Procurement]
  B --> E[Inventory Mgmt]
  B --> F[CRM/Field Service]
  F <--> G[External Integrations]
  B --> H[PowerBI/Analytics]
  B --> I[Security & Compliance]
  H --> J[Dashboards/Reports]
'''
diagram_path = render_mermaid(mermaid_code, 'solution_architecture', width=1600, height=900, scale=1.5)
doc.add_picture(str(diagram_path), width=Inches(5.8))
add_caption('Figure 1: Solution architecture overview')

# --- Implementation Plan & Delivery Schedule ---
doc.add_page_break()
doc.add_heading('Implementation Plan & Delivery Schedule', level=1)
doc.add_paragraph('Argano follows a proven, five-phase implementation methodology:')
doc.add_paragraph('1. Project Initiation', style='List Number')
doc.add_paragraph('2. Analyze', style='List Number')
doc.add_paragraph('3. Design/Build', style='List Number')
doc.add_paragraph('4. Deploy (Go-Live)', style='List Number')
doc.add_paragraph('5. Optimize (Post Go-Live Support)', style='List Number')
doc.add_paragraph('Each phase is supported by joint project management, change control, and executive sponsorship. Our approach emphasizes rapid knowledge transfer, best-practice configuration, and iterative validation.')

# Gantt Chart
import pandas as pd
schedule = pd.DataFrame({
    'Task': [
        'Project Initiation', 'Analyze', 'Design/Build', 'Deploy', 'Optimize'],
    'Team': [
        'PMO', 'Business Analysts', 'Functional/Technical', 'PMO/Client', 'Support'],
    'Start': pd.to_datetime([
        '2024-01-08', '2024-01-29', '2024-03-04', '2024-06-10', '2024-07-01']),
    'End': pd.to_datetime([
        '2024-01-26', '2024-02-23', '2024-06-07', '2024-06-28', '2024-08-30'])
})
project_start = schedule['Start'].min()
schedule['days_to_start'] = (schedule['Start'] - project_start).dt.days
schedule['duration'] = (schedule['End'] - schedule['Start']).dt.days + 1
team_colors = {
    'PMO': '#2980b9',
    'Business Analysts': '#27ae60',
    'Functional/Technical': '#8e44ad',
    'PMO/Client': '#e67e22',
    'Support': '#34495e',
}
import matplotlib.pyplot as plt
fig, ax = plt.subplots(figsize=(10, 4.5))
for idx, row in schedule.iterrows():
    ax.barh(y=row['Task'], width=row['duration'], left=row['days_to_start'],
            color=team_colors[row['Team']], edgecolor='white', linewidth=0.5)
ax.invert_yaxis()
max_days = schedule['days_to_start'].max() + schedule['duration'].max()
exticks = range(0, max_days+7, 21)
exticklabels = pd.date_range(start=project_start, periods=len(xticks), freq='21D').strftime('%b %d')
ax.set_xticks(list(xticks)[:len(xticklabels)])
ax.set_xticklabels(xticklabels, fontsize=9)
ax.set_xlabel('Timeline', fontsize=10)
ax.set_title('Implementation Schedule by Phase', fontsize=14, fontweight='bold')
ax.xaxis.grid(True, alpha=0.3, linestyle='--')
patches = [mpatches.Patch(color=team_colors[team], label=team) for team in schedule['Team'].unique()]
ax.legend(handles=patches, loc='lower right', fontsize=8, framealpha=0.9)
plt.tight_layout()
gantt_path = output_dir / 'implementation_gantt.png'
plt.savefig(gantt_path, dpi=150, bbox_inches='tight', facecolor='white')
plt.close()
doc.add_picture(str(gantt_path), width=Inches(6.0))
add_caption('Figure 2: High-level implementation schedule by phase')

doc.add_paragraph('A detailed cutover plan and resource matrix will be delivered prior to Go-Live. See Appendix B for a sample RACI and detailed timeline.')

doc.add_page_break()
# --- Project Management and Governance ---
doc.add_heading('Project Management and Governance', level=1)
doc.add_paragraph('Our governance model ensures joint accountability, transparency, and shared success. Key elements include:')
doc.add_paragraph('• Joint project management and steering committee', style='List Bullet')
doc.add_paragraph('• Weekly status reviews and dashboard reporting', style='List Bullet')
doc.add_paragraph('• Structured change control and issue escalation', style='List Bullet')
doc.add_paragraph('• Roles/responsibilities matrix and clear RACI', style='List Bullet')
doc.add_paragraph('• Executive sponsorship from Argano and Smithsonian', style='List Bullet')

mermaid_gov = '''
flowchart TD
    A[Steering Committee] --> B[Project Management]
    B --> C[Workstream Leads]
    C --> D[Functional Teams]
    B --> E[Change Control Board]
    E --> B
'''
gov_path = render_mermaid(mermaid_gov, 'gov_structure', width=1600, height=800, scale=1.5)
doc.add_picture(str(gov_path), width=Inches(5.8))
add_caption('Figure 3: Project governance structure')

doc.add_paragraph('Project reporting will be managed through Azure DevOps dashboards, providing real-time insight into sprint burndown, backlog, completion rates, and issue tracking.')

doc.add_page_break()
# --- Security, Privacy, and Compliance ---
doc.add_heading('Security, Privacy, and Compliance', level=1)
doc.add_paragraph('Argano implements D365 in full alignment with Smithsonian security, privacy, FISMA, and PCI requirements. Our approach includes:')
doc.add_paragraph('• End-to-end data encryption (in transit and at rest)', style='List Bullet')
doc.add_paragraph('• Azure AD RBAC and multi-factor authentication', style='List Bullet')
doc.add_paragraph('• Audit logs, segregation of duties, and Argano Audit & Security Manager IP', style='List Bullet')
doc.add_paragraph('• Compliance with Smithsonian Privacy and Security Clause (Exhibit B)', style='List Bullet')
doc.add_paragraph('• Regular vulnerability assessments and incident response planning', style='List Bullet')
doc.add_paragraph('• Data residency and backup in US-based Azure regions', style='List Bullet')
doc.add_paragraph('Our solution natively supports FISMA, PCI-DSS v4.0, and SOC 1/2 controls. Detailed technical documentation and security runbooks will be provided.')

doc.add_page_break()
# --- Team and Qualifications ---
doc.add_heading('Team and Qualifications', level=1)
doc.add_paragraph('Argano’s project team consists of seasoned professionals with extensive D365 experience, supported by a deep bench of technical, functional, and project management talent. The core project team will be finalized upon contract execution and includes:')
doc.add_paragraph('• Project Manager (PMP certified)', style='List Bullet')
doc.add_paragraph('• Solution Architect', style='List Bullet')
doc.add_paragraph('• Technical Architect', style='List Bullet')
doc.add_paragraph('• Functional Consultants (Finance, Supply Chain, Retail, POS, Ecommerce)', style='List Bullet')
doc.add_paragraph('• Data Migration Lead', style='List Bullet')
doc.add_paragraph('• Security Consultant', style='List Bullet')
doc.add_paragraph('• QA/Testing Lead', style='List Bullet')
doc.add_paragraph('• Executive Sponsor', style='List Bullet')
doc.add_paragraph('Resumes and bios for key team members will be provided upon request or shortlisting.')

# --- Assumptions, Risks, and Dependencies ---
doc.add_page_break()
doc.add_heading('Assumptions, Risks, and Dependencies', level=1)
doc.add_paragraph('Our proposal and estimates are based on the following key assumptions:')
doc.add_paragraph('• Timely access to required Smithsonian resources and data', style='List Bullet')
doc.add_paragraph('• Client will provide subject matter experts for all core business processes', style='List Bullet')
doc.add_paragraph('• Major integrations (Logiwa, POS) are accessible and well-documented', style='List Bullet')
doc.add_paragraph('• Blackout periods and go-live dates are honored as per RFP', style='List Bullet')
doc.add_paragraph('• Data migration scope: 3 years of financial transactions, 5 years of sales and merchandising data', style='List Bullet')
doc.add_paragraph('• Pricing is based on information provided and is subject to final SOW negotiation', style='List Bullet')
doc.add_paragraph('Risks and mitigations:')
doc.add_paragraph('• Resource constraints: mitigated by phased delivery and Argano on-site/remote flexibility', style='List Bullet')
doc.add_paragraph('• Data migration complexity: addressed via proven methodology and iterative testing', style='List Bullet')
doc.add_paragraph('• Integration risk: mitigated by early engagement with third-party vendors and use of standard protocols', style='List Bullet')
doc.add_paragraph('• Change management: addressed through joint training, leadership engagement, and clear communications', style='List Bullet')

# --- Appendices ---
doc.add_page_break()
doc.add_heading('Appendices', level=1)
doc.add_paragraph('Appendix A: Detailed Requirements Traceability Matrix (provided as spreadsheet)')
doc.add_paragraph('Appendix B: Sample Project Timeline & RACI')
doc.add_paragraph('Appendix C: Sample Training Materials and Change Management Packages')
doc.add_paragraph('Appendix D: Master Services Agreement (provided upon request)')

# --- End of Document ---
