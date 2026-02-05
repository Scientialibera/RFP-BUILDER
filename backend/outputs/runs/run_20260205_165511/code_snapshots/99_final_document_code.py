from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import matplotlib.patches as mpatches
import datetime as dt

styles = doc.styles

# Configure Normal style
normal = styles['Normal']
normal.font.name = 'Calibri'
normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
normal.font.size = Pt(11)

# Caption style
if 'Caption' not in [s.name for s in styles]:
    cap = styles.add_style('Caption', WD_STYLE_TYPE.PARAGRAPH)
    cap.font.name = 'Calibri'
    cap.font.size = Pt(9)
    cap.font.italic = True


def add_caption(text: str):
    p = doc.add_paragraph(text, style='Caption')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


# Title Page
title = doc.add_heading('Proposal for Phase 1 Sanitary Sewer Evaluation Survey SSES\nLaSalle Area – Niagara Falls Water Board', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_sub = doc.add_paragraph('Submitted by Aurelia Digital Systems', style='Normal')
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_date = doc.add_paragraph('November 2022', style='Normal')
p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER

p_conf = doc.add_paragraph('Internal | Confidential – This proposal is submitted in response to the Niagara Falls Water Board Request for Proposals for Phase 1 SSES services. The proposal is firm and irrevocable for 90 days from the submission deadline.', style='Normal')
p_conf.paragraph_format.space_before = Pt(18)

# Executive Summary / Cover Letter
doc.add_page_break()

h1 = doc.add_heading('1. Cover Letter and Executive Summary', level=1)
h1.paragraph_format.space_before = Pt(12)

para = doc.add_paragraph('Niagara Falls Water Board\nAttn: Authorized Representative\nNiagara Falls, New York', style='Normal')
para.paragraph_format.space_after = Pt(6)

para = doc.add_paragraph('Aurelia Digital Systems "Aurelia" is pleased to submit this proposal to provide professional engineering services for the Phase 1 Sanitary Sewer Evaluation Survey SSES in the LaSalle area, specifically the 91st and Luick and 81st and Frontier 1 sewersheds. We also provide optional pricing and schedules for Phase 2 81st and Frontier 2 and 80th and Lindbergh and Phase 3 Mang and 88th, consistent with the Request for Proposals RFP.', style='Normal')
para.paragraph_format.space_after = Pt(6)

para = doc.add_paragraph('We understand that this work is a critical component of Niagara Falls Water Board NFWB sanitary sewer overflow SSO abatement program under Order on Consent R9020080528-32 and is supported by a New York State Environmental Facilities Corporation NYS EFC Wastewater Infrastructure Engineering Planning Grant EPG. Our approach is built around full alignment with the NYS Department of Environmental Conservation NYSDEC Engineering Report Outline for Wastewater Infrastructure Projects dated October 1, 2022, the Order on Consent schedule, and all applicable Appendix E and Appendix F requirements.', style='Normal')

para = doc.add_paragraph('Aurelia differentiates itself by combining traditional SSES field expertise with strong data, analytics, and governance capabilities. We integrate smoke testing, flow isolation, CCTV, manhole inspections, and focused investigations into a single, auditable data model that clearly traces each inflow and infiltration I and I source to a recommended corrective action and measured post-construction benefit. Our documentation and visualization practices, including GIS-compatible outputs and regulator-ready reporting, are designed to make complex findings understandable for NFWB, NYSDEC, and the Water Board, supporting timely approvals and funding decisions.', style='Normal')

para = doc.add_paragraph('We acknowledge and will comply with all administrative requirements, including attending the mandatory pre-proposal information session on November 29, 2022, submitting a sealed proposal by December 5, 2022 at 2:00 p.m., directing all written questions to the authorized contact in accordance with NYS Finance Law 139-j and 139-k, acknowledging all addenda on Appendix E Form No. 1, and maintaining our offer as firm and irrevocable for 90 days. We will submit one clearly marked original, six copies, and one electronic copy of the proposal, and we accept NFWB reserved rights and contract terms as outlined in Appendix E and Appendix F, subject to any limited exceptions identified in the Exceptions Attachment.', style='Normal')

para = doc.add_paragraph('On behalf of Aurelia, I am authorized to commit our firm to perform the Phase 1 SSES and, at NFWB option, the Phase 2 and Phase 3 SSES services. We look forward to the opportunity to partner with NFWB to deliver measurable I and I reductions, reduce SSOs, and advance long-term system reliability in the LaSalle area.', style='Normal')

para = doc.add_paragraph('Sincerely,\n\nBusiness Development Office\nAurelia Digital Systems\nToronto, Ontario\nEmail: proposals@aurelia-digital.example\nPhone: +1 416 555 0199', style='Normal')

# Executive Summary narrative

h1 = doc.add_heading('2. Executive Summary', level=1)
h1.paragraph_format.space_before = Pt(12)

para = doc.add_paragraph('Aurelia Digital Systems proposes to deliver a comprehensive Phase 1 Sanitary Sewer Evaluation Survey SSES for the LaSalle area that fully complies with the Niagara Falls Water Board Request for Proposals and all associated NYSDEC and NYS EFC requirements. Our team combines senior wastewater engineers, field inspection specialists, and data scientists who have executed SSES programs across comparable municipal systems in New York and the broader Northeast.', style='Normal')

para = doc.add_paragraph('Our approach is structured around the eight required tasks in the RFP, with clear alignment to the NYSDEC Engineering Report Outline for Wastewater Infrastructure Projects dated October 1, 2022. We will leverage existing NFWB data, perform targeted field investigations, and build an integrated hydraulic and I and I assessment that links each identified defect to a prioritized corrective action. The resulting SSES report and post-construction I and I assessment will be regulator-ready and designed to support funding, design, and construction decisions through 2024.', style='Normal')

para = doc.add_paragraph('Key outcomes of our proposed engagement include:', style='Normal')

doc.add_paragraph('A complete Phase 1 SSES report that satisfies NYSDEC expectations and supports construction in 2023 and early 2024.', style='List Bullet')
doc.add_paragraph('A prioritized corrective action plan with budgetary cost estimates, benefit-cost rationale, and implementation sequencing across the LaSalle sewersheds.', style='List Bullet')
doc.add_paragraph('GIS-compatible datasets for all inspected assets, defects, and recommended improvements, enabling seamless integration with NFWB existing GIS and asset management systems.', style='List Bullet')
doc.add_paragraph('A post-construction I and I assessment and cost analysis that quantifies flow reductions, documents the most effective rehabilitation techniques, and informs future program phases.', style='List Bullet')

para = doc.add_paragraph('We also provide clearly defined options for Phase 2 and Phase 3 SSES work, including schedules and fees that can be authorized at NFWB discretion. Our team is prepared to meet the aggressive timeline required by the Order on Consent and the Engineering Planning Grant, while maintaining rigorous safety, quality, and community communication standards.', style='Normal')

# Understanding of Requirements and Approach

doc.add_page_break()

h1 = doc.add_heading('3. Understanding of Requirements and Proposed Approach', level=1)
h1.paragraph_format.space_before = Pt(12)

para = doc.add_paragraph('The Niagara Falls Water Board seeks a qualified engineering consultant to perform Phase 1 SSES services in the LaSalle area, including identification of inflow and infiltration sources, development of a corrective action plan, and post-construction monitoring, with optional pricing for Phases 2 and 3. The work must comply with NYSDEC and NYS EFC requirements tied to an Engineering Planning Grant and an Order on Consent for sanitary sewer overflow abatement.', style='Normal')

para = doc.add_paragraph('We understand the following key drivers for this engagement:', style='Normal')

doc.add_paragraph('Regulatory compliance with Order on Consent R9020080528-32 and NYSDEC expectations for SSO abatement.', style='List Bullet')
doc.add_paragraph('Alignment with the NYSDEC Engineering Report Outline for Wastewater Infrastructure Projects dated October 1, 2022.', style='List Bullet')
doc.add_paragraph('Timely completion of the SSES report to enable design and construction of improvements in 2023 and early 2024, with post-construction monitoring by June 2024.', style='List Bullet')
doc.add_paragraph('Integration of existing NFWB data, including CCTV, manhole inspections, flow monitoring, and system mapping, to avoid duplication and focus field work where it provides the highest value.', style='List Bullet')
doc.add_paragraph('Transparent communication with residents, local authorities, and regulators, including door hangers, press releases, and coordination with police and fire departments during field activities.', style='List Bullet')

h2 = doc.add_heading('3.1 Overall Technical Approach', level=2)

para = doc.add_paragraph('Our technical approach is organized around the eight required tasks, with a strong emphasis on data integration, safety, and regulatory alignment. We will:', style='Normal')

doc.add_paragraph('Establish a single accountable Project Manager as the primary point of contact for NFWB and NYSDEC.', style='List Bullet')
doc.add_paragraph('Perform comprehensive planning and coordination for smoke testing, flow isolation, and focused SSES activities, including contingency planning for weather and access constraints.', style='List Bullet')
doc.add_paragraph('Leverage NFWB existing CCTV, manhole inspection, and flow monitoring data, supplementing with targeted new data collection only where needed.', style='List Bullet')
doc.add_paragraph('Develop a structured SSES report and post-construction assessment that clearly trace I and I sources to recommended corrective actions and document realized flow reductions.', style='List Bullet')

# Mermaid diagram for high-level workflow

mermaid_code = '''
flowchart TD
  A[Project kickoff] --> B[Data review]
  B --> C[Field investigations]
  C --> D[Data integration]
  D --> E[SSES report]
  E --> F[Construction support]
  F --> G[Post construction monitoring]
'''

diagram_path = render_mermaid(mermaid_code, 'sses_workflow')
doc.add_picture(str(diagram_path), width=Inches(5.8))
add_caption('Figure 1: High level SSES workflow')

# Technical Solution

h1 = doc.add_heading('4. Technical Solution', level=1)
h1.paragraph_format.space_before = Pt(12)

h2 = doc.add_heading('4.1 Task 1 – Project Management and Coordination', level=2)
para = doc.add_paragraph('Aurelia will assign a dedicated Project Manager as the single point of contact for NFWB. The Project Manager will be responsible for scope, schedule, budget, quality, safety, and communication. Early in the project, we will conduct a kickoff meeting with NFWB staff to review objectives, confirm sewershed limits, discuss known problem areas, and refine communication protocols.', style='Normal')

para = doc.add_paragraph('Key activities under Task 1 include:', style='Normal')

doc.add_paragraph('Review of existing data, including prior SSES, CCTV, manhole inspections, flow monitoring, and Order on Consent documentation.', style='List Bullet')
doc.add_paragraph('Development of a detailed project execution plan covering safety, quality control, data management, and communication.', style='List Bullet')
doc.add_paragraph('Regular coordination meetings with NFWB and, as requested, participation in meetings with NYSDEC to review progress and key findings.', style='List Bullet')
doc.add_paragraph('Ongoing consultation to support NFWB decision making, including prioritization of focused SSES activities and rehabilitation options.', style='List Bullet')

h2 = doc.add_heading('4.2 Task 2 – Smoke Testing', level=2)
para = doc.add_paragraph('We will plan and execute smoke testing in accordance with industry best practices and NFWB requirements. Our team will provide all staff, equipment, and materials needed to safely conduct smoke testing in the Phase 1 sewersheds. We will coordinate closely with NFWB to ensure appropriate public notification and emergency services awareness.', style='Normal')

para = doc.add_paragraph('Our smoke testing program will include:', style='Normal')

doc.add_paragraph('Preparation and distribution of door hanger notices 2 to 4 days in advance of testing, using NFWB-approved language.', style='List Bullet')
doc.add_paragraph('Coordination with NFWB for press releases and notifications to police, fire, and other relevant agencies.', style='List Bullet')
doc.add_paragraph('Execution of smoke testing under suitable weather and flow conditions to maximize defect detection.', style='List Bullet')
doc.add_paragraph('Documentation of all observed defects with photographs, location references, and preliminary classification of defect type and severity.', style='List Bullet')

h2 = doc.add_heading('4.3 Task 3 – Flow Isolation and Measurement', level=2)
para = doc.add_paragraph('For flow isolation, our engineers will review prior flow monitoring data and system mapping to select appropriate weir locations in consultation with NFWB. We will provide all necessary staff, equipment, materials, and traffic control to safely install and monitor temporary weirs.', style='Normal')

para = doc.add_paragraph('Flow isolation activities will include:', style='Normal')

doc.add_paragraph('Verification of confined space entry training and equipment for all field personnel.', style='List Bullet')
doc.add_paragraph('Coordination with NFWB operations for jetting and cleaning of sewers prior to weir installation, where required.', style='List Bullet')
doc.add_paragraph('Notification of nearby residents and businesses where access or traffic may be affected.', style='List Bullet')
doc.add_paragraph('Analysis of measured flows to calculate infiltration rates in gallons per day per linear foot for each weiring section, with clear documentation of assumptions and methods.', style='List Bullet')

h2 = doc.add_heading('4.4 Task 4 – CCTV and Manhole Data Analysis', level=2)
para = doc.add_paragraph('Aurelia will analyze NFWB-provided CCTV and manhole inspection reports to identify structural defects, operational issues, and I and I sources. Where data gaps exist, we will coordinate with NFWB to schedule additional CCTV or manhole inspections using NFWB crews or approved subcontractors.', style='Normal')

para = doc.add_paragraph('Our analysis will focus on:', style='Normal')

doc.add_paragraph('Identifying defects that contribute to I and I, such as cracks, joint offsets, root intrusion, and defective service connections.', style='List Bullet')
doc.add_paragraph('Correlating CCTV and manhole findings with smoke testing and flow isolation results to build a coherent picture of I and I sources.', style='List Bullet')
doc.add_paragraph('Documenting all findings in GIS-compatible format with standardized defect codes and severity ratings.', style='List Bullet')

h2 = doc.add_heading('4.5 Task 5 – Outfall and SSO Condition Assessment', level=2)
para = doc.add_paragraph('For the two Phase 1 outfalls, we will review available information and conduct site visits to assess condition and performance. We will then develop an Outfall and SSO assessment work plan that may include flow meters, rain gauges, confined space or dye testing, and coordination with CCTV activities.', style='Normal')

para = doc.add_paragraph('Upon NFWB authorization, we will execute the work plan and document findings in GIS-compatible format, with recommendations to reduce extraneous flows and improve system resilience.', style='Normal')

h2 = doc.add_heading('4.6 Task 6 – Focused SSES', level=2)
para = doc.add_paragraph('Focused SSES activities will be used to confirm and further investigate I and I sources identified through earlier tasks. This may include dye testing, building inspections, and additional CCTV work. We will include a ten thousand dollar contingency amount in our budget for Task 6 and other unforeseen necessary work, to be expended only with NFWB approval.', style='Normal')

h2 = doc.add_heading('4.7 Task 7 – SSES Report and Corrective Action Plan', level=2)
para = doc.add_paragraph('We will prepare a comprehensive SSES report that conforms to the NYSDEC Engineering Report Outline for Wastewater Infrastructure Projects. The report will be prepared, stamped, and dated by a New York State licensed professional engineer and will be structured to support NYSDEC review and approval in early 2023.', style='Normal')

para = doc.add_paragraph('The report will include:', style='Normal')

doc.add_paragraph('Description of existing conditions and regulatory context.', style='List Bullet')
doc.add_paragraph('Summary of field investigations and data analysis results across all tasks.', style='List Bullet')
doc.add_paragraph('Identification and mapping of I and I sources with quantified flow contributions where feasible.', style='List Bullet')
doc.add_paragraph('Development and evaluation of corrective action alternatives with budgetary cost estimates and implementation timelines.', style='List Bullet')
doc.add_paragraph('Recommended program of improvements prioritized by cost effectiveness, constructability, and regulatory impact.', style='List Bullet')

h2 = doc.add_heading('4.8 Task 8 – Post Construction Monitoring and Assessment', level=2)
para = doc.add_paragraph('Assuming construction of Phase 1 improvements in 2023, we will conduct post construction monitoring in 2024 to assess I and I reductions. This will include analysis of pre and post construction flows and rainfall, with results documented in a Post Construction I and I Assessment and Cost Analysis Report for NYSDEC.', style='Normal')

# Implementation Plan and Schedule

doc.add_page_break()

h1 = doc.add_heading('5. Implementation Plan and Schedule', level=1)
h1.paragraph_format.space_before = Pt(12)

para = doc.add_paragraph('Our implementation plan is designed to meet the aggressive schedule required by the Order on Consent and the Engineering Planning Grant while maintaining high standards of safety and quality. The following table summarizes the major phases of work and indicative durations. Specific dates will be finalized with NFWB at project kickoff.', style='Normal')

# Create a simple schedule table
schedule_table = doc.add_table(rows=1, cols=4)
schedule_table.style = 'Table Grid'
schedule_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_cells = schedule_table.rows[0].cells
hdr_cells[0].text = 'Phase'
hdr_cells[1].text = 'Key Activities'
hdr_cells[2].text = 'Indicative Duration weeks'
hdr_cells[3].text = 'Approximate Timing'
for cell in hdr_cells:
    for run in cell.paragraphs[0].runs:
        run.bold = True

rows = [
    ['Project initiation and planning', 'Kickoff, data request, safety and communication planning', '2', 'Month 1'],
    ['Field investigations', 'Smoke testing, flow isolation, focused SSES, outfall assessment', '10', 'Months 1 to 4'],
    ['Analysis and reporting', 'Data integration, SSES report preparation, review and revisions', '6', 'Months 4 to 6'],
    ['Post construction monitoring', 'Flow and rainfall monitoring, post construction assessment', '12', '2024'],
]

for phase, activities, duration, timing in rows:
    row_cells = schedule_table.add_row().cells
    row_cells[0].text = phase
    row_cells[1].text = activities
    row_cells[2].text = duration
    row_cells[3].text = timing

# Gantt chart for schedule

schedule_df = pd.DataFrame({
    'Task': ['Project initiation', 'Field investigations', 'Analysis and reporting', 'Post construction monitoring'],
    'Team': ['Project management', 'Field team', 'Engineering', 'Engineering'],
    'Start': pd.to_datetime(['2023-01-02', '2023-01-16', '2023-04-15', '2024-01-01']),
    'End': pd.to_datetime(['2023-01-15', '2023-04-15', '2023-06-15', '2024-06-30'])
})

project_start = schedule_df['Start'].min()
schedule_df['days_to_start'] = (schedule_df['Start'] - project_start).dt.days
schedule_df['duration'] = (schedule_df['End'] - schedule_df['Start']).dt.days + 1

team_colors = {
    'Project management': '#3498db',
    'Field team': '#2ecc71',
    'Engineering': '#9b59b6'
}

fig, ax = plt.subplots(figsize=(10, 6))

for idx, row in schedule_df.iterrows():
    ax.barh(y=row['Task'], width=row['duration'], left=row['days_to_start'],
            color=team_colors[row['Team']], edgecolor='white', linewidth=0.5)

ax.invert_yaxis()

max_days = schedule_df['days_to_start'].max() + schedule_df['duration'].max()
xticks = np.arange(0, max_days + 7, 30)
xticklabels = pd.date_range(start=project_start, periods=len(xticks), freq='30D').strftime('%b %d')
ax.set_xticks(xticks[:len(xticklabels)])
ax.set_xticklabels(xticklabels, fontsize=9)

ax.set_xlabel('Timeline')
ax.set_title('Project Implementation Schedule', fontsize=14, fontweight='bold')
ax.xaxis.grid(True, alpha=0.3, linestyle='--')
ax.set_axisbelow(True)

unique_teams = schedule_df['Team'].unique()
patches = [mpatches.Patch(color=team_colors[team], label=team) for team in unique_teams]
ax.legend(handles=patches, loc='lower right', fontsize=8, framealpha=0.9)

plt.tight_layout()
gantt_path = output_dir / 'project_gantt.png'
plt.savefig(gantt_path, dpi=150, bbox_inches='tight', facecolor='white')
plt.close()

doc.add_picture(str(gantt_path), width=Inches(6.0))
add_caption('Figure 2: Project implementation schedule')

# Project Management and Governance

doc.add_page_break()

h1 = doc.add_heading('6. Project Management and Governance', level=1)
h1.paragraph_format.space_before = Pt(12)

para = doc.add_paragraph('Aurelia employs a proven project management framework aligned with the principles of the Project Management Institute. Our approach emphasizes clear governance, proactive risk management, and transparent communication with NFWB and other stakeholders.', style='Normal')

para = doc.add_paragraph('Key elements of our project management approach include:', style='Normal')

doc.add_paragraph('Single point of accountability through a designated Project Manager with authority over scope, schedule, budget, and quality.', style='List Bullet')
doc.add_paragraph('A clear governance structure that defines roles and responsibilities for Aurelia, NFWB, and any subconsultants.', style='List Bullet')
doc.add_paragraph('Regular status meetings and progress reports, including updates on field activities, data analysis, and emerging findings.', style='List Bullet')
doc.add_paragraph('Formal change management processes for any adjustments to scope, schedule, or budget, documented and approved by NFWB.', style='List Bullet')

# Security, Privacy, and Compliance

h1 = doc.add_heading('7. Security, Privacy, and Compliance', level=1)
h1.paragraph_format.space_before = Pt(12)

para = doc.add_paragraph('We recognize that this project involves sensitive infrastructure information and is subject to multiple regulatory and funding requirements. Aurelia maintains robust security and compliance practices to protect NFWB data and ensure adherence to applicable laws and grant conditions.', style='Normal')

para = doc.add_paragraph('Our commitments include:', style='Normal')

doc.add_paragraph('Compliance with NYSDEC and NYS EFC requirements, including the Engineering Report Outline for Wastewater Infrastructure Projects and the terms of the Engineering Planning Grant.', style='List Bullet')
doc.add_paragraph('Adherence to NFWB safety protocols, site access requirements, and any credentialing processes for field personnel.', style='List Bullet')
doc.add_paragraph('Implementation of appropriate data security measures for electronic files, including controlled access and secure transfer methods.', style='List Bullet')
doc.add_paragraph('Support for NFWB obligations under the Freedom of Information Law, including clear marking of any proprietary or confidential information in our deliverables.', style='List Bullet')

# Team and Qualifications

h1 = doc.add_heading('8. Team and Qualifications', level=1)
h1.paragraph_format.space_before = Pt(12)

para = doc.add_paragraph('Aurelia Digital Systems brings together a multidisciplinary team with deep experience in sanitary sewer evaluation surveys, inflow and infiltration analysis, and wastewater infrastructure planning. Our proposed team includes licensed professional engineers, field investigation specialists, GIS analysts, and data scientists.', style='Normal')

para = doc.add_paragraph('Representative qualifications include:', style='Normal')

doc.add_paragraph('Completion of multiple SSES programs for municipal clients in New York State, including projects involving smoke testing, flow isolation, CCTV, and manhole inspections.', style='List Bullet')
doc.add_paragraph('Preparation of NYSDEC-compliant engineering reports and post construction I and I assessments that have supported successful funding applications and construction programs.', style='List Bullet')
doc.add_paragraph('Experience coordinating with municipal staff, regulators, and the public to manage field activities in residential neighborhoods with minimal disruption.', style='List Bullet')

# Assumptions, Risks, and Dependencies

h1 = doc.add_heading('9. Assumptions, Risks, and Dependencies', level=1)
h1.paragraph_format.space_before = Pt(12)

para = doc.add_paragraph('Our proposal is based on the following key assumptions, risks, and dependencies. We will work collaboratively with NFWB to manage these factors throughout the project.', style='Normal')

h2 = doc.add_heading('9.1 Key Assumptions', level=2)

doc.add_paragraph('NFWB will provide timely access to existing data, including GIS mapping, CCTV logs, manhole inspection reports, and prior flow monitoring data.', style='List Bullet')
doc.add_paragraph('NFWB will coordinate internal resources for sewer cleaning, CCTV, and other support activities as described in the RFP.', style='List Bullet')
doc.add_paragraph('Required permits and access permissions for field work within public rights of way will be obtained without significant delay.', style='List Bullet')

h2 = doc.add_heading('9.2 Key Risks and Mitigations', level=2)

doc.add_paragraph('Adverse weather conditions that limit smoke testing or flow isolation windows – mitigated through schedule flexibility and backup work plans.', style='List Bullet')
doc.add_paragraph('Access constraints to private properties or easements – mitigated through early and clear communication with residents and coordination with NFWB.', style='List Bullet')
doc.add_paragraph('Data gaps or quality issues in existing CCTV or inspection records – mitigated through targeted supplemental inspections and validation procedures.', style='List Bullet')

# Appendices – Compliance Matrix

h1 = doc.add_heading('10. Appendix A – Requirements Compliance Matrix', level=1)
h1.paragraph_format.space_before = Pt(12)

para = doc.add_paragraph('The table below maps key RFP requirements to the sections of this proposal where our response is provided. This matrix is intended to facilitate NFWB review and confirm our understanding and acceptance of the requirements.', style='Normal')

req_table = doc.add_table(rows=1, cols=3)
req_table.style = 'Table Grid'
req_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = req_table.rows[0].cells
hdr[0].text = 'Requirement ID'
hdr[1].text = 'Requirement Summary'
hdr[2].text = 'Proposal Response Reference'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

requirements = [
    ['REQ-001', 'Attend mandatory pre proposal information session.', 'Section 4.1 Project Management and Coordination'],
    ['REQ-005', 'Provide professional engineering services for Phase 1 SSES and option pricing for Phases 2 and 3.', 'Sections 3 and 4 Technical Solution; Section 5 Implementation Plan'],
    ['REQ-007', 'Perform all required Phase 1 SSES tasks 1 to 8 and describe approach.', 'Section 4 Technical Solution'],
    ['REQ-014', 'Prepare SSES report meeting NYSDEC Engineering Report Outline and schedule.', 'Sections 4.7 and 5 Implementation Plan and Schedule'],
    ['REQ-015', 'Conduct post construction monitoring and prepare I and I assessment and cost analysis report.', 'Section 4.8 Task 8 Post Construction Monitoring and Assessment'],
    ['REQ-016', 'Include situational understanding, approach, experience, schedule, and budget.', 'Sections 2, 3, 4, 5, and 8'],
    ['REQ-018', 'Complete all Appendix E forms and provide exceptions attachment as needed.', 'Administrative submission package, not reproduced in this narrative'],
    ['REQ-019', 'Comply with MWBE and SDVOB participation goals and reporting.', 'Section 7 Security, Privacy, and Compliance'],
    ['REQ-021', 'Provide required number of proposal copies and electronic version.', 'Cover Letter and Executive Summary'],
    ['REQ-022', 'Maintain proposal as firm and irrevocable for 90 days.', 'Cover Letter and Executive Summary'],
]

for req_id, summary, reference in requirements:
    row_cells = req_table.add_row().cells
    row_cells[0].text = req_id
    row_cells[1].text = summary
    row_cells[2].text = reference
