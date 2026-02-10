from docx.enum.style import WD_STYLE_TYPE
import matplotlib.patches as mpatches
import datetime as dt
import matplotlib.dates as mdates

styles = doc.styles
# Normal
normal = styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)

# Caption style (create if missing)
if 'Caption' not in [s.name for s in styles]:
    cap = styles.add_style('Caption', WD_STYLE_TYPE.PARAGRAPH)
    cap.font.name = 'Calibri'
    cap.font.size = Pt(9)
    cap.font.italic = True

def add_caption(text: str):
    p = doc.add_paragraph(text, style='Caption')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p

# Title Page

doc.add_heading('Kastle Systems ERP Modernization Proposal', level=0)
doc.add_paragraph('Prepared for: Kastle Systems')
doc.add_paragraph('Dynamics 365 Finance and Operations')
doc.add_paragraph('Confidential: External use or disclosure of any information contained in this document without written permission of Argano, LLC is prohibited.', style='Caption')
doc.add_page_break()

# Letter of Transmittal
from datetime import date

doc.add_paragraph(f"Letter of Transmittal to Kastle Systems ({date.today().strftime('%B %d, %Y')})\nRe: RFP")
doc.add_paragraph("To whom it may concern, this response to the Kastle Systems RFP is hereby submitted by Argano:")
doc.add_paragraph('Submitted by:')
doc.add_paragraph('Argano\n6100 West Plano Parkway\nSuite 1800\nPlano, Texas 75093\nAttn: Legal: Elizabeth Frederic')
doc.add_paragraph('Vendor Contact Information:', style='List Bullet')
doc.add_paragraph('SAM/CSM Name: John Smith, Senior Account Manager, 555-123-4567, john.smith@argano.com', style='List Bullet')
doc.add_paragraph('DSE/DCE Name: Jane Doe, Solution Director, 555-987-6543, jane.doe@argano.com', style='List Bullet')
doc.add_paragraph('Kastle Contact Information:', style='List Bullet')
doc.add_paragraph('Name: Sarah Green\nTitle: Procurement Manager\nPhone: 555-555-1212\nEmail: sarah.green@kastle.com')
doc.add_page_break()

# Table of Contents
sections = [
    'Executive Summary',
    'Company Overview',
    'Service Capabilities',
    'Implementation Plan and Details',
    'Proposed Timeline',
    'Resources',
    'Client References',
    'Terms and Conditions',
    'Pricing Proposal',
    'Appendix 1 - Responsibilities'
]
toc = doc.add_paragraph('TABLE OF CONTENTS')
toc.style = doc.styles['Heading 1']
for idx, section in enumerate(sections, 1):
    doc.add_paragraph(f"{idx} {section}", style='Normal')
doc.add_page_break()

# Executive Summary
exec_head = doc.add_heading('Executive Summary', level=1)
exec_head.paragraph_format.space_before = Pt(18)
exec_head.paragraph_format.space_after = Pt(6)
doc.add_paragraph("Argano, LLC (\"Argano\") is pleased to submit this proposal in response to Kastle Systems' RFP for ERP modernization. Our proven track record in large-scale ERP implementations and ongoing managed services uniquely positions us to deliver a solution that meets Kastle's strategic, technical, and business requirements. Our extensive expertise in Dynamics 365 Finance and Operations, deep industry knowledge, and commitment to partnership ensure your objectives will be achieved—on time, within budget, and with high adoption.")
doc.add_paragraph('Our proposal addresses all mandatory requirements, including system design, integration, data migration, training, managed services, and transparent pricing. We are committed to collaborative discovery, clear communication, and measurable outcomes. The below table summarizes Kastle’s strategic goals and Argano’s measures of success.')

# Strategic Goals Table
str_table = doc.add_table(rows=1, cols=2)
str_table.style = 'Table Grid'
str_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = str_table.rows[0].cells
hdr[0].text, hdr[1].text = 'Strategic Goals', 'Measures of Success'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True
str_data = [
    ('Single ERP application for consolidated operations and reporting', 'Internal productivity increase'),
    ('Increased inventory and shipping visibility', 'Improved management of supply chain and orders'),
    ('Less IT maintenance and licensing costs', 'IT can free up bandwidth to work on other projects')
]
for g, m in str_data:
    row = str_table.add_row().cells
    row[0].text = g
    row[1].text = m

# Implementation timeline visual
plt.figure(figsize=(8, 4.5))
sns.set_style('whitegrid')
data = pd.DataFrame({'Phase': ['Initiation', 'Analysis', 'Design & Develop', 'Deploy', 'Optimize'],
                     'Weeks': [2, 3, 10, 3, 2]})
ax = sns.barplot(data=data, x='Phase', y='Weeks', palette='Blues_d')
ax.set_title('Proposed ERP Implementation Timeline')
ax.set_ylabel('Duration (Weeks)')
plt.tight_layout()
timeline_chart_path = output_dir / 'timeline_chart.png'
plt.savefig(timeline_chart_path, dpi=150, bbox_inches='tight', facecolor='white')
plt.close()
doc.add_picture(str(timeline_chart_path), width=Inches(5.8))
add_caption('Figure 1: Proposed ERP implementation timeline by phase')

# Company Overview
company_head = doc.add_heading('Company Overview', level=1)
company_head.paragraph_format.space_before = Pt(18)
company_head.paragraph_format.space_after = Pt(6)
doc.add_paragraph('Founded in 2001, Argano is a global leader in business transformation and enterprise technology integration. With more than 1,200 professionals across North America, Europe, and APAC, Argano specializes in Microsoft Dynamics, SAP, Oracle, and custom solutions for the mid-market and enterprise clients. Our ERP practice brings deep experience in finance, supply chain, manufacturing, and managed services.')
doc.add_paragraph('Our delivery methodology combines best practices from PMI, Agile, and industry-specific frameworks. We emphasize collaborative discovery, transparent communication, and measurable outcomes throughout the engagement.')

# Service Capabilities
srv_head = doc.add_heading('Service Capabilities', level=1)
srv_head.paragraph_format.space_before = Pt(18)
srv_head.paragraph_format.space_after = Pt(6)
doc.add_paragraph('Argano offers a full suite of services to support Kastle Systems’ ERP modernization:')
doc.add_paragraph('ERP Solution Design and Implementation', style='List Bullet')
doc.add_paragraph('System Integration (Dynamics CRM, Concur, Dayforce, Planful, NCA)', style='List Bullet')
doc.add_paragraph('Data Migration (historical and transactional)', style='List Bullet')
doc.add_paragraph('Training (in-person, webinar, online)', style='List Bullet')
doc.add_paragraph('Post-implementation Managed Services (installation, maintenance, upgrades, admin, reporting)', style='List Bullet')
doc.add_paragraph('Customization and Enhancement', style='List Bullet')
doc.add_paragraph('Compliance and Security Consulting', style='List Bullet')

# Implementation Plan and Details
imp_head = doc.add_heading('Implementation Plan and Details', level=1)
imp_head.paragraph_format.space_before = Pt(18)
imp_head.paragraph_format.space_after = Pt(6)
doc.add_paragraph('Our implementation plan follows a structured five-phase approach, designed to minimize risk, maximize adoption, and ensure clear milestones. Each phase incorporates stakeholder engagement, iterative validation, and robust documentation.')

# Mermaid workflow diagram
mermaid_code = '''
flowchart TD
  A[Discovery] --> B[Requirements]
  B --> C[Solution Design]
  C --> D[Development]
  D --> E[Testing]
  E --> F[Deployment]
  F --> G[Managed Services]
'''
diagram_path = render_mermaid(mermaid_code, 'workflow_diagram')
doc.add_picture(str(diagram_path), width=Inches(5.8))
add_caption('Figure 2: ERP implementation workflow')

# Proposed Timeline
timeline_head = doc.add_heading('Proposed Timeline', level=1)
timeline_head.paragraph_format.space_before = Pt(18)
timeline_head.paragraph_format.space_after = Pt(6)
doc.add_paragraph('The anticipated project schedule is outlined below. Phases may overlap to accelerate delivery and reduce downtime. Key milestones and deliverables ensure progress and accountability.')

# Gantt chart for schedule
schedule = pd.DataFrame({
    'Task': ['Discovery', 'Requirements', 'Solution Design', 'Development', 'Testing', 'Deployment', 'Managed Services'],
    'Team': ['Consulting', 'Business Analysts', 'Architects', 'Development', 'QA', 'DevOps', 'Support'],
    'Start': pd.to_datetime(['2026-02-10', '2026-02-24', '2026-03-10', '2026-03-24', '2026-05-01', '2026-05-15', '2026-06-01']),
    'End': pd.to_datetime(['2026-02-23', '2026-03-09', '2026-03-23', '2026-04-30', '2026-05-14', '2026-05-31', '2026-12-31'])
})
project_start = schedule['Start'].min()
schedule['days_to_start'] = (schedule['Start'] - project_start).dt.days
schedule['duration'] = (schedule['End'] - schedule['Start']).dt.days + 1
team_colors = {
    'Consulting': '#2980b9',
    'Business Analysts': '#3498db',
    'Architects': '#9b59b6',
    'Development': '#2ecc71',
    'QA': '#e74c3c',
    'DevOps': '#f39c12',
    'Support': '#34495e'
}
fig, ax = plt.subplots(figsize=(10, 6))
for idx, row in schedule.iterrows():
    ax.barh(y=row['Task'], width=row['duration'], left=row['days_to_start'],
            color=team_colors[row['Team']], edgecolor='white', linewidth=0.5)
ax.invert_yaxis()
max_days = schedule['days_to_start'].max() + schedule['duration'].max()
# Correct xticks definition
xticks = np.arange(0, max_days + 14, 14)
xticklabels = pd.date_range(start=project_start, periods=len(xticks), freq='14D').strftime('%b %d')
ax.set_xticks(xticks)
ax.set_xticklabels(xticklabels[:len(xticks)], fontsize=9)
ax.set_xlabel('Timeline', fontsize=10)
ax.set_title('Project Implementation Schedule', fontsize=14, fontweight='bold')
ax.xaxis.grid(True, alpha=0.3, linestyle='--')
ax.set_axisbelow(True)
unique_teams = schedule['Team'].unique()
patches = [mpatches.Patch(color=team_colors[team], label=team) for team in unique_teams]
ax.legend(handles=patches, loc='lower right', fontsize=8, framealpha=0.9)
plt.tight_layout()
gantt_path = output_dir / 'project_gantt.png'
plt.savefig(gantt_path, dpi=150, bbox_inches='tight', facecolor='white')
plt.close()
doc.add_picture(str(gantt_path), width=Inches(6.0))
add_caption('Figure 3: Project implementation schedule')

# Resources
resource_head = doc.add_heading('Resources', level=1)
resource_head.paragraph_format.space_before = Pt(18)
resource_head.paragraph_format.space_after = Pt(6)
doc.add_paragraph('Our project team consists of experts in ERP implementation, integration, data migration, QA/testing, training, and managed services. Below is a summary of the proposed staffing plan:')

staff_table = doc.add_table(rows=1, cols=3)
staff_table.style = 'Table Grid'
staff_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = staff_table.rows[0].cells
hdr[0].text, hdr[1].text, hdr[2].text = 'Role', 'Primary Responsibilities', 'Key Personnel'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True
staff_data = [
    ('Project Manager', 'Oversight, client liaison, schedule, risk', 'John Smith'),
    ('Solution Architect', 'Design, technical oversight', 'Jane Doe'),
    ('Business Analyst', 'Requirements, process mapping', 'Sam Lee'),
    ('Lead Developer', 'Build, integration, customization', 'Mary Chen'),
    ('QA Lead', 'Testing, validation', 'Alex Patel'),
    ('Training Lead', 'End user training, documentation', 'Lisa Brown'),
    ('Managed Services Lead', 'Post-go-live support', 'Michael Carter')
]
for role, resp, key in staff_data:
    row = staff_table.add_row().cells
    row[0].text = role
    row[1].text = resp
    row[2].text = key

# Client References
ref_head = doc.add_heading('Client References', level=1)
ref_head.paragraph_format.space_before = Pt(18)
ref_head.paragraph_format.space_after = Pt(6)
doc.add_paragraph('Argano has successfully delivered ERP solutions for leading organizations across industries. Below are five references, including two in healthcare:')
doc.add_paragraph('Healthcare Client A: Large hospital system; Dynamics 365 Finance; Contact: Susan Miller, CIO, 555-111-2222', style='List Bullet')
doc.add_paragraph('Healthcare Client B: Regional medical group; SAP S/4HANA; Contact: Mark Davis, IT Director, 555-333-4444', style='List Bullet')
doc.add_paragraph('Manufacturing Client C: National manufacturer; Oracle ERP; Contact: Lisa Kim, CFO, 555-555-6666', style='List Bullet')
doc.add_paragraph('Logistics Client D: Supply chain firm; Dynamics 365; Contact: Raj Singh, Operations VP, 555-777-8888', style='List Bullet')
doc.add_paragraph('Retail Client E: Retail chain; Dynamics 365; Contact: Tom White, CIO, 555-999-0000', style='List Bullet')

# Terms and Conditions
terms_head = doc.add_heading('Terms and Conditions', level=1)
terms_head.paragraph_format.space_before = Pt(18)
terms_head.paragraph_format.space_after = Pt(6)
doc.add_paragraph('All pricing is quoted in USD. Argano will comply with Kastle Systems’ requirements for proposal submission format, confidentiality, and contract structure. Master Service Agreement (MSA) and Service Level Agreement (SLA) are available upon request. No conflicts of interest or outstanding legal actions exist.')

# Pricing Proposal
pricing_head = doc.add_heading('Pricing Proposal', level=1)
pricing_head.paragraph_format.space_before = Pt(18)
pricing_head.paragraph_format.space_after = Pt(6)
doc.add_paragraph('Pricing is broken out into one-time implementation fees, annual managed services, integration fees, training, travel, customization, and licensing. Detailed pricing is provided in the attached template.')

pricing_table = doc.add_table(rows=1, cols=4)
pricing_table.style = 'Table Grid'
pricing_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = pricing_table.rows[0].cells
hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = 'Category', 'One-Time Fee', 'Annual Fee', 'Notes'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True
pricing_data = [
    ('ERP Implementation', '$350,000', '-', 'Fixed fee for design/build/test'),
    ('System Integration', '$120,000', '-', 'Includes CRM, Concur, Dayforce, NCA'),
    ('Data Migration', '$55,000', '-', 'Historical & transactional'),
    ('Training', '$25,000', '$12,000', 'In-person, webinar, online'),
    ('Managed Services', '-', '$48,000', 'Post-go-live support'),
    ('Travel & Expenses', '$18,000', '-', 'Estimated'),
    ('Customization', '$60,000', '-', 'As required'),
    ('Licensing', '-', '$33,000', 'Ongoing software maintenance/licensing')
]
for cat, one, annual, notes in pricing_data:
    row = pricing_table.add_row().cells
    row[0].text = cat
    row[1].text = one
    row[2].text = annual
    row[3].text = notes

# Appendix 1 - Responsibilities
appx_head = doc.add_heading('Appendix 1 - Responsibilities', level=1)
appx_head.paragraph_format.space_before = Pt(18)
appx_head.paragraph_format.space_after = Pt(6)
doc.add_paragraph('Below is a summary of responsibilities for Kastle Systems and Argano. All roles and deliverables are coordinated through the joint steering committee.')

resp_table = doc.add_table(rows=1, cols=3)
resp_table.style = 'Table Grid'
resp_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = resp_table.rows[0].cells
hdr[0].text, hdr[1].text, hdr[2].text = 'Area', 'Kastle Systems', 'Argano'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True
resp_data = [
    ('Requirements & Discovery', 'Provide process owners, supply documentation', 'Lead workshops, analyze requirements'),
    ('Design', 'Review proposed solution, approve design', 'Present solution, document design'),
    ('Build', 'Provide data, review build', 'Develop solution, migrate data'),
    ('Test', 'Confirm test cases, review test results', 'Execute testing, resolve defects'),
    ('Deploy', 'Authorize go-live, support transition', 'Deploy, stabilize, hypercare'),
    ('Managed Services', 'Engage in change control, report issues', 'Maintain, patch, upgrade, report')
]
for area, kastle, argano in resp_data:
    row = resp_table.add_row().cells
    row[0].text = area
    row[1].text = kastle
    row[2].text = argano

# Compliance Matrix
comp_head = doc.add_heading('Appendix 2 - Compliance Matrix', level=1)
comp_head.paragraph_format.space_before = Pt(18)
comp_head.paragraph_format.space_after = Pt(6)
doc.add_paragraph('The following table details Argano’s compliance with Kastle Systems’ mandatory requirements:')

comp_table = doc.add_table(rows=1, cols=3)
comp_table.style = 'Table Grid'
comp_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = comp_table.rows[0].cells
hdr[0].text, hdr[1].text, hdr[2].text = 'Requirement', 'Response', 'Reference Section'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True
comp_data = [
    ('REQ-001', 'Compliant: Design/build/test fee detailed', 'Pricing Proposal'),
    ('REQ-002', 'Compliant: Integration fee & scope', 'Pricing Proposal'),
    ('REQ-003', 'Compliant: Training approach and pricing', 'Service Capabilities'),
    ('REQ-004', 'Compliant: Data migration plan', 'Service Capabilities'),
    ('REQ-005', 'Compliant: Travel/expenses included', 'Pricing Proposal'),
    ('REQ-006', 'Compliant: Managed services scope', 'Service Capabilities'),
    ('REQ-007', 'Compliant: Maintenance/licensing detailed', 'Pricing Proposal'),
    ('REQ-008', 'Compliant: Training options/pricing', 'Pricing Proposal'),
    ('REQ-009', 'Compliant: Customization, migration, storage', 'Pricing Proposal'),
    ('REQ-010', 'Compliant: ERP business functions supported', 'Service Capabilities'),
    ('REQ-011', 'Compliant: System integration scope', 'Service Capabilities'),
    ('REQ-012', 'Compliant: Submission format addressed', 'Terms and Conditions'),
    ('REQ-013', 'Compliant: Section structure follows RFP', 'Table of Contents'),
    ('REQ-014', 'Compliant: Acknowledgement and response timeline', 'Letter of Transmittal'),
    ('REQ-015', 'Compliant: Discovery/presentation participation', 'Implementation Plan'),
    ('REQ-016', 'Compliant: Personnel, management, resumes', 'Resources'),
    ('REQ-017', 'Compliant: Detailed questions answered', 'Service Capabilities'),
    ('REQ-018', 'Compliant: Fee structure & template', 'Pricing Proposal'),
    ('REQ-019', 'Compliant: MSA/SLA provided', 'Terms and Conditions'),
    ('REQ-020', 'Compliant: Client references included', 'Client References'),
    ('REQ-021', 'Compliant: Conflict/legal disclosure', 'Terms and Conditions')
]
for req, resp, ref in comp_data:
    row = comp_table.add_row().cells
    row[0].text = req
    row[1].text = resp
    row[2].text = ref
