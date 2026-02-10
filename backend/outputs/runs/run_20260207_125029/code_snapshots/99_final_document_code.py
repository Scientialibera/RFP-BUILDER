from docx.enum.style import WD_STYLE_TYPE

styles = doc.styles

# Configure base styles
normal = styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)

# Caption style helper
if 'Caption' not in [s.name for s in styles]:
    cap = styles.add_style('Caption', WD_STYLE_TYPE.PARAGRAPH)
    cap.font.name = 'Calibri'
    cap.font.size = Pt(9)
    cap.font.italic = True


def add_caption(text: str):
    p = doc.add_paragraph(text, style='Caption')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


figure_counter = {'n': 0}


def next_figure(label: str) -> str:
    figure_counter['n'] += 1
    return f"Figure {figure_counter['n']}: {label}"


# -------------------------
# Title Page
# -------------------------

title = doc.add_heading('Proposal for Phase 1 Sanitary Sewer Evaluation Survey (SSES)\nLaSalle Sewersheds – Niagara Falls Water Board', level=0)
para = doc.add_paragraph('Submitted by Aurelia Digital Systems')
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
para = doc.add_paragraph('Professional Engineering Services for Phase 1 SSES, Engineering Report, Corrective Action Plan, and Post-Construction Monitoring\nwith Options for Phase 2 and Phase 3')
para.alignment = WD_ALIGN_PARAGRAPH.CENTER

# -------------------------
# Executive Summary and Win Themes
# -------------------------

doc.add_heading('Executive Summary and Win Themes', level=1)

p = doc.add_paragraph()
p.add_run(
    'The Niagara Falls Water Board (NFWB) is implementing a multi-phase sanitary sewer '
    'evaluation survey (SSES) program in the LaSalle sewersheds to eliminate sanitary sewer '
    'overflows (SSOs), reduce inflow and infiltration (I&I), and demonstrate compliance with '
    'its Order on Consent and New York State Department of Environmental Conservation '
    '(NYSDEC) expectations. Phase 1 focuses on the 91st & Luick and 81st & Frontier 1 '
    'sewersheds, with subsequent Phases 2 and 3 extending the program to additional LaSalle '
    'areas.'
)

p = doc.add_paragraph(
    'Aurelia Digital Systems (Aurelia) proposes to provide comprehensive professional '
    'engineering services to plan and execute Phase 1 SSES, prepare an Engineering Report '
    'and corrective action plan acceptable to NYSDEC, and perform post-construction '
    'monitoring to document I&I reduction by June 2024. Our approach is explicitly aligned '
    'to NYSDEC’s Engineering Report Outline for Wastewater Infrastructure Projects '
    '(October 1, 2022), NFWB’s Order on Consent milestones, and all requirements defined '
    'in this Request for Proposals (RFP).'
)

p = doc.add_paragraph(
    'Our win strategy is to be NFWB’s lowest-risk, highest-value partner by combining deep '
    'technical capability in sewer system evaluation with rigorous project management, '
    'strong governance, and advanced data and analytics. We will integrate flow, rainfall, '
    'CCTV, smoke testing, and weiring data in a unified analytical framework to target the '
    'most cost-effective corrective actions and to clearly demonstrate I&I reduction in the '
    'post-construction I&I Assessment and Cost Analysis Report.'
)

# Win themes bullets

doc.add_paragraph(
    'NYSDEC-compliant, audit-ready documentation mapped to the Engineering Report Outline '
    'and EPG requirements, including PE-stamped deliverables.',
    style='List Bullet'
)
doc.add_paragraph(
    'On-time delivery of Phase 1 SSES, Engineering Report, and corrective action plan to '
    'support early 2023 NYSDEC review, 2023–early 2024 construction, and June 2024 '
    'post-construction monitoring.',
    style='List Bullet'
)
doc.add_paragraph(
    'Data-driven I&I identification and reduction using integrated analysis of flow, '
    'rainfall, CCTV, smoke testing, and nighttime weiring results.',
    style='List Bullet'
)
doc.add_paragraph(
    'Strong project governance and communication leveraging Aurelia’s Initiate–Discover–'
    'Design–Build–Deploy–Operate lifecycle with clear stage gates and RACI.',
    style='List Bullet'
)
doc.add_paragraph(
    'Robust safety, security, and information governance practices that reduce operational '
    'and reputational risk for NFWB.',
    style='List Bullet'
)
doc.add_paragraph(
    'Concrete plan to meet or exceed MWBE (30% combined) and SDVOB (6%) participation '
    'goals with transparent utilization tracking and reporting.',
    style='List Bullet'
)
doc.add_paragraph(
    'Transparent pricing and effort estimation with not-to-exceed fees, clearly defined '
    'contingency for focused SSES activities, and options for Phases 2 and 3.',
    style='List Bullet'
)
doc.add_paragraph(
    'Experienced, multidisciplinary team with proven success delivering sewer system '
    'evaluations and public-sector infrastructure projects.',
    style='List Bullet'
)

# High-level value table

doc.add_paragraph('High-Level Value to NFWB', style='Heading 2')

value_table = doc.add_table(rows=1, cols=2)
value_table.style = 'Table Grid'
value_hdr = value_table.rows[0].cells
value_hdr[0].text = 'Value Theme'
value_hdr[1].text = 'Benefit to NFWB'
for cell in value_hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

rows = [
    ('Schedule certainty', 'Aligned to early 2023 NYSDEC review, 2023–early 2024 construction, and June 2024 post-construction monitoring milestones.'),
    ('NYSDEC compliance', 'Engineering Report and Post-Construction Report structured to the NYS Engineering Report Outline and EPG guidance, PE-stamped and audit-ready.'),
    ('I&I reduction', 'Targeted corrective actions prioritized by cost per gallon removed and ability to eliminate SSOs at 91st & Luick and 81st & Frontier 1.'),
    ('MWBE/SDVOB participation', 'Credible utilization plan with identified roles for certified partners and monthly reporting cadence.'),
]

for theme, benefit in rows:
    row_cells = value_table.add_row().cells
    row_cells[0].text = theme
    row_cells[1].text = benefit

# Compliance cross-walk (selected key requirements)

doc.add_paragraph('Summary of Key Mandatory Requirements and Proposal Cross-Reference', style='Heading 2')

comp_table = doc.add_table(rows=1, cols=3)
comp_table.style = 'Table Grid'
comp_hdr = comp_table.rows[0].cells
comp_hdr[0].text = 'Requirement ID'
comp_hdr[1].text = 'Requirement Summary'
comp_hdr[2].text = 'Proposal Section'
for cell in comp_hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

comp_rows = [
    ('REQ-001', 'Phase 1 SSES in 91st & Luick and 81st & Frontier 1 sewersheds', 'Phase 1 SSES Technical Approach and Scope of Services'),
    ('REQ-002', 'Option prices, schedules, and scopes for Phases 2 and 3', 'Options for Phase 2 and Phase 3 SSES (Scope, Schedule, and Fees)'),
    ('REQ-003', 'Compliance with NYSDEC, EPG, and Engineering Report Outline', 'Regulatory and Standards Compliance Approach'),
    ('REQ-016', 'Phase 1 SSES report with corrective action plan', 'Phase 1 Engineering Analysis, Corrective Action Planning, and SSES Report'),
    ('REQ-017', 'Schedule aligned to early 2023 NYSDEC review and June 2024 monitoring', 'Schedule, Level of Effort, and Budget'),
    ('REQ-020', 'Comprehensive proposal content including schedule, effort, and budget', 'Multiple sections including Approach, Schedule, Team, and Options'),
]

for rid, summary, section in comp_rows:
    row_cells = comp_table.add_row().cells
    row_cells[0].text = rid
    row_cells[1].text = summary
    row_cells[2].text = section

# Visual 1: Flowchart linking Phase 1 SSES to construction and monitoring (Mermaid)

mermaid_code_exec = '''
flowchart LR
  A[Phase 1 SSES field work] --> B[Phase 1 engineering analysis]
  B --> C[Phase 1 Engineering Report]
  C --> D[NYSDEC review]
  D --> E[Design and construction of improvements]
  E --> F[Post construction monitoring]
  F --> G[Post construction I and I report]
'''
workflow_path_exec = render_mermaid(mermaid_code_exec, 'phase1_lifecycle')
doc.add_picture(str(workflow_path_exec), width=Inches(5.8))
add_caption(next_figure('Linkage between Phase 1 SSES, corrective actions, and post-construction monitoring'))

# Visual 2: High-level bar chart comparing planned timelines for Phases 1–3

import matplotlib.pyplot as plt
import seaborn as sns

sns.set_style('whitegrid')

phase_timeline = pd.DataFrame({
    'Phase': ['Phase 1', 'Phase 2', 'Phase 3'],
    'Months': [18, 16, 16],
})

plt.figure(figsize=(6, 3.5))
ax = sns.barplot(data=phase_timeline, x='Phase', y='Months', color='#4c72b0')
ax.set_ylabel('Approximate Duration (months)')
ax.set_xlabel('')
ax.set_title('Indicative Phase Durations (Design through Monitoring)')
for container in ax.containers:
    ax.bar_label(container, fmt='%d', padding=3)
plt.tight_layout()
phase_timeline_path = output_dir / 'phase_timeline_overview.png'
plt.savefig(phase_timeline_path, dpi=150, bbox_inches='tight', facecolor='white')
plt.close()

doc.add_picture(str(phase_timeline_path), width=Inches(5.8))
add_caption(next_figure('Indicative durations for Phases 1, 2, and 3'))

# -------------------------
# Understanding of NFWB Objectives, LaSalle Sewersheds, and Regulatory Context
# -------------------------

doc.add_heading('Understanding of NFWB Objectives, LaSalle Sewersheds, and Regulatory Context', level=1)

p = doc.add_paragraph(
    'NFWB’s overarching objective for the LaSalle sewersheds is to eliminate SSOs and '
    'achieve sustainable reductions in I&I such that the collection system and treatment '
    'facilities can reliably convey and treat wet-weather flows while complying with the '
    'Order on Consent and NYSDEC requirements. The LaSalle area includes multiple '
    'sewersheds that discharge to the Niagara Falls Wastewater Treatment Plant and '
    'historically have experienced wet-weather SSOs at 91st & Luick and 81st & Frontier.'
)

p = doc.add_paragraph(
    'Phase 1 focuses on the 91st & Luick and 81st & Frontier 1 sewersheds, which have the '
    'most acute SSO history and where prior I&I reduction efforts have already been '
    'implemented. Phases 2 and 3 extend similar SSES activities to the 81st & Frontier 2, '
    '80th & Lindbergh, and Mang & 88th sewersheds to complete a comprehensive LaSalle '
    'program. Each phase must be sequenced to support design and construction of '
    'improvements and subsequent post-construction monitoring within NYSDEC’s expected '
    'timeframes.'
)

p = doc.add_paragraph(
    'We understand that NYSDEC expects Engineering Reports prepared under this program '
    'to follow the October 1, 2022 Engineering Report Outline, including documentation of '
    'existing conditions, regulatory drivers, alternatives analysis, recommended alternative, '
    'implementation schedule, and financial considerations. Our Phase 1 SSES report and '
    'Post-Construction I&I Assessment and Cost Analysis Report will be structured to this '
    'outline and will clearly demonstrate how recommended corrective actions address '
    'LaSalle SSO drivers.'
)

# Sewershed summary table

doc.add_paragraph('Summary of LaSalle Sewersheds by Phase', style='Heading 2')

sewer_table = doc.add_table(rows=1, cols=5)
sewer_table.style = 'Table Grid'
sh = sewer_table.rows[0].cells
sh[0].text = 'Phase'
sh[1].text = 'Sewershed(s)'
sh[2].text = 'Representative Outfall/SSO'
sh[3].text = 'Key Issues'
sh[4].text = 'Prior I&I Work (Illustrative)'
for cell in sh:
    for run in cell.paragraphs[0].runs:
        run.bold = True

sewer_rows = [
    ('Phase 1', '91st & Luick; 81st & Frontier 1', 'SSOs at 91st & Luick and 81st & Frontier 1', 'Wet-weather SSOs, high RDII, surcharging in local mains.', 'Targeted repairs, limited CCTV and manhole inspections, pump station control adjustments.'),
    ('Phase 2', '81st & Frontier 2; 80th & Lindbergh', 'Downstream of Frontier and Lindbergh areas', 'Capacity constraints, suspected private inflow sources.', 'Spot repairs, previous smoke testing in limited areas.'),
    ('Phase 3', 'Mang & 88th', 'Outfall near Mang & 88th', 'Localized basement backups, high groundwater I&I.', 'Isolated manhole rehabilitation, limited data on private inflow.'),
]

for phase, sewersheds, outfall, issues, prior in sewer_rows:
    row_cells = sewer_table.add_row().cells
    row_cells[0].text = phase
    row_cells[1].text = sewersheds
    row_cells[2].text = outfall
    row_cells[3].text = issues
    row_cells[4].text = prior

# Visual 3: Conceptual I&I vs rainfall line chart

rain_data = pd.DataFrame({
    'Storm Event': [f'E{i}' for i in range(1, 8)],
    'Rainfall_inches': [0.3, 0.8, 1.2, 0.5, 1.5, 0.9, 1.8],
    'Peak_Flow_MGD': [8.0, 10.5, 13.0, 9.0, 15.5, 11.0, 17.0],
})

plt.figure(figsize=(6, 3.5))
ax = plt.gca()
ax2 = ax.twinx()
ax.plot(rain_data['Storm Event'], rain_data['Rainfall_inches'], marker='o', color='#4c72b0', label='Rainfall (in)')
ax2.plot(rain_data['Storm Event'], rain_data['Peak_Flow_MGD'], marker='s', color='#dd8452', label='Peak flow (MGD)')
ax.set_xlabel('Storm event')
ax.set_ylabel('Rainfall (inches)', color='#4c72b0')
ax2.set_ylabel('Peak flow (MGD)', color='#dd8452')
ax.set_title('Illustrative Relationship Between Rainfall and Peak Flow')
lines, labels = ax.get_legend_handles_labels()
lines2, labels2 = ax2.get_legend_handles_labels()
ax2.legend(lines + lines2, labels + labels2, loc='upper left', fontsize=8)
plt.tight_layout()
rii_path = output_dir / 'conceptual_rainfall_flow.png'
plt.savefig(rii_path, dpi=150, bbox_inches='tight', facecolor='white')
plt.close()

doc.add_picture(str(rii_path), width=Inches(5.8))
add_caption(next_figure('Conceptual relationship between rainfall and peak wet-weather flow'))

# -------------------------
# Regulatory and Standards Compliance Approach
# -------------------------

doc.add_heading('Regulatory and Standards Compliance Approach (NYSDEC, EPG, Engineering Report Outline)', level=1)

p = doc.add_paragraph(
    'Aurelia will structure all Phase 1 deliverables to fully comply with NYSDEC and '
    'Engineering Planning Guidance (EPG) requirements, including the Engineering Report '
    'Outline for NYS Wastewater Infrastructure Projects dated October 1, 2022. The Phase 1 '
    'SSES report and the Post-Construction I&I Assessment and Cost Analysis Report will be '
    'prepared, stamped, and dated by a New York State licensed Professional Engineer and '
    'will include sufficient documentation and analysis to support NYSDEC review and '
    'funding decisions.'
)

p = doc.add_paragraph(
    'Our internal quality management process includes structured technical reviews, '
    'compliance checklists mapped to the Engineering Report Outline, and formal PE review '
    'and approval prior to submittal. We will coordinate with NFWB and NYSDEC at key '
    'milestones to confirm expectations and minimize the risk of rework or schedule delay.'
)

# Compliance matrix mapping Engineering Report Outline elements

doc.add_paragraph('Engineering Report Outline Compliance Matrix (Excerpt)', style='Heading 2')

eng_table = doc.add_table(rows=1, cols=3)
eng_table.style = 'Table Grid'
eh = eng_table.rows[0].cells
eh[0].text = 'Engineering Report Outline Element'
eh[1].text = 'Planned Report Section'
eh[2].text = 'Responsible Role'
for cell in eh:
    for run in cell.paragraphs[0].runs:
        run.bold = True

eng_rows = [
    ('Project background and purpose', 'Introduction and Regulatory Context', 'Project Manager / Lead PE'),
    ('Description of existing facilities and service area', 'Existing System and LaSalle Sewershed Description', 'Collection System Engineer'),
    ('Regulatory and compliance requirements', 'Regulatory and Standards Compliance Approach', 'Lead PE'),
    ('I&I analysis and SSES findings', 'Phase 1 Engineering Analysis and Findings', 'Hydraulic / I&I Specialist'),
    ('Alternatives analysis and recommended improvements', 'Corrective Action Planning and Alternatives Evaluation', 'Lead PE with Cost Estimator'),
    ('Implementation schedule', 'Implementation Plan and Schedule', 'Project Manager'),
]

for element, section, role in eng_rows:
    row_cells = eng_table.add_row().cells
    row_cells[0].text = element
    row_cells[1].text = section
    row_cells[2].text = role

# Visual 4: Internal quality and compliance review flowchart (Mermaid)

mermaid_code_comp = '''
flowchart TD
  A[Draft technical content] --> B[Internal technical review]
  B --> C[Compliance checklist]
  C --> D[PE review and revisions]
  D --> E[NFWB review]
  E --> F[Final PE stamp]
  F --> G[Submittal to NYSDEC]
'''
comp_path = render_mermaid(mermaid_code_comp, 'compliance_flow')
doc.add_picture(str(comp_path), width=Inches(5.8))
add_caption(next_figure('Internal quality and compliance review process for PE-stamped deliverables'))

# -------------------------
# Phase 1 SSES Technical Approach and Scope of Services
# -------------------------

doc.add_heading('Phase 1 SSES Technical Approach and Scope of Services', level=1)

p = doc.add_paragraph(
    'Our Phase 1 SSES approach is structured into clear tasks that align with the RFP '
    'requirements and NFWB’s Order on Consent milestones. We will perform a comprehensive '
    'investigation of the 91st & Luick and 81st & Frontier 1 sewersheds, including data '
    'review, field investigations (smoke testing, nighttime flow isolation and weiring, '
    'CCTV and manhole inspections, focused SSES activities), outfall/SSO assessment, '
    'engineering analysis and corrective action planning, and post-construction monitoring '
    'and reporting.'
)

# Task-by-task scope table (high level)

doc.add_paragraph('Phase 1 Tasks and Scope Summary', style='Heading 2')

scope_table = doc.add_table(rows=1, cols=4)
scope_table.style = 'Table Grid'
shdr = scope_table.rows[0].cells
shdr[0].text = 'Task'
shdr[1].text = 'Objective'
shdr[2].text = 'Key Activities'
shdr[3].text = 'Primary Deliverables'
for cell in shdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

scope_rows = [
    ('Task 1 – Project initiation and data review', 'Confirm scope and gain understanding of LaSalle sewersheds.', 'Kickoff meeting, review Order on Consent, prior reports, CCTV, manhole, SCADA, and NYSDEC correspondence.', 'Kickoff meeting notes, data gap assessment, refined work plan.'),
    ('Task 2 – Smoke testing', 'Identify inflow sources such as cross-connections and illicit connections.', 'Plan test areas, notifications, field smoke testing, defect logging with photos and GIS coordinates.', 'Smoke testing field logs, GIS feature classes, summary memo.'),
    ('Task 3 – Nighttime flow isolation and weiring', 'Quantify groundwater I&I by sewer reach.', 'Select weir locations, perform nighttime isolation, measure flows, calculate gpd/lf.', 'Weir measurement database, summary figures, infiltration rates by reach.'),
    ('Task 4 – CCTV and manhole assessment', 'Assess structural and maintenance defects and confirm I&I sources.', 'Review existing PACP data, coordinate additional CCTV by NFWB crews, update condition ratings.', 'Updated PACP database, defect summaries, rehabilitation candidates.'),
    ('Task 5 – Outfall/SSO assessment', 'Characterize performance of 91st & Luick and 81st & Frontier 1 outfalls.', 'Develop and execute Work Plan, monitor flows, document conditions.', 'Outfall assessment technical memo and GIS-compatible data.'),
    ('Task 6 – Focused SSES (contingency)', 'Investigate specific I&I sources using targeted methods.', 'Dye testing, house inspections, targeted CCTV using $10,000 contingency upon NFWB approval.', 'Focused SSES logs, findings, and recommendations.'),
    ('Task 7 – Engineering analysis and SSES report', 'Develop corrective action plan and PE-stamped SSES report.', 'Integrate data, evaluate alternatives, develop budgetary costs and schedule.', 'Phase 1 SSES Engineering Report and corrective action plan.'),
    ('Task 8 – Post-construction monitoring and reporting', 'Document I&I reduction following construction.', 'Flow and rainfall monitoring, data normalization, cost-effectiveness analysis.', 'Post-Construction I&I Assessment and Cost Analysis Report.'),
]

for task, obj, acts, deliv in scope_rows:
    row_cells = scope_table.add_row().cells
    row_cells[0].text = task
    row_cells[1].text = obj
    row_cells[2].text = acts
    row_cells[3].text = deliv

# Visual 5: High-level Gantt-style chart for Phase 1 tasks

import matplotlib.patches as mpatches
import numpy as np
import pandas as pd
import datetime as dt

schedule = pd.DataFrame({
    'Task': [
        'Initiation and data review',
        'Smoke testing',
        'Nighttime weiring',
        'CCTV and manhole assessment',
        'Outfall and SSO assessment',
        'Focused SSES (contingency)',
        'Engineering analysis and SSES report',
        'Post construction monitoring',
    ],
    'Team': [
        'PM and PE',
        'Field crew',
        'Field crew',
        'Field and analysis',
        'Field and analysis',
        'Field crew',
        'Engineering',
        'Field and analysis',
    ],
    'Start': pd.to_datetime([
        '2022-12-01',
        '2023-02-01',
        '2023-03-01',
        '2023-02-15',
        '2023-03-15',
        '2023-04-01',
        '2023-04-15',
        '2024-03-01',
    ]),
    'End': pd.to_datetime([
        '2023-01-15',
        '2023-03-15',
        '2023-03-31',
        '2023-04-15',
        '2023-04-30',
        '2023-06-15',
        '2023-06-30',
        '2024-06-30',
    ]),
})

project_start = schedule['Start'].min()
schedule['days_to_start'] = (schedule['Start'] - project_start).dt.days
schedule['duration'] = (schedule['End'] - schedule['Start']).dt.days + 1

team_colors = {
    'PM and PE': '#4c72b0',
    'Field crew': '#55a868',
    'Field and analysis': '#c44e52',
    'Engineering': '#8172b3',
}

fig, ax = plt.subplots(figsize=(8, 4.5))
for idx, row in schedule.iterrows():
    ax.barh(y=row['Task'], width=row['duration'], left=row['days_to_start'],
            color=team_colors[row['Team']], edgecolor='white', linewidth=0.5)

ax.invert_yaxis()
max_days = schedule['days_to_start'].max() + schedule['duration'].max()
xticks = np.arange(0, max_days + 30, 30)
xticklabels = pd.date_range(start=project_start, periods=len(xticks), freq='30D').strftime('%b %Y')
ax.set_xticks(xticks[:len(xticklabels)])
ax.set_xticklabels(xticklabels, fontsize=8)
ax.set_xlabel('Timeline')
ax.set_title('Phase 1 SSES – High-Level Implementation Schedule')
ax.xaxis.grid(True, alpha=0.3, linestyle='--')
ax.set_axisbelow(True)

unique_teams = schedule['Team'].unique()
patches = [mpatches.Patch(color=team_colors[t], label=t) for t in unique_teams]
ax.legend(handles=patches, loc='lower right', fontsize=8, framealpha=0.9)

plt.tight_layout()
phase1_gantt_path = output_dir / 'phase1_gantt.png'
plt.savefig(phase1_gantt_path, dpi=150, bbox_inches='tight', facecolor='white')
plt.close()

doc.add_picture(str(phase1_gantt_path), width=Inches(6.0))
add_caption(next_figure('High-level implementation schedule for Phase 1 SSES'))

# -------------------------
# Data Review, Existing Information Assessment, and System Familiarization
# -------------------------

doc.add_heading('Data Review, Existing Information Assessment, and System Familiarization', level=1)

p = doc.add_paragraph(
    'Aurelia will begin Phase 1 with a structured data review and system familiarization '
    'process. We will compile and review all available information, including the Order on '
    'Consent, prior engineering reports, NYSDEC correspondence, existing CCTV and manhole '
    'inspection data (NASSCO PACP), SCADA data for lift stations, and any previous I&I '
    'studies in the LaSalle sewersheds. This review will identify data gaps and inform the '
    'targeting of field investigations.'
)

p = doc.add_paragraph(
    'We will hold a kickoff and technical workshops with NFWB staff to discuss system '
    'history, operational practices (including wet-weather operation of lift stations), '
    'known problem areas, and candidate locations for flow isolation and weiring. Our team '
    'will use GIS to integrate available mapping, asset data, and field results to support '
    'efficient planning and clear communication with NFWB and NYSDEC.'
)

# Data inventory table

doc.add_paragraph('Existing Data Inventory and Intended Use', style='Heading 2')

inv_table = doc.add_table(rows=1, cols=3)
inv_table.style = 'Table Grid'
inv_hdr = inv_table.rows[0].cells
inv_hdr[0].text = 'Data Type'
inv_hdr[1].text = 'Source'
inv_hdr[2].text = 'Intended Use in Analysis'
for cell in inv_hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

inv_rows = [
    ('Order on Consent and NYSDEC correspondence', 'NFWB / NYSDEC', 'Define regulatory drivers, milestones, and required reporting content.'),
    ('System mapping and GIS data', 'NFWB', 'Define sewershed boundaries, asset locations, and connectivity for analysis and mapping.'),
    ('CCTV inspection reports and videos (PACP)', 'NFWB / prior consultants', 'Assess structural and O&M defects, confirm I&I sources, prioritize rehabilitation.'),
    ('Manhole inspection data', 'NFWB / prior consultants', 'Identify defects such as inflow sources, frame and cover issues, and groundwater infiltration.'),
    ('SCADA data for lift stations', 'NFWB', 'Understand wet-weather operations, pump run times, and system response to storms.'),
    ('Flow monitoring and rainfall data', 'NFWB / third-party meters', 'Quantify RDII, calibrate I&I models, and support pre- and post-construction comparisons.'),
]

for dtype, source, use in inv_rows:
    row_cells = inv_table.add_row().cells
    row_cells[0].text = dtype
    row_cells[1].text = source
    row_cells[2].text = use

# -------------------------
# Smoke Testing Plan and Methodology
# -------------------------

doc.add_heading('Smoke Testing Plan and Methodology', level=1)

p = doc.add_paragraph(
    'Smoke testing is a core component of the Phase 1 SSES and will be used to identify '
    'direct inflow sources such as cross-connections, defective cleanouts, roof and yard '
    'drains, and other illicit connections. Aurelia will provide all staff, equipment, and '
    'materials necessary to perform smoke testing in the 91st & Luick and 81st & Frontier 1 '
    'sewersheds, in coordination with NFWB.'
)

p = doc.add_paragraph(
    'We will select test areas based on sewershed hydraulics, known SSO locations, and '
    'existing data. Testing will be conducted under dry weather conditions with advance '
    'resident notifications (2–4 days), coordination with local police and fire departments, '
    'and clear public information materials. Field crews will document each observed defect '
    'with photographs, GPS coordinates, and standardized defect codes, and will enter '
    'results into a GIS database for use in the engineering analysis.'
)

# Smoke testing procedures table

doc.add_paragraph('Smoke Testing Procedures and Controls', style='Heading 2')

smoke_table = doc.add_table(rows=1, cols=3)
smoke_table.style = 'Table Grid'
smh = smoke_table.rows[0].cells
smh[0].text = 'Step'
smh[1].text = 'Description'
smh[2].text = 'Key Controls'
for cell in smh:
    for run in cell.paragraphs[0].runs:
        run.bold = True

smoke_rows = [
    ('Planning', 'Define test areas and schedule based on sewershed hydraulics and weather.', 'Confirm dry weather window, coordinate with NFWB operations.'),
    ('Public notification', 'Distribute door hangers and public notices 2–4 days in advance.', 'Use clear messaging, provide NFWB contact, document addresses notified.'),
    ('Agency coordination', 'Notify police, fire, and other stakeholders prior to testing.', 'Provide maps and schedule, maintain contact list in field vehicles.'),
    ('Field setup', 'Stage blowers, smoke, and safety equipment.', 'Conduct tailgate safety briefing, verify PPE, review confined space procedures as applicable.'),
    ('Testing', 'Introduce smoke into mains and observe emergence points.', 'Document each observation with photo, GPS, and standardized defect code.'),
    ('Data entry', 'Upload field logs and photos to GIS database.', 'QA/QC by supervisor, cross-check with mapping and address lists.'),
]

for step, desc, ctrl in smoke_rows:
    row_cells = smoke_table.add_row().cells
    row_cells[0].text = step
    row_cells[1].text = desc
    row_cells[2].text = ctrl

# -------------------------
# Night-Time Flow Isolation and Measurement (Weiring) Approach
# -------------------------

doc.add_heading('Night-Time Flow Isolation and Measurement (Weiring) Approach', level=1)

p = doc.add_paragraph(
    'Nighttime flow isolation and weiring will be used to quantify groundwater infiltration '
    'in gallons per day per linear foot (gpd/lf) for key sewer reaches. Aurelia will review '
    'prior flow data and system configuration to select weir locations that allow effective '
    'isolation of subareas, while maintaining safe operations and access.'
)

p = doc.add_paragraph(
    'Our field teams will provide all necessary staff, equipment, traffic control, and '
    'confined space entry capabilities to perform nighttime measurements. We will coordinate '
    'with NFWB on resident notifications where nighttime work may affect access or create '
    'visible activity. Flow measurements will be recorded at high temporal resolution and '
    'analyzed to compute infiltration rates by reach, supporting prioritization of '
    'rehabilitation measures.'
)

# Candidate weir locations table (conceptual)

doc.add_paragraph('Conceptual Candidate Weir Location Criteria', style='Heading 2')

weir_table = doc.add_table(rows=1, cols=3)
weir_table.style = 'Table Grid'
wh = weir_table.rows[0].cells
wh[0].text = 'Location Type'
wh[1].text = 'Selection Criteria'
wh[2].text = 'Constraints'
for cell in wh:
    for run in cell.paragraphs[0].runs:
        run.bold = True

weir_rows = [
    ('Upstream subarea manhole', 'Isolates a discrete subcatchment, accessible for nighttime work.', 'Traffic control requirements, depth and confined space entry limitations.'),
    ('Downstream trunk manhole', 'Aggregates flows from multiple upstream reaches for screening.', 'Hydraulic grade line, risk of surcharge during unexpected storms.'),
    ('Near SSO location', 'Directly relates infiltration to observed SSO performance.', 'Proximity to residences and noise/light considerations.'),
]

for loc, crit, con in weir_rows:
    row_cells = weir_table.add_row().cells
    row_cells[0].text = loc
    row_cells[1].text = crit
    row_cells[2].text = con

# -------------------------
# CCTV, Manhole, and Focused SSES Activities
# -------------------------

doc.add_heading('CCTV, Manhole, and Focused SSES Activities (Dye Testing, House Inspections)', level=1)

p = doc.add_paragraph(
    'Aurelia will review all existing CCTV and manhole inspection data, including NASSCO '
    'PACP condition coding, to understand structural condition and potential I&I sources. '
    'We will coordinate with NFWB to prioritize additional CCTV work to be performed by '
    'NFWB crews, focusing on reaches with high infiltration or incomplete data.'
)

p = doc.add_paragraph(
    'Focused SSES activities such as dye testing and house inspections will be used '
    'selectively where data indicate likely private inflow sources or complex connectivity. '
    'These activities will be funded from the $10,000 contingency only upon NFWB approval, '
    'with clear documentation of scope, cost, and findings for each authorization.'
)

# NASSCO PACP summary table

doc.add_paragraph('NASSCO PACP Condition Grades and Recommended Actions (Summary)', style='Heading 2')

pacc_table = doc.add_table(rows=1, cols=3)
pacc_table.style = 'Table Grid'
ph = pacc_table.rows[0].cells
ph[0].text = 'PACP Grade'
ph[1].text = 'Condition Description'
ph[2].text = 'Typical Recommended Action'
for cell in ph:
    for run in cell.paragraphs[0].runs:
        run.bold = True

pacc_rows = [
    ('1–2', 'Minor defects, low risk of failure.', 'Routine maintenance, monitor in future inspections.'),
    ('3', 'Moderate defects, potential for future issues.', 'Plan rehabilitation within program horizon, coordinate with other work.'),
    ('4', 'Significant defects, likely structural or I&I issues.', 'Prioritize rehabilitation or replacement in near term.'),
    ('5', 'Severe defects, imminent failure or major I&I.', 'Immediate repair or replacement, potential emergency response.'),
]

for grade, desc, act in pacc_rows:
    row_cells = pacc_table.add_row().cells
    row_cells[0].text = grade
    row_cells[1].text = desc
    row_cells[2].text = act

# -------------------------
# Outfall/SSO Condition Assessment for 91st & Luick and 81st & Frontier 1
# -------------------------

doc.add_heading('Outfall/SSO Condition Assessment for 91st & Luick and 81st & Frontier 1', level=1)

p = doc.add_paragraph(
    'The two Phase 1 outfalls at 91st & Luick and 81st & Frontier 1 are focal points for '
    'LaSalle SSOs and will be assessed in detail. Aurelia will develop a site-specific '
    'Outfall and SSO Assessment Work Plan, including review of existing data, site visits, '
    'and potential deployment of meters, gauges, dye testing, and CCTV in coordination with '
    'NFWB crews.'
)

p = doc.add_paragraph(
    'Assessment activities will be documented in GIS-compatible formats, including '
    'photographs, condition ratings, and flow monitoring results. We will develop '
    'recommendations to reduce extraneous flows and improve performance at each outfall, '
    'coordinated with upstream SSES findings and proposed corrective actions.'
)

# Outfall summary table

doc.add_paragraph('Phase 1 Outfall Assessment Summary', style='Heading 2')

outfall_table = doc.add_table(rows=1, cols=4)
outfall_table.style = 'Table Grid'
oth = outfall_table.rows[0].cells
oth[0].text = 'Outfall'
oth[1].text = 'Existing Issues (Illustrative)'
oth[2].text = 'Assessment Techniques'
oth[3].text = 'Expected Outputs'
for cell in oth:
    for run in cell.paragraphs[0].runs:
        run.bold = True

outfall_rows = [
    ('91st & Luick', 'Wet-weather SSOs, surcharging, evidence of high RDII.', 'Flow monitoring, dye testing, CCTV of connecting sewers, site inspections.', 'Quantified wet-weather performance, identification of key inflow sources.'),
    ('81st & Frontier 1', 'SSOs during moderate to large storms, limited storage.', 'Level/flow monitoring, SCADA data review, CCTV, site inspections.', 'Performance curves, recommended storage or conveyance improvements.'),
]

for name, issues, tech, outputs in outfall_rows:
    row_cells = outfall_table.add_row().cells
    row_cells[0].text = name
    row_cells[1].text = issues
    row_cells[2].text = tech
    row_cells[3].text = outputs

# -------------------------
# Phase 1 Engineering Analysis, Corrective Action Planning, and SSES Report
# -------------------------

doc.add_heading('Phase 1 Engineering Analysis, Corrective Action Planning, and SSES Report', level=1)

p = doc.add_paragraph(
    'Aurelia will integrate all Phase 1 data—including smoke testing, nighttime weiring, '
    'CCTV and manhole inspections, outfall assessments, and focused SSES results—into a '
    'coherent engineering analysis. We will quantify RDII contributions by subarea, '
    'identify root causes of SSOs, and develop a prioritized corrective action plan that '
    'maximizes I&I reduction per dollar invested.'
)

p = doc.add_paragraph(
    'For each candidate corrective action (e.g., cured-in-place lining, point repairs, '
    'manhole rehabilitation, targeted inflow removal, outfall modifications), we will '
    'estimate I&I reduction, capital cost, constructability considerations, and schedule '
    'implications. Alternatives will be evaluated using transparent criteria and documented '
    'in the Phase 1 SSES Engineering Report, which will be structured to the NYS Engineering '
    'Report Outline and stamped by a New York State licensed Professional Engineer.'
)

# Corrective action alternatives matrix (excerpt)

doc.add_paragraph('Corrective Action Alternatives Evaluation (Excerpt)', style='Heading 2')

alt_table = doc.add_table(rows=1, cols=5)
alt_table.style = 'Table Grid'
alth = alt_table.rows[0].cells
alth[0].text = 'Alternative'
alth[1].text = 'Primary Application'
alth[2].text = 'Relative Effectiveness'
alth[3].text = 'Relative Cost'
alth[4].text = 'Typical Use in LaSalle'
for cell in alth:
    for run in cell.paragraphs[0].runs:
        run.bold = True

alt_rows = [
    ('CIPP lining', 'Long segments with multiple defects and high infiltration.', 'High', 'High', 'Targeted trunk and collector sewers with Grade 4–5 defects.'),
    ('Point repairs', 'Isolated structural defects.', 'Moderate', 'Moderate', 'Specific joints or cracks identified on CCTV.'),
    ('Manhole rehabilitation', 'Leaking frames, covers, and walls.', 'Moderate', 'Low to Moderate', 'High groundwater areas with defective manholes.'),
    ('Private inflow removal', 'Roof drains, yard drains, cross-connections.', 'High', 'Variable', 'Residential areas with significant smoke testing defects.'),
]

for alt, app, eff, cost, use in alt_rows:
    row_cells = alt_table.add_row().cells
    row_cells[0].text = alt
    row_cells[1].text = app
    row_cells[2].text = eff
    row_cells[3].text = cost
    row_cells[4].text = use

# -------------------------
# Post-Construction Monitoring and I&I Reduction Assessment
# -------------------------

doc.add_heading('Post-Construction Monitoring and I&I Reduction Assessment', level=1)

p = doc.add_paragraph(
    'Following construction of Phase 1 improvements in 2023–early 2024, Aurelia will '
    'conduct post-construction monitoring to quantify I&I reduction and confirm that SSOs '
    'have been effectively mitigated. Monitoring will include flow and rainfall data '
    'collection at key locations within the 91st & Luick and 81st & Frontier 1 sewersheds, '
    'with data normalized for rainfall and seasonal conditions.'
)

p = doc.add_paragraph(
    'We will compare pre- and post-construction flow metrics (e.g., base infiltration, '
    'RDII volume, peak wet-weather flows) and calculate cost per gallon of I&I removed for '
    'major corrective action categories. Results will be summarized in the Post-Construction '
    'I&I Assessment and Cost Analysis Report for NYSDEC, including identification of the '
    'most effective techniques to inform Phases 2 and 3.'
)

# Pre- and post-construction comparison table

doc.add_paragraph('Illustrative Pre- and Post-Construction Flow Metrics (Conceptual)', style='Heading 2')

post_table = doc.add_table(rows=1, cols=4)
post_table.style = 'Table Grid'
phd = post_table.rows[0].cells
phd[0].text = 'Metric'
phd[1].text = 'Pre-Construction'
phd[2].text = 'Post-Construction'
phd[3].text = 'Change'
for cell in phd:
    for run in cell.paragraphs[0].runs:
        run.bold = True

post_rows = [
    ('Average dry-weather base infiltration (MGD)', '2.0', '1.2', '40% reduction'),
    ('RDII volume for 1-inch storm (MG)', '10.0', '5.5', '45% reduction'),
    ('Peak flow at 91st & Luick (MGD)', '16.0', '9.5', '40% reduction'),
    ('Number of SSOs per year', '6', '0–1', 'Substantial reduction'),
]

for metric, pre, post, change in post_rows:
    row_cells = post_table.add_row().cells
    row_cells[0].text = metric
    row_cells[1].text = pre
    row_cells[2].text = post
    row_cells[3].text = change

# -------------------------
# Options for Phase 2 and Phase 3 SSES (Scope, Schedule, and Fees)
# -------------------------

doc.add_heading('Options for Phase 2 and Phase 3 SSES (Scope, Schedule, and Fees)', level=1)

p = doc.add_paragraph(
    'In addition to the base Phase 1 scope, Aurelia proposes optional services for Phase 2 '
    '(81st & Frontier 2 and 80th & Lindbergh sewersheds) and Phase 3 (Mang & 88th '
    'sewershed). Both options leverage the methods, tools, and lessons learned from Phase 1 '
    'to provide a consistent, efficient multi-phase program for NFWB.'
)

p = doc.add_paragraph(
    'Option 1 (Phase 2) includes all applicable SSES tasks except outfall/SSO assessment, '
    'as no outfall locations are identified in the RFP for Phase 2. Option 2 (Phase 3) '
    'includes all applicable SSES tasks, including outfall/SSO assessment for one outfall. '
    'For each option, we will provide a detailed scope, schedule aligned with NFWB’s '
    'overall LaSalle program, and not-to-exceed fees with task/hour estimate matrices.'
)

# Phase comparison table

doc.add_paragraph('Phase 1, Phase 2, and Phase 3 Scope Comparison', style='Heading 2')

ph_table = doc.add_table(rows=1, cols=4)
ph_table.style = 'Table Grid'
phh = ph_table.rows[0].cells
phh[0].text = 'Element'
phh[1].text = 'Phase 1'
phh[2].text = 'Phase 2 (Option 1)'
phh[3].text = 'Phase 3 (Option 2)'
for cell in phh:
    for run in cell.paragraphs[0].runs:
        run.bold = True

phase_rows = [
    ('Sewersheds', '91st & Luick; 81st & Frontier 1', '81st & Frontier 2; 80th & Lindbergh', 'Mang & 88th'),
    ('Smoke testing', 'Included', 'Included', 'Included'),
    ('Nighttime weiring', 'Included', 'Included', 'Included'),
    ('CCTV and manhole review', 'Included', 'Included', 'Included'),
    ('Outfall/SSO assessment', 'Two outfalls', 'Not applicable', 'One outfall'),
    ('Focused SSES contingency', '$10,000', 'Similar contingency proposed', 'Similar contingency proposed'),
]

for elem, p1, p2, p3 in phase_rows:
    row_cells = ph_table.add_row().cells
    row_cells[0].text = elem
    row_cells[1].text = p1
    row_cells[2].text = p2
    row_cells[3].text = p3

# -------------------------
# Project Management, Governance, and Communication Plan
# -------------------------

doc.add_heading('Project Management, Governance, and Communication Plan', level=1)

p = doc.add_paragraph(
    'Aurelia will designate a Project Manager (PM) as the primary point of contact for '
    'NFWB. The PM will be responsible for overall coordination, schedule and budget '
    'management, risk management, and communication with NFWB and NYSDEC. Our project '
    'governance model follows Aurelia’s Initiate–Discover–Design–Build–Deploy–Operate '
    'lifecycle, with clearly defined stage gates and decision points.'
)

p = doc.add_paragraph(
    'We will attend the mandatory pre-proposal information session, submit any questions '
    'through the authorized contact by the stated deadline, and acknowledge all addenda in '
    'Appendix E Form No. 1. If selected, we are prepared to participate in interviews or '
    'presentations with NFWB staff or the Board to discuss our concepts and methods.'
)

# Project organization table

doc.add_paragraph('Project Organization and Key Roles', style='Heading 2')

org_table = doc.add_table(rows=1, cols=3)
org_table.style = 'Table Grid'
ogh = org_table.rows[0].cells
ogh[0].text = 'Role'
ogh[1].text = 'Primary Responsibilities'
ogh[2].text = 'Representative Individual'
for cell in ogh:
    for run in cell.paragraphs[0].runs:
        run.bold = True

org_rows = [
    ('Project Manager', 'Primary contact, schedule and budget management, coordination with NFWB and NYSDEC.', 'Senior civil engineer with SSES program experience.'),
    ('Lead PE', 'Technical oversight, Engineering Report authorship, PE stamping.', 'New York licensed professional engineer.'),
    ('I&I / Hydraulics Specialist', 'Flow and rainfall analysis, RDII quantification, alternatives evaluation.', 'Senior hydraulic engineer.'),
    ('Field Operations Lead', 'Manage smoke testing, weiring, and focused SSES crews.', 'Field supervisor with SSES experience.'),
    ('GIS / Data Analyst', 'Data integration, mapping, dashboards, and reporting.', 'GIS and data specialist.'),
]

for role, resp, person in org_rows:
    row_cells = org_table.add_row().cells
    row_cells[0].text = role
    row_cells[1].text = resp
    row_cells[2].text = person

# Communication plan table

doc.add_paragraph('Communication and Meeting Plan', style='Heading 2')

comm_table = doc.add_table(rows=1, cols=4)
comm_table.style = 'Table Grid'
ch = comm_table.rows[0].cells
ch[0].text = 'Meeting / Communication'
ch[1].text = 'Frequency'
ch[2].text = 'Participants'
ch[3].text = 'Purpose'
for cell in ch:
    for run in cell.paragraphs[0].runs:
        run.bold = True

comm_rows = [
    ('Project kickoff', 'Once at project start', 'NFWB, Aurelia PM and Lead PE', 'Confirm scope, schedule, roles, and communication protocols.'),
    ('Biweekly progress meetings', 'Every two weeks during active work', 'NFWB project team, Aurelia PM and task leads', 'Review progress, risks, and upcoming activities.'),
    ('Technical workshops', 'At key milestones', 'Subject matter experts from NFWB and Aurelia', 'Review findings, alternatives, and recommendations.'),
    ('NYSDEC coordination calls', 'As needed at major milestones', 'NFWB, Aurelia PM/Lead PE, NYSDEC staff', 'Align on expectations, discuss report content and schedule.'),
]

for mtg, freq, part, purpose in comm_rows:
    row_cells = comm_table.add_row().cells
    row_cells[0].text = mtg
    row_cells[1].text = freq
    row_cells[2].text = part
    row_cells[3].text = purpose

# -------------------------
# Schedule, Level of Effort, and Budget (Including Contingency)
# -------------------------

doc.add_heading('Schedule, Level of Effort, and Budget (Including Contingency)', level=1)

p = doc.add_paragraph(
    'Our schedule is designed to deliver the Phase 1 SSES report in time for NYSDEC review '
    'in early 2023, supporting construction of recommended improvements in 2023–early 2024 '
    'and post-construction monitoring completed by June 2024. We will provide detailed bar '
    'chart schedules and task/hour estimate matrices as part of the final proposal package.'
)

p = doc.add_paragraph(
    'Fees will be proposed on a not-to-exceed basis, with a clearly defined $10,000 '
    'contingency for Task 6 focused SSES activities, to be used only upon NFWB approval. '
    'We will provide monthly invoices with supporting documentation and accept the Water '
    'Board’s 60-day payment terms. Our pricing will exclude sales tax in recognition of '
    'NFWB’s tax-exempt status.'
)

# Task/hour estimate matrix (conceptual)

doc.add_paragraph('Illustrative Task and Level of Effort Matrix (Phase 1)', style='Heading 2')

loetable = doc.add_table(rows=1, cols=5)
loetable.style = 'Table Grid'
leh = loetable.rows[0].cells
leh[0].text = 'Task'
leh[1].text = 'PM Hours'
leh[2].text = 'PE / Engineer Hours'
leh[3].text = 'Field Crew Hours'
leh[4].text = 'GIS / Analyst Hours'
for cell in leh:
    for run in cell.paragraphs[0].runs:
        run.bold = True

loe_rows = [
    ('Task 1 – Initiation and data review', '40', '80', '0', '40'),
    ('Task 2 – Smoke testing', '24', '40', '240', '24'),
    ('Task 3 – Nighttime weiring', '24', '40', '200', '24'),
    ('Task 4 – CCTV and manhole assessment', '24', '80', '80', '40'),
    ('Task 5 – Outfall/SSO assessment', '16', '40', '80', '24'),
    ('Task 6 – Focused SSES (contingency)', '16', '40', '120', '16'),
    ('Task 7 – Engineering analysis and SSES report', '40', '160', '0', '80'),
    ('Task 8 – Post-construction monitoring', '32', '80', '80', '40'),
]

for task, pmh, peh, fch, gish in loe_rows:
    row_cells = loetable.add_row().cells
    row_cells[0].text = task
    row_cells[1].text = pmh
    row_cells[2].text = peh
    row_cells[3].text = fch
    row_cells[4].text = gish

# -------------------------
# Team Qualifications, Experience, and Staffing Plan
# -------------------------

doc.add_heading('Team Qualifications, Experience, and Staffing Plan', level=1)

p = doc.add_paragraph(
    'Aurelia Digital Systems is a privately held engineering and technology firm founded in '
    '2016 with more than 120 professionals across Canada and the United States. Our team '
    'includes civil and environmental engineers, data and analytics specialists, project '
    'managers, and field staff experienced in SSES, I&I reduction programs, and public '
    'sector infrastructure projects.'
)

p = doc.add_paragraph(
    'We will provide a comprehensive statement of qualifications including our business '
    'structure, licensing, years in business, similar project experience with references, '
    'and details for any subcontractors performing more than 10 percent of the work. Key '
    'project professionals will be identified by name, with resumes included in the '
    'Appendices.'
)

# Team qualifications table

doc.add_paragraph('Key Personnel Qualifications (Summary)', style='Heading 2')

team_table = doc.add_table(rows=1, cols=4)
team_table.style = 'Table Grid'
th = team_table.rows[0].cells
th[0].text = 'Role'
th[1].text = 'Education / Licenses'
th[2].text = 'Relevant Experience'
th[3].text = 'Representative Responsibilities'
for cell in th:
    for run in cell.paragraphs[0].runs:
        run.bold = True

team_rows = [
    ('Project Manager', 'B.S. Civil Engineering, NYS PE', '15+ years managing SSES and I&I reduction programs for municipal clients.', 'Overall project management, client coordination, schedule and budget control.'),
    ('Lead PE', 'M.S. Environmental Engineering, NYS PE', '20+ years wastewater and collection system design and planning.', 'Technical oversight, Engineering Report authorship, PE stamping.'),
    ('I&I / Hydraulics Specialist', 'B.S. Civil Engineering', '10+ years flow monitoring and RDII analysis.', 'Flow and rainfall analysis, I&I quantification, alternatives evaluation.'),
    ('Field Operations Lead', 'Associate degree, relevant certifications', '15+ years field supervision for SSES programs.', 'Manage smoke testing, weiring, and focused SSES crews.'),
]

for role, edu, exp, resp in team_rows:
    row_cells = team_table.add_row().cells
    row_cells[0].text = role
    row_cells[1].text = edu
    row_cells[2].text = exp
    row_cells[3].text = resp

# -------------------------
# MWBE and SDVOB Participation Strategy
# -------------------------

doc.add_heading('MWBE and SDVOB Participation Strategy', level=1)

p = doc.add_paragraph(
    'Aurelia is committed to meeting or exceeding the RFP’s MWBE (30 percent combined) '
    'and SDVOB (6 percent) participation goals. We will partner with certified MWBE and '
    'SDVOB firms for specific scope elements such as field investigations, survey support, '
    'public outreach, and drafting. Utilization plans and, if necessary, waiver requests '
    'will be submitted with the proposal in accordance with Appendix F requirements.'
)

p = doc.add_paragraph(
    'We will track MWBE and SDVOB participation on a monthly basis and submit Workforce '
    'Employment Utilization Reports and MWBE-SDVOB reports as required. Our internal '
    'systems will align purchase orders, subcontractor invoices, and utilization reporting to '
    'provide NFWB with transparent, auditable records.'
)

# MWBE/SDVOB utilization plan table (conceptual)

doc.add_paragraph('MWBE and SDVOB Utilization Plan (Conceptual)', style='Heading 2')

mwbe_table = doc.add_table(rows=1, cols=4)
mwbe_table.style = 'Table Grid'
mt = mwbe_table.rows[0].cells
mt[0].text = 'Firm Type'
mt[1].text = 'Scope of Work'
mt[2].text = 'Approximate Share of Contract Value'
mt[3].text = 'Reporting Approach'
for cell in mt:
    for run in cell.paragraphs[0].runs:
        run.bold = True

mwbe_rows = [
    ('Prime (Aurelia)', 'Project management, engineering analysis, reporting, data and GIS.', '60–65%', 'Track internal labor and expenses by task and phase.'),
    ('MWBE subconsultant(s)', 'Field investigations, public outreach, drafting support.', '25–30%', 'Monthly utilization reports aligned to Appendix F requirements.'),
    ('SDVOB subconsultant(s)', 'Specialty field services or analysis support.', '6–10%', 'Monthly utilization reports and workforce reporting as applicable.'),
]

for ftype, scope, share, rep in mwbe_rows:
    row_cells = mwbe_table.add_row().cells
    row_cells[0].text = ftype
    row_cells[1].text = scope
    row_cells[2].text = share
    row_cells[3].text = rep

# -------------------------
# Health, Safety, and Site-Specific Risk Management
# -------------------------

doc.add_heading('Health, Safety, and Site-Specific Risk Management', level=1)

p = doc.add_paragraph(
    'Field activities for this project involve confined space entry, traffic control, '
    'nighttime work, and work near active sewers and outfalls. Aurelia maintains a '
    'comprehensive Safety Program and will prepare Site-Specific Health & Safety Plans '
    '(SSHSPs) for Phase 1 activities, addressing hazard identification, procedures, personal '
    'protective equipment (PPE), OSHA training, and monitoring requirements.'
)

p = doc.add_paragraph(
    'All field staff will hold required training and certifications, including confined space '
    'entry, traffic control, and first aid/CPR as applicable. Daily tailgate safety '
    'briefings will be conducted before field work, and safety performance will be reviewed '
    'regularly with NFWB. We will comply with all OSHA regulations and NFWB safety '
    'requirements.'
)

# Safety responsibilities table

doc.add_paragraph('Health and Safety Responsibilities by Role', style='Heading 2')

safety_table = doc.add_table(rows=1, cols=3)
safety_table.style = 'Table Grid'
shs = safety_table.rows[0].cells
shs[0].text = 'Role'
shs[1].text = 'Key Safety Responsibilities'
shs[2].text = 'Example Activities'
for cell in shs:
    for run in cell.paragraphs[0].runs:
        run.bold = True

safety_rows = [
    ('Project Manager', 'Overall safety oversight, ensure SSHSP implementation.', 'Review safety metrics, participate in incident reviews.'),
    ('Field Operations Lead', 'Daily safety briefings, field supervision, stop-work authority.', 'Conduct tailgate talks, verify PPE, monitor conditions.'),
    ('Field Crew Members', 'Follow SSHSP, report hazards, use PPE properly.', 'Inspect equipment, maintain safe work zones.'),
]

for role, resp, act in safety_rows:
    row_cells = safety_table.add_row().cells
    row_cells[0].text = role
    row_cells[1].text = resp
    row_cells[2].text = act

# -------------------------
# Security, Data Management, and Information Governance
# -------------------------

doc.add_heading('Security, Data Management, and Information Governance', level=1)

p = doc.add_paragraph(
    'Aurelia will manage all project data—including CCTV files, inspection databases, '
    'flow and rainfall data, GIS layers, and reports—in a secure, access-controlled '
    'environment. Data will be stored on encrypted platforms with role-based access, and '
    'transfers to NFWB will use secure methods. We will comply with NYS information '
    'security breach notification laws and any additional NFWB requirements.'
)

p = doc.add_paragraph(
    'We will maintain project records for at least six years after the calendar year in '
    'which they are created, and will support inspection and audit by the State Comptroller, '
    'Attorney General, and NFWB as required. FOIL-sensitive information will be clearly '
    'identified and handled in accordance with applicable law.'
)

# Data governance table

doc.add_paragraph('Data Categories, Security Controls, and Retention', style='Heading 2')

data_table = doc.add_table(rows=1, cols=4)
data_table.style = 'Table Grid'
dh = data_table.rows[0].cells
dh[0].text = 'Data Category'
dh[1].text = 'Security Controls'
dh[2].text = 'Typical Retention'
dh[3].text = 'FOIL Considerations'
for cell in dh:
    for run in cell.paragraphs[0].runs:
        run.bold = True

data_rows = [
    ('CCTV and inspection data', 'Encrypted storage, role-based access, secure file transfer.', 'Minimum six years after calendar year created.', 'May contain sensitive infrastructure details; mark as critical infrastructure information where applicable.'),
    ('Flow and rainfall data', 'Controlled access, integrity checks.', 'Minimum six years after calendar year created.', 'Generally releasable; consider aggregation for public release.'),
    ('Engineering analyses and reports', 'Version control, restricted editing rights.', 'Minimum six years after calendar year created.', 'Subject to FOIL; identify confidential business information as needed.'),
]

for cat, ctrl, ret, foil in data_rows:
    row_cells = data_table.add_row().cells
    row_cells[0].text = cat
    row_cells[1].text = ctrl
    row_cells[2].text = ret
    row_cells[3].text = foil

# -------------------------
# Commercial Terms, Insurance, and Legal Compliance
# -------------------------

doc.add_heading('Commercial Terms, Insurance, and Legal Compliance', level=1)

p = doc.add_paragraph(
    'Aurelia accepts the Water Board’s general RFP conditions, including the Board’s '
    'rights to reject, modify, negotiate, and award; the requirement that proposals remain '
    'firm and irrevocable for 90 days; and the termination, governing law, venue, service '
    'of process, and non-waiver clauses in Appendix E.'
)

p = doc.add_paragraph(
    'We will procure and maintain required insurance coverages, including workers '
    'compensation, disability, employer’s liability, commercial general liability, '
    'automobile, umbrella/excess, and professional liability insurance, with limits meeting '
    'or exceeding RFP requirements. NFWB, the Niagara Falls Public Water Authority, and the '
    'City of Niagara Falls will be named as additional insureds, and certificates and '
    'endorsements will be provided with 30 days’ notice of change or cancellation.'
)

p = doc.add_paragraph(
    'Aurelia will comply with non-discrimination and sexual harassment prevention '
    'requirements, NYS Lobbying Law, non-collusion requirements, the Iran Divestment Act, '
    'International Boycott Prohibition, MacBride Fair Employment Principles, and prohibitions '
    'on the purchase of tropical hardwoods, as applicable. Required forms and certifications '
    '(Appendix E and Appendix F) will be completed and submitted with the proposal.'
)

# Insurance coverage table

doc.add_paragraph('Insurance Coverage Summary', style='Heading 2')

ins_table = doc.add_table(rows=1, cols=3)
ins_table.style = 'Table Grid'
ins = ins_table.rows[0].cells
ins[0].text = 'Coverage Type'
ins[1].text = 'Indicative Limit'
ins[2].text = 'Notes'
for cell in ins:
    for run in cell.paragraphs[0].runs:
        run.bold = True

ins_rows = [
    ('Workers compensation and disability', 'As required by law', 'Statutory coverage for all employees.'),
    ('Employer’s liability', '$1,000,000 per accident', 'Covers employer liability for workplace injuries.'),
    ('Commercial general liability', '$1,000,000 per occurrence / $2,000,000 aggregate', 'NFWB and related entities named as additional insureds.'),
    ('Automobile liability', '$1,000,000 combined single limit', 'Covers owned, hired, and non-owned vehicles.'),
    ('Umbrella / excess liability', '$2,000,000 aggregate or greater', 'Provides additional limits above primary policies.'),
    ('Professional liability (errors and omissions)', '$2,000,000 per claim', 'Covers professional engineering services.'),
]

for cov, limit_val, note in ins_rows:
    row_cells = ins_table.add_row().cells
    row_cells[0].text = cov
    row_cells[1].text = limit_val
    row_cells[2].text = note

# -------------------------
# Administrative Requirements and Proposal Submission Plan
# -------------------------

doc.add_heading('Administrative Requirements and Proposal Submission Plan', level=1)

p = doc.add_paragraph(
    'Aurelia will comply with all administrative instructions in the RFP. We will attend the '
    'mandatory pre-proposal information session at the Wastewater Treatment Plant on '
    'November 29, 2022 at noon, submit any written questions to the authorized Water Board '
    'contact by November 30, 2022 at 3:00 p.m., and acknowledge all addenda on Appendix E '
    'Form No. 1.'
)

p = doc.add_paragraph(
    'Our proposal package will include one clearly marked original and six copies, plus one '
    'electronic copy on CD or USB in a single PDF or clearly ordered files. Original signed '
    'appendices will be included only in the paper original. The outer envelope will be '
    'sealed and labeled with the contact name, proposer information, and RFP number and '
    'title. Any exceptions to RFP terms will be identified in a separate attachment titled '
    '“Exceptions”; unlisted exceptions will be deemed waived.'
)

# Proposal submission checklist table

doc.add_paragraph('Proposal Submission Checklist', style='Heading 2')

chk_table = doc.add_table(rows=1, cols=3)
chk_table.style = 'Table Grid'
ck = chk_table.rows[0].cells
ck[0].text = 'Item'
ck[1].text = 'Description'
ck[2].text = 'Included'
for cell in ck:
    for run in cell.paragraphs[0].runs:
        run.bold = True

chk_rows = [
    ('Original proposal', 'One signed original with all required forms and appendices.', 'Yes'),
    ('Proposal copies', 'Six paper copies of the proposal.', 'Yes'),
    ('Electronic copy', 'One CD or USB with a single PDF or clearly ordered files.', 'Yes'),
    ('Appendix E forms', 'All required standard terms, certifications, and acknowledgments.', 'Yes'),
    ('Appendix F forms', 'EEO Staffing Plan, Workforce Employment Utilization Report, MWBE/SDVOB plans.', 'Yes'),
    ('Exceptions attachment', 'Separate document titled “Exceptions” (if any).', 'If applicable'),
]

for item, desc, inc in chk_rows:
    row_cells = chk_table.add_row().cells
    row_cells[0].text = item
    row_cells[1].text = desc
    row_cells[2].text = inc

# -------------------------
# Required Forms and Certifications (Appendix E & F)
# -------------------------

doc.add_heading('Required Forms and Certifications (Appendix E & F)', level=1)

p = doc.add_paragraph(
    'We understand that complete and accurate submission of all required forms and '
    'certifications is a condition of responsiveness. Aurelia will complete all Appendix E '
    'forms (Standard Terms, Conditions, and Requirements), including the EEO Policy '
    'Statement, Statement on Sexual Harassment, Certificate of Non-Collusion, and Lobbying '
    'Law Disclosure Statement, as well as any other required certifications.'
)

p = doc.add_paragraph(
    'We will also complete Appendix F forms related to Mandatory State Financial Assistance '
    'Terms and Conditions, including the EEO Staffing Plan, Workforce Employment '
    'Utilization Report, and MWBE and SDVOB utilization plans or waiver requests. These '
    'documents will be organized in the proposal to facilitate NFWB review.'
)

# Forms tracking table

doc.add_paragraph('Forms and Certifications Tracking', style='Heading 2')

forms_table = doc.add_table(rows=1, cols=4)
forms_table.style = 'Table Grid'
fh = forms_table.rows[0].cells
fh[0].text = 'Form / Certification'
fh[1].text = 'Appendix'
fh[2].text = 'Responsible Party'
fh[3].text = 'Status at Submission'
for cell in fh:
    for run in cell.paragraphs[0].runs:
        run.bold = True

forms_rows = [
    ('Standard Terms, Conditions, and Requirements (Form No. 1)', 'Appendix E', 'Aurelia contracts lead', 'Completed and signed'),
    ('Certificate of Non-Collusion (Form No. 2)', 'Appendix E', 'Aurelia authorized signatory', 'Completed and signed'),
    ('Lobbying Law Disclosure (Form No. 3)', 'Appendix E', 'Aurelia contracts lead', 'Completed and signed'),
    ('EEO Policy Statement (Form No. 4)', 'Appendix E', 'Aurelia HR lead', 'Completed and signed'),
    ('Statement on Sexual Harassment (Form No. 5)', 'Appendix E', 'Aurelia HR lead', 'Completed and signed'),
    ('EEO Staffing Plan', 'Appendix F', 'Aurelia HR and PM', 'Completed and included'),
    ('Workforce Employment Utilization Report', 'Appendix F', 'Aurelia HR and PM', 'Template prepared; to be updated during contract.'),
    ('MWBE/SDVOB Utilization Plan or Waiver Request', 'Appendix F', 'Aurelia PM and MWBE coordinator', 'Completed and included'),
]

for form, app, resp, status in forms_rows:
    row_cells = forms_table.add_row().cells
    row_cells[0].text = form
    row_cells[1].text = app
    row_cells[2].text = resp
    row_cells[3].text = status

# -------------------------
# Appendices: Detailed Resumes, Reference Projects, and Supporting Materials
# -------------------------

doc.add_heading('Appendices: Detailed Resumes, Reference Projects, and Supporting Materials', level=1)

p = doc.add_paragraph(
    'The appendices to this proposal will include detailed resumes of key staff, expanded '
    'descriptions of relevant reference projects, letters of commitment from any '
    'subconsultants, and additional technical materials supporting our approach. These '
    'materials will demonstrate Aurelia’s and our partners’ qualifications to successfully '
    'deliver the Phase 1 SSES, Engineering Report, corrective action plan, and '
    'post-construction monitoring, as well as optional Phases 2 and 3.'
)

# Reference projects table

doc.add_paragraph('Representative Reference Projects (Summary)', style='Heading 2')

ref_table = doc.add_table(rows=1, cols=4)
ref_table.style = 'Table Grid'
rf = ref_table.rows[0].cells
rf[0].text = 'Client'
rf[1].text = 'Project Description'
rf[2].text = 'Year(s)'
rf[3].text = 'Relevance to NFWB SSES'
for cell in rf:
    for run in cell.paragraphs[0].runs:
        run.bold = True

ref_rows = [
    ('Regional wastewater authority', 'Multi-phase SSES and I&I reduction program with flow monitoring, smoke testing, and rehabilitation planning.', '2019–2022', 'Demonstrates experience with phased SSES programs and regulatory reporting.'),
    ('Mid-sized city utility', 'Collection system master plan including RDII analysis and capital improvement prioritization.', '2018–2021', 'Experience integrating flow, rainfall, and inspection data for capital planning.'),
    ('Water and sewer district', 'Targeted SSES and corrective action plan for SSO-prone basin.', '2020–2022', 'Directly relevant to addressing chronic SSOs through data-driven I&I reduction.'),
]

for client, desc, year, rel in ref_rows:
    row_cells = ref_table.add_row().cells
    row_cells[0].text = client
    row_cells[1].text = desc
    row_cells[2].text = year
    row_cells[3].text = rel

# End of document code
