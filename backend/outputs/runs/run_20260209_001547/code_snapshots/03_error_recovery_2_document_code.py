from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import matplotlib.patches as mpatches
import datetime as dt
import matplotlib.dates as mdates

styles = doc.styles
# --- Style Bootstrap ---
normal = styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)

if 'Caption' not in [s.name for s in styles]:
    cap = styles.add_style('Caption', WD_STYLE_TYPE.PARAGRAPH)
    cap.font.name = 'Calibri'
    cap.font.size = Pt(9)
    cap.font.italic = True

def add_caption(text: str):
    p = doc.add_paragraph(text, style='Caption')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p

# --- Title Page ---
doc.add_heading('Dynamics 365 Finance and Operations\nComprehensive ERP Solution Proposal', level=0)
p = doc.add_paragraph()
p.add_run('Prepared for: ').bold = True
p.add_run('Smithsonian Enterprises')
p = doc.add_paragraph()
p.add_run('Submitted by: ').bold = True
p.add_run('Argano, LLC')
p = doc.add_paragraph()
p.add_run('Confidentiality Notice').bold = True
p = doc.add_paragraph('External use or disclosure of any information contained in this document without written permission of Argano, LLC is prohibited.')
p.paragraph_format.space_before = Pt(24)
doc.add_page_break()

# --- Table of Contents ---
doc.add_heading('Table of Contents', level=1)
toc_items = [
    'Executive Summary',
    'Company Overview',
    'Requirements Response & Compliance Matrix',
    'Solution Architecture and Technical Approach',
    'Implementation Plan & Delivery Schedule',
    'Project Management and Governance',
    'Security, Privacy, and Compliance',
    'Team and Qualifications',
    'Assumptions, Risks, and Dependencies',
    'Appendices',
]
for item in toc_items:
    doc.add_paragraph(item, style='List Number')
doc.add_page_break()

# --- Executive Summary ---
doc.add_heading('Executive Summary', level=1)
doc.add_paragraph(
    'Argano, LLC is excited to submit this proposal for the implementation of Microsoft Dynamics 365 Finance and Operations for Smithsonian Enterprises. Our approach is designed to support Smithsonian Enterprises in digitally transforming its core finance, operations, retail, and e-commerce processes. We leverage deep industry experience, a proven agile methodology, and a commitment to partnership to ensure a successful and sustainable ERP implementation.'
)

sum_table = doc.add_table(rows=1, cols=2)
sum_table.style = 'Table Grid'
sum_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = sum_table.rows[0].cells
hdr[0].text, hdr[1].text = 'Strategic Goal', 'Measure of Success'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True
sum_data = [
    ('Unified ERP for consolidated operations & reporting', 'Increased productivity, improved insight'),
    ('Increased inventory and shipping visibility', 'Improved management of supply chain and orders'),
    ('Cost reduction through system consolidation', 'Lower IT maintenance and licensing costs'),
    ('Enhanced compliance and security', 'Sustained audit readiness, FISMA/PCI support'),
    ('Scalable foundation for growth', 'Future-proofed platform, seamless integrations'),
]
for goal, measure in sum_data:
    row = sum_table.add_row().cells
    row[0].text = goal
    row[1].text = measure

doc.add_paragraph('Argano is confident in our ability to deliver the project scope and meet all requirements based on our extensive experience in nonprofit, retail, and public sector ERP implementations.')
doc.add_page_break()

# --- Company Overview ---
doc.add_heading('Company Overview', level=1)
doc.add_paragraph(
    'Argano is a global consulting firm and a top-ranked Microsoft Dynamics Gold Certified Partner, specializing in large-scale ERP, CRM, and analytics solutions for public sector, nonprofit, retail, and complex multi-entity organizations. As a member of Microsoft’s Inner Circle and a US Partner of the Year for Dynamics 365 Finance, Argano brings over two decades of experience and an award-winning team of 1,900+ professionals.'
)
doc.add_paragraph('Key differentiators:')
doc.add_paragraph('Agile, outcome-driven delivery methodology', style='List Bullet')
doc.add_paragraph('Deep nonprofit and public sector experience', style='List Bullet')
doc.add_paragraph('Commitment to business process transformation', style='List Bullet')
doc.add_paragraph('Executive sponsorship for every implementation', style='List Bullet')
doc.add_paragraph('Extensive experience with legacy migrations and integrations', style='List Bullet')
doc.add_paragraph('Proprietary IP to accelerate value (Retail Dashboard, Audit & Security Manager)', style='List Bullet')
doc.add_page_break()

# --- Requirements Response & Compliance Matrix ---
doc.add_heading('Requirements Response & Compliance Matrix', level=1)
req_table = doc.add_table(rows=1, cols=3)
req_table.style = 'Table Grid'
req_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = req_table.rows[0].cells
hdr[0].text, hdr[1].text, hdr[2].text = 'Requirement', 'Response', 'Reference'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True
req_data = [
    ('GA-1', 'Supported via multi-entity configuration and rollup reporting.', 'Dynamics 365 Core'),
    ('GA-2', 'Supported through flexible organizational structures and reporting.', 'Dynamics 365 Core'),
    ('GA-3', 'Supported with scalable legal entity addition.', 'Dynamics 365 Core'),
    ('GA-5', 'Configurable approval workflows for entity changes.', 'Power Automate'),
    ('GA-15', 'Attach supporting docs to journal entries.', 'Dynamics 365 Finance'),
    ('GA-19', 'Automated anomaly flagging and notification.', 'Dynamics 365 AI'),
    ('O2C-2', 'CRM to ERP integration for customer account data.', 'Dynamics 365 Integration'),
    ('O2C-10', 'Enforce credit limits during order/project activation.', 'Dynamics 365 Credit Mgmt'),
    ('P2P-5', 'Approval workflows for vendor onboarding.', 'Dynamics 365 Workflow'),
    ('INV-17', 'Multi-warehouse inventory tracking.', 'Dynamics 365 Supply Chain'),
]
for req, resp, ref in req_data:
    row = req_table.add_row().cells
    row[0].text = req
    row[1].text = resp
    row[2].text = ref

# --- Solution Architecture and Technical Approach ---
doc.add_heading('Solution Architecture and Technical Approach', level=1)
doc.add_paragraph(
    'Our solution leverages Microsoft Dynamics 365 Finance and Operations as the centralized ERP platform, integrated with CRM, Power Automate, and Azure for end-to-end business process management. The technical design includes scalable multi-entity support, robust workflow configuration, and secure integrations.'
)
doc.add_paragraph('Core components:')
doc.add_paragraph('Dynamics 365 Finance & Operations', style='List Bullet')
doc.add_paragraph('Dynamics 365 CRM', style='List Bullet')
doc.add_paragraph('Power Automate for workflow orchestration', style='List Bullet')
doc.add_paragraph('Azure Active Directory for security', style='List Bullet')
doc.add_paragraph('Azure Data Lake for reporting and analytics', style='List Bullet')
doc.add_paragraph('Integration with third-party tax and compliance engines', style='List Bullet')

# Mermaid architecture diagram
mermaid_code = '''
flowchart TD
  A[User] --> B[Azure AD]
  B --> C[Dynamics 365 Finance]
  C --> D[Power Automate]
  C --> E[CRM]
  C --> F[Azure Data Lake]
  E --> F
'''
diagram_path = render_mermaid(mermaid_code, 'solution_architecture')
doc.add_picture(str(diagram_path), width=Inches(5.8))
add_caption('Figure 1: Solution architecture overview')

doc.add_page_break()

# --- Implementation Plan & Delivery Schedule ---
doc.add_heading('Implementation Plan & Delivery Schedule', level=1)
doc.add_paragraph(
    'Our implementation plan follows a proven project lifecycle, beginning with requirements discovery, solution design, iterative build, testing, deployment, and hypercare. We use agile methods to maximize stakeholder involvement and minimize risk.'
)

# Gantt chart data
schedule = pd.DataFrame({
    'Task': ['Requirements Analysis', 'Solution Design', 'Development Phase 1', 
             'Development Phase 2', 'Integration Testing', 'User Acceptance Testing',
             'Training & Documentation', 'Deployment & Go-Live', 'Hypercare Support'],
    'Team': ['Business Analysts', 'Architects', 'Development', 'Development', 
             'QA Team', 'QA Team', 'Training', 'DevOps', 'Support'],
    'Start': pd.to_datetime(['2024-07-01', '2024-07-15', '2024-08-01', '2024-09-01',
                             '2024-10-01', '2024-10-15', '2024-10-20', '2024-11-01', '2024-11-08']),
    'End': pd.to_datetime(['2024-07-14', '2024-07-31', '2024-08-31', '2024-09-30',
                           '2024-10-14', '2024-10-31', '2024-11-05', '2024-11-07', '2024-11-30'])
})
project_start = schedule['Start'].min()
schedule['days_to_start'] = (schedule['Start'] - project_start).dt.days
schedule['duration'] = (schedule['End'] - schedule['Start']).dt.days + 1
team_colors = {
    'Business Analysts': '#3498db',  # Blue
    'Architects': '#9b59b6',         # Purple
    'Development': '#2ecc71',        # Green
    'QA Team': '#e74c3c',            # Red
    'Training': '#f39c12',           # Orange
    'DevOps': '#1abc9c',             # Teal
    'Support': '#34495e'             # Dark gray
}
fig, ax = plt.subplots(figsize=(10, 6))
for idx, row in schedule.iterrows():
    ax.barh(y=row['Task'], width=row['duration'], left=row['days_to_start'],
            color=team_colors[row['Team']], edgecolor='white', linewidth=0.5)
ax.invert_yaxis()
max_days = schedule['days_to_start'].max() + schedule['duration'].max()
exticks = np.arange(0, max_days + 7, 14)  # Every 2 weeks
xticklabels = pd.date_range(start=project_start, periods=len(xticks), freq='14D').strftime('%b %d')
ax.set_xticks(xticks[:len(xticklabels)])
ax.set_xticklabels(xticklabels, fontsize=9)
ax.set_xlabel('Timeline', fontsize=10)
ax.set_title('Project Implementation Schedule', fontsize=14, fontweight='bold')
ax.xaxis.grid(True, alpha=0.3, linestyle='--')
ax.set_axisbelow(True)
unique_teams = schedule['Team'].unique()
patches = [mpatches.Patch(color=team_colors[team], label=team) for team in unique_teams]
ax.legend(handles=patches, loc='lower right', fontsize=8, framealpha=0.9)
plt.tight_layout()
gantt_path = output_dir / 'project_gantt.png'
plt.savefig(gantt_path, dpi=150, bbox_inches='tight', facecolor='white')
plt.close()
doc.add_picture(str(gantt_path), width=Inches(6.0))
add_caption('Figure 2: Project Implementation Schedule')

doc.add_page_break()

# --- Project Management and Governance ---
doc.add_heading('Project Management and Governance', level=1)
doc.add_paragraph(
    'Our project management approach is based on PMI standards, agile principles, and transparent stakeholder engagement. Key governance measures include weekly status reviews, milestone tracking, risk management, and executive sponsorship.'
)
doc.add_paragraph('Governance framework:')
doc.add_paragraph('Steering committee with executive sponsor', style='List Bullet')
doc.add_paragraph('Weekly status and risk review', style='List Bullet')
doc.add_paragraph('Milestone-based reporting', style='List Bullet')
doc.add_paragraph('Change control process for scope and requirements', style='List Bullet')
doc.add_paragraph('Stakeholder engagement and feedback loops', style='List Bullet')

# Project workflow Mermaid diagram
mermaid_code = '''
flowchart TD
  A[Kickoff] --> B[Discovery]
  B --> C[Design]
  C --> D[Build]
  D --> E[Test]
  E --> F[Deploy]
  F --> G[Hypercare]
'''
diagram_path2 = render_mermaid(mermaid_code, 'project_workflow')
doc.add_picture(str(diagram_path2), width=Inches(5.8))
add_caption('Figure 3: Project delivery workflow')

doc.add_page_break()

# --- Security, Privacy, and Compliance ---
doc.add_heading('Security, Privacy, and Compliance', level=1)
doc.add_paragraph(
    'Our solution implements Microsoft-recommended security practices, including role-based access control, audit logging, data encryption, and compliance with FISMA, PCI, and SOC standards. All integrations and workflows are subject to robust authentication and authorization via Azure Active Directory.'
)
doc.add_paragraph('Security measures include:')
doc.add_paragraph('Role-based access control (RBAC)', style='List Bullet')
doc.add_paragraph('Audit logging of all critical actions', style='List Bullet')
doc.add_paragraph('Data encryption in transit and at rest', style='List Bullet')
doc.add_paragraph('Multi-factor authentication', style='List Bullet')
doc.add_paragraph('FISMA, PCI, SOC 1/2 compliance support', style='List Bullet')
doc.add_paragraph('Automated compliance workflows for approvals and notifications', style='List Bullet')

# Compliance matrix table
comp_table = doc.add_table(rows=1, cols=3)
comp_table.style = 'Table Grid'
comp_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = comp_table.rows[0].cells
hdr[0].text, hdr[1].text, hdr[2].text = 'Control Area', 'Approach', 'Compliance'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True
comp_data = [
    ('Access Control', 'RBAC, Azure AD', 'FISMA, SOC 1/2'),
    ('Audit Trail', 'Activity logging', 'PCI, SOC 2'),
    ('Data Encryption', 'TLS, Azure encryption', 'FISMA, PCI'),
    ('Approval Workflow', 'Power Automate', 'SOC 2'),
    ('Reporting & Analytics', 'Azure Data Lake', 'SOC 1/2'),
]
for area, approach, comp in comp_data:
    row = comp_table.add_row().cells
    row[0].text = area
    row[1].text = approach
    row[2].text = comp

doc.add_page_break()

# --- Team and Qualifications ---
doc.add_heading('Team and Qualifications', level=1)
doc.add_paragraph(
    'Argano’s project team includes certified Dynamics 365 solution architects, finance subject matter experts, technical integrators, and project managers. All team members have deep sector experience and a track record of delivering complex ERP projects.'
)
team_table = doc.add_table(rows=1, cols=3)
team_table.style = 'Table Grid'
team_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = team_table.rows[0].cells
hdr[0].text, hdr[1].text, hdr[2].text = 'Role', 'Name', 'Qualifications'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True
team_data = [
    ('Project Manager', 'Jessica Lee', 'PMP, 12+ years ERP delivery'),
    ('Solution Architect', 'Samuel Jones', 'MS Certified, 18+ years Dynamics'),
    ('Finance SME', 'Priya Patel', 'CPA, 15+ years finance systems'),
    ('Technical Lead', 'Brandon Smith', 'Azure, Integration specialist'),
    ('QA Lead', 'Lisa Chen', 'ISTQB, 8+ years testing'),
]
for role, name, qual in team_data:
    row = team_table.add_row().cells
    row[0].text = role
    row[1].text = name
    row[2].text = qual

doc.add_page_break()

# --- Assumptions, Risks, and Dependencies ---
doc.add_heading('Assumptions, Risks, and Dependencies', level=1)
doc.add_paragraph(
    'The proposal assumes timely access to key stakeholders, data, and legacy systems. Major risks include data migration complexity, change management, and external dependencies on third-party integrations. Risk mitigation includes detailed planning, regular communication, and phased rollout.'
)
doc.add_paragraph('Top risks and mitigation:')
doc.add_paragraph('Data migration complexity – Mitigated via phased validation and testing', style='List Bullet')
doc.add_paragraph('Change management – Mitigated via training, clear communication, and stakeholder engagement', style='List Bullet')
doc.add_paragraph('External integration dependencies – Mitigated via early engagement and fallback planning', style='List Bullet')
doc.add_paragraph('Resource availability – Mitigated via executive sponsorship and clear escalation channels', style='List Bullet')

# --- Appendices ---
doc.add_heading('Appendices', level=1)
doc.add_paragraph('Full requirements list, extended compliance matrix, references available upon request.')
