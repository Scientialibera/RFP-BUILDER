from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import matplotlib.patches as mpatches
import datetime as dt

styles = doc.styles

normal = styles['Normal']
normal.font.name = 'Calibri'
normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
normal.font.size = Pt(11)

if 'Caption' not in [s.name for s in styles]:
    cap = styles.add_style('Caption', WD_STYLE_TYPE.PARAGRAPH)
    cap.font.name = 'Calibri'
    cap.font.size = Pt(9)
    cap.font.italic = True
else:
    cap = styles['Caption']
    cap.font.name = 'Calibri'
    cap.font.size = Pt(9)
    cap.font.italic = True

figure_counter = {'n': 0}


def add_caption(text: str):
    p = doc.add_paragraph(text, style='Caption')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


section = doc.sections[0]
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

# Title Page
title = doc.add_heading('Phase 1 Sanitary Sewer Evaluation Survey (SSES)\nLaSalle Sewersheds – Engineering Services Proposal', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('Submitted to: Niagara Falls Water Board', style='Normal')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle2 = doc.add_paragraph('Submitted by: Aurelia Digital Systems – Water & Infrastructure Practice', style='Normal')
subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle3 = doc.add_paragraph('Date: November 2022', style='Normal')
subtitle3.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# 1. Executive Summary
h1 = doc.add_heading('1. Executive Summary', level=1)
h1.paragraph_format.space_before = Pt(6)
h1.paragraph_format.space_after = Pt(6)

p = doc.add_paragraph(
    'The Niagara Falls Water Board (NFWB) is undertaking a focused Phase 1 Sanitary Sewer '
    'Evaluation Survey (SSES) in the LaSalle sewersheds to meet the requirements of Order on '
    'Consent R9020080528-32, reduce inflow and infiltration (I&I), and enable reliable compliance '
    'with NYSDEC regulatory expectations. Aurelia Digital Systems (Aurelia) proposes a clear, '
    'technically rigorous, and fully compliant approach that delivers an NYSDEC-acceptable SSES '
    'Engineering Report and Corrective Action Plan on the Consent Order schedule, while de-risking '
    'NFWB’s obligations and setting a repeatable playbook for Phases 2 and 3.'
)

p = doc.add_paragraph(
    'Our team will perform all eight Phase 1 SSES tasks in the specified LaSalle sewersheds, '
    'including project management, smoke testing, flow isolation and measurement, CCTV and '
    'manhole report review, outfall and SSO condition assessment, focused SSES, reporting, and '
    'post-construction monitoring (REQ-001). A dedicated Project Manager will serve as the single '
    'point of contact (REQ-002), coordinating closely with NFWB staff and NYSDEC to maintain '
    'alignment with the Engineering Report Outline for NYS Wastewater Infrastructure Projects '
    'and the Consent Order timelines (REQ-019, REQ-020, REQ-021, REQ-022).'
)

p = doc.add_paragraph(
    'The proposed solution emphasizes regulatory reliability, data-driven decision making, and '
    'constructible recommendations. By integrating flow, rainfall, CCTV, manhole, and SCADA data, '
    'we will prioritize corrective actions that maximize I&I reduction per dollar invested, provide '
    'transparent assumptions, and clearly connect each recommendation to Consent Order '
    'milestones and sewershed R-value prioritization. Option pricing for Phases 2 and 3 is '
    'developed using the same methodology, enabling NFWB to scale the program efficiently '
    '(REQ-023).'
)

p = doc.add_paragraph(
    'Aurelia’s approach reduces risk for NFWB by: (1) tightly aligning the Phase 1 SSES scope, '
    'schedule, and report structure to NYSDEC expectations; (2) implementing robust health, safety, '
    'and public communication measures for smoke testing and flow isolation; (3) applying proven '
    'project management and governance practices; and (4) proactively addressing MWBE/SDVOB '
    'participation, prevailing wage, and statutory compliance requirements (REQ-027, REQ-030, '
    'REQ-031, REQ-038, REQ-042).'
)

# Executive Summary – Mermaid timeline diagram
mermaid_code_exec = '''
flowchart LR
  A[Notice to proceed] --> B[Phase 1 SSES field work]
  B --> C[Draft SSES report]
  C --> D[NYSDEC review]
  D --> E[Construction 2023]
  E --> F[Post construction monitoring 2024]
  F --> G[Final I and I assessment June 2024]
'''
diagram_path_exec = render_mermaid(mermaid_code_exec, 'exec_timeline')
doc.add_picture(str(diagram_path_exec), width=Inches(5.8))
figure_counter['n'] += 1
add_caption(f"Figure {figure_counter['n']}: Phase 1 SSES, construction, and post-construction monitoring timeline")

# Executive Summary – I&I reduction bar chart
plt.figure(figsize=(6, 4))
sns.set_style('whitegrid')
sewersheds = ['LaSalle A', 'LaSalle B', 'LaSalle C', 'LaSalle D']
reduction_mgd = [0.7, 0.5, 0.4, 0.3]
exec_data = pd.DataFrame({'Sewershed': sewersheds, 'Projected I&I reduction MGD': reduction_mgd})
ax = sns.barplot(data=exec_data, x='Sewershed', y='Projected I&I reduction MGD', color='#4c72b0')
ax.set_title('Illustrative projected I&I reduction by Phase 1 sewershed')
ax.set_ylabel('Reduction (MGD)')
plt.tight_layout()
chart_path_exec = output_dir / 'exec_iandi_reduction.png'
plt.savefig(chart_path_exec, dpi=150, bbox_inches='tight', facecolor='white')
plt.close()

doc.add_picture(str(chart_path_exec), width=Inches(5.8))
figure_counter['n'] += 1
add_caption(f"Figure {figure_counter['n']}: Illustrative projected I&I reduction by Phase 1 sewershed")

# Executive Summary – benefits table
table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'NFWB objective'
hdr[1].text = 'Aurelia response'
hdr[2].text = 'Resulting benefit'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

rows = [
    (
        'Meet Order on Consent and NYSDEC expectations',
        'Align SSES scope, report structure, and schedule with Engineering Report Outline and Consent Order milestones',
        'Reduced regulatory risk; clear path to NYSDEC acceptance (REQ-019, REQ-020, REQ-021, REQ-022)'
    ),
    (
        'Maximize I&I reduction for limited capital funds',
        'Use integrated flow, rainfall, CCTV, and SCADA analysis to prioritize highest return corrective actions',
        'Higher I&I reduction per dollar; transparent prioritization and budgeting (REQ-019)'
    ),
    (
        'Protect residents and field staff during SSES',
        'Implement robust health and safety program, traffic control, and public notification for smoke testing and flow isolation',
        'Minimized disruption, improved public trust, and safe field operations (REQ-007, REQ-008, REQ-009, REQ-011, REQ-012, REQ-033, REQ-034)'
    ),
    (
        'Deliver on time to enable 2023–early 2024 construction',
        'Use detailed bar-chart schedule, critical path tracking, and contingency planning',
        'Schedule certainty and ability to complete post-construction monitoring by June 2024 (REQ-020, REQ-022, REQ-028)'
    ),
    (
        'Meet MWBE/SDVOB and workforce requirements',
        'Implement a proactive utilization and reporting strategy with qualified MWBE/SDVOB partners',
        'Compliance with participation goals while maintaining technical quality (REQ-027, REQ-030, REQ-031, REQ-038, REQ-047, REQ-048)'
    ),
]

for obj, resp, ben in rows:
    row_cells = table.add_row().cells
    row_cells[0].text = obj
    row_cells[1].text = resp
    row_cells[2].text = ben

# 2. Understanding of Project, System, and Regulatory Context
doc.add_page_break()
h1 = doc.add_heading('2. Understanding of Project, System, and Regulatory Context', level=1)
h1.paragraph_format.space_before = Pt(6)

p = doc.add_paragraph(
    'NFWB’s LaSalle area has experienced recurring sanitary sewer overflows (SSOs) and '
    'wet-weather surcharging that are subject to Order on Consent R9020080528-32. The Consent '
    'Order requires a structured SSES program, prioritized by sewershed R-values, to identify and '
    'mitigate I&I sources contributing to outfalls 013, 014, and 017. Phase 1 focuses on a subset of '
    'LaSalle sewersheds with the highest regulatory and hydraulic significance, while Phase 3 will '
    'address additional contributing areas.'
)

p = doc.add_paragraph(
    'We understand that NYSDEC expects the Phase 1 SSES to produce an Engineering Report and '
    'Corrective Action Plan that clearly demonstrate how recommended projects will reduce '
    'extraneous flows, protect receiving waters, and eliminate or significantly reduce SSOs at the '
    'LaSalle outfalls. This requires a defensible linkage between field observations, flow and '
    'rainfall data, sewershed R-values, and the proposed implementation schedule (REQ-003, '
    'REQ-004, REQ-019, REQ-020, REQ-049, REQ-050).'
)

p = doc.add_paragraph(
    'Aurelia’s team has worked with utilities facing similar Consent Orders and understands the '
    'importance of transparent documentation, traceability from data to decisions, and early '
    'coordination with NYSDEC staff. We will review existing records, prior SSES work, SCADA '
    'data, maintenance history, and correspondence with NYSDEC to ensure our methodology is '
    'grounded in NFWB’s system knowledge and past investments (REQ-003, REQ-004, REQ-005, '
    'REQ-019, REQ-027).'
)

# System context Mermaid diagram
mermaid_code_sys = '''
flowchart LR
  LS1[LaSalle sewershed 1] --> O013[Outfall 013]
  LS2[LaSalle sewershed 2] --> O014[Outfall 014]
  LS3[LaSalle sewershed 3] --> O017[Outfall 017]
  LS4[Phase 3 sewersheds] --> O017
  O013 --> WWTP[Wastewater treatment plant]
  O014 --> WWTP
  O017 --> WWTP
'''
diagram_path_sys = render_mermaid(mermaid_code_sys, 'system_context')
doc.add_picture(str(diagram_path_sys), width=Inches(5.8))
figure_counter['n'] += 1
add_caption(f"Figure {figure_counter['n']}: LaSalle Phase 1 sewersheds and outfalls in system context")

# R-value prioritization chart
plt.figure(figsize=(6, 4))
sns.set_style('whitegrid')
rv_data = pd.DataFrame({
    'Sewershed': ['LaSalle A', 'LaSalle B', 'LaSalle C', 'LaSalle D'],
    'R value': [3.2, 2.8, 2.1, 1.7]
})
ax = sns.barplot(data=rv_data, x='Sewershed', y='R value', color='#55a868')
ax.set_title('Illustrative R values for Phase 1 sewersheds')
ax.set_ylabel('R value (dimensionless)')
plt.tight_layout()
rv_chart_path = output_dir / 'r_values.png'
plt.savefig(rv_chart_path, dpi=150, bbox_inches='tight', facecolor='white')
plt.close()

doc.add_picture(str(rv_chart_path), width=Inches(5.8))
figure_counter['n'] += 1
add_caption(f"Figure {figure_counter['n']}: Illustrative R values supporting Phase 1 prioritization")

# Sewershed summary table
table = doc.add_table(rows=1, cols=5)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'Sewershed ID'
hdr[1].text = 'Approximate size (acres)'
hdr[2].text = 'Known I&I issues'
hdr[3].text = 'Past repairs / SCADA management'
hdr[4].text = 'Consent Order relevance'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

rows = [
    (
        'LaSalle A', '220',
        'Basement backups and wet-weather surcharging near outfall 013',
        'Targeted lining and manhole sealing; SCADA-based pump control adjustments',
        'High priority contributor to overflow at outfall 013 (REQ-049, REQ-050)'
    ),
    (
        'LaSalle B', '180',
        'Inflow from cross-connections and storm inflow through pickholes',
        'Spot repairs and inflow source investigations; limited prior SSES',
        'Contributes to outfall 014 performance and Consent Order milestones'
    ),
    (
        'LaSalle C', '150',
        'High groundwater infiltration in older clay mains',
        'Historical CCTV and partial main replacement; ongoing SCADA monitoring',
        'Linked to outfall 017 and Phase 3 coordination'
    ),
    (
        'LaSalle D', '130',
        'Localized inflow in commercial areas and private property connections',
        'Manhole frame and cover improvements; limited private property investigations',
        'Supports phased I&I reduction strategy for LaSalle system'
    ),
]

for row in rows:
    cells = table.add_row().cells
    for i, val in enumerate(row):
        cells[i].text = val

# 3. Compliance Matrix and RFP Requirements Traceability
doc.add_page_break()
h1 = doc.add_heading('3. Compliance Matrix and RFP Requirements Traceability', level=1)

p = doc.add_paragraph(
    'To make evaluation straightforward and to demonstrate full compliance, this section maps '
    'each RFP requirement (REQ-001 through REQ-050) to our response and the section(s) where '
    'it is addressed. This traceability ensures that all mandatory technical, management, cost, '
    'and compliance obligations are explicitly covered and easy to verify.'
)

# Requirements flow Mermaid diagram
mermaid_code_req = '''
flowchart TD
  RFP[Requirements REQ 001 to REQ 050] --> S1[Phase 1 SSES tasks]
  RFP --> S2[Project management]
  RFP --> S3[Reporting and compliance]
  S1 --> D1[Engineering report]
  S2 --> D2[Governance and communication]
  S3 --> D3[Appendix E forms and certifications]
'''
req_diagram_path = render_mermaid(mermaid_code_req, 'requirements_flow')
doc.add_picture(str(req_diagram_path), width=Inches(5.8))
figure_counter['n'] += 1
add_caption(f"Figure {figure_counter['n']}: Flow of RFP requirements into workstreams and deliverables")

# Compliance matrix table (summarized but covering all IDs)
reqs = [
    ('REQ-001', 'Perform all Phase 1 SSES tasks (Tasks 1–8) in LaSalle sewersheds.', 'Mandatory', 'Sections 5–13, 14, 19'),
    ('REQ-002', 'Identify a dedicated Project Manager as primary contact.', 'Mandatory', 'Sections 4, 15'),
    ('REQ-003', 'Review existing data, records, and Consent Order correspondence.', 'Mandatory', 'Sections 2, 6, 7, 9, 10, 12'),
    ('REQ-004', 'Gain familiarity with Phase 1 sewersheds and I&I history.', 'Mandatory', 'Sections 2, 5–8, 10–13'),
    ('REQ-005', 'Meet with NFWB staff to discuss scope and system response.', 'Mandatory', 'Sections 4, 6'),
    ('REQ-006', 'Provide ongoing consultation to NFWB and NYSDEC.', 'Mandatory', 'Sections 4, 5, 12, 13, 14'),
    ('REQ-007', 'Provide staff, equipment, and materials for smoke testing.', 'Mandatory', 'Sections 5, 7, 18'),
    ('REQ-008', 'Notify residents and agencies prior to smoke testing.', 'Mandatory', 'Sections 4, 7, 18'),
    ('REQ-009', 'Conduct smoke testing in proper weather and document defects.', 'Mandatory', 'Sections 5, 7'),
    ('REQ-010', 'Review prior flow data before night-time weiring.', 'Mandatory', 'Sections 5, 6, 8'),
    ('REQ-011', 'Select adequate weir locations for accurate analysis.', 'Mandatory', 'Sections 5, 8'),
    ('REQ-012', 'Provide staff, equipment, traffic control, and confined space support for flow isolation.', 'Mandatory', 'Sections 5, 8, 18'),
    ('REQ-013', 'Notify residents before disruptive flow isolation work.', 'Mandatory', 'Sections 4, 8, 18'),
    ('REQ-014', 'Analyze flow data and calculate infiltration (gpd/lf).', 'Mandatory', 'Sections 5, 8, 12'),
    ('REQ-015', 'Analyze CCTV/manhole reports and identify data gaps.', 'Mandatory', 'Sections 5, 6, 9, 12'),
    ('REQ-016', 'Review outfall information and develop Work Plan.', 'Mandatory', 'Sections 5, 10, 12'),
    ('REQ-017', 'Execute Outfall/SSO Work Plan and document findings.', 'Mandatory', 'Sections 5, 10, 12'),
    ('REQ-018', 'Conduct focused SSES using $10,000 contingency.', 'Mandatory', 'Sections 5, 11, 12, 19'),
    ('REQ-019', 'Prepare NYSDEC-acceptable SSES report and corrective action plan.', 'Mandatory', 'Sections 1, 2, 5, 12, 14'),
    ('REQ-020', 'Complete SSES report in time for early 2023 NYSDEC review.', 'Mandatory', 'Sections 1, 5, 12, 14'),
    ('REQ-021', 'Ensure report is PE-stamped and follows Engineering Report Outline.', 'Mandatory', 'Sections 5, 12, 20'),
    ('REQ-022', 'Conduct post-construction monitoring and reporting.', 'Mandatory', 'Sections 5, 13, 14'),
    ('REQ-023', 'Provide option pricing for Phase 2 and Phase 3 SSES.', 'Mandatory', 'Sections 1, 14, 19'),
    ('REQ-024', 'Comply with Appendix E terms unless exceptions listed.', 'Mandatory', 'Sections 20, 22'),
    ('REQ-025', 'Attend mandatory pre-proposal information session.', 'Mandatory', 'Sections 4, 21'),
    ('REQ-026', 'Submit questions per communication restrictions.', 'Mandatory', 'Sections 20, 21'),
    ('REQ-027', 'Submit comprehensive proposal including scope, schedule, and budget.', 'Mandatory', 'All sections, especially 1–5, 14–19, 22'),
    ('REQ-028', 'Provide bar-chart schedules and task/hour estimate matrix.', 'Mandatory', 'Sections 5, 14, 19'),
    ('REQ-029', 'Complete and return all Appendix E forms.', 'Mandatory', 'Sections 20, 21'),
    ('REQ-030', 'Submit MWBE/SDVOB utilization plans and meet goals.', 'Mandatory', 'Section 17, 20'),
    ('REQ-031', 'Submit EEO Staffing Plan and utilization reports.', 'Mandatory', 'Sections 17, 20'),
    ('REQ-032', 'Maintain required insurance coverages and additional insureds.', 'Mandatory', 'Section 20'),
    ('REQ-033', 'Provide adequate and competent personnel.', 'Mandatory', 'Sections 4, 7–11, 15, 18'),
    ('REQ-034', 'Develop Safety Program and Site-Specific H&S Plan if requested.', 'Noted', 'Sections 18, 20'),
    ('REQ-035', 'Agree to indemnification and waiver of subrogation.', 'Mandatory', 'Section 20'),
    ('REQ-036', 'Maintain and provide access to records for six years.', 'Mandatory', 'Section 20'),
    ('REQ-037', 'Comply with NYS Information Security Breach laws.', 'Mandatory', 'Section 20'),
    ('REQ-038', 'Pay prevailing wages and submit utilization reports.', 'Mandatory', 'Sections 17, 18, 20'),
    ('REQ-039', 'Provide monthly invoices and comply to receive payment.', 'Mandatory', 'Sections 4, 19, 20, 21'),
    ('REQ-040', 'Certify non-discrimination and sexual harassment policy.', 'Mandatory', 'Section 20'),
    ('REQ-041', 'Certify statutory requirements including Iran Divestment and boycott prohibitions.', 'Mandatory', 'Section 20'),
    ('REQ-042', 'Acknowledge NFWB rights regarding evaluation and contract.', 'Mandatory', 'Sections 1, 20, 22'),
    ('REQ-043', 'Provide Statement of Qualifications and disclosures.', 'Mandatory', 'Sections 15, 16, 20'),
    ('REQ-044', 'Deliver required copies and electronic version by deadline.', 'Mandatory', 'Section 21'),
    ('REQ-045', 'Acknowledge receipt of addenda.', 'Mandatory', 'Sections 20, 21'),
    ('REQ-046', 'Agree to lobbying law restrictions and disclosure.', 'Mandatory', 'Sections 20, 21'),
    ('REQ-047', 'Comply with MWBE/SDVOB monthly reporting and proof of payment.', 'Mandatory', 'Sections 17, 19, 20'),
    ('REQ-048', 'Ensure no suspended or debarred entities are used.', 'Mandatory', 'Sections 15, 17, 20'),
    ('REQ-049', 'Demonstrate understanding of Consent Order and R-value methodology.', 'Noted', 'Sections 2, 5, 6, 10, 12'),
    ('REQ-050', 'Document outfall locations 013, 014, and 017 and relevance.', 'Noted', 'Sections 2, 5, 10, 12, 13'),
]

# Create table with multiple rows
req_table = doc.add_table(rows=1, cols=4)
req_table.style = 'Table Grid'
req_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = req_table.rows[0].cells
hdr[0].text = 'Requirement ID'
hdr[1].text = 'Requirement summary'
hdr[2].text = 'Mandatory / priority'
hdr[3].text = 'Proposal section reference'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

for rid, desc, mand, ref in reqs:
    row_cells = req_table.add_row().cells
    row_cells[0].text = rid
    row_cells[1].text = desc
    row_cells[2].text = mand
    row_cells[3].text = ref

# 4. Project Management and Governance Approach
doc.add_page_break()
h1 = doc.add_heading('4. Project Management and Governance Approach', level=1)

p = doc.add_paragraph(
    'Aurelia will manage the Phase 1 SSES as a structured, governance-driven engagement with a '
    'dedicated Project Manager (PM) as the primary point of contact for NFWB (REQ-002). Our PM '
    'will coordinate all technical disciplines, subconsultants, and field crews; manage scope, '
    'schedule, and budget; and facilitate decision making with NFWB and NYSDEC (REQ-005, '
    'REQ-006, REQ-020, REQ-025, REQ-027, REQ-039, REQ-044).'
)

p = doc.add_paragraph(
    'Governance will include a formal kickoff, biweekly progress meetings during active field work, '
    'monthly status reports aligned with invoicing, and targeted workshops for key milestones such '
    'as selection of flow isolation locations, review of preliminary findings, and development of the '
    'Corrective Action Plan. Communication protocols will respect NYS lobbying law restrictions and '
    'ensure that all official communications during the Restricted Period occur through the authorized '
    'NFWB contact (REQ-026, REQ-046).'
)

# Governance org chart Mermaid diagram
mermaid_code_org = '''
flowchart TD
  NFWB[NFWB leadership] --> PM[Aurelia project manager]
  PM --> ENG[Technical leads]
  PM --> FIELD[Field crews]
  PM --> MWBE[MWBE and SDVOB partners]
  PM --> NYSDEC[NYSDEC coordination]
'''
org_diagram_path = render_mermaid(mermaid_code_org, 'governance_org')
doc.add_picture(str(org_diagram_path), width=Inches(5.8))
figure_counter['n'] += 1
add_caption(f"Figure {figure_counter['n']}: Project governance structure")

# Communication flow Mermaid diagram
mermaid_code_comm = '''
flowchart LR
  NFWB[NFWB staff] <---> PM[Aurelia project manager]
  PM <---> TECH[Discipline leads]
  PM --> RES[Residents]
  PM --> EMG[Police and fire]
  PM --> NYSDEC[NYSDEC]
'''
comm_diagram_path = render_mermaid(mermaid_code_comm, 'communication_flow')
doc.add_picture(str(comm_diagram_path), width=Inches(5.8))
figure_counter['n'] += 1
add_caption(f"Figure {figure_counter['n']}: Communication pathways for the Phase 1 SSES")

# RACI table
table = doc.add_table(rows=1, cols=5)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'Activity'
hdr[1].text = 'NFWB'
hdr[2].text = 'Aurelia PM'
hdr[3].text = 'Aurelia technical leads'
hdr[4].text = 'Subconsultants / partners'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

rows = [
    ('Project initiation and scope confirmation', 'A', 'R', 'C', 'I'),
    ('Consent Order and data review workshops', 'C', 'R', 'R', 'I'),
    ('Smoke testing planning and notifications', 'C', 'R', 'R', 'C'),
    ('Flow isolation and weiring operations', 'C', 'R', 'R', 'C'),
    ('Outfall/SSO Work Plan approval', 'A', 'R', 'C', 'I'),
    ('Corrective Action Plan development', 'A', 'R', 'R', 'C'),
    ('MWBE/SDVOB reporting', 'I', 'R', 'I', 'C'),
    ('Monthly invoicing and progress reporting', 'A', 'R', 'C', 'I'),
]

for act, nfw, pm, tech, subs in rows:
    cells = table.add_row().cells
    cells[0].text = act
    cells[1].text = nfw
    cells[2].text = pm
    cells[3].text = tech
    cells[4].text = subs

# Risk register table
table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'Risk'
hdr[1].text = 'Likelihood / impact'
hdr[2].text = 'Mitigation strategy'
hdr[3].text = 'Owner'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

rows = [
    (
        'Adverse weather delaying smoke testing and flow isolation',
        'Medium / High',
        'Plan seasonal windows, maintain schedule float, and resequence tasks to maximize productive days.',
        'Aurelia PM'
    ),
    (
        'Data gaps in historical CCTV or flow records',
        'Medium / Medium',
        'Identify gaps early during Task 1 and use focused SSES contingency to fill critical gaps.',
        'Aurelia technical leads'
    ),
    (
        'NYSDEC review taking longer than anticipated',
        'Low / High',
        'Engage NYSDEC during report development, provide clear traceability, and submit a high-quality draft.',
        'NFWB and Aurelia PM'
    ),
    (
        'Community concerns about smoke testing or night work',
        'Medium / Medium',
        'Implement robust public communication, coordinate with police and fire, and provide NFWB with FAQs.',
        'Aurelia PM and NFWB'
    ),
]

for r, li, mit, own in rows:
    cells = table.add_row().cells
    cells[0].text = r
    cells[1].text = li
    cells[2].text = mit
    cells[3].text = own

# 5. Technical Approach – Phase 1 SSES Overview
doc.add_page_break()
h1 = doc.add_heading('5. Technical Approach – Phase 1 SSES Overview', level=1)

p = doc.add_paragraph(
    'Our Phase 1 SSES methodology is structured around the eight tasks defined in the RFP and is '
    'designed to satisfy NYSDEC’s Engineering Report Outline and the Consent Order schedule '
    '(REQ-001, REQ-003, REQ-004, REQ-005, REQ-006, REQ-007, REQ-009, REQ-010, REQ-011, '
    'REQ-012, REQ-014, REQ-015, REQ-016, REQ-017, REQ-018, REQ-019, REQ-020, REQ-021, '
    'REQ-022, REQ-049, REQ-050). The tasks are tightly integrated so that findings from early work '
    'inform later field investigations and the final Corrective Action Plan.'
)

p = doc.add_paragraph(
    'We begin with project initiation, data review, and planning (Task 1), followed by smoke testing '
    '(Task 2) and flow isolation and measurement (Task 3) to identify inflow and infiltration patterns. '
    'Task 4 focuses on CCTV and manhole report review, while Task 5 addresses outfall and SSO '
    'condition assessment. Task 6 provides a focused SSES contingency for targeted investigations. '
    'Task 7 synthesizes all findings into an NYSDEC-acceptable SSES Engineering Report and '
    'Corrective Action Plan, and Task 8 provides post-construction monitoring and I&I assessment.'
)

# High-level process flow Mermaid diagram
mermaid_code_overview = '''
flowchart LR
  T1[Task 1 initiation and planning] --> T2[Task 2 smoke testing]
  T2 --> T3[Task 3 flow isolation]
  T3 --> T4[Task 4 CCTV review]
  T4 --> T5[Task 5 outfall assessment]
  T5 --> T6[Task 6 focused SSES]
  T6 --> T7[Task 7 engineering report]
  T7 --> T8[Task 8 post construction monitoring]
'''
overview_diagram_path = render_mermaid(mermaid_code_overview, 'tasks_overview')
doc.add_picture(str(overview_diagram_path), width=Inches(5.8))
figure_counter['n'] += 1
add_caption(f"Figure {figure_counter['n']}: Phase 1 SSES task sequence and dependencies")

# Gantt style bar chart for tasks
plt.figure(figsize=(8, 4.5))
sns.set_style('whitegrid')
phases = ['Task 1', 'Task 2', 'Task 3', 'Task 4', 'Task 5', 'Task 6', 'Task 7', 'Task 8']
weeks = [4, 6, 6, 5, 4, 3, 8, 6]
start_weeks = [0, 4, 6, 8, 10, 12, 14, 22]

df_tasks = pd.DataFrame({'Task': phases, 'Start': start_weeks, 'Duration': weeks})
fig, ax = plt.subplots(figsize=(8, 4.5))
colors = ['#4c72b0', '#55a868', '#c44e52', '#8172b3', '#ccb974', '#64b5cd', '#8da0cb', '#66c2a5']
for i, row in df_tasks.iterrows():
    ax.barh(row['Task'], row['Duration'], left=row['Start'], color=colors[i], edgecolor='white')

ax.invert_yaxis()
ax.set_xlabel('Weeks from notice to proceed')
ax.set_title('Illustrative Phase 1 SSES task durations')
ax.xaxis.grid(True, linestyle='--', alpha=0.3)
plt.tight_layout()
sseschart_path = output_dir / 'tasks_timeline.png'
plt.savefig(sseschart_path, dpi=150, bbox_inches='tight', facecolor='white')
plt.close()

doc.add_picture(str(sseschart_path), width=Inches(6.0))
figure_counter['n'] += 1
add_caption(f"Figure {figure_counter['n']}: Illustrative Phase 1 SSES task durations and sequencing")

# Task summary table
table = doc.add_table(rows=1, cols=5)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'Task'
hdr[1].text = 'Objectives'
hdr[2].text = 'Key activities'
hdr[3].text = 'Key inputs / outputs'
hdr[4].text = 'Primary roles'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

rows = [
    ('Task 1 – Initiation and planning', 'Establish scope, schedule, and data foundation.', 'Kickoff, data requests, workshops, Consent Order review.', 'Inputs: RFP, Consent Order, existing data. Outputs: Project plan, data inventory.', 'PM, system analyst.'),
    ('Task 2 – Smoke testing', 'Identify inflow sources and cross-connections.', 'Field smoke testing, notifications, documentation.', 'Inputs: sewershed maps. Outputs: defect logs, photos, GIS layers.', 'Field lead, safety officer.'),
    ('Task 3 – Flow isolation', 'Quantify infiltration by segment.', 'Night-time weiring, flow measurement, traffic control.', 'Inputs: prior flow data. Outputs: gpd/lf calculations.', 'Hydraulic engineer, field crew.'),
    ('Task 4 – CCTV/manhole review', 'Characterize structural and I&I defects.', 'Review CCTV reports, code defects, identify gaps.', 'Inputs: CCTV footage/reports. Outputs: defect database.', 'Condition assessment lead.'),
    ('Task 5 – Outfall/SSO assessment', 'Understand outfall performance and drivers.', 'Site visits, Work Plan, monitoring, dye testing.', 'Inputs: outfall records. Outputs: assessment results.', 'Hydraulic engineer, field crew.'),
    ('Task 6 – Focused SSES', 'Resolve uncertainties and confirm sources.', 'Targeted dye testing, house inspections, additional CCTV.', 'Inputs: earlier task findings. Outputs: refined defect locations.', 'PM, field team.'),
    ('Task 7 – Report and CAP', 'Deliver NYSDEC-acceptable SSES report and plan.', 'Synthesis, alternatives analysis, cost estimating, scheduling.', 'Inputs: all task data. Outputs: PE-stamped report and CAP.', 'Lead engineer, PM.'),
    ('Task 8 – Post-construction monitoring', 'Verify I&I reduction and cost-effectiveness.', 'Flow and rainfall monitoring, before/after analysis.', 'Inputs: pre and post data. Outputs: Post-Construction I&I Assessment.', 'Hydraulic engineer.'),
]

for row in rows:
    cells = table.add_row().cells
    for i, val in enumerate(row):
        cells[i].text = val

# 6. Detailed Methodology – Task 1
doc.add_page_break()
h1 = doc.add_heading('6. Detailed Methodology – Task 1: Project Initiation, Data Review, and Planning', level=1)

p = doc.add_paragraph(
    'Task 1 establishes the technical and governance foundation for the SSES. We will conduct a '
    'formal kickoff meeting with NFWB to confirm scope, deliverables, and success criteria; review '
    'the Consent Order and revised work schedule; and finalize the project communication and '
    'safety expectations (REQ-001, REQ-003, REQ-004, REQ-005, REQ-010, REQ-015, REQ-019, '
    'REQ-027, REQ-049).'
)

p = doc.add_paragraph(
    'Our team will compile and review available records, including prior SSES reports, SCADA '
    'data, flow monitoring results, CCTV and manhole inspection reports, maintenance logs, and '
    'NYSDEC correspondence related to LaSalle SSO abatement. We will use structured data '
    'checklists and GIS mapping to identify gaps and prioritize any additional data collection.'
)

# Task 1 sequence diagram
mermaid_code_t1 = '''
sequenceDiagram
  participant NFWB
  participant PM
  participant Eng
  NFWB->>PM: Provide existing data and Consent Order
  PM->>Eng: Distribute data and assign reviews
  Eng->>PM: Summaries and data gaps
  PM->>NFWB: Planning workshop and confirmation of priorities
'''
seq_t1_path = render_mermaid(mermaid_code_t1, 'task1_sequence')
doc.add_picture(str(seq_t1_path), width=Inches(5.8))
figure_counter['n'] += 1
add_caption(f"Figure {figure_counter['n']}: Task 1 initiation and planning interactions")

# SCADA flow vs rainfall illustrative chart
dates = pd.date_range(start='2022-03-01', periods=14, freq='D')
flow = [4.2, 4.1, 4.3, 5.0, 5.5, 4.6, 4.3, 4.2, 5.1, 5.6, 4.5, 4.3, 4.2, 4.1]
rain = [0.0, 0.0, 0.1, 0.8, 1.2, 0.0, 0.0, 0.0, 0.6, 1.0, 0.0, 0.0, 0.0, 0.0]

df_scada = pd.DataFrame({'Date': dates, 'Flow MGD': flow, 'Rainfall inches': rain})
fig, ax1 = plt.subplots(figsize=(7, 4))
color_flow = '#4c72b0'
color_rain = '#dd8452'
ax1.plot(df_scada['Date'], df_scada['Flow MGD'], color=color_flow, marker='o')
ax1.set_xlabel('Date')
ax1.set_ylabel('Flow (MGD)', color=color_flow)
ax1.tick_params(axis='y', labelcolor=color_flow)
ax2 = ax1.twinx()
ax2.bar(df_scada['Date'], df_scada['Rainfall inches'], color=color_rain, alpha=0.4)
ax2.set_ylabel('Rainfall (inches)', color=color_rain)
ax2.tick_params(axis='y', labelcolor=color_rain)
plt.title('Illustrative SCADA flow and rainfall response')
fig.autofmt_xdate(rotation=45)
plt.tight_layout()
scada_chart_path = output_dir / 'scada_flow_rain.png'
plt.savefig(scada_chart_path, dpi=150, bbox_inches='tight', facecolor='white')
plt.close()

doc.add_picture(str(scada_chart_path), width=Inches(5.8))
figure_counter['n'] += 1
add_caption(f"Figure {figure_counter['n']}: Example analysis of SCADA flow response to rainfall")

# Task 1 data checklist table
table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'Data source'
hdr[1].text = 'Description'
hdr[2].text = 'Use in SSES'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

rows = [
    ('Consent Order and NYSDEC correspondence', 'Order on Consent R9020080528-32 and related letters.', 'Define regulatory context, milestones, and reporting expectations (REQ-003, REQ-049).'),
    ('Previous SSES reports', 'Any prior I&I studies in LaSalle or related sewersheds.', 'Leverage past findings and avoid duplication.'),
    ('SCADA and flow data', 'Pump station flows, wet-weather trends, alarms.', 'Understand system response and target monitoring (REQ-004, REQ-010).'),
    ('CCTV and manhole reports', 'Existing condition assessments for mains and structures.', 'Identify structural and I&I defects and data gaps (REQ-015).'),
    ('Maintenance and repair history', 'Work orders, lining projects, spot repairs.', 'Account for completed work and residual issues.'),
]

for row in rows:
    cells = table.add_row().cells
    for i, val in enumerate(row):
        cells[i].text = val

# 7. Detailed Methodology – Task 2: Smoke Testing Program
doc.add_page_break()
h1 = doc.add_heading('7. Detailed Methodology – Task 2: Smoke Testing Program', level=1)

p = doc.add_paragraph(
    'Smoke testing is a core tool for identifying inflow sources such as cross-connections, roof '
    'leaders, and defective cleanouts. Aurelia will provide all staff, equipment, and materials needed '
    'to conduct smoke testing in the Phase 1 sewersheds (REQ-001, REQ-004, REQ-007). We will '
    'implement a robust public notification and safety program, including door hangers 2–4 days '
    'before testing, assistance with NFWB press releases, and coordination with police and fire '
    'officials (REQ-008, REQ-033, REQ-034).'
)

p = doc.add_paragraph(
    'Testing will be scheduled only during appropriate weather conditions to ensure meaningful '
    'results and minimize nuisance impacts. All observed defects will be documented and '
    'photographed with GIS-referenced locations for integration into the defect database and '
    'Corrective Action Plan (REQ-009).'
)

# Smoke testing workflow Mermaid diagram
mermaid_code_smoke = '''
flowchart LR
  P[Plan smoke testing] --> N[Notify residents and agencies]
  N --> F[Field setup and safety briefing]
  F --> T[Conduct smoke testing]
  T --> D[Document and photograph defects]
  D --> G[Update GIS and defect database]
'''
smoke_diagram_path = render_mermaid(mermaid_code_smoke, 'smoke_workflow')
doc.add_picture(str(smoke_diagram_path), width=Inches(5.8))
figure_counter['n'] += 1
add_caption(f"Figure {figure_counter['n']}: Smoke testing workflow from planning to documentation")

# Smoke testing segments bar chart
plt.figure(figsize=(6, 4))
segments_data = pd.DataFrame({
    'Sewershed': ['LaSalle A', 'LaSalle B', 'LaSalle C', 'LaSalle D'],
    'Segments planned': [85, 70, 60, 55]
})
ax = sns.barplot(data=segments_data, x='Sewershed', y='Segments planned', color='#4c72b0')
ax.set_title('Illustrative smoke testing segments by sewershed')
ax.set_ylabel('Number of segments')
plt.tight_layout()
smoke_chart_path = output_dir / 'smoke_segments.png'
plt.savefig(smoke_chart_path, dpi=150, bbox_inches='tight', facecolor='white')
plt.close()

doc.add_picture(str(smoke_chart_path), width=Inches(5.8))
figure_counter['n'] += 1
add_caption(f"Figure {figure_counter['n']}: Illustrative smoke testing coverage by sewershed")

# Notification plan table
table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'Audience'
hdr[1].text = 'Notification method'
hdr[2].text = 'Timing'
hdr[3].text = 'Key content'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

rows = [
    ('Residents and businesses', 'Hand-delivered door hangers.', '2–4 days before testing (REQ-008).', 'Purpose of testing, safety information, contact details.'),
    ('Police and fire', 'Email and phone coordination.', 'At least 3 days before and day-of reminder.', 'Schedule, locations, and emergency contact protocol.'),
    ('NFWB communications team', 'Draft press release and FAQs.', 'One week before testing.', 'Plain-language explanation and key messages.'),
]

for row in rows:
    cells = table.add_row().cells
    for i, val in enumerate(row):
        cells[i].text = val

# Equipment and staffing table
table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'Role / equipment'
hdr[1].text = 'Quantity'
hdr[2].text = 'Key responsibilities'
hdr[3].text = 'Safety considerations'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

rows = [
    ('Field supervisor', '1', 'Oversee daily operations and safety.', 'Ensures adherence to H&S Plan (REQ-033, REQ-034).'),
    ('Smoke testing technicians', '3–4', 'Set up blowers, inject smoke, observe defects.', 'Use PPE and follow traffic control procedures.'),
    ('Blowers and smoke generators', '2–3', 'Provide visible smoke in sewer mains.', 'Proper placement and monitoring to avoid nuisance.'),
    ('Vehicle and traffic control devices', '1 truck, cones, signs', 'Safe access to manholes in streets.', 'Comply with MUTCD and local requirements.'),
]

for row in rows:
    cells = table.add_row().cells
    for i, val in enumerate(row):
        cells[i].text = val

# 8. Detailed Methodology – Task 3: Flow Isolation and Measurement
doc.add_page_break()
h1 = doc.add_heading('8. Detailed Methodology – Task 3: Flow Isolation and Measurement', level=1)

p = doc.add_paragraph(
    'Flow isolation and night-time weiring will quantify groundwater infiltration and guide targeted '
    'corrective actions. Before installing weirs, we will review prior flow measurement and analysis '
    'for each relevant sewershed (REQ-010) and work with NFWB staff to select adequate '
    'locations and number of weirs to support accurate infiltration analysis (REQ-011).'
)

p = doc.add_paragraph(
    'Aurelia will provide all staff, equipment, and materials necessary for flow isolation and '
    'measurement, including traffic control and confined space entry-trained personnel and '
    'equipment (REQ-012, REQ-033, REQ-034). Residents will be notified in advance of any '
    'disruptive night work via door hangers and coordination with NFWB’s communications team '
    '(REQ-013). We will calculate infiltration in gallons per day per linear foot (gpd/lf) for each '
    'weiring section and integrate these results with CCTV and smoke testing findings (REQ-014).'
)

# Flow isolation process Mermaid diagram
mermaid_code_flow = '''
flowchart LR
  R[Review prior flow data] --> S[Select weir locations]
  S --> P[Plan night work and notifications]
  P --> I[Install weirs and measure flow]
  I --> C[Calculate infiltration gpd per lf]
  C --> U[Update hydraulic analysis and CAP]
'''
flow_diagram_path = render_mermaid(mermaid_code_flow, 'flow_isolation_process')
doc.add_picture(str(flow_diagram_path), width=Inches(5.8))
figure_counter['n'] += 1
add_caption(f"Figure {figure_counter['n']}: Flow isolation and weiring process")

# Infiltration results template chart
sections = ['Section 1', 'Section 2', 'Section 3', 'Section 4']
infiltration = [1200, 800, 600, 1500]
plt.figure(figsize=(6, 4))
inf_data = pd.DataFrame({'Weir section': sections, 'Infiltration gpd per lf': infiltration})
ax = sns.barplot(data=inf_data, x='Weir section', y='Infiltration gpd per lf', color='#55a868')
ax.set_title('Illustrative infiltration rates by weir section')
ax.set_ylabel('Infiltration (gpd/lf)')
plt.tight_layout()
inf_chart_path = output_dir / 'infiltration_rates.png'
plt.savefig(inf_chart_path, dpi=150, bbox_inches='tight', facecolor='white')
plt.close()

doc.add_picture(str(inf_chart_path), width=Inches(5.8))
figure_counter['n'] += 1
add_caption(f"Figure {figure_counter['n']}: Example presentation of infiltration results by section")

# Flow isolation results table template
table = doc.add_table(rows=1, cols=5)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'Location ID'
hdr[1].text = 'Conditions'
hdr[2].text = 'Measured flow (gpm)'
hdr[3].text = 'Infiltration (gpd/lf)'
hdr[4].text = 'Notes'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

rows = [
    ('FI-01', 'Dry weather, stable groundwater.', '12', '1,200', 'High infiltration, candidate for lining.'),
    ('FI-02', 'Dry weather, low groundwater.', '6', '600', 'Moderate infiltration, monitor after repairs upstream.'),
]

for row in rows:
    cells = table.add_row().cells
    for i, val in enumerate(row):
        cells[i].text = val

# 9. Detailed Methodology – Task 4: CCTV and Manhole Report Review
doc.add_page_break()
h1 = doc.add_heading('9. Detailed Methodology – Task 4: CCTV and Manhole Report Review', level=1)

p = doc.add_paragraph(
    'Task 4 leverages existing CCTV and manhole inspection data provided by NFWB to '
    'characterize structural and I&I-related defects and identify data gaps (REQ-001, REQ-003, '
    'REQ-004, REQ-015, REQ-018, REQ-033). Aurelia will ingest digital CCTV footage and reports, '
    'standardize defect coding, and integrate results into a GIS-based asset and defect database.'
)

p = doc.add_paragraph(
    'We will coordinate with NFWB to understand repairs already completed, confirm access to '
    'footage, and prioritize any additional CCTV needed under Task 6 focused SSES. This ensures '
    'that the Corrective Action Plan reflects current conditions and avoids recommending repairs '
    'that have already been addressed.'
)

# Data flow Mermaid diagram for CCTV
mermaid_code_cctv = '''
flowchart LR
  C1[CCTV footage and reports] --> P1[Review and code defects]
  P1 --> G1[GIS and database]
  G1 --> A1[Analysis and prioritization]
  A1 --> R1[Corrective Action Plan]
'''
cctv_diagram_path = render_mermaid(mermaid_code_cctv, 'cctv_data_flow')
doc.add_picture(str(cctv_diagram_path), width=Inches(5.8))
figure_counter['n'] += 1
add_caption(f"Figure {figure_counter['n']}: CCTV and manhole data integration workflow")

# Defect distribution bar chart
plt.figure(figsize=(6, 4))
defect_data = pd.DataFrame({
    'Defect type': ['Structural', 'I&I related', 'Maintenance', 'Other'],
    'Count': [120, 95, 60, 15]
})
ax = sns.barplot(data=defect_data, x='Defect type', y='Count', color='#4c72b0')
ax.set_title('Illustrative distribution of defect types')
ax.set_ylabel('Number of defects')
plt.tight_layout()
defect_chart_path = output_dir / 'defect_distribution.png'
plt.savefig(defect_chart_path, dpi=150, bbox_inches='tight', facecolor='white')
plt.close()

doc.add_picture(str(defect_chart_path), width=Inches(5.8))
figure_counter['n'] += 1
add_caption(f"Figure {figure_counter['n']}: Example distribution of CCTV-identified defects")

# Defect summary table
table = doc.add_table(rows=1, cols=5)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'Segment / manhole ID'
hdr[1].text = 'Defect type'
hdr[2].text = 'Severity'
hdr[3].text = 'Recommended action'
hdr[4].text = 'Notes'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

rows = [
    ('MH-101', 'I&I at frame and cover', 'High', 'Raise and seal frame, replace cover.', 'Near low-lying area with ponding.'),
    ('SEG-205', 'Longitudinal crack', 'Medium', 'Monitor and consider lining with adjacent segments.', 'Coordinate with infiltration data.'),
]

for row in rows:
    cells = table.add_row().cells
    for i, val in enumerate(row):
        cells[i].text = val

# Additional CCTV needs table
table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'Area'
hdr[1].text = 'Reason for additional CCTV'
hdr[2].text = 'Estimated effort'
hdr[3].text = 'Coordination with NFWB crews'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

rows = [
    ('LaSalle B subcatchment 2', 'No recent CCTV and high infiltration from Task 3.', '2 days.', 'Schedule with NFWB CCTV crew under Task 6.'),
]

for row in rows:
    cells = table.add_row().cells
    for i, val in enumerate(row):
        cells[i].text = val

# 10. Detailed Methodology – Task 5: Outfall/SSO Condition Assessment
doc.add_page_break()
h1 = doc.add_heading('10. Detailed Methodology – Task 5: Outfall/SSO Condition Assessment', level=1)

p = doc.add_paragraph(
    'Task 5 focuses on understanding the condition and performance of the two Phase 1 outfalls '
    'and their relationship to Phase 1 and Phase 3 sewersheds (REQ-001, REQ-004, REQ-016, '
    'REQ-017, REQ-050). We will review existing information for outfalls 013, 014, and 017, '
    'conduct site visits to confirm physical conditions, and develop an Outfall/SSO Work Plan in '
    'coordination with NFWB.'
)

p = doc.add_paragraph(
    'The Work Plan may include temporary flow meters, rain gauges, dye testing, and coordination '
    'for additional CCTV as needed. Upon NFWB authorization, we will execute the Work Plan and '
    'document findings in GIS-compatible format, including recommendations for reducing '
    'extraneous flows and mitigating SSOs (REQ-017).'
)

# Outfall assessment workflow Mermaid diagram
mermaid_code_outfall = '''
flowchart LR
  O1[Review outfall records] --> O2[Site visits and condition checks]
  O2 --> O3[Develop Outfall SSO Work Plan]
  O3 --> O4[Install meters and collect data]
  O4 --> O5[Analyze flows and rainfall]
  O5 --> O6[Recommend outfall improvements]
'''
outfall_diagram_path = render_mermaid(mermaid_code_outfall, 'outfall_workflow')
doc.add_picture(str(outfall_diagram_path), width=Inches(5.8))
figure_counter['n'] += 1
add_caption(f"Figure {figure_counter['n']}: Outfall and SSO condition assessment workflow")

# Outfall inventory table
table = doc.add_table(rows=1, cols=5)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'Outfall ID'
hdr[1].text = 'Location description'
hdr[2].text = 'Current condition (illustrative)'
hdr[3].text = 'Associated sewersheds'
hdr[4].text = 'Consent Order relevance'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

rows = [
    ('013', 'LaSalle area outfall to receiving water.', 'Operational, evidence of surcharge.', 'LaSalle A, LaSalle B.', 'Key SSO location under Consent Order (REQ-050).'),
    ('014', 'Downstream outfall near commercial corridor.', 'Operational, minor debris.', 'LaSalle B.', 'Supports R-value-based prioritization.'),
    ('017', 'Outfall serving Phase 1 and Phase 3 areas.', 'Operational, subject to high flows.', 'LaSalle C, Phase 3 sewersheds.', 'Critical to long-term LaSalle strategy.'),
]

for row in rows:
    cells = table.add_row().cells
    for i, val in enumerate(row):
        cells[i].text = val

# Work Plan summary table
table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'Work Plan element'
hdr[1].text = 'Purpose'
hdr[2].text = 'Trigger'
hdr[3].text = 'Notes'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

rows = [
    ('Temporary flow meters', 'Measure outfall discharge under various conditions.', 'Evidence of unexplained SSOs or high flows.', 'Sized and installed per manufacturer guidance.'),
    ('Rain gauges', 'Correlate rainfall with outfall response.', 'Needed to refine R-values or validate models.', 'Prefer co-location with existing gauges.'),
    ('Dye testing', 'Confirm suspected cross-connections.', 'Indications from smoke testing or CCTV.', 'Conducted with NFWB approval and safety controls.'),
]

for row in rows:
    cells = table.add_row().cells
    for i, val in enumerate(row):
        cells[i].text = val

# 11. Detailed Methodology – Task 6: Focused SSES (Contingency)
doc.add_page_break()
h1 = doc.add_heading('11. Detailed Methodology – Task 6: Focused SSES (Contingency)', level=1)

p = doc.add_paragraph(
    'Task 6 provides a structured mechanism to resolve uncertainties and confirm specific I&I '
    'sources using a $10,000 contingency budget (REQ-001, REQ-004, REQ-018, REQ-027, '
    'REQ-028, REQ-033). Aurelia will only use this contingency with explicit NFWB approval, '
    'based on documented triggers from Tasks 2–5.'
)

p = doc.add_paragraph(
    'Focused SSES activities may include dye testing of suspected cross-connections, house '
    'inspections for improper connections, and additional CCTV via NFWB crews. Each activity '
    'will be scoped, estimated, and approved before execution, with results integrated into the '
    'overall defect database and Corrective Action Plan.'
)

# Decision tree Mermaid diagram
mermaid_code_focus = '''
flowchart TD
  F0[Findings from Tasks 2 to 5] --> Q1[Is source location uncertain]
  Q1 -->|Yes| A1[Define focused SSES activity]
  A1 --> A2[Estimate effort and cost]
  A2 --> A3[NFWB approval]
  A3 -->|Approved| A4[Execute focused SSES]
  A4 --> A5[Update findings and CAP]
  Q1 -->|No| A6[No focused SSES needed]
'''
focus_diagram_path = render_mermaid(mermaid_code_focus, 'focused_sses_decision')
doc.add_picture(str(focus_diagram_path), width=Inches(5.8))
figure_counter['n'] += 1
add_caption(f"Figure {figure_counter['n']}: Decision process for using focused SSES contingency")

# Contingency tracking table
table = doc.add_table(rows=1, cols=5)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'Activity'
hdr[1].text = 'Estimated hours'
hdr[2].text = 'Estimated cost'
hdr[3].text = 'Approval status'
hdr[4].text = 'Notes'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

rows = [
    ('Dye testing in LaSalle B subcatchment 2', '16', '$3,200', 'To be approved by NFWB.', 'Triggered by high infiltration and limited CCTV.'),
]

for row in rows:
    cells = table.add_row().cells
    for i, val in enumerate(row):
        cells[i].text = val

# 12. Detailed Methodology – Task 7: SSES Engineering Report and Corrective Action Plan
doc.add_page_break()
h1 = doc.add_heading('12. Detailed Methodology – Task 7: SSES Engineering Report and Corrective Action Plan', level=1)

p = doc.add_paragraph(
    'Task 7 culminates in an SSES Engineering Report and Corrective Action Plan (CAP) that '
    'meet NYSDEC requirements and the Engineering Report Outline dated October 1, 2022 '
    '(REQ-001, REQ-003, REQ-014, REQ-015, REQ-016, REQ-017, REQ-018, REQ-019, REQ-020, '
    'REQ-021, REQ-027, REQ-028, REQ-049, REQ-050). The report will be prepared, stamped, and '
    'dated by a New York State licensed Professional Engineer (REQ-021).'
)

p = doc.add_paragraph(
    'We will synthesize findings from Tasks 1–6 into a transparent set of recommendations, '
    'prioritized by I&I reduction per dollar invested, constructability, and alignment with Consent '
    'Order milestones. Budgetary costs and implementation timelines will be developed for each '
    'project, with a clear bar-chart schedule to enable construction in 2023–early 2024 and '
    'post-construction monitoring by June 2024 (REQ-019, REQ-020).'
)

# Dependency diagram Mermaid
mermaid_code_report = '''
flowchart TD
  D1[Task 1 data review] --> RPT[SSES report and CAP]
  D2[Task 2 smoke testing] --> RPT
  D3[Task 3 flow isolation] --> RPT
  D4[Task 4 CCTV review] --> RPT
  D5[Task 5 outfall assessment] --> RPT
  D6[Task 6 focused SSES] --> RPT
'''
report_diagram_path = render_mermaid(mermaid_code_report, 'report_dependencies')
doc.add_picture(str(report_diagram_path), width=Inches(5.8))
figure_counter['n'] += 1
add_caption(f"Figure {figure_counter['n']}: Data dependencies feeding the SSES report and CAP")

# Corrective actions ranking chart
actions = ['Lining mains', 'Manhole rehab', 'Private inflow removal', 'Targeted repairs']
reduction_per_dollar = [1.8, 1.2, 2.1, 1.0]
plt.figure(figsize=(6, 4))
ca_data = pd.DataFrame({'Action': actions, 'I&I reduction per 1000 dollars': reduction_per_dollar})
ax = sns.barplot(data=ca_data, x='Action', y='I&I reduction per 1000 dollars', color='#4c72b0')
ax.set_title('Illustrative I&I reduction per 1,000 dollars invested')
ax.set_ylabel('Reduction (MGD per 1,000 dollars)')
plt.tight_layout()
ca_chart_path = output_dir / 'corrective_action_value.png'
plt.savefig(ca_chart_path, dpi=150, bbox_inches='tight', facecolor='white')
plt.close()

doc.add_picture(str(ca_chart_path), width=Inches(5.8))
figure_counter['n'] += 1
add_caption(f"Figure {figure_counter['n']}: Example ranking of corrective actions by I&I reduction per dollar")

# Corrective action matrix table
table = doc.add_table(rows=1, cols=6)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'Defect / location'
hdr[1].text = 'Recommended fix'
hdr[2].text = 'Estimated cost'
hdr[3].text = 'Expected I&I reduction'
hdr[4].text = 'Priority'
hdr[5].text = 'Implementation window'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

rows = [
    ('High infiltration in SEG-205', 'CIPP lining of 400 lf.', '$220,000', '0.25 MGD.', 'High.', '2023 construction.'),
    ('I&I at MH-101', 'Frame and cover raising and sealing.', '$15,000', '0.05 MGD.', 'Medium.', '2023 construction.'),
]

for row in rows:
    cells = table.add_row().cells
    for i, val in enumerate(row):
        cells[i].text = val

# Implementation schedule table
table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'Milestone'
hdr[1].text = 'Target date'
hdr[2].text = 'Dependency'
hdr[3].text = 'Notes'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

rows = [
    ('Draft SSES report to NFWB', 'January 2023', 'Completion of Tasks 1–6.', 'Enables internal review.'),
    ('Final SSES report to NYSDEC', 'March 2023', 'NFWB review and revisions.', 'Supports early NYSDEC review (REQ-020).'),
    ('Construction window', 'Mid 2023–early 2024', 'NYSDEC acceptance and funding.', 'Implements high-priority projects.'),
]

for row in rows:
    cells = table.add_row().cells
    for i, val in enumerate(row):
        cells[i].text = val

# 13. Detailed Methodology – Task 8: Post-Construction Monitoring and I&I Assessment
doc.add_page_break()
h1 = doc.add_heading('13. Detailed Methodology – Task 8: Post-Construction Monitoring and I&I Assessment', level=1)

p = doc.add_paragraph(
    'Task 8 verifies the effectiveness of the construction program through post-construction '
    'monitoring in 2024 (assuming 2023 construction) (REQ-001, REQ-004, REQ-020, REQ-022, '
    'REQ-027). Aurelia will design a monitoring plan that collects flow and rainfall data sufficient to '
    'compare pre- and post-construction conditions and normalize for rainfall variability.'
)

p = doc.add_paragraph(
    'We will prepare a Post-Construction I&I Assessment and Cost Analysis Report for NYSDEC '
    'documenting changes in base and wet-weather flows, confirming I&I reductions, and relating '
    'these outcomes to construction costs and Consent Order commitments.'
)

# Monitoring timeline Mermaid diagram
mermaid_code_monitor = '''
flowchart LR
  C1[Construction complete] --> M1[Post construction baseline monitoring]
  M1 --> M2[Wet weather monitoring]
  M2 --> M3[Data analysis and normalization]
  M3 --> M4[Post construction I and I report]
  M4 --> M5[NYSDEC review]
'''
monitor_diagram_path = render_mermaid(mermaid_code_monitor, 'monitor_timeline')
doc.add_picture(str(monitor_diagram_path), width=Inches(5.8))
figure_counter['n'] += 1
add_caption(f"Figure {figure_counter['n']}: Post-construction monitoring and reporting timeline")

# Before/after flow comparison chart
periods = ['Pre construction', 'Post construction']
flows = [5.0, 3.5]
plt.figure(figsize=(5, 4))
ba_data = pd.DataFrame({'Period': periods, 'Average wet weather flow MGD': flows})
ax = sns.barplot(data=ba_data, x='Period', y='Average wet weather flow MGD', color='#4c72b0')
ax.set_title('Illustrative pre and post construction flow comparison')
ax.set_ylabel('Flow (MGD)')
plt.tight_layout()
ba_chart_path = output_dir / 'before_after_flow.png'
plt.savefig(ba_chart_path, dpi=150, bbox_inches='tight', facecolor='white')
plt.close()

doc.add_picture(str(ba_chart_path), width=Inches(4.5))
figure_counter['n'] += 1
add_caption(f"Figure {figure_counter['n']}: Example pre- and post-construction flow comparison")

# Pre vs post I&I metrics table
table = doc.add_table(rows=1, cols=5)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'Sewershed'
hdr[1].text = 'Pre construction I&I (MGD)'
hdr[2].text = 'Post construction I&I (MGD)'
hdr[3].text = 'Reduction (MGD)'
hdr[4].text = 'Cost effectiveness (MGD per million dollars)'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

rows = [
    ('LaSalle A', '0.90', '0.55', '0.35', '0.70'),
    ('LaSalle B', '0.70', '0.45', '0.25', '0.60'),
]

for row in rows:
    cells = table.add_row().cells
    for i, val in enumerate(row):
        cells[i].text = val

# 14. Schedule and Key Milestones
doc.add_page_break()
h1 = doc.add_heading('14. Schedule and Key Milestones', level=1)

p = doc.add_paragraph(
    'Our schedule is designed to deliver the Phase 1 SSES report in time for early 2023 NYSDEC '
    'review, support construction in 2023–early 2024, and complete post-construction monitoring by '
    'June 2024 (REQ-001, REQ-019, REQ-020, REQ-022, REQ-023, REQ-027, REQ-028). We use a '
    'bar-chart schedule to clearly represent task durations, overlaps, and critical path activities.'
)

# Gantt chart using dates
schedule = pd.DataFrame({
    'Task': [
        'Task 1 initiation', 'Task 2 smoke testing', 'Task 3 flow isolation',
        'Task 4 CCTV review', 'Task 5 outfall assessment', 'Task 6 focused SSES',
        'Task 7 report and CAP', 'Task 8 post construction monitoring'
    ],
    'Team': [
        'PM and analysts', 'Field crews', 'Field crews',
        'Analysts', 'Field crews', 'Field crews',
        'Engineers', 'Engineers'
    ],
    'Start': pd.to_datetime([
        '2022-12-01', '2023-01-02', '2023-02-01',
        '2023-02-15', '2023-03-01', '2023-03-15',
        '2023-04-01', '2024-01-15'
    ]),
    'End': pd.to_datetime([
        '2022-12-31', '2023-02-15', '2023-03-10',
        '2023-03-31', '2023-03-31', '2023-04-15',
        '2023-06-30', '2024-06-15'
    ])
})

project_start = schedule['Start'].min()
schedule['days_to_start'] = (schedule['Start'] - project_start).dt.days
schedule['duration'] = (schedule['End'] - schedule['Start']).dt.days + 1

team_colors = {
    'PM and analysts': '#4c72b0',
    'Field crews': '#55a868',
    'Analysts': '#c44e52',
    'Engineers': '#8172b3'
}

fig, ax = plt.subplots(figsize=(10, 5))
for idx, row in schedule.iterrows():
    ax.barh(y=row['Task'], width=row['duration'], left=row['days_to_start'],
            color=team_colors[row['Team']], edgecolor='white', linewidth=0.5)

ax.invert_yaxis()
max_days = schedule['days_to_start'].max() + schedule['duration'].max()
xticks = np.arange(0, max_days + 30, 30)
xticklabels = pd.date_range(start=project_start, periods=len(xticks), freq='30D').strftime('%b %Y')
ax.set_xticks(xticks[:len(xticklabels)])
ax.set_xticklabels(xticklabels, fontsize=8)
ax.set_xlabel('Timeline')
ax.set_title('Project implementation schedule')
ax.xaxis.grid(True, alpha=0.3, linestyle='--')
ax.set_axisbelow(True)

unique_teams = schedule['Team'].unique()
patches = [mpatches.Patch(color=team_colors[team], label=team) for team in unique_teams]
ax.legend(handles=patches, loc='lower right', fontsize=8, framealpha=0.9)

plt.tight_layout()
gantt_path = output_dir / 'project_gantt_schedule.png'
plt.savefig(gantt_path, dpi=150, bbox_inches='tight', facecolor='white')
plt.close()

doc.add_picture(str(gantt_path), width=Inches(6.0))
figure_counter['n'] += 1
add_caption(f"Figure {figure_counter['n']}: Project implementation schedule from initiation through June 2024")

# Milestone table
table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'Milestone'
hdr[1].text = 'Target date'
hdr[2].text = 'Responsible party'
hdr[3].text = 'Notes'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

rows = [
    ('Notice to proceed', 'November 2022', 'NFWB and Aurelia.', 'Triggers Task 1 initiation.'),
    ('Kickoff meeting and data workshop', 'Early December 2022', 'Aurelia PM and NFWB.', 'Confirms scope and data needs (REQ-005).'),
    ('Completion of field work (Tasks 2–5)', 'March 2023', 'Aurelia field leads.', 'Supports report development.'),
    ('Draft SSES report to NFWB', 'January 2023', 'Aurelia lead engineer.', 'Enables internal review and comment.'),
    ('Final SSES report to NYSDEC', 'March 2023', 'Aurelia and NFWB.', 'Supports early NYSDEC review (REQ-020).'),
    ('Construction window', 'Mid 2023–early 2024', 'NFWB and contractors.', 'Implements CAP projects.'),
    ('Post-construction monitoring', 'January–June 2024', 'Aurelia and NFWB.', 'Supports final I&I Assessment (REQ-022).'),
]

for row in rows:
    cells = table.add_row().cells
    for i, val in enumerate(row):
        cells[i].text = val

# 15. Team, Key Personnel, and Organizational Qualifications
doc.add_page_break()
h1 = doc.add_heading('15. Team, Key Personnel, and Organizational Qualifications', level=1)

p = doc.add_paragraph(
    'Aurelia Digital Systems is a multidisciplinary engineering and delivery firm with deep '
    'experience supporting utilities and public-sector clients on complex, compliance-driven '
    'projects (REQ-002, REQ-027, REQ-033, REQ-034, REQ-043, REQ-048). For this engagement, '
    'we will provide a dedicated Project Manager and a team of engineers, analysts, and field staff '
    'with the licensing, training, and experience required to safely and effectively execute the Phase '
    '1 SSES.'
)

p = doc.add_paragraph(
    'Our Project Manager will be responsible for day-to-day coordination, schedule and budget '
    'management, and primary communication with NFWB. Key technical leads will cover '
    'hydraulics and hydrology, condition assessment, field operations, and MWBE/SDVOB '
    'compliance. Subconsultants, including MWBE and SDVOB partners, will be selected to '
    'provide specialized services such as survey, traffic control, or additional field crews as needed.'
)

# Team org chart Mermaid diagram
mermaid_code_team = '''
flowchart TD
  PM[Aurelia project manager] --> H[Hydraulic engineer]
  PM --> CA[Condition assessment lead]
  PM --> FO[Field operations lead]
  PM --> MW[MWBE and SDVOB coordinator]
  FO --> FC[Field crews]
'''
team_diagram_path = render_mermaid(mermaid_code_team, 'team_org')
doc.add_picture(str(team_diagram_path), width=Inches(5.8))
figure_counter['n'] += 1
add_caption(f"Figure {figure_counter['n']}: Project team organization")

# Team qualifications table
table = doc.add_table(rows=1, cols=5)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'Name / role'
hdr[1].text = 'Years of experience'
hdr[2].text = 'Relevant project experience'
hdr[3].text = 'Licenses / certifications'
hdr[4].text = 'Responsibilities'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

rows = [
    ('Project Manager', '15+', 'Multiple Consent Order-driven SSES and I&I reduction programs.', 'NYS PE, PMP.', 'Primary contact, governance, schedule, and budget.'),
    ('Hydraulic engineer', '12', 'Collection system modeling and flow monitoring analysis.', 'NYS PE.', 'Flow isolation design, R-value review, I&I analysis.'),
    ('Condition assessment lead', '10', 'CCTV and manhole inspection programs.', 'NASSCO PACP/MACP.', 'CCTV review, defect coding, and CAP input.'),
    ('Field operations lead', '12', 'Smoke testing, flow isolation, and field safety.', 'OSHA 30, Confined Space, Traffic Control.', 'Field planning, safety oversight, and coordination.'),
]

for row in rows:
    cells = table.add_row().cells
    for i, val in enumerate(row):
        cells[i].text = val

# Subconsultant summary table
table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'Firm'
hdr[1].text = 'Certification / status'
hdr[2].text = 'Scope of services'
hdr[3].text = 'Approximate share of work'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

rows = [
    ('MWBE field services partner', 'NYS-certified MWBE.', 'Supplemental smoke testing and CCTV crews.', '20–25%.'),
    ('SDVOB traffic control partner', 'NYS-certified SDVOB.', 'Traffic control and MOT design support.', '6–8%.'),
]

for row in rows:
    cells = table.add_row().cells
    for i, val in enumerate(row):
        cells[i].text = val

# 16. Relevant Experience and References
doc.add_page_break()
h1 = doc.add_heading('16. Relevant Experience and References', level=1)

p = doc.add_paragraph(
    'Aurelia and its key personnel have delivered multiple SSES and I&I reduction projects for '
    'public utilities under regulatory oversight (REQ-027, REQ-043). Our experience includes '
    'developing NYSDEC-acceptable engineering reports, implementing data-driven corrective '
    'action plans, and performing post-construction verification of I&I reductions.'
)

# Past project timeline Mermaid diagram
mermaid_code_past = '''
flowchart LR
  P1[Assessment] --> P2[Engineering report]
  P2 --> P3[Construction]
  P3 --> P4[Post construction monitoring]
  P4 --> P5[Regulatory closeout]
'''
past_diagram_path = render_mermaid(mermaid_code_past, 'past_project_timeline')
doc.add_picture(str(past_diagram_path), width=Inches(5.8))
figure_counter['n'] += 1
add_caption(f"Figure {figure_counter['n']}: Representative SSES project lifecycle")

# Before/after metrics chart for prior project
plt.figure(figsize=(6, 4))
proj_data = pd.DataFrame({
    'Metric': ['I&I volume', 'Overflow events'],
    'Before': [100, 12],
    'After': [55, 3]
})
proj_data_melt = proj_data.melt(id_vars='Metric', value_vars=['Before', 'After'], var_name='Period', value_name='Value')
ax = sns.barplot(data=proj_data_melt, x='Metric', y='Value', hue='Period')
ax.set_title('Illustrative outcomes from prior SSES project')
plt.tight_layout()
prior_chart_path = output_dir / 'prior_project_outcomes.png'
plt.savefig(prior_chart_path, dpi=150, bbox_inches='tight', facecolor='white')
plt.close()

doc.add_picture(str(prior_chart_path), width=Inches(5.8))
figure_counter['n'] += 1
add_caption(f"Figure {figure_counter['n']}: Example I&I and overflow reductions from a prior project")

# References table
table = doc.add_table(rows=1, cols=5)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'Client'
hdr[1].text = 'Project description'
hdr[2].text = 'Services provided'
hdr[3].text = 'Dates'
hdr[4].text = 'Reference contact'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

rows = [
    ('Regional wastewater authority', 'SSES and CAP for consent order-driven I&I program.', 'SSES, flow monitoring, report preparation, post-construction assessment.', '2019–2021.', 'Program manager, phone and email available upon request.'),
]

for row in rows:
    cells = table.add_row().cells
    for i, val in enumerate(row):
        cells[i].text = val

# 17. MWBE/SDVOB Strategy and Compliance
doc.add_page_break()
h1 = doc.add_heading('17. MWBE/SDVOB Strategy and Compliance', level=1)

p = doc.add_paragraph(
    'Aurelia is committed to meeting or exceeding NFWB’s MWBE (30%) and SDVOB (6%) '
    'participation goals and complying with all related reporting requirements (REQ-030, REQ-031, '
    'REQ-038, REQ-047, REQ-048). We will submit MWBE and SDVOB utilization plans with the '
    'proposal, and, if necessary, properly documented waiver requests.'
)

p = doc.add_paragraph(
    'Our subcontracting strategy includes early engagement with certified MWBE and SDVOB '
    'firms for field services, traffic control, and specialized tasks. We will track utilization monthly '
    'and provide NFWB with the required Monthly MWBE-SDVOB Contractor Compliance Reports '
    'and proof of timely payments to subcontractors.'
)

# MWBE/SDVOB process Mermaid diagram
mermaid_code_mwbe = '''
flowchart LR
  S1[Identify scope packages] --> S2[Engage MWBE and SDVOB firms]
  S2 --> S3[Finalize utilization plan]
  S3 --> S4[Execute work and track participation]
  S4 --> S5[Submit monthly compliance reports]
'''
mwbe_diagram_path = render_mermaid(mermaid_code_mwbe, 'mwbe_process')
doc.add_picture(str(mwbe_diagram_path), width=Inches(5.8))
figure_counter['n'] += 1
add_caption(f"Figure {figure_counter['n']}: MWBE and SDVOB engagement and reporting process")

# Contract value distribution pie chart
plt.figure(figsize=(5, 5))
shares = [64, 30, 6]
labels = ['Prime Aurelia', 'MWBE', 'SDVOB']
colors = ['#4c72b0', '#55a868', '#c44e52']
plt.pie(shares, labels=labels, colors=colors, autopct='%1.0f%%', startangle=90)
plt.title('Illustrative distribution of contract value')
plt.tight_layout()
share_chart_path = output_dir / 'mwbe_sdvob_shares.png'
plt.savefig(share_chart_path, dpi=150, bbox_inches='tight', facecolor='white')
plt.close()

doc.add_picture(str(share_chart_path), width=Inches(4.0))
figure_counter['n'] += 1
add_caption(f"Figure {figure_counter['n']}: Planned distribution of contract value among prime, MWBE, and SDVOB firms")

# MWBE/SDVOB utilization plan table
table = doc.add_table(rows=1, cols=5)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'Firm'
hdr[1].text = 'Certification type'
hdr[2].text = 'Scope of work'
hdr[3].text = 'Estimated contract value share'
hdr[4].text = 'Reporting responsibility'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

rows = [
    ('Aurelia Digital Systems', 'Prime.', 'Project management, engineering, reporting.', '64%.', 'Submit utilization and compliance reports.'),
    ('MWBE field services partner', 'MWBE.', 'Smoke testing, CCTV support, focused SSES.', '30%.', 'Provide utilization data to Aurelia.'),
    ('SDVOB traffic control partner', 'SDVOB.', 'Traffic control and MOT support.', '6%.', 'Provide utilization data to Aurelia.'),
]

for row in rows:
    cells = table.add_row().cells
    for i, val in enumerate(row):
        cells[i].text = val

# EEO staffing plan summary table
table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'Labor category'
hdr[1].text = 'Estimated FTEs'
hdr[2].text = 'EEO and utilization reporting'
hdr[3].text = 'Notes'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

rows = [
    ('Professional engineers', '2.0', 'Included in EEO Staffing Plan and utilization reports.', 'Includes NYS PE responsible for stamping report.'),
    ('Field technicians', '3.0', 'Tracked by demographic group per requirements.', 'Includes MWBE and SDVOB staff.'),
]

for row in rows:
    cells = table.add_row().cells
    for i, val in enumerate(row):
        cells[i].text = val

# 18. Health, Safety, and Risk Management
doc.add_page_break()
h1 = doc.add_heading('18. Health, Safety, and Risk Management', level=1)

p = doc.add_paragraph(
    'Aurelia’s health and safety program is central to our field operations. We will develop and '
    'maintain a Site-Specific Health & Safety Plan (HASP) that meets OSHA and applicable '
    'federal, state, and local requirements, including hazard analysis, PPE requirements, training '
    'documentation, and monitoring plans (REQ-007, REQ-008, REQ-009, REQ-011, REQ-012, '
    'REQ-013, REQ-033, REQ-034, REQ-038).'
)

p = doc.add_paragraph(
    'Field activities such as smoke testing, night-time flow isolation, and confined space entry will be '
    'subject to job hazard analyses, daily safety briefings, and documented permits where required. '
    'Traffic control plans will follow MUTCD and local requirements, and resident safety will be '
    'protected through clear notifications and coordination with emergency services.'
)

# Safety workflow Mermaid diagram
mermaid_code_safety = '''
flowchart LR
  H1[Identify field task] --> H2[Job hazard analysis]
  H2 --> H3[Develop site specific HASP]
  H3 --> H4[Daily safety briefing]
  H4 --> H5[Execute work with monitoring]
  H5 --> H6[Incident review and lessons]
'''
safety_diagram_path = render_mermaid(mermaid_code_safety, 'safety_workflow')
doc.add_picture(str(safety_diagram_path), width=Inches(5.8))
figure_counter['n'] += 1
add_caption(f"Figure {figure_counter['n']}: Safety planning and field execution workflow")

# Hazards and mitigation table
table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'Hazard'
hdr[1].text = 'Potential impact'
hdr[2].text = 'Mitigation measures'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

rows = [
    ('Smoke testing near residences', 'Resident concern or respiratory irritation.', 'Advance notification, use of non-toxic smoke, clear instructions, and hotline.'),
    ('Night-time flow isolation', 'Reduced visibility and traffic conflicts.', 'Traffic control plans, high-visibility PPE, lighting, and coordination with police.'),
    ('Confined space entry', 'Exposure to hazardous atmospheres or engulfment.', 'Permit-required confined space procedures, gas monitoring, rescue plan.'),
]

for row in rows:
    cells = table.add_row().cells
    for i, val in enumerate(row):
        cells[i].text = val

# Training and certification matrix
table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'Role'
hdr[1].text = 'Required training'
hdr[2].text = 'Certification'
hdr[3].text = 'Verification method'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

rows = [
    ('Field supervisor', 'OSHA 30, confined space, traffic control.', 'Current certifications.', 'Training records reviewed before mobilization.'),
    ('Field technicians', 'OSHA 10, task-specific training.', 'Documented completion.', 'Maintained in Aurelia training database.'),
]

for row in rows:
    cells = table.add_row().cells
    for i, val in enumerate(row):
        cells[i].text = val

# 19. Pricing, Level of Effort, and Option Phases 2 & 3
doc.add_page_break()
h1 = doc.add_heading('19. Pricing, Level of Effort, and Option Phases 2 & 3', level=1)

p = doc.add_paragraph(
    'Aurelia will provide a detailed task/hour estimate matrix and not-to-exceed fee for Phase 1, '
    'including a $10,000 contingency for Task 6 focused SSES (REQ-018, REQ-023, REQ-027, '
    'REQ-028, REQ-039). Separate option pricing, schedules, and fees will be provided for Phase 2 '
    'and Phase 3 SSES, applying applicable Phase 1 tasks and omitting Task 5 for Phase 2 where '
    'no outfalls are present.'
)

p = doc.add_paragraph(
    'Our pricing structure is transparent and aligned with the work plan, enabling NFWB to clearly '
    'see how resources are allocated by task and labor category. Monthly invoices will include '
    'sufficient supporting documentation to facilitate timely payment (REQ-039).'
)

# Mini Gantt timelines Mermaid for Phase 2 and 3
mermaid_code_phases = '''
flowchart LR
  P1[Phase 1 SSES] --> P2[Phase 2 SSES option]
  P2 --> P3[Phase 3 SSES option]
'''
phases_diagram_path = render_mermaid(mermaid_code_phases, 'phase_options_timeline')
doc.add_picture(str(phases_diagram_path), width=Inches(5.0))
figure_counter['n'] += 1
add_caption(f"Figure {figure_counter['n']}: Relationship of Phase 1 to Phase 2 and Phase 3 SSES options")

# Cost distribution stacked bar chart
plt.figure(figsize=(7, 4))
phases = ['Phase 1', 'Phase 2', 'Phase 3']
fees_task1 = [80, 60, 65]
fees_task2 = [120, 90, 100]
fees_task3 = [110, 85, 95]
fees_task4 = [70, 60, 65]
fees_task5 = [50, 0, 45]
fees_task6 = [10, 10, 10]
fees_task7 = [90, 70, 80]
fees_task8 = [40, 35, 40]

ind = np.arange(len(phases))
width = 0.6
bottom = np.zeros(len(phases))
for fees, label in [
    (fees_task1, 'Task 1'), (fees_task2, 'Task 2'), (fees_task3, 'Task 3'),
    (fees_task4, 'Task 4'), (fees_task5, 'Task 5'), (fees_task6, 'Task 6'),
    (fees_task7, 'Task 7'), (fees_task8, 'Task 8')
]:
    plt.bar(ind, fees, width, bottom=bottom, label=label)
    bottom += np.array(fees)

plt.xticks(ind, phases)
plt.ylabel('Illustrative cost (thousand dollars)')
plt.title('Illustrative distribution of costs by task and phase')
plt.legend(fontsize=7, ncol=4)
plt.tight_layout()
fees_chart_path = output_dir / 'cost_distribution.png'
plt.savefig(fees_chart_path, dpi=150, bbox_inches='tight', facecolor='white')
plt.close()

doc.add_picture(str(fees_chart_path), width=Inches(5.8))
figure_counter['n'] += 1
add_caption(f"Figure {figure_counter['n']}: Illustrative cost distribution by task for Phase 1 and option phases")

# Task/hour estimate matrix table
table = doc.add_table(rows=1, cols=6)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'Task'
hdr[1].text = 'PM hours'
hdr[2].text = 'Engineer hours'
hdr[3].text = 'Field hours'
hdr[4].text = 'Support hours'
hdr[5].text = 'Total hours'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

rows = [
    ('Task 1', '40', '80', '0', '16', '136'),
    ('Task 2', '30', '40', '220', '24', '314'),
    ('Task 3', '30', '60', '200', '24', '314'),
    ('Task 4', '20', '80', '0', '16', '116'),
    ('Task 5', '16', '40', '80', '8', '144'),
    ('Task 6 (contingency)', '8', '16', '40', '8', '72'),
    ('Task 7', '60', '160', '0', '24', '244'),
    ('Task 8', '24', '80', '40', '16', '160'),
]

for row in rows:
    cells = table.add_row().cells
    for i, val in enumerate(row):
        cells[i].text = val

# Pricing breakdown table
table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'Phase'
hdr[1].text = 'Scope summary'
hdr[2].text = 'Not-to-exceed fee (illustrative)'
hdr[3].text = 'Notes'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

rows = [
    ('Phase 1', 'Tasks 1–8 with $10,000 contingency in Task 6.', '$XXX,XXX.', 'Detailed fee schedule provided in commercial attachment.'),
    ('Phase 2 option', 'Tasks 1–4, 6–8 (no Task 5).', '$XXX,XXX.', 'Pricing assumes similar sewershed scale.'),
    ('Phase 3 option', 'Tasks 1–8 for additional LaSalle sewersheds.', '$XXX,XXX.', 'To be refined based on Consent Order updates.'),
]

for row in rows:
    cells = table.add_row().cells
    for i, val in enumerate(row):
        cells[i].text = val

# 20. Administrative, Legal, and Compliance Commitments
doc.add_page_break()
h1 = doc.add_heading('20. Administrative, Legal, and Compliance Commitments', level=1)

p = doc.add_paragraph(
    'Aurelia agrees to comply with all applicable terms and conditions in Appendix E, including '
    'forms, insurance, indemnification, records retention, FOIL, information security, prevailing '
    'wage, non-discrimination, sexual harassment, Iran Divestment, boycott and MacBride '
    'certifications, and communication restrictions (REQ-021, REQ-024, REQ-026, REQ-029, '
    'REQ-030, REQ-031, REQ-032, REQ-035, REQ-036, REQ-037, REQ-038, REQ-040, REQ-041, '
    'REQ-042, REQ-045, REQ-046, REQ-047, REQ-048). Completed Appendix E forms will be '
    'submitted with the proposal.'
)

p = doc.add_paragraph(
    'We maintain required insurance coverages and will name Niagara Falls Water Board, Niagara '
    'Falls Public Water Authority, and the City of Niagara Falls as additional insureds on applicable '
    'policies (REQ-032). Aurelia certifies that no suspended or debarred entities will be used as '
    'contractors or subcontractors (REQ-048) and that we will comply with prevailing wage and '
    'workforce reporting requirements (REQ-038).'
)

# Compliance checklist table
table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'Requirement / form'
hdr[1].text = 'Description'
hdr[2].text = 'Status in this proposal'
hdr[3].text = 'Location / reference'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

rows = [
    ('Appendix E Form No. 1', 'Acknowledgement of Addenda.', 'Completed and included.', 'Administrative attachment (REQ-029, REQ-045).'),
    ('Appendix E Form No. 2', 'Certificate of Non-Collusion.', 'Completed and included.', 'Administrative attachment (REQ-029, REQ-041).'),
    ('Appendix E Form No. 3', 'Lobbying Law Disclosure Statement.', 'Completed and included.', 'Administrative attachment (REQ-026, REQ-046).'),
    ('EEO Policy Statement', 'Affirmation of non-discrimination and EEO.', 'Completed and included.', 'Administrative attachment (REQ-029, REQ-040).'),
    ('Sexual Harassment Statement', 'Compliance with NY Labor Law 201-g.', 'Completed and included.', 'Administrative attachment (REQ-029, REQ-040).'),
]

for row in rows:
    cells = table.add_row().cells
    for i, val in enumerate(row):
        cells[i].text = val

# Insurance coverage table
table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'Coverage type'
hdr[1].text = 'Limit (illustrative)'
hdr[2].text = 'Carrier rating'
hdr[3].text = 'Additional insureds / notes'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

rows = [
    ('Workers Compensation and Disability', 'Statutory.', 'A rated.', 'Covers all project staff.'),
    ('Commercial General Liability', '$1,000,000 per occurrence / $2,000,000 aggregate.', 'A rated.', 'NFWB, NFPWA, and City named as additional insureds.'),
    ('Automobile Liability', '$1,000,000 combined single limit.', 'A rated.', 'Covers all project vehicles.'),
    ('Professional Liability', '$2,000,000 per claim / aggregate.', 'A rated.', 'Covers engineering services.'),
]

for row in rows:
    cells = table.add_row().cells
    for i, val in enumerate(row):
        cells[i].text = val

# 21. Proposal Administration and Submission Plan
doc.add_page_break()
h1 = doc.add_heading('21. Proposal Administration and Submission Plan', level=1)

p = doc.add_paragraph(
    'Aurelia will manage proposal preparation and submission to fully comply with NFWB’s '
    'requirements (REQ-025, REQ-026, REQ-027, REQ-029, REQ-039, REQ-044, REQ-045, '
    'REQ-046). We will attend the mandatory pre-proposal information session at the Wastewater '
    'Treatment Plant on November 29, 2022, and direct any written questions to the authorized '
    'Water Board contact only, in accordance with NYS lobbying law restrictions.'
)

p = doc.add_paragraph(
    'The final submission will include one clearly marked original and six copies of the proposal, '
    'plus one electronic copy on CD or USB as a single PDF or logically ordered files, in a sealed '
    'outer envelope labeled with proposer name, address, RFP number, and title (REQ-044). All '
    'Appendix E forms and required signatures will be included.'
)

# Internal proposal timeline Mermaid diagram
mermaid_code_prop = '''
flowchart LR
  R1[RFP receipt] --> R2[Pre proposal meeting]
  R2 --> R3[Draft development]
  R3 --> R4[Internal review]
  R4 --> R5[Finalize forms and signatures]
  R5 --> R6[Production and delivery]
'''
prop_diagram_path = render_mermaid(mermaid_code_prop, 'proposal_timeline')
doc.add_picture(str(prop_diagram_path), width=Inches(5.8))
figure_counter['n'] += 1
add_caption(f"Figure {figure_counter['n']}: Internal proposal development and submission timeline")

# Submission checklist table
table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'Item'
hdr[1].text = 'Requirement'
hdr[2].text = 'Responsibility'
hdr[3].text = 'Status tracking'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

rows = [
    ('Pre-proposal meeting attendance', 'November 29, 2022 at noon (REQ-025).', 'Aurelia PM.', 'Sign-in sheet and meeting notes.'),
    ('Original and six copies', 'Printed and bound per RFP (REQ-044).', 'Proposal coordinator.', 'Checklist before shipment.'),
    ('Electronic copy on CD or USB', 'Single PDF or organized files (REQ-044).', 'Proposal coordinator.', 'Media verification.'),
    ('Appendix E forms', 'All required forms completed and signed (REQ-029).', 'Aurelia leadership.', 'Signature tracking.'),
]

for row in rows:
    cells = table.add_row().cells
    for i, val in enumerate(row):
        cells[i].text = val

# 22. Exceptions, Assumptions, and Clarifications
doc.add_page_break()
h1 = doc.add_heading('22. Exceptions, Assumptions, and Clarifications', level=1)

p = doc.add_paragraph(
    'Aurelia has structured this proposal to comply with all material requirements of the RFP and '
    'Appendix E. Any exceptions are limited and intended to clarify responsibilities and avoid '
    'misunderstandings (REQ-024, REQ-027, REQ-042). We welcome the opportunity to discuss '
    'these items with NFWB during contract negotiations.'
)

# Assumptions and clarifications table
table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'Assumption / clarification'
hdr[1].text = 'Description'
hdr[2].text = 'Potential impact'
hdr[3].text = 'NFWB responsibility'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

rows = [
    ('Access to data and facilities', 'NFWB will provide timely access to existing data and facilitate access to manholes and sites.', 'Delays could extend schedule.', 'Provide data and coordinate access.'),
    ('Construction by others', 'NFWB will procure and manage construction contractors.', 'Our scope focuses on SSES and reporting.', 'Manage construction contracts.'),
]

for row in rows:
    cells = table.add_row().cells
    for i, val in enumerate(row):
        cells[i].text = val

# Exception log table (if any)
table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'Reference'
hdr[1].text = 'Exception summary'
hdr[2].text = 'Rationale / proposed resolution'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

row = table.add_row().cells
row[0].text = 'Appendix E general terms'
row[1].text = 'No substantive exceptions are requested at this time.'
row[2].text = 'Aurelia will work with NFWB to address any clarifications during contract finalization.'

# Appendices heading
h1 = doc.add_heading('Appendices', level=1)

p = doc.add_paragraph(
    'The following appendices are provided under separate cover and are incorporated by '
    'reference: detailed fee schedules, full resumes of key personnel, additional project case '
    'studies, completed Appendix E forms, and any required insurance certificates.'
)
