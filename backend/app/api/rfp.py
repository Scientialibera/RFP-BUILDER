"""
RFP Generation API endpoints.
"""

import logging
import json
import time
import base64
import ast
import html
import re
import shutil
from datetime import datetime
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional, TypeVar

from fastapi import APIRouter, File, UploadFile, HTTPException, Depends, Form
from fastapi.responses import FileResponse, StreamingResponse
from pydantic import BaseModel, ValidationError

from app.core.config import get_config, FeaturesConfig
from app.core.auth import verify_api_token
from app.core.llm_client import create_llm_client
from app.models.schemas import (
    GeneratedCodePackage,
    GeneratedCodeSnippet,
    GenerateRFPResponse,
    ExtractReqsResponse,
    PlanStepRequest,
    PlanStepResponse,
    GenerateRFPStepResponse,
    CritiqueStepRequest,
    CritiqueStepResponse,
    RFPAnalysis,
    ProposalPlan,
    RFPRequirement,
)
from app.services.pdf_service import PDFService
from app.services.blob_storage import get_blob_storage
from app.workflows.rfp_workflow import create_rfp_workflow
from app.workflows.executors import (
    RFPAnalyzerExecutor,
    PlannerExecutor,
    SectionGeneratorExecutor,
    CritiquerExecutor,
    CodeInterpreterExecutor,
)
from app.workflows.run_dirs import RUN_SUBDIRECTORIES, create_unique_run_directory
from app.workflows.state import WorkflowInput, DocumentInfo


logger = logging.getLogger(__name__)
router = APIRouter(prefix="/api/rfp", tags=["RFP"])
ModelT = TypeVar("ModelT", bound=BaseModel)


@dataclass
class ProcessedDocuments:
    """Container for processed document data."""
    rfp_text: str
    rfp_images: list[dict] | None
    example_texts: list[str]
    example_images: list[list[dict]] | None
    context_text: str | None
    context_images: list[dict] | None
    # Document info for storage
    documents: list[DocumentInfo] = field(default_factory=list)


def _validate_pdf_files(files: list[UploadFile]) -> None:
    """Validate that all uploads are PDF files."""
    for file in files:
        if file.filename and not file.filename.lower().endswith(".pdf"):
            raise HTTPException(
                status_code=400,
                detail=f"Only PDF files are allowed. Got: {file.filename}",
            )


def _safe_filename(filename: str) -> str:
    """Return a path-safe basename."""
    return Path(filename).name


def _validate_simple_path_part(value: str, field_name: str) -> str:
    """Reject path traversal by forcing a single path segment."""
    safe = Path(value).name
    if safe != value:
        raise HTTPException(status_code=400, detail=f"Invalid {field_name}: {value}")
    return safe


def _create_step_run_directory() -> Path:
    """Create a run directory for step endpoints."""
    config = get_config()
    output_root = Path(config.app.output_dir)
    return create_unique_run_directory(output_root)


def _ensure_run_subdirectories(run_dir: Path) -> None:
    """Ensure all standard run subdirectories exist."""
    for subdir_name in RUN_SUBDIRECTORIES:
        (run_dir / subdir_name).mkdir(parents=True, exist_ok=True)


def _resolve_step_run_directory(run_id: Optional[str]) -> Path:
    """Resolve an existing run directory or create a new one."""
    if not run_id or not run_id.strip():
        return _create_step_run_directory()

    safe_run_id = _validate_simple_path_part(run_id.strip(), "run_id")
    config = get_config()
    output_root = Path(config.app.output_dir).resolve()
    run_dir = (output_root / safe_run_id).resolve()
    if not run_dir.is_relative_to(output_root):
        raise HTTPException(status_code=400, detail="Invalid run_id path")

    if not run_dir.exists():
        storage = get_blob_storage()
        if storage.enabled:
            try:
                storage.hydrate_run_directory(safe_run_id, run_dir)
            except Exception as exc:
                logger.warning("Failed to hydrate run %s from blob: %s", safe_run_id, exc)

    _ensure_run_subdirectories(run_dir)
    return run_dir


def _save_documents(run_dir: Path, documents: list[DocumentInfo]) -> None:
    """Persist uploaded source documents in the run directory."""
    if not documents:
        return

    docs_dir = run_dir / "documents"
    manifest_path = docs_dir / "documents_manifest.json"
    for doc_info in documents:
        if not doc_info.file_bytes:
            continue
        safe_name = _safe_filename(doc_info.filename)
        try:
            with open(docs_dir / safe_name, "wb") as f:
                f.write(doc_info.file_bytes)
        except Exception as exc:
            logger.warning(f"Failed to save source document {safe_name}: {exc}")

    existing_documents: list[dict[str, str]] = []
    if manifest_path.exists():
        try:
            with open(manifest_path, "r", encoding="utf-8") as f:
                loaded_manifest = json.load(f)
            if isinstance(loaded_manifest, dict):
                loaded_docs = loaded_manifest.get("documents", [])
                if isinstance(loaded_docs, list):
                    for item in loaded_docs:
                        if not isinstance(item, dict):
                            continue
                        filename = item.get("filename")
                        file_type = item.get("type")
                        if isinstance(filename, str) and isinstance(file_type, str):
                            existing_documents.append(
                                {"filename": _safe_filename(filename), "type": file_type}
                            )
        except Exception as exc:
            logger.warning(f"Failed to load existing source documents manifest: {exc}")

    merged_docs: list[dict[str, str]] = []
    seen: set[tuple[str, str]] = set()
    for item in existing_documents:
        key = (item["filename"], item["type"])
        if key in seen:
            continue
        seen.add(key)
        merged_docs.append(item)

    for doc_info in documents:
        item = {"filename": _safe_filename(doc_info.filename), "type": doc_info.file_type}
        key = (item["filename"], item["type"])
        if key in seen:
            continue
        seen.add(key)
        merged_docs.append(item)

    manifest = {"documents": merged_docs}
    try:
        with open(manifest_path, "w", encoding="utf-8") as f:
            json.dump(manifest, f, indent=2)
    except Exception as exc:
        logger.warning(f"Failed to save source documents manifest: {exc}")


def _save_code_snapshot(run_dir: Path, stage: str, code: str) -> None:
    """Save generated code snapshot."""
    snapshot_file = run_dir / "code_snapshots" / f"{stage}_document_code.py"
    try:
        with open(snapshot_file, "w", encoding="utf-8") as f:
            f.write(code)
    except Exception as exc:
        logger.warning(f"Failed to save code snapshot {snapshot_file.name}: {exc}")


def _save_json_artifact(run_dir: Path, relative_path: str, payload: object) -> None:
    """Save a JSON payload under the run directory."""
    target = run_dir / relative_path
    target.parent.mkdir(parents=True, exist_ok=True)
    try:
        with open(target, "w", encoding="utf-8") as f:
            json.dump(payload, f, indent=2, default=str)
    except Exception as exc:
        logger.warning("Failed to save JSON artifact %s: %s", target, exc)


def _sync_run_to_blob(run_dir: Path) -> None:
    """Best-effort sync of run artifacts to Azure Blob storage."""
    storage = get_blob_storage()
    if not storage.enabled:
        return
    try:
        storage.sync_run_directory(run_dir)
        if not get_config().storage.use_local_storage:
            try:
                shutil.rmtree(run_dir)
                logger.info("Removed local run directory after blob sync: %s", run_dir)
            except Exception as cleanup_exc:
                logger.warning("Failed to remove local run directory %s: %s", run_dir, cleanup_exc)
    except Exception as exc:
        logger.warning("Blob sync failed for run %s: %s", run_dir.name, exc)


def _save_analysis_metadata(run_dir: Path, analysis: RFPAnalysis) -> None:
    _save_json_artifact(run_dir, "metadata/analysis.json", analysis.model_dump())


def _save_plan_metadata(run_dir: Path, plan: ProposalPlan) -> None:
    _save_json_artifact(run_dir, "metadata/plan.json", plan.model_dump())


def _load_json_list_artifact(run_dir: Path, relative_path: str) -> list[dict]:
    """Load a list JSON artifact; return empty list when missing or invalid."""
    target = run_dir / relative_path
    if not target.exists():
        return []
    try:
        with open(target, "r", encoding="utf-8") as f:
            loaded = json.load(f)
        if isinstance(loaded, list):
            return [item for item in loaded if isinstance(item, dict)]
    except Exception as exc:
        logger.warning("Failed to load JSON artifact %s: %s", target, exc)
    return []


def _append_analysis_version(run_dir: Path, analysis: RFPAnalysis, comment: Optional[str] = None) -> None:
    """Append a new analysis version entry to metadata history."""
    versions = _load_json_list_artifact(run_dir, "metadata/analysis_versions.json")
    version_id = f"analysis_v{len(versions) + 1:03d}"
    entry = {
        "version_id": version_id,
        "created_at": datetime.now().isoformat(),
        "comment": (comment or "").strip() or None,
        "analysis": analysis.model_dump(),
    }
    versions.append(entry)
    _save_json_artifact(run_dir, "metadata/analysis_versions.json", versions)


def _append_plan_version(run_dir: Path, plan: ProposalPlan, comment: Optional[str] = None) -> None:
    """Append a new plan version entry to metadata history."""
    versions = _load_json_list_artifact(run_dir, "metadata/plan_versions.json")
    version_id = f"plan_v{len(versions) + 1:03d}"
    entry = {
        "version_id": version_id,
        "created_at": datetime.now().isoformat(),
        "comment": (comment or "").strip() or None,
        "plan": plan.model_dump(),
    }
    versions.append(entry)
    _save_json_artifact(run_dir, "metadata/plan_versions.json", versions)


def _extract_target_names(target: ast.expr) -> set[str]:
    if isinstance(target, ast.Name):
        return {target.id}
    if isinstance(target, (ast.Tuple, ast.List)):
        names: set[str] = set()
        for elt in target.elts:
            names.update(_extract_target_names(elt))
        return names
    return set()


def _extract_assigned_names(stmt: ast.stmt) -> set[str]:
    names: set[str] = set()
    if isinstance(stmt, ast.Assign):
        for target in stmt.targets:
            names.update(_extract_target_names(target))
    elif isinstance(stmt, ast.AnnAssign):
        names.update(_extract_target_names(stmt.target))
    elif isinstance(stmt, ast.AugAssign):
        names.update(_extract_target_names(stmt.target))
    return names


def _stmt_uses_names(stmt: ast.stmt, names: set[str]) -> bool:
    if not names:
        return False
    for node in ast.walk(stmt):
        if isinstance(node, ast.Name) and node.id in names:
            return True
    return False


def _has_call(stmt: ast.stmt, owner: str | None, attr: str) -> bool:
    for node in ast.walk(stmt):
        if not isinstance(node, ast.Call):
            continue
        func = node.func
        if isinstance(func, ast.Attribute):
            if owner is None:
                continue
            if isinstance(func.value, ast.Name) and func.value.id == owner and func.attr == attr:
                return True
        elif isinstance(func, ast.Name):
            if owner is None and func.id == attr:
                return True
    return False


def _extract_call_string_arg(stmt: ast.stmt, func_name: str, arg_index: int) -> str | None:
    for node in ast.walk(stmt):
        if not isinstance(node, ast.Call):
            continue
        if not isinstance(node.func, ast.Name) or node.func.id != func_name:
            continue
        if len(node.args) > arg_index and isinstance(node.args[arg_index], ast.Constant):
            value = node.args[arg_index].value
            return value if isinstance(value, str) else None
        return None
    return None


def _extract_savefig_target(stmt: ast.stmt) -> str | None:
    for node in ast.walk(stmt):
        if not isinstance(node, ast.Call):
            continue
        func = node.func
        if not (
            isinstance(func, ast.Attribute)
            and isinstance(func.value, ast.Name)
            and func.value.id == "plt"
            and func.attr == "savefig"
        ):
            continue
        if node.args:
            first = node.args[0]
            if isinstance(first, ast.Constant) and isinstance(first.value, str):
                return first.value
            if isinstance(first, ast.Name):
                return first.id
        for keyword in node.keywords:
            if keyword.arg == "fname":
                if isinstance(keyword.value, ast.Constant) and isinstance(keyword.value.value, str):
                    return keyword.value.value
                if isinstance(keyword.value, ast.Name):
                    return keyword.value.id
        return None
    return None


def _slice_statement_block(lines: list[str], statements: list[ast.stmt], start_idx: int, end_idx: int) -> str:
    start_line = statements[start_idx].lineno
    end_line = statements[end_idx].end_lineno or statements[end_idx].lineno
    return "\n".join(lines[start_line - 1:end_line]).rstrip()


def _is_doc_add_table_assign(stmt: ast.stmt) -> str | None:
    if not isinstance(stmt, ast.Assign) or len(stmt.targets) != 1:
        return None
    target = stmt.targets[0]
    if not isinstance(target, ast.Name):
        return None
    value = stmt.value
    if not isinstance(value, ast.Call):
        return None
    func = value.func
    if (
        isinstance(func, ast.Attribute)
        and isinstance(func.value, ast.Name)
        and func.value.id == "doc"
        and func.attr == "add_table"
    ):
        return target.id
    return None


def _is_mermaid_code_assign(stmt: ast.stmt) -> str | None:
    if not isinstance(stmt, ast.Assign) or len(stmt.targets) != 1:
        return None
    target = stmt.targets[0]
    if not isinstance(target, ast.Name):
        return None
    if not isinstance(stmt.value, ast.Constant) or not isinstance(stmt.value.value, str):
        return None

    text = stmt.value.value.strip().lower()
    name = target.id.lower()
    if name.startswith("mermaid_code"):
        return target.id

    mermaid_prefixes = (
        "graph",
        "flowchart",
        "sequencediagram",
        "classdiagram",
        "statediagram",
        "erdiagram",
        "journey",
        "gantt",
        "pie",
        "mindmap",
        "timeline",
        "gitgraph",
        "quadrantchart",
        "requirementdiagram",
        "xychart",
    )
    if any(text.startswith(prefix) for prefix in mermaid_prefixes):
        return target.id
    return None


def _is_chart_start(statements: list[ast.stmt], idx: int) -> bool:
    stmt = statements[idx]
    if _has_call(stmt, "sns", "set_style") or _has_call(stmt, "plt", "figure") or _has_call(stmt, "plt", "subplots"):
        return True
    if _has_call(stmt, "sns", "barplot") or _has_call(stmt, "sns", "lineplot") or _has_call(stmt, "sns", "scatterplot"):
        return True
    if _has_call(stmt, "sns", "heatmap") or _has_call(stmt, "sns", "histplot") or _has_call(stmt, "sns", "boxplot"):
        return True
    if _has_call(stmt, "sns", "violinplot"):
        return True

    if not isinstance(stmt, ast.Assign) or not isinstance(stmt.value, ast.Call):
        return False
    call = stmt.value
    if not (
        isinstance(call.func, ast.Attribute)
        and isinstance(call.func.value, ast.Name)
        and call.func.value.id == "pd"
        and call.func.attr == "DataFrame"
    ):
        return False

    lookahead = statements[idx + 1: idx + 5]
    return any(
        _has_call(candidate, "plt", "figure")
        or _has_call(candidate, "plt", "subplots")
        or _has_call(candidate, "sns", "barplot")
        or _has_call(candidate, "sns", "lineplot")
        or _has_call(candidate, "sns", "scatterplot")
        or _has_call(candidate, "sns", "heatmap")
        for candidate in lookahead
    )


def _build_code_package(document_code: str) -> GeneratedCodePackage:
    package = GeneratedCodePackage()
    try:
        tree = ast.parse(document_code)
    except SyntaxError:
        return package

    statements = tree.body
    lines = document_code.splitlines()

    mermaid_items: list[GeneratedCodeSnippet] = []
    for idx, stmt in enumerate(statements):
        mermaid_var = _is_mermaid_code_assign(stmt)
        if not mermaid_var:
            continue

        end_idx = idx
        block_vars = {mermaid_var}
        output_name: str | None = None
        for look_idx in range(idx + 1, len(statements)):
            candidate = statements[look_idx]
            if _stmt_uses_names(candidate, block_vars):
                block_vars.update(_extract_assigned_names(candidate))
                output_name = output_name or _extract_call_string_arg(candidate, "render_mermaid", 1)
                end_idx = look_idx
                continue
            if _has_call(candidate, None, "add_caption") and end_idx > idx:
                end_idx = look_idx
            break

        title = output_name or mermaid_var
        mermaid_items.append(
            GeneratedCodeSnippet(
                snippet_id=f"mermaid_{len(mermaid_items) + 1}",
                title=title.replace("_", " ").strip().title(),
                code=_slice_statement_block(lines, statements, idx, end_idx),
            )
        )

    table_items: list[GeneratedCodeSnippet] = []
    for idx, stmt in enumerate(statements):
        table_var = _is_doc_add_table_assign(stmt)
        if not table_var:
            continue

        end_idx = idx
        block_vars = {table_var}
        for look_idx in range(idx + 1, len(statements)):
            candidate = statements[look_idx]
            if not _stmt_uses_names(candidate, block_vars):
                break
            block_vars.update(_extract_assigned_names(candidate))
            end_idx = look_idx

        table_items.append(
            GeneratedCodeSnippet(
                snippet_id=f"table_{len(table_items) + 1}",
                title=table_var.replace("_", " ").strip().title(),
                code=_slice_statement_block(lines, statements, idx, end_idx),
            )
        )

    chart_items: list[GeneratedCodeSnippet] = []
    used_chart_indexes: set[int] = set()
    idx = 0
    while idx < len(statements):
        if idx in used_chart_indexes or not _is_chart_start(statements, idx):
            idx += 1
            continue

        start_idx = idx
        for back in range(1, 4):
            prev_idx = idx - back
            if prev_idx < 0 or prev_idx in used_chart_indexes:
                break
            prev_stmt = statements[prev_idx]
            if _is_doc_add_table_assign(prev_stmt) or _is_mermaid_code_assign(prev_stmt):
                break
            if _has_call(prev_stmt, "doc", "add_heading") or _has_call(prev_stmt, "doc", "add_page_break"):
                break
            if isinstance(prev_stmt, ast.Assign):
                start_idx = prev_idx
                continue
            break

        close_idx: int | None = None
        for look_idx in range(idx, len(statements)):
            candidate = statements[look_idx]
            if look_idx > idx and _is_chart_start(statements, look_idx):
                break
            if _is_doc_add_table_assign(candidate) or _is_mermaid_code_assign(candidate):
                break
            if _has_call(candidate, "plt", "close"):
                close_idx = look_idx
                break

        if close_idx is None:
            idx += 1
            continue

        end_idx = close_idx
        chart_title: str | None = None
        for look_idx in range(start_idx, close_idx + 1):
            chart_title = chart_title or _extract_savefig_target(statements[look_idx])

        block_vars: set[str] = set()
        for look_idx in range(start_idx, close_idx + 1):
            block_vars.update(_extract_assigned_names(statements[look_idx]))

        for look_idx in range(close_idx + 1, min(close_idx + 5, len(statements))):
            candidate = statements[look_idx]
            if _has_call(candidate, "doc", "add_picture") or _has_call(candidate, None, "add_caption"):
                end_idx = look_idx
                block_vars.update(_extract_assigned_names(candidate))
                continue
            if _stmt_uses_names(candidate, block_vars):
                end_idx = look_idx
                block_vars.update(_extract_assigned_names(candidate))
                continue
            break

        for used_idx in range(start_idx, end_idx + 1):
            used_chart_indexes.add(used_idx)

        chart_items.append(
            GeneratedCodeSnippet(
                snippet_id=f"diagram_{len(chart_items) + 1}",
                title=(chart_title or f"Chart {len(chart_items) + 1}").replace("_", " ").strip().title(),
                code=_slice_statement_block(lines, statements, start_idx, end_idx),
            )
        )

        idx = end_idx + 1

    package.mermaid = mermaid_items
    package.tables = table_items
    package.diagrams = chart_items
    return package


def _extract_png_names_from_code(code: str) -> list[str]:
    matches = re.findall(r"['\"]([^'\"]+\.png)['\"]", code, flags=re.IGNORECASE)
    names: list[str] = []
    seen: set[str] = set()
    for match in matches:
        name = Path(match).name
        lowered = name.lower()
        if lowered in seen:
            continue
        seen.add(lowered)
        names.append(name)
    return names


def _extract_mermaid_output_filename(code: str) -> Optional[str]:
    mermaid_call = re.search(
        r"render_mermaid\s*\(\s*.+?,\s*['\"]([^'\"]+)['\"]",
        code,
        flags=re.IGNORECASE | re.DOTALL,
    )
    if not mermaid_call:
        return None
    return f"{Path(mermaid_call.group(1)).name}.png"


def _collect_png_assets(image_dir: Path) -> dict[str, tuple[str, str]]:
    assets: dict[str, tuple[str, str]] = {}
    if not image_dir.exists():
        return assets

    for file_path in sorted(image_dir.iterdir()):
        if not file_path.is_file() or file_path.suffix.lower() != ".png":
            continue
        try:
            encoded = base64.b64encode(file_path.read_bytes()).decode("ascii")
            assets[file_path.name.lower()] = (file_path.name, encoded)
        except Exception as exc:
            logger.warning("Failed to encode PNG asset %s: %s", file_path.name, exc)
    return assets


def _normalize_hint(value: str) -> str:
    return re.sub(r"[^a-z0-9]+", "", value.lower())


def _attach_asset_to_snippet(
    snippet: GeneratedCodeSnippet,
    assets: dict[str, tuple[str, str]],
    used_assets: set[str],
) -> None:
    candidate_names: list[str] = []
    if snippet.snippet_id.startswith("mermaid_"):
        mermaid_output = _extract_mermaid_output_filename(snippet.code)
        if mermaid_output:
            candidate_names.append(mermaid_output)
    candidate_names.extend(_extract_png_names_from_code(snippet.code))

    for candidate in candidate_names:
        key = candidate.lower()
        if key not in assets or key in used_assets:
            continue
        filename, encoded = assets[key]
        snippet.asset_filename = filename
        snippet.asset_base64 = encoded
        snippet.asset_content_type = "image/png"
        used_assets.add(key)
        return

    title_hint = _normalize_hint(snippet.title)
    if title_hint:
        for key, (filename, encoded) in assets.items():
            if key in used_assets:
                continue
            filename_hint = _normalize_hint(filename)
            if title_hint and title_hint in filename_hint:
                snippet.asset_filename = filename
                snippet.asset_base64 = encoded
                snippet.asset_content_type = "image/png"
                used_assets.add(key)
                return

    for key, (filename, encoded) in assets.items():
        if key in used_assets:
            continue
        snippet.asset_filename = filename
        snippet.asset_base64 = encoded
        snippet.asset_content_type = "image/png"
        used_assets.add(key)
        return


def _table_to_html(table) -> str:
    rows: list[list[str]] = []
    for row in table.rows:
        rows.append([cell.text.strip() for cell in row.cells])
    if not rows:
        return ""

    header = rows[0]
    body = rows[1:]
    parts = [
        '<table style="border-collapse:collapse;width:100%;font-size:12px;">',
        "<thead><tr>",
    ]
    for cell in header:
        parts.append(
            '<th style="border:1px solid #d1d5db;padding:6px;background:#f9fafb;text-align:left;">'
            f"{html.escape(cell)}"
            "</th>"
        )
    parts.append("</tr></thead><tbody>")

    for row in body:
        parts.append("<tr>")
        for cell in row:
            parts.append(
                '<td style="border:1px solid #e5e7eb;padding:6px;vertical-align:top;">'
                f"{html.escape(cell)}"
                "</td>"
            )
        parts.append("</tr>")

    parts.append("</tbody></table>")
    return "".join(parts)


def _extract_table_html_from_docx(docx_path: Path) -> list[str]:
    if not docx_path.exists():
        return []

    try:
        from docx import Document as WordDocument
    except Exception as exc:
        logger.warning("python-docx unavailable for table HTML extraction: %s", exc)
        return []

    try:
        document = WordDocument(str(docx_path))
    except Exception as exc:
        logger.warning("Failed to read DOCX for table HTML extraction: %s", exc)
        return []

    previews: list[str] = []
    for table in document.tables:
        html_table = _table_to_html(table)
        if html_table:
            previews.append(html_table)
    return previews


def _enrich_code_package_with_previews(
    package: GeneratedCodePackage,
    image_dir: Path,
    docx_path: Path,
) -> GeneratedCodePackage:
    assets = _collect_png_assets(image_dir)
    used_assets: set[str] = set()
    for snippet in package.mermaid:
        _attach_asset_to_snippet(snippet, assets, used_assets)
    for snippet in package.diagrams:
        _attach_asset_to_snippet(snippet, assets, used_assets)

    table_html = _extract_table_html_from_docx(docx_path)
    for idx, snippet in enumerate(package.tables):
        if idx < len(table_html):
            snippet.html_code = table_html[idx]

    return package


def _chunk_sections(sections: list, max_sections: int) -> list[list]:
    """Chunk sections list by max size."""
    if max_sections <= 0:
        return [sections]
    return [sections[i:i + max_sections] for i in range(0, len(sections), max_sections)]


def _parse_json_model(value: Optional[str], model_cls: type[ModelT], field_name: str) -> Optional[ModelT]:
    """Parse JSON from form field and validate against a pydantic model."""
    if value is None or not value.strip():
        return None
    try:
        parsed = json.loads(value)
        return model_cls.model_validate(parsed)
    except json.JSONDecodeError as exc:
        raise HTTPException(status_code=400, detail=f"Invalid JSON for {field_name}: {exc}") from exc
    except ValidationError as exc:
        raise HTTPException(status_code=400, detail=f"Invalid schema for {field_name}: {exc}") from exc


def _parse_previous_requirements(value: Optional[str]) -> Optional[list[RFPRequirement]]:
    """Parse previous requirements JSON payload."""
    if value is None or not value.strip():
        return None
    try:
        parsed = json.loads(value)
    except json.JSONDecodeError as exc:
        raise HTTPException(status_code=400, detail=f"Invalid JSON for previous_requirements: {exc}") from exc
    if not isinstance(parsed, list):
        raise HTTPException(status_code=400, detail="previous_requirements must be a JSON array.")
    try:
        return [RFPRequirement.model_validate(item) for item in parsed]
    except ValidationError as exc:
        raise HTTPException(status_code=400, detail=f"Invalid requirement schema: {exc}") from exc


async def _extract_context_from_files(
    company_context: Optional[list[UploadFile]],
    pdf_service: PDFService,
) -> tuple[Optional[str], list[DocumentInfo]]:
    """Extract text from optional context PDFs for planning."""
    if not company_context:
        return None, []

    _validate_pdf_files(company_context)

    context_parts: list[str] = []
    documents: list[DocumentInfo] = []
    for idx, ctx in enumerate(company_context):
        ctx_bytes = await ctx.read()
        context_parts.append(pdf_service.extract_text_from_bytes(ctx_bytes))
        documents.append(
            DocumentInfo(
                filename=ctx.filename or f"context_{idx + 1}.pdf",
                file_type="context",
                file_bytes=ctx_bytes,
            )
        )

    context_text = "\n\n---\n\n".join(context_parts) if context_parts else None
    return context_text, documents


def _resolve_docx_download_path(run_id: str, filename: str) -> Path:
    """Resolve and validate docx download path."""
    safe_run_id = _validate_simple_path_part(run_id, "run_id")
    safe_filename = _validate_simple_path_part(filename, "filename")

    if not safe_filename.lower().endswith(".docx"):
        raise HTTPException(status_code=400, detail="Only DOCX downloads allowed")

    config = get_config()
    output_root = Path(config.app.output_dir).resolve()
    file_path = (output_root / safe_run_id / "word_document" / safe_filename).resolve()

    if not file_path.is_relative_to(output_root):
        raise HTTPException(status_code=400, detail="Invalid download path")

    return file_path


async def _process_documents(
    rfp: UploadFile,
    example_rfps: list[UploadFile],
    company_context: list[UploadFile] | None,
    features: FeaturesConfig,
    pdf_service: PDFService,
    store_documents: bool = True,
) -> ProcessedDocuments:
    """
    Process uploaded PDF documents and extract text/images.
    
    Shared logic between streaming and non-streaming endpoints.
    
    Args:
        store_documents: If True, include file bytes in DocumentInfo for storage.
    """
    image_budgets = _allocate_image_budgets(
        features,
        len(example_rfps),
        1,
        len(company_context or []),
    )
    
    # Collect document info for storage
    documents: list[DocumentInfo] = []
    
    # Read RFP
    logger.info(f"Processing RFP: {rfp.filename}")
    rfp_bytes = await rfp.read()
    rfp_text = pdf_service.extract_text_from_bytes(rfp_bytes)
    
    if store_documents:
        documents.append(DocumentInfo(
            filename=rfp.filename or "rfp.pdf",
            file_type="rfp",
            file_bytes=rfp_bytes
        ))
    
    rfp_images = None
    if features.enable_images and image_budgets["rfp"] > 0:
        rfp_images = pdf_service.pdf_to_base64_images(
            rfp_bytes,
            max_pages=image_budgets["rfp"],
            min_table_rows=features.min_table_rows,
            min_table_cols=features.min_table_cols,
        )
    
    # Read example RFPs
    example_texts = []
    example_images = []
    for idx, ex in enumerate(example_rfps):
        logger.info(f"Processing example: {ex.filename}")
        ex_bytes = await ex.read()
        example_texts.append(pdf_service.extract_text_from_bytes(ex_bytes))
        
        if store_documents:
            documents.append(DocumentInfo(
                filename=ex.filename or f"example_{idx+1}.pdf",
                file_type="example",
                file_bytes=ex_bytes
            ))
        
        if features.enable_images and image_budgets["examples_per_doc"] > 0:
            example_images.append(
                pdf_service.pdf_to_base64_images(
                    ex_bytes,
                    max_pages=image_budgets["examples_per_doc"],
                    min_table_rows=features.min_table_rows,
                    min_table_cols=features.min_table_cols,
                )
            )
    
    # Read company context
    context_text = None
    context_images = None
    if company_context:
        context_parts = []
        context_imgs = []
        for idx, ctx in enumerate(company_context):
            logger.info(f"Processing context: {ctx.filename}")
            ctx_bytes = await ctx.read()
            context_parts.append(pdf_service.extract_text_from_bytes(ctx_bytes))
            
            if store_documents:
                documents.append(DocumentInfo(
                    filename=ctx.filename or f"context_{idx+1}.pdf",
                    file_type="context",
                    file_bytes=ctx_bytes
                ))
            
            if features.enable_images and image_budgets["context_per_doc"] > 0:
                context_imgs.extend(
                    pdf_service.pdf_to_base64_images(
                        ctx_bytes,
                        max_pages=image_budgets["context_per_doc"],
                        min_table_rows=features.min_table_rows,
                        min_table_cols=features.min_table_cols,
                    )
                )
        
        context_text = "\n\n---\n\n".join(context_parts)
        context_images = context_imgs if context_imgs else None
    
    return ProcessedDocuments(
        rfp_text=rfp_text,
        rfp_images=rfp_images,
        example_texts=example_texts,
        example_images=example_images if example_images else None,
        context_text=context_text,
        context_images=context_images,
        documents=documents,
    )


def _build_workflow_input(
    docs: ProcessedDocuments,
    *,
    enable_planner: Optional[bool] = None,
    enable_critiquer: Optional[bool] = None,
    generator_formatting_injection: Optional[str] = None,
    generator_intro_pages: Optional[int] = None,
    generation_page_overlap: Optional[int] = None,
    toggle_generation_chunking: Optional[bool] = None,
    max_tokens_generation_chunking: Optional[int] = None,
    max_sections_per_chunk: Optional[int] = None,
    extract_reqs_comment: Optional[str] = None,
    previous_requirements: Optional[list[RFPRequirement]] = None,
    plan_comment: Optional[str] = None,
    previous_plan: Optional[ProposalPlan] = None,
    generate_rfp_comment: Optional[str] = None,
    previous_document_code: Optional[str] = None,
    critique_comment: Optional[str] = None,
) -> WorkflowInput:
    """Build a WorkflowInput object from processed document content."""
    return WorkflowInput(
        rfp_text=docs.rfp_text,
        rfp_images=docs.rfp_images,
        example_rfps_text=docs.example_texts,
        example_rfps_images=docs.example_images,
        company_context_text=docs.context_text,
        company_context_images=docs.context_images,
        enable_planner=enable_planner,
        enable_critiquer=enable_critiquer,
        generator_formatting_injection=generator_formatting_injection,
        generator_intro_pages=generator_intro_pages,
        generation_page_overlap=generation_page_overlap,
        toggle_generation_chunking=toggle_generation_chunking,
        max_tokens_generation_chunking=max_tokens_generation_chunking,
        max_sections_per_chunk=max_sections_per_chunk,
        extract_reqs_comment=extract_reqs_comment,
        previous_requirements=previous_requirements,
        plan_comment=plan_comment,
        previous_plan=previous_plan,
        generate_rfp_comment=generate_rfp_comment,
        previous_document_code=previous_document_code,
        critique_comment=critique_comment,
        documents=docs.documents,
    )


async def _run_orchestrated_generation(
    rfp: UploadFile,
    example_rfps: list[UploadFile],
    company_context: Optional[list[UploadFile]],
    enable_planner: Optional[bool],
    enable_critiquer: Optional[bool],
    generator_formatting_injection: Optional[str],
    generator_intro_pages: Optional[int],
    generation_page_overlap: Optional[int],
    toggle_generation_chunking: Optional[bool],
    max_tokens_generation_chunking: Optional[int],
    max_sections_per_chunk: Optional[int],
) -> GenerateRFPResponse:
    """Shared orchestrated flow implementation."""
    start_time = time.time()
    config = get_config()
    pdf_service = PDFService()

    all_files = [rfp] + example_rfps + (company_context or [])
    _validate_pdf_files(all_files)

    docs = await _process_documents(
        rfp,
        example_rfps,
        company_context,
        config.features,
        pdf_service,
    )

    workflow_input = _build_workflow_input(
        docs,
        enable_planner=enable_planner,
        enable_critiquer=enable_critiquer,
        generator_formatting_injection=generator_formatting_injection,
        generator_intro_pages=generator_intro_pages,
        generation_page_overlap=generation_page_overlap,
        toggle_generation_chunking=toggle_generation_chunking,
        max_tokens_generation_chunking=max_tokens_generation_chunking,
        max_sections_per_chunk=max_sections_per_chunk,
    )

    logger.info("Starting orchestrated RFP workflow")
    workflow = create_rfp_workflow()
    result = await workflow.run(workflow_input)

    processing_time = time.time() - start_time
    docx_path = Path(result.docx_path) if result.docx_path else None
    if docx_path:
        run_id = docx_path.parent.parent.name
        docx_filename = docx_path.name
        _sync_run_to_blob(docx_path.parent.parent)
    else:
        run_id = "unknown"
        docx_filename = "proposal.docx"

    return GenerateRFPResponse(
        success=True,
        message="RFP response generated successfully",
        rfp_response=result.response,
        analysis=result.analysis,
        docx_download_url=f"/api/rfp/download/{run_id}/{docx_filename}",
        processing_time_seconds=round(processing_time, 2),
    )


@router.post("/generate", response_model=GenerateRFPResponse)
async def generate_rfp(
    rfp: UploadFile = File(..., description="The RFP document to respond to (PDF)"),
    example_rfps: list[UploadFile] = File(
        ...,
        description="Example RFP responses for style reference (PDF)",
    ),
    company_context: Optional[list[UploadFile]] = File(
        None,
        description="Company context documents (PDF)",
    ),
    enable_planner: Optional[bool] = Form(None),
    enable_critiquer: Optional[bool] = Form(None),
    generator_formatting_injection: Optional[str] = Form(None),
    generator_intro_pages: Optional[int] = Form(None),
    generation_page_overlap: Optional[int] = Form(None),
    toggle_generation_chunking: Optional[bool] = Form(None),
    max_tokens_generation_chunking: Optional[int] = Form(None),
    max_sections_per_chunk: Optional[int] = Form(None),
    _: Optional[str] = Depends(verify_api_token),
):
    """Legacy orchestrated endpoint kept for compatibility."""
    try:
        return await _run_orchestrated_generation(
            rfp=rfp,
            example_rfps=example_rfps,
            company_context=company_context,
            enable_planner=enable_planner,
            enable_critiquer=enable_critiquer,
            generator_formatting_injection=generator_formatting_injection,
            generator_intro_pages=generator_intro_pages,
            generation_page_overlap=generation_page_overlap,
            toggle_generation_chunking=toggle_generation_chunking,
            max_tokens_generation_chunking=max_tokens_generation_chunking,
            max_sections_per_chunk=max_sections_per_chunk,
        )
    except HTTPException:
        raise
    except Exception as exc:
        logger.exception("RFP generation failed")
        raise HTTPException(status_code=500, detail=f"RFP generation failed: {exc}") from exc


@router.post("/orchestrate", response_model=GenerateRFPResponse)
async def orchestrate_rfp(
    source_rfp: UploadFile = File(..., description="The RFP document to respond to (PDF)"),
    example_rfps: list[UploadFile] = File(
        ...,
        description="Example RFP responses for style reference (PDF)",
    ),
    company_context: Optional[list[UploadFile]] = File(
        None,
        description="Company context documents (PDF)",
    ),
    enable_planner: Optional[bool] = Form(None),
    enable_critiquer: Optional[bool] = Form(None),
    generator_formatting_injection: Optional[str] = Form(None),
    generator_intro_pages: Optional[int] = Form(None),
    generation_page_overlap: Optional[int] = Form(None),
    toggle_generation_chunking: Optional[bool] = Form(None),
    max_tokens_generation_chunking: Optional[int] = Form(None),
    max_sections_per_chunk: Optional[int] = Form(None),
    _: Optional[str] = Depends(verify_api_token),
):
    """New orchestrated endpoint using snake_case parameter names."""
    try:
        return await _run_orchestrated_generation(
            rfp=source_rfp,
            example_rfps=example_rfps,
            company_context=company_context,
            enable_planner=enable_planner,
            enable_critiquer=enable_critiquer,
            generator_formatting_injection=generator_formatting_injection,
            generator_intro_pages=generator_intro_pages,
            generation_page_overlap=generation_page_overlap,
            toggle_generation_chunking=toggle_generation_chunking,
            max_tokens_generation_chunking=max_tokens_generation_chunking,
            max_sections_per_chunk=max_sections_per_chunk,
        )
    except HTTPException:
        raise
    except Exception as exc:
        logger.exception("RFP orchestration failed")
        raise HTTPException(status_code=500, detail=f"RFP orchestration failed: {exc}") from exc


def _allocate_image_budgets(features, example_count: int, rfp_count: int, context_count: int) -> dict[str, int]:
    total = max(0, int(features.max_images))
    ratios = features.normalized_image_ratios(
        include_examples=example_count > 0,
        include_rfp=rfp_count > 0,
        include_context=context_count > 0,
    )
    targets = {
        "examples": ratios["examples"] * total,
        "rfp": ratios["rfp"] * total,
        "context": ratios["context"] * total,
    }
    floors = {key: int(value) for key, value in targets.items()}
    remainder = total - sum(floors.values())
    if remainder > 0:
        order = sorted(
            targets.items(),
            key=lambda item: (item[1] - floors[item[0]]),
            reverse=True,
        )
        for key, _ in order:
            if remainder <= 0:
                break
            floors[key] += 1
            remainder -= 1
    examples_per_doc = floors["examples"] // example_count if example_count else 0
    context_per_doc = floors["context"] // context_count if context_count else 0
    return {
        "examples": floors["examples"],
        "rfp": floors["rfp"],
        "context": floors["context"],
        "examples_per_doc": examples_per_doc,
        "context_per_doc": context_per_doc,
    }


@router.post("/extract-reqs", response_model=ExtractReqsResponse)
async def extract_reqs(
    source_rfp: UploadFile = File(..., description="RFP PDF to extract requirements from"),
    company_context: Optional[list[UploadFile]] = File(None, description="Optional company context PDFs"),
    comment: Optional[str] = Form(None, description="Optional guidance for regeneration"),
    run_id: Optional[str] = Form(None, description="Optional run identifier to reuse the same run directory"),
    previous_requirements: Optional[str] = Form(
        None,
        description="Optional JSON array of previously generated requirements",
    ),
    _: Optional[str] = Depends(verify_api_token),
):
    """Extract requirements from an RFP with optional regeneration context."""
    config = get_config()
    pdf_service = PDFService()
    _validate_pdf_files([source_rfp] + (company_context or []))

    previous_reqs = _parse_previous_requirements(previous_requirements)
    docs = await _process_documents(
        source_rfp,
        [],
        company_context,
        config.features,
        pdf_service,
    )

    run_dir = _resolve_step_run_directory(run_id)
    _save_documents(run_dir, docs.documents)

    client = create_llm_client()
    analyzer = RFPAnalyzerExecutor(client, run_dir)
    input_data = _build_workflow_input(
        docs,
        extract_reqs_comment=comment,
        previous_requirements=previous_reqs,
    )

    try:
        analysis_result = await analyzer.execute(input_data)
        _save_analysis_metadata(run_dir, analysis_result.analysis)
        _append_analysis_version(run_dir, analysis_result.analysis, comment)
        _sync_run_to_blob(run_dir)
        return ExtractReqsResponse(
            success=True,
            message="Requirements extracted successfully",
            analysis=analysis_result.analysis,
            run_id=run_dir.name,
        )
    except Exception as exc:
        logger.exception("Requirement extraction failed")
        raise HTTPException(status_code=500, detail=f"Requirement extraction failed: {exc}") from exc


@router.post("/plan", response_model=PlanStepResponse)
async def plan_rfp(
    request: PlanStepRequest,
    _: Optional[str] = Depends(verify_api_token),
):
    """Generate a proposal plan from extracted requirements."""
    run_dir = _resolve_step_run_directory(request.run_id)
    client = create_llm_client()
    planner = PlannerExecutor(client, run_dir)

    input_data = WorkflowInput(
        rfp_text="",
        company_context_text=request.company_context_text,
        plan_comment=request.comment,
        previous_plan=request.previous_plan,
    )

    try:
        planning_result = await planner.execute(input_data, request.analysis)
        _save_analysis_metadata(run_dir, request.analysis)
        _save_plan_metadata(run_dir, planning_result.plan)
        _append_plan_version(run_dir, planning_result.plan, request.comment)
        _sync_run_to_blob(run_dir)
        return PlanStepResponse(
            success=True,
            message="Plan generated successfully",
            plan=planning_result.plan,
            run_id=run_dir.name,
        )
    except Exception as exc:
        logger.exception("Plan generation failed")
        raise HTTPException(status_code=500, detail=f"Plan generation failed: {exc}") from exc


@router.post("/plan-with-context", response_model=PlanStepResponse)
async def plan_rfp_with_context(
    analysis_json: str = Form(..., description="JSON for RFPAnalysis"),
    company_context_text: Optional[str] = Form(None, description="Optional free-text planning context"),
    company_context: Optional[list[UploadFile]] = File(None, description="Optional context PDFs"),
    comment: Optional[str] = Form(None, description="Optional guidance for planning/regeneration"),
    run_id: Optional[str] = Form(None, description="Optional run identifier to reuse the same run directory"),
    previous_plan_json: Optional[str] = Form(None, description="Optional JSON for ProposalPlan"),
    _: Optional[str] = Depends(verify_api_token),
):
    """Generate a proposal plan using analysis + optional context PDFs + optional context text."""
    try:
        analysis = RFPAnalysis.model_validate_json(analysis_json)
    except ValidationError as exc:
        raise HTTPException(status_code=400, detail=f"Invalid analysis_json schema: {exc}") from exc
    except json.JSONDecodeError as exc:
        raise HTTPException(status_code=400, detail=f"Invalid analysis_json: {exc}") from exc

    previous_plan = _parse_json_model(previous_plan_json, ProposalPlan, "previous_plan_json")

    pdf_service = PDFService()
    context_pdf_text, context_docs = await _extract_context_from_files(company_context, pdf_service)

    context_parts = []
    if company_context_text and company_context_text.strip():
        context_parts.append(company_context_text.strip())
    if context_pdf_text and context_pdf_text.strip():
        context_parts.append(context_pdf_text.strip())
    combined_context = "\n\n---\n\n".join(context_parts) if context_parts else None

    run_dir = _resolve_step_run_directory(run_id)
    _save_documents(run_dir, context_docs)

    client = create_llm_client()
    planner = PlannerExecutor(client, run_dir)
    input_data = WorkflowInput(
        rfp_text="",
        company_context_text=combined_context,
        plan_comment=comment,
        previous_plan=previous_plan,
        documents=context_docs,
    )

    try:
        planning_result = await planner.execute(input_data, analysis)
        _save_analysis_metadata(run_dir, analysis)
        _save_plan_metadata(run_dir, planning_result.plan)
        _append_plan_version(run_dir, planning_result.plan, comment)
        _sync_run_to_blob(run_dir)
        return PlanStepResponse(
            success=True,
            message="Plan generated successfully",
            plan=planning_result.plan,
            run_id=run_dir.name,
        )
    except Exception as exc:
        logger.exception("Plan generation (with context) failed")
        raise HTTPException(status_code=500, detail=f"Plan generation failed: {exc}") from exc


@router.post("/generate-rfp", response_model=GenerateRFPStepResponse)
async def generate_rfp_step(
    source_rfp: UploadFile = File(..., description="RFP PDF to generate against"),
    example_rfps: list[UploadFile] = File(..., description="Style reference PDFs"),
    company_context: Optional[list[UploadFile]] = File(None, description="Optional company context PDFs"),
    analysis_json: str = Form(..., description="JSON for RFPAnalysis"),
    plan_json: Optional[str] = Form(None, description="Optional JSON for ProposalPlan"),
    comment: Optional[str] = Form(None, description="Optional guidance for regeneration"),
    run_id: Optional[str] = Form(None, description="Optional run identifier to reuse the same run directory"),
    previous_document_code: Optional[str] = Form(
        None,
        description="Optional previous generated code for iterative regeneration",
    ),
    generator_formatting_injection: Optional[str] = Form(None),
    generator_intro_pages: Optional[int] = Form(None),
    generation_page_overlap: Optional[int] = Form(None),
    toggle_generation_chunking: Optional[bool] = Form(None),
    max_tokens_generation_chunking: Optional[int] = Form(None),
    max_sections_per_chunk: Optional[int] = Form(None),
    _: Optional[str] = Depends(verify_api_token),
):
    """
    Generate and execute proposal code.

    Returns both executable code and a base64 DOCX payload.
    """
    config = get_config()
    pdf_service = PDFService()
    _validate_pdf_files([source_rfp] + example_rfps + (company_context or []))

    try:
        analysis = RFPAnalysis.model_validate_json(analysis_json)
    except ValidationError as exc:
        raise HTTPException(status_code=400, detail=f"Invalid analysis_json schema: {exc}") from exc
    except json.JSONDecodeError as exc:
        raise HTTPException(status_code=400, detail=f"Invalid analysis_json: {exc}") from exc

    plan = _parse_json_model(plan_json, ProposalPlan, "plan_json")
    docs = await _process_documents(
        source_rfp,
        example_rfps,
        company_context,
        config.features,
        pdf_service,
    )

    run_dir = _resolve_step_run_directory(run_id)
    _save_documents(run_dir, docs.documents)
    _save_analysis_metadata(run_dir, analysis)
    if plan:
        _save_plan_metadata(run_dir, plan)

    client = create_llm_client()
    generator = SectionGeneratorExecutor(client, run_dir)
    code_interpreter = CodeInterpreterExecutor(client, run_dir)

    input_data = _build_workflow_input(
        docs,
        generator_formatting_injection=generator_formatting_injection,
        generator_intro_pages=generator_intro_pages,
        generation_page_overlap=generation_page_overlap,
        toggle_generation_chunking=toggle_generation_chunking,
        max_tokens_generation_chunking=max_tokens_generation_chunking,
        max_sections_per_chunk=max_sections_per_chunk,
        generate_rfp_comment=comment,
        previous_document_code=previous_document_code,
    )

    try:
        if plan:
            chunking_enabled = (
                toggle_generation_chunking
                if toggle_generation_chunking is not None
                else config.features.toggle_generation_chunking
            )

            if chunking_enabled and plan.sections:
                max_sections = (
                    max_sections_per_chunk
                    if max_sections_per_chunk is not None
                    else config.features.max_sections_per_chunk
                )
                section_chunks = _chunk_sections(plan.sections, max_sections)
                chunk_codes: list[str] = []
                for idx, chunk in enumerate(section_chunks, start=1):
                    chunk_result = await generator.execute_chunk_with_plan(
                        input_data,
                        analysis,
                        chunk,
                        part_number=idx,
                        total_parts=len(section_chunks),
                    )
                    _save_code_snapshot(run_dir, f"01_chunk_{idx}", chunk_result.response.document_code)
                    chunk_codes.append(chunk_result.response.document_code)
                generation_result = await generator.synthesize_from_chunks(input_data, chunk_codes)
                _save_code_snapshot(run_dir, "02_synthesized", generation_result.response.document_code)
            else:
                generation_result = await generator.execute_with_plan(input_data, analysis, plan)
                _save_code_snapshot(run_dir, "01_initial", generation_result.response.document_code)
        else:
            generation_result = await generator.execute(input_data, analysis)
            _save_code_snapshot(run_dir, "01_initial", generation_result.response.document_code)

        max_error_loops = config.workflow.max_error_loops
        current_response = generation_result.response
        docx_dir = run_dir / "word_document"
        image_dir = run_dir / "image_assets"
        docx_path = docx_dir / "proposal.docx"
        code_stats: dict = {}

        for error_loop in range(max_error_loops + 1):
            docx_path, code_stats = await code_interpreter.execute(
                current_response,
                image_dir=image_dir,
                docx_dir=docx_dir,
            )

            if code_stats.get("document_success"):
                break

            if error_loop < max_error_loops:
                error_msg = code_stats.get("errors", ["Unknown error"])[0]
                generation_result = await generator.execute_with_error(
                    input_data,
                    analysis,
                    current_response.document_code,
                    error_msg,
                )
                current_response = generation_result.response
                _save_code_snapshot(run_dir, f"03_error_recovery_{error_loop + 1}", current_response.document_code)

        _save_code_snapshot(run_dir, "99_final", current_response.document_code)

        docx_bytes = docx_path.read_bytes() if docx_path.exists() else b""
        docx_base64 = base64.b64encode(docx_bytes).decode("ascii")
        execution_success = bool(code_stats.get("document_success", False))
        code_package = _enrich_code_package_with_previews(
            _build_code_package(current_response.document_code),
            image_dir=image_dir,
            docx_path=docx_path,
        )
        _sync_run_to_blob(run_dir)

        return GenerateRFPStepResponse(
            success=execution_success,
            message="RFP generated and executed successfully" if execution_success else "RFP code generated, but execution failed",
            document_code=current_response.document_code,
            docx_base64=docx_base64,
            docx_filename="proposal.docx",
            execution_stats=code_stats,
            code_package=code_package,
            run_id=run_dir.name,
            docx_download_url=f"/api/rfp/download/{run_dir.name}/proposal.docx",
        )
    except HTTPException:
        raise
    except Exception as exc:
        logger.exception("Generate RFP step failed")
        raise HTTPException(status_code=500, detail=f"Generate RFP step failed: {exc}") from exc


@router.post("/critique", response_model=CritiqueStepResponse)
async def critique_rfp(
    request: CritiqueStepRequest,
    _: Optional[str] = Depends(verify_api_token),
):
    """Critique generated code with optional user guidance."""
    run_dir = _resolve_step_run_directory(request.run_id)
    client = create_llm_client()
    critiquer = CritiquerExecutor(client, run_dir)
    try:
        critique_result = await critiquer.execute(
            request.analysis,
            request.document_code,
            comment=request.comment,
        )
        _save_analysis_metadata(run_dir, request.analysis)
        if request.run_id:
            existing_critiques: list[dict] = []
            critiques_path = run_dir / "metadata" / "critiques.json"
            if critiques_path.exists():
                try:
                    with open(critiques_path, "r", encoding="utf-8") as f:
                        loaded = json.load(f)
                    if isinstance(loaded, list):
                        existing_critiques = loaded
                except Exception:
                    existing_critiques = []
            if critique_result.critique:
                existing_critiques.append(critique_result.critique.model_dump())
                _save_json_artifact(run_dir, "metadata/critiques.json", existing_critiques)
        _sync_run_to_blob(run_dir)
        return CritiqueStepResponse(
            success=True,
            message="Critique completed successfully",
            critique=critique_result.critique,
            run_id=run_dir.name,
        )
    except Exception as exc:
        logger.exception("Critique step failed")
        raise HTTPException(status_code=500, detail=f"Critique step failed: {exc}") from exc


@router.get("/download/{run_id}/{filename}")
async def download_file(run_id: str, filename: str):
    """Download a generated DOCX file from a run directory."""
    file_path = _resolve_docx_download_path(run_id, filename)

    if not file_path.exists():
        storage = get_blob_storage()
        if storage.enabled:
            relative = f"word_document/{file_path.name}"
            data = storage.download_blob_bytes(run_id, relative)
            if data:
                file_path.parent.mkdir(parents=True, exist_ok=True)
                with open(file_path, "wb") as f:
                    f.write(data)

    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File not found")
    
    return FileResponse(
        path=file_path,
        filename=file_path.name,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


@router.post("/generate/stream")
async def generate_rfp_stream(
    rfp: UploadFile = File(...),
    example_rfps: list[UploadFile] = File(...),
    company_context: Optional[list[UploadFile]] = File(None),
    enable_planner: Optional[bool] = Form(None),
    enable_critiquer: Optional[bool] = Form(None),
    generator_formatting_injection: Optional[str] = Form(None),
    generator_intro_pages: Optional[int] = Form(None),
    generation_page_overlap: Optional[int] = Form(None),
    toggle_generation_chunking: Optional[bool] = Form(None),
    max_tokens_generation_chunking: Optional[int] = Form(None),
    max_sections_per_chunk: Optional[int] = Form(None),
):
    """
    Generate an RFP response with streaming progress events.
    
    Returns Server-Sent Events (SSE) with workflow progress.
    """
    config = get_config()
    pdf_service = PDFService()
    _validate_pdf_files([rfp] + example_rfps + (company_context or []))
    
    async def event_generator():
        try:
            # Process files using shared function
            yield f"data: {json.dumps({'event': 'processing', 'message': 'Reading documents...'})}\n\n"
            
            docs = await _process_documents(
                rfp, example_rfps, company_context,
                config.features, pdf_service
            )
            
            workflow_input = _build_workflow_input(
                docs,
                enable_planner=enable_planner,
                enable_critiquer=enable_critiquer,
                generator_formatting_injection=generator_formatting_injection,
                generator_intro_pages=generator_intro_pages,
                generation_page_overlap=generation_page_overlap,
                toggle_generation_chunking=toggle_generation_chunking,
                max_tokens_generation_chunking=max_tokens_generation_chunking,
                max_sections_per_chunk=max_sections_per_chunk,
            )
            
            # Run workflow with streaming
            workflow = create_rfp_workflow()
            run_id = None
            async for event in workflow.run_stream(workflow_input):
                yield f"data: {json.dumps({'event': event.event_type, 'step': event.step_name, 'message': event.message, 'data': event.data})}\n\n"
                # Capture run_id from the finished event
                if event.event_type == "finished" and event.data:
                    run_id = event.data.get("run_id") or event.data.get("run_dir", "").split("/")[-1] or event.data.get("run_dir", "").split("\\")[-1]
            
            # Send final download URL using run_id from workflow
            if run_id:
                run_dir = Path(config.app.output_dir) / run_id
                _sync_run_to_blob(run_dir)
                yield f"data: {json.dumps({'event': 'complete', 'download_url': f'/api/rfp/download/{run_id}/proposal.docx', 'run_id': run_id})}\n\n"
            
        except Exception as e:
            yield f"data: {json.dumps({'event': 'error', 'message': str(e)})}\n\n"
    
    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream"
    )


@router.post("/orchestrate/stream")
async def orchestrate_rfp_stream(
    source_rfp: UploadFile = File(...),
    example_rfps: list[UploadFile] = File(...),
    company_context: Optional[list[UploadFile]] = File(None),
    enable_planner: Optional[bool] = Form(None),
    enable_critiquer: Optional[bool] = Form(None),
    generator_formatting_injection: Optional[str] = Form(None),
    generator_intro_pages: Optional[int] = Form(None),
    generation_page_overlap: Optional[int] = Form(None),
    toggle_generation_chunking: Optional[bool] = Form(None),
    max_tokens_generation_chunking: Optional[int] = Form(None),
    max_sections_per_chunk: Optional[int] = Form(None),
):
    """Snake_case alias for orchestrated streaming generation."""
    return await generate_rfp_stream(
        rfp=source_rfp,
        example_rfps=example_rfps,
        company_context=company_context,
        enable_planner=enable_planner,
        enable_critiquer=enable_critiquer,
        generator_formatting_injection=generator_formatting_injection,
        generator_intro_pages=generator_intro_pages,
        generation_page_overlap=generation_page_overlap,
        toggle_generation_chunking=toggle_generation_chunking,
        max_tokens_generation_chunking=max_tokens_generation_chunking,
        max_sections_per_chunk=max_sections_per_chunk,
    )
