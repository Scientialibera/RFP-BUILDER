from docx.enum.style import WD_STYLE_TYPE
import matplotlib.patches as mpatches
import matplotlib.dates as mdates
import datetime as dt

styles = doc.styles

# Configure Normal style
normal = styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)

# Create Caption style if missing
if 'Caption' not in [s.name for s in styles]:
    cap = styles.add_style('Caption', WD_STYLE_TYPE.PARAGRAPH)
    cap.font.name = 'Calibri'
    cap.font.size = Pt(9)
    cap.font.italic = True


def add_caption(text: str):
    p = doc.add_paragraph(text, style='Caption')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


figure_counter = {'n': 0}


def next_figure_label(description: str) -> str:
    figure_counter['n'] += 1
    return f"Figure {figure_counter['n']}: {description}"


# ------------------------
# Title Page
# ------------------------
doc.add_heading('Proposal for Phase 1 Sanitary Sewer Evaluation Survey (SSES)\nLaSalle Area Sewersheds', level=0)
subtitle = doc.add_paragraph('Submitted to: Niagara Falls Water Board')
subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT
subtitle.paragraph_format.space_after = Pt(6)

doc.add_paragraph('Submitted by: Aurelia Digital Systems').paragraph_format.space_after = Pt(0)
doc.add_paragraph('Date: February 2026').paragraph_format.space_after = Pt(0)

doc.add_paragraph('This proposal outlines Aurelia Digital Systems’ comprehensive, compliance-focused approach to delivering Phase 1 SSES services in the LaSalle area sewersheds, with scalable options for Phases 2 and 3.').paragraph_format.space_before = Pt(12)

doc.add_page_break()

# ------------------------
# Executive Summary and Win Themes
# ------------------------
sec = doc.add_heading('Executive Summary and Win Themes', level=1)
sec.paragraph_format.space_before = Pt(6)

p = doc.add_paragraph()
run = p.add_run('Overview of NFWB objectives')
run.bold = True
p.add_run(': The Niagara Falls Water Board (NFWB) seeks a qualified engineering consultant to perform a Phase 1 sanitary sewer evaluation survey (SSES) in the LaSalle area sewersheds, identify inflow and infiltration (I&I) sources, and develop a corrective action plan that satisfies NYSDEC expectations and the Wastewater Infrastructure Engineering Planning Grant (EPG) conditions. Aurelia Digital Systems will deliver a methodical yet adaptable Phase 1 program, with clearly priced options for Phases 2 and 3, that minimizes regulatory, schedule, and funding risk for the Board.')

p = doc.add_paragraph()
run = p.add_run('Our value to NFWB')
run.bold = True
p.add_run(': We combine deep public-sector delivery experience with rigorous governance, safety, and documentation practices. Our approach is explicitly aligned to the October 1, 2022 Engineering Report Outline for NYS Wastewater Infrastructure Projects and to the LaSalle Order on Consent objectives. We emphasize traceable compliance, transparent decision-making, and data-driven I&I reduction strategies that can be defended to NYSDEC and other stakeholders.')

p = doc.add_paragraph()
run = p.add_run('Key benefits of selecting Aurelia')
run.bold = True
p.add_run(' include:')

doc.add_paragraph('Regulatory and grant compliance excellence – explicit alignment to NYSDEC and EPG requirements, with a structured compliance matrix and clear documentation of how each requirement is satisfied.', style='List Bullet')
doc.add_paragraph('Reduced risk for NFWB – robust project governance, QA/QC, and safety management across all SSES tasks, including smoke testing, flow isolation, and confined space work.', style='List Bullet')
doc.add_paragraph('Defensible I&I mitigation roadmap – a transparent methodology to identify, quantify, and prioritize I&I sources, resulting in a corrective action plan that maximizes I&I reduction per dollar invested.', style='List Bullet')
doc.add_paragraph('Schedule certainty – a realistic yet efficient schedule that supports NYSDEC review early in 2023, construction in 2023–2024, and post-construction monitoring by June 2024.', style='List Bullet')
doc.add_paragraph('Scalable options for Phases 2 and 3 – clearly priced options following the same proven methodology, minimizing incremental mobilization and learning curve.', style='List Bullet')
doc.add_paragraph('Strong MWBE/SDVOB utilization – a proactive strategy to meet the 30% MWBE and 6% SDVOB goals, with clear reporting and documentation to protect grant funding.', style='List Bullet')

# Objectives vs outcomes table
obj_table = doc.add_table(rows=1, cols=3)
obj_table.style = 'Table Grid'
obj_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = obj_table.rows[0].cells
hdr[0].text = 'NFWB Objective'
hdr[1].text = 'Our Proposed Outcome'
hdr[2].text = 'Evidence of Ability to Deliver'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

rows = [
    ('Comply with NYSDEC and EPG requirements for Phase 1 SSES and reporting',
     'Engineering Report and Post-Construction I&I Assessment that conform to the October 1, 2022 outline and all grant conditions',
     'Experience preparing NYSDEC-accepted wastewater planning reports and grant-funded deliverables; PE-led QA/QC and compliance matrix.'),
    ('Identify and prioritize I&I sources in LaSalle sewersheds',
     'Defensible inventory of I&I defects, ranked by cost-effectiveness and feasibility of correction',
     'Structured SSES methodology including data review, smoke testing, flow isolation, CCTV and manhole analysis, and focused investigations.'),
    ('Support timely construction and monitoring windows',
     'Schedule that enables NYSDEC review early in 2023, construction in 2023–early 2024, and monitoring through June 2024',
     'Realistic Gantt schedule, active risk management, and experienced project management with utilities and public-sector programs.'),
    ('Maximize value of EPG funding and protect state assistance',
     'Audit-ready documentation of compliance, MWBE/SDVOB utilization, and financial reporting',
     'Established internal controls for grant-funded work, monthly reporting routines, and experience with NYS assistance terms.'),
]

for obj, outcome, evidence in rows:
    row_cells = obj_table.add_row().cells
    row_cells[0].text = obj
    row_cells[1].text = outcome
    row_cells[2].text = evidence

# Mermaid diagram linking Phase 1 SSES to outcomes
mermaid_code = '''
flowchart TD
  A[Phase 1 SSES] --> B[Identify I and I sources]
  B --> C[Corrective action plan]
  C --> D[Design and construction]
  D --> E[Post construction monitoring]
  E --> F[Documented I and I reduction]
  F --> G[NYSDEC and EPG compliance]
'''
mermaid_path = render_mermaid(mermaid_code, 'exec_summary_flow')
doc.add_picture(str(mermaid_path), width=Inches(5.8))
add_caption(next_figure_label('Linkage from Phase 1 SSES to corrective actions and NYSDEC outcomes'))

# Conceptual I&I reduction impact chart
sns.set_style('whitegrid')
plt.figure(figsize=(6, 3.5))
impact_data = pd.DataFrame({
    'Intervention': ['CIPP lining', 'Spot repairs', 'Inflow removal', 'Manhole rehab'],
    'Relative I and I Reduction': [40, 20, 25, 15]
})
ax = sns.barplot(data=impact_data, x='Intervention', y='Relative I and I Reduction', color='#4c72b0')
ax.set_ylabel('Relative reduction index')
ax.set_xlabel('Intervention category')
ax.set_title('Conceptual I and I reduction by intervention category')
plt.tight_layout()
impact_path = output_dir / 'conceptual_ii_reduction.png'
plt.savefig(impact_path, dpi=150, bbox_inches='tight', facecolor='white')
plt.close()

doc.add_picture(str(impact_path), width=Inches(5.8))
add_caption(next_figure_label('Conceptual I and I reduction impact by intervention category'))

# ------------------------
# Understanding of Project Context and Objectives
# ------------------------
sec = doc.add_heading('Understanding of Project Context and Objectives', level=1)
sec.paragraph_format.space_before = Pt(12)

p = doc.add_paragraph('NFWB is implementing a multi-phase sanitary sewer evaluation survey program in the LaSalle area to reduce sanitary sewer overflows (SSOs), protect water quality, and comply with an Order on Consent and related NYSDEC directives. Phase 1 focuses on selected sewersheds and two outfalls/SSOs, with subsequent Phases 2 and 3 extending similar work to additional basins.')

p = doc.add_paragraph('We understand that the project is funded in part through a Wastewater Infrastructure Engineering Planning Grant. As a result, the work must strictly adhere to EPG terms, the October 1, 2022 Engineering Report Outline, and Mandatory State Financial Assistance Terms and Conditions. The schedule must also support NYSDEC review early in 2023, construction in 2023–early 2024, and post-construction monitoring through June 2024.')

p = doc.add_paragraph('The LaSalle sewersheds have experienced recurring wet-weather SSOs driven by excessive I&I. Previous maintenance and I&I reduction efforts have been undertaken, but a systematic, data-driven SSES is required to fully characterize remaining sources, prioritize remedial measures, and demonstrate measurable reduction in extraneous flows.')

# Regulatory/EPG drivers mapping table
reg_table = doc.add_table(rows=1, cols=3)
reg_table.style = 'Table Grid'
reg_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = reg_table.rows[0].cells
hdr[0].text = 'Regulatory / Funding Driver'
hdr[1].text = 'Implication for This Project'
hdr[2].text = 'Associated Deliverables / Checkpoints'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

reg_rows = [
    ('Order on Consent and NYSDEC directives',
     'Demonstrate sustained SSO abatement and I&I reduction in LaSalle sewersheds.',
     'Phase 1 SSES Engineering Report, Corrective Action Plan, and Post-Construction I and I Assessment and Cost Analysis Report.'),
    ('EPG Wastewater Infrastructure Engineering Planning Grant',
     'Meet all grant terms, including eligible work scope, schedule, MWBE/SDVOB goals, and documentation.',
     'Grant-compliant project files, MWBE/SDVOB utilization plans and reports, and audit-ready financial records.'),
    ('Engineering Report Outline (October 1, 2022)',
     'Prepare an engineering report that conforms to required sections, level of analysis, and PE certification.',
     'Structured report crosswalked to outline sections, with PE stamp and NYSDEC-ready figures and appendices.'),
    ('Mandatory State Financial Assistance Terms',
     'Comply with MWBE 30 percent goal, SDVOB 6 percent goal, EEO, reporting, and recordkeeping.',
     'Approved utilization plans, monthly MWBE-SDVOB reports, workforce utilization reports, and maintained records.'),
]

for drv, imp, deliv in reg_rows:
    row_cells = reg_table.add_row().cells
    row_cells[0].text = drv
    row_cells[1].text = imp
    row_cells[2].text = deliv

# ------------------------
# Regulatory, Grant, and Standards Compliance Approach
# ------------------------
sec = doc.add_heading('Regulatory, Grant, and Standards Compliance Approach', level=1)
sec.paragraph_format.space_before = Pt(12)

p = doc.add_paragraph('Our compliance approach is designed to provide NFWB and NYSDEC with clear, audit-ready evidence that every applicable requirement has been met. We integrate EPG terms, NYSDEC Engineering Report expectations, Mandatory State Financial Assistance Terms, and all Appendix E conditions directly into our project controls, templates, and QA/QC processes.')

p = doc.add_paragraph()
run = p.add_run('Alignment to NYSDEC and EPG requirements')
run.bold = True
p.add_run(': We will establish a compliance matrix at project initiation that maps each requirement (including the October 1, 2022 Engineering Report Outline, Order on Consent provisions, and EPG conditions) to specific tasks, deliverables, and responsible staff. This matrix will be maintained throughout the project and used as a checklist prior to each submission to NFWB and NYSDEC.')

p = doc.add_paragraph()
run = p.add_run('Statutory and Appendix E compliance')
run.bold = True
p.add_run(': We will complete and submit all Appendix E forms (Acknowledgement of Addenda, Certificate of Non-Collusion, Lobbying Law Disclosure, EEO Policy Statement, Statement on Sexual Harassment, RFP Acknowledgement and Certification) in accordance with the RFP. Our internal legal and compliance team will review all certifications related to Lobbying Law, FOIL, Iran Divestment, boycott prohibitions, MacBride principles, tropical hardwoods, conflicts of interest, and debarment status to ensure accuracy and completeness.')

# Compliance workflow Mermaid diagram
mermaid_code = '''
flowchart TD
  A[Project initiation] --> B[Compliance matrix setup]
  B --> C[Form and certification completion]
  C --> D[Internal legal review]
  D --> E[Submission to NFWB]
  E --> F[NYSDEC and EPG review]
  F --> G[Ongoing reporting and recordkeeping]
'''
comp_flow_path = render_mermaid(mermaid_code, 'compliance_workflow')
doc.add_picture(str(comp_flow_path), width=Inches(5.8))
add_caption(next_figure_label('Compliance and approvals workflow from initiation through final reporting'))

# MWBE/SDVOB utilization chart
util_data = pd.DataFrame({
    'Category': ['Prime', 'MWBE', 'SDVOB'],
    'Planned Percent': [64, 30, 6]
})
plt.figure(figsize=(5.5, 3.2))
ax = sns.barplot(data=util_data, x='Category', y='Planned Percent', palette=['#4c72b0', '#55a868', '#c44e52'])
ax.axhline(30, color='#55a868', linestyle='--', linewidth=1, label='MWBE goal')
ax.axhline(6, color='#c44e52', linestyle=':', linewidth=1, label='SDVOB goal')
ax.set_ylabel('Percent of total contract value')
ax.set_title('Planned MWBE and SDVOB utilization relative to goals')
ax.legend(fontsize=8)
plt.tight_layout()
util_path = output_dir / 'mwbe_sdvob_utilization.png'
plt.savefig(util_path, dpi=150, bbox_inches='tight', facecolor='white')
plt.close()

doc.add_picture(str(util_path), width=Inches(5.8))
add_caption(next_figure_label('Planned MWBE and SDVOB utilization compared to required goals'))

# Compliance matrix table (high level)
comp_table = doc.add_table(rows=1, cols=3)
comp_table.style = 'Table Grid'
comp_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = comp_table.rows[0].cells
hdr[0].text = 'Requirement Area'
hdr[1].text = 'How Requirement Is Met'
hdr[2].text = 'Primary Evidence / Location in Proposal'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

comp_rows = [
    ('EPG and NYSDEC Engineering Report Outline',
     'Report structure, content, and level of analysis explicitly follow the October 1, 2022 outline.',
     'Section "Task 7 – SSES Engineering Report and Corrective Action Plan" and Appendix crosswalk table.'),
    ('Appendix E forms and certifications',
     'All required forms completed, signed where required, and included in the original proposal.',
     'Appendices – Required Forms, Insurance Certificates, and Supporting Materials.'),
    ('Mandatory State Financial Assistance Terms (MWBE/SDVOB, EEO, reporting)',
     'MWBE/SDVOB utilization plan, monthly reporting routines, and workforce reporting integrated into project governance.',
     'Sections "MWBE and SDVOB Utilization Strategy" and "Administrative Requirements and Deliverables Checklist".'),
    ('FOIL and confidentiality',
     'Sensitive pages marked with confidentiality legends while recognizing potential disclosure.',
     'Section "Administrative Requirements and Deliverables Checklist" – FOIL compliance discussion.'),
]

for area, how, ev in comp_rows:
    row_cells = comp_table.add_row().cells
    row_cells[0].text = area
    row_cells[1].text = how
    row_cells[2].text = ev

# ------------------------
# Project Management and Governance Plan
# ------------------------
sec = doc.add_heading('Project Management and Governance Plan', level=1)
sec.paragraph_format.space_before = Pt(12)

p = doc.add_paragraph('Aurelia will apply a structured yet flexible project management framework that combines clear governance controls with responsive communication. Our Project Manager (PM) will be NFWB’s primary point of contact and will coordinate all technical, field, and reporting activities across Aurelia staff and subcontractors.')

p = doc.add_paragraph()
run = p.add_run('Governance structure and decision-making')
run.bold = True
p.add_run(': Governance will be anchored by a joint NFWB–Aurelia project steering group, regular progress meetings, and defined decision gates aligned to key milestones such as SSES work plan approval, completion of field investigations, draft report submittal, and final NYSDEC submittals.')

# RACI matrix table (high-level)
raci_table = doc.add_table(rows=1, cols=6)
raci_table.style = 'Table Grid'
raci_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = raci_table.rows[0].cells
hdr[0].text = 'Major Task'
hdr[1].text = 'NFWB'
hdr[2].text = 'Project Manager'
hdr[3].text = 'Field Lead'
hdr[4].text = 'Technical Lead'
hdr[5].text = 'MWBE / Admin'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

raci_rows = [
    ('Data and historical information review', 'A', 'R', 'C', 'R', 'I'),
    ('Smoke testing program', 'C', 'R', 'R', 'C', 'I'),
    ('Flow isolation and weiring', 'C', 'R', 'R', 'C', 'I'),
    ('CCTV and manhole analysis', 'C', 'R', 'C', 'R', 'I'),
    ('Outfall/SSO assessment', 'C', 'R', 'R', 'C', 'I'),
    ('Focused SSES and contingency use', 'A', 'R', 'R', 'C', 'C'),
    ('SSES engineering report', 'A', 'R', 'C', 'R', 'I'),
    ('Post-construction monitoring and I and I assessment', 'A', 'R', 'R', 'R', 'I'),
    ('MWBE/SDVOB reporting and invoicing', 'I', 'A', 'I', 'I', 'R'),
]

for task, nfw, pm, fld, tech, mwbe in raci_rows:
    row_cells = raci_table.add_row().cells
    row_cells[0].text = task
    row_cells[1].text = nfw
    row_cells[2].text = pm
    row_cells[3].text = fld
    row_cells[4].text = tech
    row_cells[5].text = mwbe

# ------------------------
# Technical Approach – Phase 1 SSES Overview
# ------------------------
sec = doc.add_heading('Technical Approach – Phase 1 SSES Overview', level=1)
sec.paragraph_format.space_before = Pt(12)

p = doc.add_paragraph('Our Phase 1 SSES approach follows a logical sequence from data review through field investigations, analysis, and reporting, with explicit checkpoints to confirm NYSDEC and NFWB expectations are being met. Each task is designed to produce quantitative and qualitative information that feeds directly into the corrective action plan and engineering report.')

# Phase 1 process flow Mermaid diagram
mermaid_code = '''
flowchart TD
  A[Kickoff and planning] --> B[Data and historical review]
  B --> C[Smoke testing]
  C --> D[Flow isolation and weiring]
  D --> E[CCTV and manhole analysis]
  E --> F[Outfall and SSO assessment]
  F --> G[Focused SSES activities]
  G --> H[Integrated analysis]
  H --> I[Engineering report and corrective plan]
'''
phase1_flow_path = render_mermaid(mermaid_code, 'phase1_overview_flow')
doc.add_picture(str(phase1_flow_path), width=Inches(5.8))
add_caption(next_figure_label('High level Phase 1 SSES workflow'))

# Level of effort chart by SSES task
loe_data = pd.DataFrame({
    'Task': ['Data review', 'Smoke testing', 'Flow isolation', 'CCTV analysis', 'Outfall assessment', 'Focused SSES'],
    'Relative Effort': [10, 20, 20, 20, 15, 15]
})
plt.figure(figsize=(6, 3.5))
ax = sns.barplot(data=loe_data, x='Task', y='Relative Effort', color='#4c72b0')
ax.set_ylabel('Relative effort index')
ax.set_xlabel('Phase 1 task')
ax.set_title('Relative level of effort by Phase 1 SSES task')
plt.xticks(rotation=25, ha='right')
plt.tight_layout()
loe_path = output_dir / 'phase1_loe.png'
plt.savefig(loe_path, dpi=150, bbox_inches='tight', facecolor='white')
plt.close()

doc.add_picture(str(loe_path), width=Inches(5.8))
add_caption(next_figure_label('Relative level of effort by Phase 1 SSES task'))

# Phase 1 task summary table
phase1_table = doc.add_table(rows=1, cols=4)
phase1_table.style = 'Table Grid'
phase1_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = phase1_table.rows[0].cells
hdr[0].text = 'Task'
hdr[1].text = 'Objective'
hdr[2].text = 'Key Methods'
hdr[3].text = 'Primary Deliverables'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

phase1_rows = [
    ('Task 1 – Data and historical review', 'Establish baseline understanding and target high-value SSES areas.', 'Review reports, SCADA, CCTV logs, maintenance history, Order on Consent.', 'Data gaps log, preliminary SSES targeting memo.'),
    ('Task 2 – Smoke testing', 'Identify inflow sources such as cross-connections and illicit connections.', 'Smoke testing under dry conditions, notifications, defect logging and photos.', 'Smoke test maps, defect inventory, recommended follow-up actions.'),
    ('Task 3 – Flow isolation and weiring', 'Quantify infiltration rates by sewer segment.', 'Night-time weiring, flow isolation, safety controls, data analysis in gpd per lf.', 'Segment-level infiltration estimates and prioritized list of high-infiltration sections.'),
    ('Task 4 – CCTV and manhole analysis', 'Characterize structural and maintenance defects contributing to I and I.', 'PACP-based review, condition rating, integration with prior repairs.', 'Defect database, condition ratings, and rehabilitation recommendations.'),
    ('Task 5 – Outfall/SSO assessment', 'Assess condition and function of Phase 1 outfalls and SSOs.', 'Site visits, dye testing, instrumentation, GIS documentation.', 'Outfall condition summaries and GIS-compatible data layers.'),
    ('Task 6 – Focused SSES and contingency', 'Investigate complex or unresolved I and I sources.', 'Dye testing, house inspections, targeted CCTV using contingency budget.', 'Focused investigation findings and recommendations.'),
    ('Task 7 – Engineering report and corrective plan', 'Produce NYSDEC-ready SSES report and action plan.', 'Integrated analysis, cost-effectiveness ranking, PE review and stamping.', 'SSES Engineering Report and Corrective Action Plan.'),
    ('Task 8 – Post-construction monitoring', 'Demonstrate I and I reduction and effectiveness of measures.', 'Flow and rainfall monitoring, pre/post analysis, cost-effectiveness assessment.', 'Post-Construction I and I Assessment and Cost Analysis Report.'),
]

for t, obj, meth, deliv in phase1_rows:
    row_cells = phase1_table.add_row().cells
    row_cells[0].text = t
    row_cells[1].text = obj
    row_cells[2].text = meth
    row_cells[3].text = deliv

# ------------------------
# Task 1 – Data and Historical Information Review
# ------------------------
sec = doc.add_heading('Task 1 – Data and Historical Information Review', level=1)
sec.paragraph_format.space_before = Pt(12)

p = doc.add_paragraph('We will begin by assembling and reviewing all available information related to the LaSalle sewersheds, including prior SSES efforts, maintenance records, SCADA data, flow monitoring reports, NYSDEC correspondence, and Order on Consent documentation. This review will identify data gaps, confirm locations of known problem areas, and inform the targeting of field investigations.')

# Data sources table
data_table = doc.add_table(rows=1, cols=5)
data_table.style = 'Table Grid'
data_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = data_table.rows[0].cells
hdr[0].text = 'Dataset'
hdr[1].text = 'Owner'
hdr[2].text = 'Coverage'
hdr[3].text = 'Data Quality / Format'
hdr[4].text = 'Use in SSES Planning'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

data_rows = [
    ('SCADA flow and level data', 'NFWB operations', 'Key pump stations and interceptors', 'Digital time series, recent years', 'Baseline flow characterization, wet-weather response, targeting of weiring locations.'),
    ('CCTV inspection reports and videos', 'NFWB collections', 'Phase 1 sewersheds', 'PACP-coded reports and video files', 'Structural/maintenance defect assessment and identification of additional inspection needs.'),
    ('Manhole inspection records', 'NFWB collections', 'Phase 1 manholes', 'Inspection forms and GIS attributes', 'Inflow sources, surcharge evidence, and prioritization of rehabilitation needs.'),
    ('Order on Consent and NYSDEC correspondence', 'NFWB management', 'System-wide', 'PDF letters and orders', 'Regulatory context, milestones, and constraints on schedule and scope.'),
    ('Prior I and I and SSO studies', 'NFWB / prior consultants', 'Relevant basins', 'Reports and GIS layers', 'Lessons learned, previously implemented measures, and remaining gaps.'),
]

for ds, owner, cov, qual, use in data_rows:
    row_cells = data_table.add_row().cells
    row_cells[0].text = ds
    row_cells[1].text = owner
    row_cells[2].text = cov
    row_cells[3].text = qual
    row_cells[4].text = use

# ------------------------
# Task 2 – Smoke Testing Program
# ------------------------
sec = doc.add_heading('Task 2 – Smoke Testing Program', level=1)
sec.paragraph_format.space_before = Pt(12)

p = doc.add_paragraph('Our smoke testing program will be executed in close coordination with NFWB operations, public safety agencies, and the community. We will provide all staff, equipment, and materials necessary to conduct testing under appropriate weather conditions, ensuring thorough documentation and safe operations.')

# Notification and outreach plan table
notif_table = doc.add_table(rows=1, cols=4)
notif_table.style = 'Table Grid'
notif_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = notif_table.rows[0].cells
hdr[0].text = 'Activity'
hdr[1].text = 'Description'
hdr[2].text = 'Responsible Party'
hdr[3].text = 'Timing'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

notif_rows = [
    ('Resident notifications', 'Door hangers describing smoke testing, safety information, and contact details.', 'Aurelia field team with NFWB review of materials.', '2–4 days prior to testing.'),
    ('Press release support', 'Draft press release content for NFWB to issue.', 'Aurelia PM and NFWB communications.', '1 week prior to first testing window.'),
    ('Police and fire coordination', 'Provide maps and schedule to police and fire to avoid unnecessary responses.', 'Aurelia PM with NFWB operations.', 'At least 3 business days prior to testing.'),
    ('On-site signage and traffic control', 'Temporary signage and cones as needed to alert motorists and pedestrians.', 'Aurelia field lead and subcontractors.', 'Day of testing.'),
]

for act, desc, resp, time in notif_rows:
    row_cells = notif_table.add_row().cells
    row_cells[0].text = act
    row_cells[1].text = desc
    row_cells[2].text = resp
    row_cells[3].text = time

# ------------------------
# Task 3 – Flow Isolation and Night-Time Weiring
# ------------------------
sec = doc.add_heading('Task 3 – Flow Isolation and Night-Time Weiring', level=1)
sec.paragraph_format.space_before = Pt(12)

p = doc.add_paragraph('We will design and implement a flow isolation and night-time weiring program to quantify groundwater infiltration by sewer segment. Our team will review existing flow data, consult with NFWB to select safe and hydraulically appropriate weir locations, and provide all staff, equipment, and traffic control needed to perform work during low-use periods. Confined space entry will be managed in full compliance with OSHA and NFWB requirements.')

# Candidate weir locations table
weir_table = doc.add_table(rows=1, cols=4)
weir_table.style = 'Table Grid'
weir_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = weir_table.rows[0].cells
hdr[0].text = 'Location Type'
hdr[1].text = 'Rationale'
hdr[2].text = 'Access / Safety Considerations'
hdr[3].text = 'Data Use'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

weir_rows = [
    ('Upstream residential submain', 'Capture localized infiltration in a residential sub-basin.', 'Confined space entry, potential traffic control on local streets.', 'Segment-level infiltration rate and prioritization of lining or repairs.'),
    ('Downstream trunk main', 'Measure cumulative infiltration for larger basin.', 'Deeper manhole, need for gas monitoring and retrieval equipment.', 'Validation of basin-level I and I reduction potential.'),
    ('Near outfall approach', 'Assess contribution of upstream sewers to SSO outfall.', 'Coordination with NFWB to avoid surcharge conditions.', 'Linkage between SSES findings and SSO mitigation.'),
]

for loc, rat, acc, use in weir_rows:
    row_cells = weir_table.add_row().cells
    row_cells[0].text = loc
    row_cells[1].text = rat
    row_cells[2].text = acc
    row_cells[3].text = use

# ------------------------
# Task 4 – CCTV and Manhole Inspection Analysis
# ------------------------
sec = doc.add_heading('Task 4 – CCTV and Manhole Inspection Analysis', level=1)
sec.paragraph_format.space_before = Pt(12)

p = doc.add_paragraph('We will analyze existing CCTV and manhole inspection data, supplemented by additional inspections where needed, using PACP standards and structured defect coding. Our analysis will identify structural defects, inflow sources, and maintenance issues that contribute to I and I, and will coordinate with NFWB on repairs already completed or planned.')

# Defect categories table
cctv_table = doc.add_table(rows=1, cols=3)
cctv_table.style = 'Table Grid'
cctv_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = cctv_table.rows[0].cells
hdr[0].text = 'Defect Category'
hdr[1].text = 'Typical PACP Rating Range'
hdr[2].text = 'Proposed Follow-Up Action'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

cctv_rows = [
    ('Structural defects (cracks, fractures)', '3–5', 'CIPP lining or segment replacement, prioritized by severity and I and I contribution.'),
    ('Joints and connections', '2–4', 'Joint sealing or spot repairs based on infiltration evidence.'),
    ('Manhole defects (covers, chimneys)', '2–4', 'Manhole rehabilitation including frame adjustment, chimney seals, and lining.'),
    ('Inflow sources (roof drains, yard drains)', '2–5', 'Disconnection and redirection of inflow sources, with property owner coordination.'),
]

for cat, rating, action in cctv_rows:
    row_cells = cctv_table.add_row().cells
    row_cells[0].text = cat
    row_cells[1].text = rating
    row_cells[2].text = action

# ------------------------
# Task 5 – Outfall/SSO Condition Assessment and GIS Integration
# ------------------------
sec = doc.add_heading('Task 5 – Outfall/SSO Condition Assessment and GIS Integration', level=1)
sec.paragraph_format.space_before = Pt(12)

p = doc.add_paragraph('For the two Phase 1 outfalls/SSOs, we will develop a Work Plan in consultation with NFWB that may include site inspections, dye testing, temporary instrumentation, and targeted CCTV. Findings will be documented in GIS-compatible format, including geometry, attributes, and photos, to support NFWB’s long-term asset management and regulatory reporting.')

# Outfall attributes table
outfall_table = doc.add_table(rows=1, cols=5)
outfall_table.style = 'Table Grid'
outfall_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = outfall_table.rows[0].cells
hdr[0].text = 'Attribute'
hdr[1].text = 'Description'
hdr[2].text = 'Example Values'
hdr[3].text = 'GIS Field'
hdr[4].text = 'Use in Analysis'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

outfall_rows = [
    ('Location', 'Georeferenced point of outfall structure.', 'Latitude, longitude.', 'geometry', 'Mapping and spatial analysis of SSO risk.'),
    ('Configuration', 'Type, size, and orientation of outfall.', '36 inch pipe to open channel.', 'config', 'Hydraulic assessment and feasibility of modifications.'),
    ('Observed condition', 'Physical condition and evidence of discharge.', 'Good, minor corrosion, evidence of surcharging.', 'condition', 'Prioritization of repair or monitoring needs.'),
    ('Recommended actions', 'Proposed rehabilitation or monitoring measures.', 'Install flap gate, regrade channel, add monitoring device.', 'recom_act', 'Corrective action planning and cost estimation.'),
    ('Related sewersheds', 'Upstream sewersheds contributing to outfall.', 'LaSalle sub-basin A.', 'sewershed_id', 'Linkage to SSES findings and I and I reductions.'),
]

for attr, desc, ex, field, use in outfall_rows:
    row_cells = outfall_table.add_row().cells
    row_cells[0].text = attr
    row_cells[1].text = desc
    row_cells[2].text = ex
    row_cells[3].text = field
    row_cells[4].text = use

# ------------------------
# Task 6 – Focused SSES Activities and Contingency Use
# ------------------------
sec = doc.add_heading('Task 6 – Focused SSES Activities and Contingency Use', level=1)
sec.paragraph_format.space_before = Pt(12)

p = doc.add_paragraph('We will reserve a $10,000 contingency allowance for focused SSES activities such as dye testing, house inspections, and targeted CCTV, to be used only with NFWB approval. Governance of this contingency will ensure that funds are applied to the highest-value investigations and that all expenditures are fully documented and traceable to specific I and I reduction objectives.')

# Contingency management table
cont_table = doc.add_table(rows=1, cols=4)
cont_table.style = 'Table Grid'
cont_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = cont_table.rows[0].cells
hdr[0].text = 'Potential Use'
hdr[1].text = 'Trigger'
hdr[2].text = 'Approval Steps'
hdr[3].text = 'Documentation'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

cont_rows = [
    ('Dye testing of suspect cross-connections', 'Conflicting evidence from smoke testing and CCTV.', 'PM prepares brief justification; NFWB approval via email.', 'Field logs, dye test maps, and photo documentation.'),
    ('House inspections in localized problem areas', 'Persistent inflow observed during storms.', 'Joint decision at progress meeting; homeowner coordination plan.', 'Inspection checklists, findings summary, and recommended corrective actions.'),
    ('Targeted CCTV on uninspected segments', 'Data gaps in high I and I segments.', 'PM submits scope and cost estimate; NFWB written authorization.', 'PACP reports, video files, and updated defect database.'),
]

for use, trig, appr, docu in cont_rows:
    row_cells = cont_table.add_row().cells
    row_cells[0].text = use
    row_cells[1].text = trig
    row_cells[2].text = appr
    row_cells[3].text = docu

# ------------------------
# Task 7 – SSES Engineering Report and Corrective Action Plan
# ------------------------
sec = doc.add_heading('Task 7 – SSES Engineering Report and Corrective Action Plan', level=1)
sec.paragraph_format.space_before = Pt(12)

p = doc.add_paragraph('We will prepare an SSES Engineering Report and Corrective Action Plan in full accordance with the October 1, 2022 Engineering Report Outline for NYS Wastewater Infrastructure Projects. The report will identify and quantify I and I sources, evaluate alternative corrective measures, and present a prioritized, cost-effective plan with budgetary estimates and implementation timelines. The report will be prepared, stamped, and dated by a New York State licensed Professional Engineer.')

# Corrective actions summary table
ca_table = doc.add_table(rows=1, cols=5)
ca_table.style = 'Table Grid'
ca_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = ca_table.rows[0].cells
hdr[0].text = 'Corrective Action Type'
hdr[1].text = 'Typical Locations'
hdr[2].text = 'Priority Basis'
hdr[3].text = 'Estimated Cost Range'
hdr[4].text = 'Expected I and I Reduction'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

ca_rows = [
    ('CIPP lining', 'Cracked or fractured gravity mains.', 'High infiltration rate and structural risk.', '$$–$$$', 'High reduction in groundwater infiltration.'),
    ('Spot repairs', 'Isolated structural defects.', 'Localized defects with limited extent.', '$–$$', 'Moderate reduction, cost-effective for small areas.'),
    ('Manhole rehabilitation', 'Deteriorated covers, chimneys, and barrels.', 'Evidence of inflow or surcharge.', '$–$$', 'Reduction of surface inflow and surcharge-related I and I.'),
    ('Inflow source removal', 'Roof drains, yard drains, sump pumps.', 'Direct inflow connections identified through smoke and dye testing.', '$', 'High reduction in wet-weather inflow for relatively low cost.'),
]

for typ, loc, pri, cost, red in ca_rows:
    row_cells = ca_table.add_row().cells
    row_cells[0].text = typ
    row_cells[1].text = loc
    row_cells[2].text = pri
    row_cells[3].text = cost
    row_cells[4].text = red

# ------------------------
# Task 8 – Post-Construction Monitoring and I&I Assessment
# ------------------------
sec = doc.add_heading('Task 8 – Post-Construction Monitoring and I&I Assessment', level=1)
sec.paragraph_format.space_before = Pt(12)

p = doc.add_paragraph('Following completion of Phase 1 construction, we will conduct post-construction monitoring in the Phase 1 sewersheds to evaluate I and I reduction. This will include collection and analysis of flow and rainfall data, comparison of pre- and post-construction conditions, and preparation of a Post-Construction I and I Assessment and Cost Analysis Report for NYSDEC.')

# Monitoring locations table
mon_table = doc.add_table(rows=1, cols=4)
mon_table.style = 'Table Grid'
mon_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = mon_table.rows[0].cells
hdr[0].text = 'Monitoring Location Type'
hdr[1].text = 'Instrumentation'
hdr[2].text = 'Data Periods'
hdr[3].text = 'Key Performance Indicators'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

mon_rows = [
    ('Upstream sub-basin manholes', 'Temporary flow meters and level sensors.', 'Pre-construction baseline and post-construction wet seasons.', 'Reduction in base infiltration and peak wet-weather flows.'),
    ('Key pump stations', 'SCADA flow and level data.', 'Continuous, with focused analysis periods.', 'Reduction in pump run times and overflow risk.'),
    ('Outfalls/SSOs', 'Visual inspections and, if needed, level sensors.', 'Targeted during wet-weather events.', 'Frequency and volume of SSO occurrences.'),
]

for loc, instr, periods, kpi in mon_rows:
    row_cells = mon_table.add_row().cells
    row_cells[0].text = loc
    row_cells[1].text = instr
    row_cells[2].text = periods
    row_cells[3].text = kpi

# ------------------------
# Schedule and Work Plan
# ------------------------
sec = doc.add_heading('Schedule and Work Plan', level=1)
sec.paragraph_format.space_before = Pt(12)

p = doc.add_paragraph('Our proposed schedule is designed to support NYSDEC review of the SSES Engineering Report early in 2023, enabling construction in 2023–early 2024 and post-construction monitoring through June 2024. The schedule also accommodates coordination with NFWB, MWBE/SDVOB reporting, and seasonal constraints on field work.')

# Define schedule data for Gantt chart
schedule = pd.DataFrame({
    'Task': [
        'Project initiation and planning',
        'Task 1 – Data review',
        'Task 2 – Smoke testing',
        'Task 3 – Flow isolation and weiring',
        'Task 4 – CCTV and manhole analysis',
        'Task 5 – Outfall and SSO assessment',
        'Task 6 – Focused SSES and contingency',
        'Task 7 – Engineering report and corrective plan',
        'Task 8 – Post construction monitoring'
    ],
    'Team': [
        'PM and NFWB',
        'Technical',
        'Field',
        'Field',
        'Technical',
        'Field',
        'Field and Technical',
        'Technical',
        'Technical'
    ],
    'Start': pd.to_datetime([
        '2022-12-06',
        '2022-12-10',
        '2023-01-10',
        '2023-02-01',
        '2023-02-15',
        '2023-03-01',
        '2023-03-20',
        '2023-04-01',
        '2023-07-01'
    ]),
    'End': pd.to_datetime([
        '2022-12-20',
        '2023-01-15',
        '2023-02-10',
        '2023-02-28',
        '2023-03-15',
        '2023-03-20',
        '2023-04-15',
        '2023-05-31',
        '2024-06-30'
    ])
})

project_start = schedule['Start'].min()
schedule['days_to_start'] = (schedule['Start'] - project_start).dt.days
schedule['duration'] = (schedule['End'] - schedule['Start']).dt.days + 1

team_colors = {
    'PM and NFWB': '#4c72b0',
    'Technical': '#55a868',
    'Field': '#c44e52',
    'Field and Technical': '#8172b2'
}

fig, ax = plt.subplots(figsize=(9, 5))
for idx, row in schedule.iterrows():
    ax.barh(y=row['Task'], width=row['duration'], left=row['days_to_start'],
            color=team_colors[row['Team']], edgecolor='white', linewidth=0.5)

ax.invert_yaxis()
max_days = schedule['days_to_start'].max() + schedule['duration'].max()
xticks = np.arange(0, max_days + 15, 30)
xticklabels = pd.date_range(start=project_start, periods=len(xticks), freq='30D').strftime('%b %d %Y')
ax.set_xticks(xticks[:len(xticklabels)])
ax.set_xticklabels(xticklabels, fontsize=8, rotation=30, ha='right')
ax.set_xlabel('Timeline', fontsize=10)
ax.set_title('Phase 1 SSES Implementation Schedule', fontsize=12, fontweight='bold')
ax.xaxis.grid(True, alpha=0.3, linestyle='--')
ax.set_axisbelow(True)
unique_teams = schedule['Team'].unique()
patches = [mpatches.Patch(color=team_colors[t], label=t) for t in unique_teams]
ax.legend(handles=patches, loc='lower right', fontsize=8, framealpha=0.9)
plt.tight_layout()
gantt_path = output_dir / 'phase1_gantt.png'
plt.savefig(gantt_path, dpi=150, bbox_inches='tight', facecolor='white')
plt.close()

doc.add_picture(str(gantt_path), width=Inches(6.0))
add_caption(next_figure_label('Phase 1 SSES implementation schedule'))

# Schedule table
sched_table = doc.add_table(rows=1, cols=4)
sched_table.style = 'Table Grid'
sched_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = sched_table.rows[0].cells
hdr[0].text = 'Task'
hdr[1].text = 'Start Date'
hdr[2].text = 'End Date'
hdr[3].text = 'Responsible Lead'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

for _, row in schedule.iterrows():
    row_cells = sched_table.add_row().cells
    row_cells[0].text = row['Task']
    row_cells[1].text = row['Start'].strftime('%b %d, %Y')
    row_cells[2].text = row['End'].strftime('%b %d, %Y')
    row_cells[3].text = row['Team']

# ------------------------
# Pricing, Level of Effort, and Options for Phases 2 and 3
# ------------------------
sec = doc.add_heading('Pricing, Level of Effort, and Options for Phases 2 and 3', level=1)
sec.paragraph_format.space_before = Pt(12)

p = doc.add_paragraph('We propose a not-to-exceed fee structure for Phase 1, supported by a detailed task/hour estimate matrix. Pricing excludes all state and local taxes and fees, consistent with NFWB’s tax-exempt status. Option prices and schedules are provided for Phase 2 and Phase 3 SSES work, following the same tasks as Phase 1, with Task 5 omitted for Phase 2.')

# Task/hour estimate matrix (conceptual, without dollar amounts)
loe_table = doc.add_table(rows=1, cols=5)
loe_table.style = 'Table Grid'
loe_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = loe_table.rows[0].cells
hdr[0].text = 'Phase / Task'
hdr[1].text = 'PM Hours'
hdr[2].text = 'Engineer Hours'
hdr[3].text = 'Field Hours'
hdr[4].text = 'Total Hours'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

loe_rows = [
    ('Phase 1 – Task 1 to Task 7', 80, 260, 320, 660),
    ('Phase 1 – Task 8', 40, 120, 80, 240),
    ('Phase 2 – Option 1', 60, 200, 280, 540),
    ('Phase 3 – Option 2', 70, 220, 300, 590),
]

for phase, pm_h, eng_h, fld_h, tot_h in loe_rows:
    row_cells = loe_table.add_row().cells
    row_cells[0].text = phase
    row_cells[1].text = str(pm_h)
    row_cells[2].text = str(eng_h)
    row_cells[3].text = str(fld_h)
    row_cells[4].text = str(tot_h)

# Cost distribution chart (conceptual, relative indices)
cost_data = pd.DataFrame({
    'Phase': ['Phase 1 base', 'Phase 1 Task 6 contingency', 'Phase 2 option', 'Phase 3 option'],
    'Relative Cost Index': [100, 10, 90, 95]
})
plt.figure(figsize=(6, 3.5))
ax = sns.barplot(data=cost_data, x='Phase', y='Relative Cost Index', color='#4c72b0')
ax.set_ylabel('Relative cost index')
ax.set_title('Relative cost distribution across base and option scopes')
plt.xticks(rotation=20, ha='right')
plt.tight_layout()
cost_path = output_dir / 'relative_cost_distribution.png'
plt.savefig(cost_path, dpi=150, bbox_inches='tight', facecolor='white')
plt.close()

doc.add_picture(str(cost_path), width=Inches(5.8))
add_caption(next_figure_label('Relative cost distribution across Phase 1 and options'))

# Assumptions and exclusions table
assump_table = doc.add_table(rows=1, cols=3)
assump_table.style = 'Table Grid'
assump_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = assump_table.rows[0].cells
hdr[0].text = 'Area'
hdr[1].text = 'Assumption'
hdr[2].text = 'Implication'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

assump_rows = [
    ('Taxes and fees', 'NFWB is tax-exempt; no state or local taxes are included.', 'Invoices will exclude sales and similar taxes; any future tax changes may require adjustment.'),
    ('Access and traffic control', 'NFWB will assist with obtaining necessary permits and approvals for access and traffic control.', 'Schedule assumes timely permit approvals and coordination with agencies.'),
    ('NFWB staff support', 'NFWB will provide timely access to data, facilities, and staff for coordination.', 'Delays in data or access may affect schedule and level of effort.'),
]

for area, assump, impl in assump_rows:
    row_cells = assump_table.add_row().cells
    row_cells[0].text = area
    row_cells[1].text = assump
    row_cells[2].text = impl

# ------------------------
# Team Organization, Key Personnel, and Safety Program
# ------------------------
sec = doc.add_heading('Team Organization, Key Personnel, and Safety Program', level=1)
sec.paragraph_format.space_before = Pt(12)

p = doc.add_paragraph('Our team combines experienced project management, SSES specialists, and field personnel with proven safety performance. The Project Manager will serve as the primary contact for NFWB, supported by technical leads for SSES analysis, GIS, and reporting, as well as MWBE/SDVOB partners for field execution and specialized services.')

# Team qualifications table
team_table = doc.add_table(rows=1, cols=4)
team_table.style = 'Table Grid'
team_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = team_table.rows[0].cells
hdr[0].text = 'Role'
hdr[1].text = 'Licenses / Credentials'
hdr[2].text = 'Years of Experience'
hdr[3].text = 'Relevant Project Experience'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

team_rows = [
    ('Project Manager', 'PE, NY; PMP', '15+', 'Lead for multiple wastewater planning and SSES projects for public utilities.'),
    ('SSES Technical Lead', 'PE, NY', '12+', 'Designed and implemented I and I reduction programs including flow monitoring and modeling.'),
    ('Field Operations Lead', 'OSHA 30, Confined Space', '10+', 'Supervised smoke testing, flow isolation, and CCTV field crews.'),
    ('GIS and Data Specialist', 'GISP', '8+', 'Developed GIS data models and dashboards for sewer system management.'),
]

for role, lic, yrs, exp in team_rows:
    row_cells = team_table.add_row().cells
    row_cells[0].text = role
    row_cells[1].text = lic
    row_cells[2].text = yrs
    row_cells[3].text = exp

# Safety program summary table
safety_table = doc.add_table(rows=1, cols=3)
safety_table.style = 'Table Grid'
safety_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = safety_table.rows[0].cells
hdr[0].text = 'Safety Area'
hdr[1].text = 'Key Procedures'
hdr[2].text = 'Regulatory Alignment'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

safety_rows = [
    ('Confined space entry', 'Entry permits, atmospheric testing, retrieval systems, attendant and entrant roles.', 'OSHA 29 CFR 1910.146 and NFWB site requirements.'),
    ('Traffic control', 'Traffic control plans, flagging, signage, and coordination with local agencies.', 'MUTCD guidelines and local permits.'),
    ('Personal protective equipment', 'Standard PPE plus task-specific equipment such as gas monitors and fall protection.', 'OSHA PPE standards and NFWB policies.'),
]

for area, proc, reg in safety_rows:
    row_cells = safety_table.add_row().cells
    row_cells[0].text = area
    row_cells[1].text = proc
    row_cells[2].text = reg

# ------------------------
# Qualifications, Experience, and References
# ------------------------
sec = doc.add_heading('Qualifications, Experience, and References', level=1)
sec.paragraph_format.space_before = Pt(12)

p = doc.add_paragraph('Aurelia Digital Systems is a privately held engineering and technology firm founded in 2016, with headquarters in Toronto and delivery teams across Canada and the United States. We specialize in complex public-sector and utilities initiatives that require strong governance, data-driven decision-making, and rigorous compliance management.')

# Statement of Qualifications table
soq_table = doc.add_table(rows=1, cols=4)
soq_table.style = 'Table Grid'
soq_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = soq_table.rows[0].cells
hdr[0].text = 'Item'
hdr[1].text = 'Description'
hdr[2].text = 'Relevance to This Project'
hdr[3].text = 'Notes'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

soq_rows = [
    ('Business structure', 'Corporation with engineering and technology services.', 'Supports multi-disciplinary SSES program with strong data and reporting.', ''),
    ('Licensing', 'Professional engineers licensed in New York and other states.', 'Enables PE-stamped reports and compliance with NYSDEC requirements.', ''),
    ('Years in business', 'In continuous operation since 2016.', 'Demonstrated stability and growth in public-sector engagements.', ''),
    ('Primary industries', 'Public sector, utilities, healthcare, financial services.', 'Direct experience with utility operations and regulatory environments.', ''),
]

for item, desc, rel, notes in soq_rows:
    row_cells = soq_table.add_row().cells
    row_cells[0].text = item
    row_cells[1].text = desc
    row_cells[2].text = rel
    row_cells[3].text = notes

# ------------------------
# MWBE and SDVOB Utilization Strategy
# ------------------------
sec = doc.add_heading('MWBE and SDVOB Utilization Strategy', level=1)
sec.paragraph_format.space_before = Pt(12)

p = doc.add_paragraph('We are committed to meeting or exceeding the combined MWBE goal of 30 percent and SDVOB goal of 6 percent for this engagement. Our strategy includes early identification and engagement of certified partners, clear work packages aligned to their capabilities, and structured reporting to demonstrate compliance with Mandatory State Financial Assistance Terms.')

# MWBE/SDVOB utilization plan table
mwbe_table = doc.add_table(rows=1, cols=4)
mwbe_table.style = 'Table Grid'
mwbe_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = mwbe_table.rows[0].cells
hdr[0].text = 'Firm Type'
hdr[1].text = 'Certification'
hdr[2].text = 'Scope of Work'
hdr[3].text = 'Estimated Share of Contract'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

mwbe_rows = [
    ('MWBE field services partner', 'NYS-certified MWBE', 'Smoke testing, flow isolation support, and traffic control.', 'Approximately 20 percent.'),
    ('MWBE data and GIS partner', 'NYS-certified MWBE', 'GIS data preparation, mapping, and documentation.', 'Approximately 10 percent.'),
    ('SDVOB specialist', 'NYS-certified SDVOB', 'Targeted CCTV, dye testing, and focused SSES.', 'Approximately 6 percent.'),
]

for firm, cert, scope, share in mwbe_rows:
    row_cells = mwbe_table.add_row().cells
    row_cells[0].text = firm
    row_cells[1].text = cert
    row_cells[2].text = scope
    row_cells[3].text = share

# ------------------------
# Risk Management and Quality Assurance
# ------------------------
sec = doc.add_heading('Risk Management and Quality Assurance', level=1)
sec.paragraph_format.space_before = Pt(12)

p = doc.add_paragraph('We will manage technical, schedule, safety, and compliance risks through a structured risk register, proactive mitigation actions, and embedded QA/QC across field work, data, and deliverables. Our governance model emphasizes early identification of issues and transparent communication with NFWB and NYSDEC.')

# Risk register table
risk_table = doc.add_table(rows=1, cols=5)
risk_table.style = 'Table Grid'
risk_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = risk_table.rows[0].cells
hdr[0].text = 'Risk'
hdr[1].text = 'Likelihood'
hdr[2].text = 'Impact'
hdr[3].text = 'Mitigation Actions'
hdr[4].text = 'Owner'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

risk_rows = [
    ('Adverse weather delaying field work', 'Medium', 'High', 'Plan flexible field windows, prioritize critical areas, and coordinate with NFWB on rescheduling.', 'Project Manager and Field Lead.'),
    ('Data gaps in existing records', 'Medium', 'Medium', 'Identify gaps early, schedule targeted field investigations, and adjust scope with NFWB approval.', 'Technical Lead.'),
    ('Safety incident during confined space entry', 'Low', 'High', 'Strict adherence to safety program, training, and permits; stop-work authority for all staff.', 'Field Operations Lead.'),
    ('MWBE/SDVOB utilization shortfall', 'Low', 'Medium', 'Early engagement, clearly defined scopes, and monthly tracking of utilization.', 'Project Manager and MWBE Coordinator.'),
]

for risk, like, imp, mit, owner in risk_rows:
    row_cells = risk_table.add_row().cells
    row_cells[0].text = risk
    row_cells[1].text = like
    row_cells[2].text = imp
    row_cells[3].text = mit
    row_cells[4].text = owner

# ------------------------
# Administrative Requirements and Deliverables Checklist
# ------------------------
sec = doc.add_heading('Administrative Requirements and Deliverables Checklist', level=1)
sec.paragraph_format.space_before = Pt(12)

p = doc.add_paragraph('We have carefully reviewed all administrative, submission, and contractual requirements in the RFP and Appendix E. The following checklist summarizes how we will comply with proposal submission, communication protocols, insurance and indemnification, records retention, and statutory certifications.')

# Deliverables and forms checklist table
adm_table = doc.add_table(rows=1, cols=4)
adm_table.style = 'Table Grid'
adm_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = adm_table.rows[0].cells
hdr[0].text = 'Requirement ID'
hdr[1].text = 'Description'
hdr[2].text = 'Our Response'
hdr[3].text = 'Internal Owner'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

adm_rows = [
    ('REQ-022 / REQ-023', 'Sealed proposal, one original and six copies, plus electronic copy.', 'We will prepare and deliver the required original, copies, and electronic media by the specified deadline.', 'Proposal Manager.'),
    ('REQ-018 / REQ-024 / REQ-054 / REQ-055', 'Appendix E forms, single signed originals, acknowledgement of addenda, and RFP certification.', 'All forms will be completed, signed, and included in the original proposal; copies as required.', 'Legal and Contracts.'),
    ('REQ-028 / REQ-029', 'Insurance and indemnification requirements.', 'We will maintain required coverages, provide certificates naming NFWB and related entities as additional insureds, and accept indemnification terms.', 'Risk Management.'),
    ('REQ-034 / REQ-035 / REQ-051', 'Records retention, security breach laws, and state assistance terms.', 'Project records will be maintained for at least six years and security incidents will be handled per law; MWBE/SDVOB and EEO requirements will be integrated into reporting.', 'Project Controls.'),
]

for req, desc, resp, owner in adm_rows:
    row_cells = adm_table.add_row().cells
    row_cells[0].text = req
    row_cells[1].text = desc
    row_cells[2].text = resp
    row_cells[3].text = owner

# ------------------------
# Appendices – Required Forms, Insurance Certificates, and Supporting Materials
# ------------------------
sec = doc.add_heading('Appendices – Required Forms, Insurance Certificates, and Supporting Materials', level=1)
sec.paragraph_format.space_before = Pt(12)

p = doc.add_paragraph('The appendices to this proposal will include all required forms, acknowledgements, certifications, insurance certificates, MWBE/SDVOB utilization plans or waivers, and any supporting technical materials requested by NFWB. Each appendix item will be clearly labeled and cross-referenced to the relevant RFP requirement ID.')

# Appendix index table
app_table = doc.add_table(rows=1, cols=3)
app_table.style = 'Table Grid'
app_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = app_table.rows[0].cells
hdr[0].text = 'Appendix Item'
hdr[1].text = 'Description'
hdr[2].text = 'Related Requirement IDs'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

app_rows = [
    ('Appendix A – Completed Appendix E Forms', 'All required forms including Acknowledgement of Addenda, Certificate of Non-Collusion, Lobbying Law Disclosure, EEO Policy, Sexual Harassment Statement, and RFP Acknowledgement and Certification.', 'REQ-018, REQ-025, REQ-039, REQ-040, REQ-046, REQ-047, REQ-054, REQ-055.'),
    ('Appendix B – Insurance Certificates', 'Certificates of insurance demonstrating required coverages and additional insured status.', 'REQ-028, REQ-029.'),
    ('Appendix C – MWBE and SDVOB Utilization Plans', 'Completed utilization plans or waiver requests and any supporting documentation.', 'REQ-051, REQ-052.'),
    ('Appendix D – Statement of Qualifications and References', 'Detailed SOQ, project experience summaries, and client references.', 'REQ-026, REQ-053.'),
]

for item, desc, reqs in app_rows:
    row_cells = app_table.add_row().cells
    row_cells[0].text = item
    row_cells[1].text = desc
    row_cells[2].text = reqs
