"""
Executors for the RFP Builder workflow.
Each executor handles a specific stage of the RFP generation process.
"""

import json
import logging
import os
import subprocess
import tempfile
from pathlib import Path
from typing import Optional

from openai import AsyncOpenAI

from app.core.config import get_config
from app.core.llm_client import create_llm_client, get_model_name
from app.core.llm_logger import create_llm_logger
from app.core.text_chunker import (
    chunk_pages_by_tokens,
    format_chunk_pages,
    parse_pages,
    truncate_to_tokens,
)
from app.prompts.system_prompts import (
    RFP_ANALYZER_SYSTEM_PROMPT,
    build_rfp_section_generator_system_prompt,
    PROPOSAL_PLANNER_SYSTEM_PROMPT,
    PROPOSAL_CRITIQUER_SYSTEM_PROMPT,
)
from app.prompts.user_prompts import (
    ANALYZE_RFP_USER_PROMPT,
    GENERATE_SECTIONS_USER_PROMPT,
    PLAN_PROPOSAL_USER_PROMPT,
    GENERATE_WITH_PLAN_USER_PROMPT,
    GENERATE_WITH_ERROR_USER_PROMPT,
    GENERATE_WITH_CRITIQUE_USER_PROMPT,
    GENERATE_SECTIONS_CHUNKED_USER_PROMPT,
    GENERATE_SECTIONS_CHUNKED_WITH_CRITIQUE_USER_PROMPT,
    SYNTHESIZE_DOCUMENT_CODE_PROMPT,
    CRITIQUE_DOCUMENT_USER_PROMPT,
)
from app.functions.rfp_functions import (
    ANALYZE_RFP_FUNCTION,
    GENERATE_RFP_RESPONSE_FUNCTION,
    PLAN_PROPOSAL_FUNCTION,
    CRITIQUE_RESPONSE_FUNCTION,
)
from app.models.schemas import (
    RFPAnalysis,
    RFPResponse,
    RFPRequirement,
    EvaluationCriterion,
    SubmissionRequirements,
    ProposalPlan,
    PlannedSection,
    CritiqueResult,
)
from .state import (
    WorkflowInput,
    AnalysisResult,
    GenerationResult,
    PlanningResult,
    CritiqueResultData,
    FinalResult,
)


logger = logging.getLogger(__name__)


def _as_page_number(text: str) -> str:
    return text.replace("PAGE TO CITE:", "PAGE NUMBER:")


class BaseExecutor:
    """Base class for workflow executors."""
    
    def __init__(self, client: Optional[AsyncOpenAI] = None, run_dir: Optional[Path] = None):
        self.config = get_config()
        self.client = client or create_llm_client()
        self.model = get_model_name()
        
        # Initialize LLM logger if enabled
        self.llm_logger = None
        if self.config.workflow.log_all_steps and run_dir:
            self.llm_logger = create_llm_logger(run_dir)
    
    def _build_messages_with_images(
        self,
        system_prompt: str,
        user_text: str,
        images: Optional[list[dict]] = None
    ) -> list[dict]:
        """Build message list with optional image content."""
        messages = [
            {"role": "system", "content": system_prompt}
        ]
        
        if images and self.config.features.enable_images:
            content = [
                {"type": "text", "text": user_text}
            ]
            content.extend(images)
            messages.append({"role": "user", "content": content})
        else:
            messages.append({"role": "user", "content": user_text})
        
        return messages
    
    async def _call_llm(
        self,
        messages: list[dict],
        functions: Optional[list[dict]] = None,
        function_call: Optional[dict] = None,
    ) -> dict:
        """Make an LLM call with optional function calling."""
        kwargs = {
            "model": self.model,
            "messages": messages,
            "temperature": 0.7
        }
        
        if functions:
            kwargs["tools"] = [
                {"type": "function", "function": f} for f in functions
            ]
            if function_call:
                kwargs["tool_choice"] = {
                    "type": "function",
                    "function": {"name": function_call["name"]}
                }
        
        response = await self.client.chat.completions.create(**kwargs)
        return response


class RFPAnalyzerExecutor(BaseExecutor):
    """Executor that analyzes the RFP and extracts requirements."""
    
    async def execute(self, input_data: WorkflowInput) -> AnalysisResult:
        """Analyze the RFP document and extract structured requirements."""
        logger.info("Starting RFP analysis")
        
        base_additional_context = ""
        if input_data.company_context_text:
            base_additional_context = f"Company Context:\n{input_data.company_context_text}"

        should_chunk = self.config.features.toggle_requirements_chunking
        pages = parse_pages(input_data.rfp_text)
        if not should_chunk or not pages:
            user_prompt = ANALYZE_RFP_USER_PROMPT.format(
                rfp_content=input_data.rfp_text,
                additional_context=base_additional_context
            )

            messages = self._build_messages_with_images(
                RFP_ANALYZER_SYSTEM_PROMPT,
                user_prompt,
                input_data.rfp_images
            )

            response = await self._call_llm(
                messages,
                functions=[ANALYZE_RFP_FUNCTION],
                function_call={"name": "analyze_rfp"}
            )

            message = response.choices[0].message
            raw_response = str(message)

            if message.tool_calls:
                func_args = json.loads(message.tool_calls[0].function.arguments)

                requirements = [
                    RFPRequirement(**req) for req in func_args.get("requirements", [])
                ]

                eval_criteria = None
                if func_args.get("evaluation_criteria"):
                    eval_criteria = [
                        EvaluationCriterion(**ec)
                        for ec in func_args["evaluation_criteria"]
                    ]

                sub_reqs = None
                if func_args.get("submission_requirements"):
                    sub_reqs = SubmissionRequirements(**func_args["submission_requirements"])

                analysis = RFPAnalysis(
                    summary=func_args.get("summary", ""),
                    requirements=requirements,
                    evaluation_criteria=eval_criteria,
                    submission_requirements=sub_reqs,
                    key_differentiators=func_args.get("key_differentiators"),
                )

                if self.llm_logger:
                    self.llm_logger.log_step(
                        step_name="analyze",
                        function_name="analyze_rfp",
                        function_args=func_args,
                        raw_response=raw_response,
                        parsed_result=json.loads(json.dumps(analysis, default=str))
                    )
            else:
                analysis = RFPAnalysis(
                    summary=message.content or "Analysis not available",
                    requirements=[]
                )

            logger.info(f"Analysis complete: {len(analysis.requirements)} requirements found")
            return AnalysisResult(analysis=analysis, raw_response=raw_response)

        max_tokens = self.config.features.max_tokens_reqs_chunking
        chunks = chunk_pages_by_tokens(pages, max_tokens=max_tokens, model=self.model)
        aggregated_requirements: list[RFPRequirement] = []
        aggregated_eval: list[EvaluationCriterion] = []
        aggregated_sub_reqs: Optional[SubmissionRequirements] = None
        aggregated_summary = ""
        aggregated_key_diff: list[str] = []
        raw_responses: list[str] = []
        previous_output = ""
        previous_output_tokens = max(500, min(2000, max_tokens // 4))

        for index, chunk in enumerate(chunks, start=1):
            chunk_text = "\n\n".join(page.text for page in chunk)
            chunk_meta = (
                "DOCUMENT CHUNKING NOTICE:\n"
                f"This RFP has been split into {len(chunks)} parts.\n"
                f"You are processing part {index} of {len(chunks)}.\n"
                f"This part contains pages: {format_chunk_pages(chunk)}.\n"
                "Use the PAGE TO CITE numbers when referencing pages.\n"
            )

            previous_text = ""
            if previous_output:
                previous_text = (
                    "Previously extracted requirements (from earlier parts). "
                    "Do NOT repeat these. Only add new items:\n"
                    f"{truncate_to_tokens(previous_output, previous_output_tokens, model=self.model)}"
                )

            additional_context = "\n\n".join(
                part for part in [base_additional_context, chunk_meta, previous_text] if part
            )

            user_prompt = ANALYZE_RFP_USER_PROMPT.format(
                rfp_content=chunk_text,
                additional_context=additional_context
            )

            messages = self._build_messages_with_images(
                RFP_ANALYZER_SYSTEM_PROMPT,
                user_prompt,
                input_data.rfp_images
            )

            response = await self._call_llm(
                messages,
                functions=[ANALYZE_RFP_FUNCTION],
                function_call={"name": "analyze_rfp"}
            )

            message = response.choices[0].message
            raw_response = str(message)
            raw_responses.append(raw_response)

            if not message.tool_calls:
                continue

            func_args = json.loads(message.tool_calls[0].function.arguments)
            previous_output = json.dumps(func_args, ensure_ascii=True)

            chunk_requirements = [
                RFPRequirement(**req) for req in func_args.get("requirements", [])
            ]
            for req in chunk_requirements:
                if any(
                    existing.id == req.id or existing.description.strip().lower() == req.description.strip().lower()
                    for existing in aggregated_requirements
                ):
                    continue
                aggregated_requirements.append(req)

            if func_args.get("evaluation_criteria"):
                for ec in func_args["evaluation_criteria"]:
                    criterion = EvaluationCriterion(**ec)
                    if any(
                        existing.criterion.strip().lower() == criterion.criterion.strip().lower()
                        for existing in aggregated_eval
                    ):
                        continue
                    aggregated_eval.append(criterion)

            if not aggregated_sub_reqs and func_args.get("submission_requirements"):
                aggregated_sub_reqs = SubmissionRequirements(**func_args["submission_requirements"])

            if not aggregated_summary:
                aggregated_summary = func_args.get("summary", "")

            if func_args.get("key_differentiators"):
                for item in func_args["key_differentiators"]:
                    if item not in aggregated_key_diff:
                        aggregated_key_diff.append(item)

            if self.llm_logger:
                self.llm_logger.log_step(
                    step_name=f"analyze_part_{index}",
                    function_name="analyze_rfp",
                    function_args=func_args,
                    raw_response=raw_response,
                    parsed_result={
                        "requirements_count": len(chunk_requirements),
                        "pages": format_chunk_pages(chunk),
                    }
                )

        analysis = RFPAnalysis(
            summary=aggregated_summary or "Analysis not available",
            requirements=aggregated_requirements,
            evaluation_criteria=aggregated_eval or None,
            submission_requirements=aggregated_sub_reqs,
            key_differentiators=aggregated_key_diff or None,
        )

        logger.info(f"Analysis complete: {len(analysis.requirements)} requirements found")
        return AnalysisResult(analysis=analysis, raw_response="\n\n".join(raw_responses))


def _format_requirements(
    requirements: list[RFPRequirement], 
    include_category: bool = True,
    include_mandatory: bool = True,
    include_priority: bool = False
) -> str:
    """Format requirements list for prompts.
    
    Args:
        requirements: List of RFP requirements
        include_category: Include category info
        include_mandatory: Include mandatory status
        include_priority: Include priority level
    """
    lines = []
    for req in requirements:
        parts = [f"- [{req.id}] {req.description}"]
        extras = []
        if include_category:
            extras.append(f"Category: {req.category}")
        if include_mandatory:
            extras.append(f"Mandatory: {req.is_mandatory}")
        if include_priority:
            extras.append(f"Priority: {req.priority or 'medium'}")
        if extras:
            parts.append(f"({', '.join(extras)})")
        lines.append(" ".join(parts))
    return "\n".join(lines)


def _collect_images(input_data: WorkflowInput) -> list[dict]:
    """Collect all images from input data for the LLM."""
    images = []
    if input_data.example_rfps_images:
        for img_list in input_data.example_rfps_images:
            images.extend(img_list)
    if input_data.company_context_images:
        images.extend(input_data.company_context_images)
    return images


def _collect_images_capped(input_data: WorkflowInput, max_images: int) -> list[dict]:
    images = _collect_images(input_data)
    if max_images > 0:
        return images[:max_images]
    return images


def _format_plan(plan: ProposalPlan) -> str:
    """Format proposal plan for prompts."""
    lines = [
        f"Overview: {plan.overview}",
        f"\nWin Strategy: {plan.win_strategy}",
        "\nKey Themes:"
    ]
    lines.extend(f"- {theme}" for theme in plan.key_themes)
    lines.append("\nPlanned Sections:")
    
    for section in plan.sections:
        lines.append(f"\n### {section.title}")
        lines.append(f"Summary: {section.summary}")
        if section.related_requirements:
            lines.append(f"Requirements: {', '.join(section.related_requirements)}")
        if section.suggested_diagrams:
            lines.append(f"Diagrams: {', '.join(section.suggested_diagrams)}")
        if section.suggested_charts:
            lines.append(f"Charts: {', '.join(section.suggested_charts)}")
        if section.suggested_tables:
            lines.append(f"Tables: {', '.join(section.suggested_tables)}")
    
    return "\n".join(lines)


def _format_critique(critique: CritiqueResult) -> str:
    """Format critique for prompts."""
    lines = [f"Needs Revision: {critique.needs_revision}", f"\nCritique: {critique.critique}"]
    if critique.weaknesses:
        lines.append("\nWeaknesses:")
        lines.extend(f"- {w}" for w in critique.weaknesses)
    if critique.priority_fixes:
        lines.append("\n\nPriority Fixes:")
        lines.extend(f"- {f}" for f in critique.priority_fixes)
    return "\n".join(lines)


class SectionGeneratorExecutor(BaseExecutor):
    """Executor that generates the proposal document code."""

    def _get_generation_settings(self, input_data: WorkflowInput) -> dict:
        cfg = self.config.features
        return {
            "formatting_injection": input_data.generator_formatting_injection,
            "intro_pages": input_data.generator_intro_pages if input_data.generator_intro_pages is not None else cfg.generator_intro_pages,
            "page_overlap": input_data.generation_page_overlap if input_data.generation_page_overlap is not None else cfg.generation_page_overlap,
            "toggle_generation_chunking": input_data.toggle_generation_chunking if input_data.toggle_generation_chunking is not None else cfg.toggle_generation_chunking,
            "max_tokens_generation_chunking": input_data.max_tokens_generation_chunking if input_data.max_tokens_generation_chunking is not None else cfg.max_tokens_generation_chunking,
            "max_sections_per_chunk": input_data.max_sections_per_chunk if input_data.max_sections_per_chunk is not None else cfg.max_sections_per_chunk,
        }

    def _system_prompt_for(self, input_data: WorkflowInput) -> str:
        settings = self._get_generation_settings(input_data)
        return build_rfp_section_generator_system_prompt(settings["formatting_injection"])

    @staticmethod
    def _extract_page_numbers(page_refs: list[str] | None) -> list[int]:
        if not page_refs:
            return []
        pages: list[int] = []
        for ref in page_refs:
            if ref is None:
                continue
            for token in str(ref).replace(",", " ").split():
                if token.isdigit():
                    pages.append(int(token))
        return pages

    def _build_page_map(self, rfp_text: str) -> dict[int, str]:
        pages = parse_pages(rfp_text)
        page_map: dict[int, str] = {}
        for page in pages:
            page_map[page.page_num] = _as_page_number(page.text)
        return page_map

    def _select_pages_text(
        self,
        page_map: dict[int, str],
        intro_pages: int,
        target_pages: list[int],
        page_overlap: int,
        max_tokens: int,
    ) -> str:
        selected: set[int] = set()
        if intro_pages > 0:
            selected.update(range(1, intro_pages + 1))
        for page in target_pages:
            if page <= 0:
                continue
            selected.add(page)
            if page_overlap > 0:
                for offset in range(1, page_overlap + 1):
                    selected.add(page - offset)
                    selected.add(page + offset)
        valid_pages = sorted(p for p in selected if p in page_map)
        combined = "\n\n".join(page_map[p] for p in valid_pages)
        if max_tokens > 0:
            combined = truncate_to_tokens(combined, max_tokens, model=self.model)
        return combined

    @staticmethod
    def _build_sections_outline(sections: list[PlannedSection]) -> str:
        lines = []
        for section in sections:
            title = section.title or "Untitled Section"
            summary = section.summary or ""
            lines.append(f"- {title}: {summary}".strip())
        return "\n".join(lines) if lines else "No sections provided."
    
    async def _generate(
        self,
        user_prompt: str,
        images: list[dict] | None,
        step_name: str,
        system_prompt: str,
        fallback_msg: str = "Generation Failed"
    ) -> GenerationResult:
        """Core generation logic - shared by all execute methods."""
        messages = self._build_messages_with_images(
            system_prompt,
            user_prompt,
            images if images else None
        )
        
        response = await self._call_llm(
            messages,
            functions=[GENERATE_RFP_RESPONSE_FUNCTION],
            function_call={"name": "generate_rfp_response"}
        )
        
        message = response.choices[0].message
        raw_response = str(message)
        
        if message.tool_calls:
            func_args = json.loads(message.tool_calls[0].function.arguments)
            rfp_response = RFPResponse(document_code=func_args.get("document_code", ""))
            
            if self.llm_logger:
                self.llm_logger.log_step(
                    step_name=step_name,
                    function_name="generate_rfp_response",
                    function_args=func_args,
                    raw_response=raw_response,
                    parsed_result={"document_code_length": len(rfp_response.document_code)}
                )
        else:
            rfp_response = RFPResponse(
                document_code=f"doc.add_heading('{fallback_msg}', level=0)\ndoc.add_paragraph('Unable to generate proposal.')"
            )
        
        logger.info(f"{step_name} complete: {len(rfp_response.document_code)} chars of document code")
        return GenerationResult(response=rfp_response, raw_response=raw_response)
    
    async def execute(
        self, 
        input_data: WorkflowInput,
        analysis: RFPAnalysis
    ) -> GenerationResult:
        """Generate Python code for the proposal document."""
        logger.info("Starting document generation")
        system_prompt = self._system_prompt_for(input_data)
        
        user_prompt = GENERATE_SECTIONS_USER_PROMPT.format(
            rfp_analysis=analysis.summary,
            example_rfps="\n\n---\n\n".join(_as_page_number(t) for t in input_data.example_rfps_text),
            company_context=_as_page_number(input_data.company_context_text) if input_data.company_context_text else "No company context provided.",
            requirements=_format_requirements(analysis.requirements)
        )
        
        return await self._generate(user_prompt, _collect_images(input_data), "generate", system_prompt)
    
    async def execute_with_plan(
        self, 
        input_data: WorkflowInput,
        analysis: RFPAnalysis,
        plan: ProposalPlan
    ) -> GenerationResult:
        """Generate Python code for the proposal document using the proposal plan."""
        logger.info("Starting document generation with plan")
        system_prompt = self._system_prompt_for(input_data)
        
        user_prompt = GENERATE_WITH_PLAN_USER_PROMPT.format(
            rfp_analysis=analysis.summary,
            proposal_plan=_format_plan(plan),
            example_rfps="\n\n---\n\n".join(_as_page_number(t) for t in input_data.example_rfps_text),
            company_context=_as_page_number(input_data.company_context_text) if input_data.company_context_text else "No company context provided.",
            requirements=_format_requirements(analysis.requirements)
        )
        
        return await self._generate(user_prompt, _collect_images(input_data), "generate_with_plan", system_prompt)

    async def execute_chunk_with_plan(
        self,
        input_data: WorkflowInput,
        analysis: RFPAnalysis,
        sections: list[PlannedSection],
        part_number: int,
        total_parts: int,
        critique_text: Optional[str] = None,
        previous_code: str | None = None,
    ) -> GenerationResult:
        settings = self._get_generation_settings(input_data)
        system_prompt = self._system_prompt_for(input_data)
        page_map = self._build_page_map(input_data.rfp_text)

        target_pages: list[int] = []
        for section in sections:
            target_pages.extend(self._extract_page_numbers(section.rfp_pages))

        rfp_pages_text = self._select_pages_text(
            page_map,
            intro_pages=settings["intro_pages"],
            target_pages=target_pages,
            page_overlap=settings["page_overlap"],
            max_tokens=settings["max_tokens_generation_chunking"],
        )

        section_req_ids: set[str] = set()
        for section in sections:
            section_req_ids.update(section.related_requirements or [])
        req_list = "\n".join([
            f"- [{req.id}] {req.description} (Category: {req.category}, Mandatory: {req.is_mandatory})"
            for req in analysis.requirements
            if req.id in section_req_ids
        ]) or "No specific requirements listed for these sections."

        example_text = "\n\n---\n\n".join(_as_page_number(t) for t in input_data.example_rfps_text)
        company_context = _as_page_number(input_data.company_context_text) if input_data.company_context_text else "No company context provided."
        sections_outline = self._build_sections_outline(sections)

        if critique_text:
            user_prompt = GENERATE_SECTIONS_CHUNKED_WITH_CRITIQUE_USER_PROMPT.format(
                critique=critique_text,
                previous_code=(previous_code or "")[:5000],
                part_number=part_number,
                total_parts=total_parts,
                sections_outline=sections_outline,
                rfp_analysis=analysis.summary,
                requirements=req_list,
                rfp_pages=rfp_pages_text or "No relevant pages found.",
                example_rfps=example_text,
                company_context=company_context,
            )
        else:
            user_prompt = GENERATE_SECTIONS_CHUNKED_USER_PROMPT.format(
                part_number=part_number,
                total_parts=total_parts,
                sections_outline=sections_outline,
                rfp_analysis=analysis.summary,
                requirements=req_list,
                rfp_pages=rfp_pages_text or "No relevant pages found.",
                example_rfps=example_text,
                company_context=company_context,
            )

        return await self._generate(
            user_prompt,
            _collect_images_capped(input_data, self.config.features.max_images),
            f"generate_part_{part_number}",
            system_prompt,
            "Generation Failed",
        )

    async def synthesize_from_chunks(
        self,
        input_data: WorkflowInput,
        chunk_codes: list[str],
    ) -> GenerationResult:
        system_prompt = self._system_prompt_for(input_data)
        combined_codes = "\n\n--- CHUNK ---\n\n".join(chunk_codes)
        user_prompt = SYNTHESIZE_DOCUMENT_CODE_PROMPT.format(chunk_codes=combined_codes)
        return await self._generate(user_prompt, None, "synthesize_document", system_prompt, "Synthesis Failed")
    
    async def execute_with_error(
        self, 
        input_data: WorkflowInput,
        analysis: RFPAnalysis,
        previous_code: str,
        error_message: str
    ) -> GenerationResult:
        """Regenerate document code after fixing an execution error."""
        logger.info(f"Regenerating document code after error: {error_message[:100]}...")
        system_prompt = self._system_prompt_for(input_data)
        
        user_prompt = GENERATE_WITH_ERROR_USER_PROMPT.format(
            error_message=error_message,
            previous_code=previous_code[:5000],
            rfp_analysis=analysis.summary,
            requirements=_format_requirements(analysis.requirements, include_category=False, include_mandatory=False)
        )
        
        return await self._generate(user_prompt, None, "error_recovery", system_prompt, "Error Recovery Failed")
    
    async def execute_with_critique(
        self, 
        input_data: WorkflowInput,
        analysis: RFPAnalysis,
        previous_code: str,
        critique: CritiqueResult
    ) -> GenerationResult:
        """Regenerate document code after critique."""
        logger.info("Regenerating document code based on critique")
        system_prompt = self._system_prompt_for(input_data)
        
        user_prompt = GENERATE_WITH_CRITIQUE_USER_PROMPT.format(
            critique=_format_critique(critique),
            previous_code=previous_code[:5000],
            rfp_analysis=analysis.summary,
            requirements=_format_requirements(analysis.requirements, include_category=False, include_mandatory=False)
        )
        
        return await self._generate(user_prompt, None, "critique_revision", system_prompt, "Critique Recovery Failed")


class PlannerExecutor(BaseExecutor):
    """Executor that creates a detailed proposal plan before generation."""
    
    async def execute(
        self, 
        input_data: WorkflowInput,
        analysis: RFPAnalysis
    ) -> PlanningResult:
        """Create a detailed proposal plan."""
        logger.info("Starting proposal planning")
        
        user_prompt = PLAN_PROPOSAL_USER_PROMPT.format(
            rfp_analysis=analysis.summary,
            requirements=_format_requirements(analysis.requirements, include_priority=True),
            company_context=input_data.company_context_text or "No company context provided."
        )
        
        messages = self._build_messages_with_images(
            PROPOSAL_PLANNER_SYSTEM_PROMPT,
            user_prompt,
            None  # No images needed for planning
        )
        
        response = await self._call_llm(
            messages,
            functions=[PLAN_PROPOSAL_FUNCTION],
            function_call={"name": "plan_proposal"}
        )
        
        message = response.choices[0].message
        raw_response = str(message)
        
        if message.tool_calls:
            func_args = json.loads(message.tool_calls[0].function.arguments)
            
            sections = [
                PlannedSection(
                    title=s.get("title", "Untitled Section"),
                    summary=s.get("summary", ""),
                    related_requirements=s.get("related_requirements", []),
                    rfp_pages=s.get("rfp_pages", []),
                    suggested_diagrams=s.get("suggested_diagrams", []),
                    suggested_charts=s.get("suggested_charts", []),
                    suggested_tables=s.get("suggested_tables", [])
                )
                for s in func_args.get("sections", [])
            ]
            
            plan = ProposalPlan(
                overview=func_args.get("overview", ""),
                sections=sections,
                key_themes=func_args.get("key_themes", []),
                win_strategy=func_args.get("win_strategy", "")
            )
            
            if self.llm_logger:
                self.llm_logger.log_step(
                    step_name="plan",
                    function_name="plan_proposal",
                    function_args=func_args,
                    raw_response=raw_response,
                    parsed_result={"sections_count": len(plan.sections)}
                )
        else:
            plan = ProposalPlan(
                overview="Planning not available",
                sections=[]
            )
        
        logger.info(f"Planning complete: {len(plan.sections)} sections planned")
        return PlanningResult(plan=plan, raw_response=raw_response)


class CritiquerExecutor(BaseExecutor):
    """Executor that reviews generated document code and provides critique."""
    
    async def execute(
        self, 
        analysis: RFPAnalysis,
        document_code: str
    ) -> CritiqueResultData:
        """Review document code and provide critique."""
        logger.info("Starting document critique")
        
        user_prompt = CRITIQUE_DOCUMENT_USER_PROMPT.format(
            requirements=_format_requirements(analysis.requirements, include_category=False),
            document_code=document_code[:10000]  # Limit for context window
        )
        
        messages = [
            {"role": "system", "content": PROPOSAL_CRITIQUER_SYSTEM_PROMPT},
            {"role": "user", "content": user_prompt}
        ]
        
        response = await self._call_llm(
            messages,
            functions=[CRITIQUE_RESPONSE_FUNCTION],
            function_call={"name": "critique_response"}
        )
        
        message = response.choices[0].message
        raw_response = str(message)
        
        if message.tool_calls:
            func_args = json.loads(message.tool_calls[0].function.arguments)
            
            critique = CritiqueResult(
                needs_revision=func_args.get("needs_revision", False),
                critique=func_args.get("critique", ""),
                strengths=func_args.get("strengths", []),
                weaknesses=func_args.get("weaknesses", []),
                priority_fixes=func_args.get("priority_fixes", [])
            )
            
            if self.llm_logger:
                self.llm_logger.log_step(
                    step_name="critique",
                    function_name="critique_response",
                    function_args=func_args,
                    raw_response=raw_response,
                    parsed_result={"needs_revision": critique.needs_revision}
                )
        else:
            critique = CritiqueResult(
                needs_revision=False,
                critique="Unable to perform critique"
            )
        
        logger.info(f"Critique complete: needs_revision={critique.needs_revision}")
        return CritiqueResultData(critique=critique, raw_response=raw_response)


class CodeInterpreterExecutor(BaseExecutor):
    """Executor that runs document code to generate the final Word document."""
    
    def __init__(self, client: Optional[AsyncOpenAI] = None, run_dir: Optional[Path] = None, output_dir: Optional[Path] = None):
        super().__init__(client, run_dir)
        self.output_dir = output_dir or Path("./output")
    
    def _find_mmdc(self) -> Optional[str]:
        """Find the mermaid CLI executable."""
        import shutil
        
        # Try to find mmdc in PATH
        mmdc = shutil.which('mmdc')
        if mmdc:
            return mmdc
        
        # Try common npm global install locations
        possible_paths = [
            Path.home() / 'AppData' / 'Roaming' / 'npm' / 'mmdc.cmd',  # Windows
            Path.home() / 'AppData' / 'Roaming' / 'npm' / 'mmdc',
            Path('/usr/local/bin/mmdc'),  # macOS/Linux
            Path('/usr/bin/mmdc'),
        ]
        
        for p in possible_paths:
            if p.exists():
                return str(p)
        
        return None
    
    async def execute(
        self,
        response: RFPResponse,
        output_dir: Optional[Path] = None
    ) -> tuple[Path, dict]:
        """
        Execute the document_code to create the Word document.
        
        The code has full access to:
        - python-docx for document creation
        - seaborn/matplotlib for charts
        - subprocess for mermaid CLI
        """
        import matplotlib
        matplotlib.use('Agg')  # Non-interactive backend
        import matplotlib.pyplot as plt
        import seaborn as sns
        import numpy as np
        import pandas as pd
        from docx import Document
        from docx.shared import Inches, Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        
        out_dir = output_dir or self.output_dir
        out_dir = Path(out_dir)
        out_dir.mkdir(parents=True, exist_ok=True)
        
        logger.info("Starting code interpreter execution")
        
        stats = {
            "document_success": False,
            "errors": []
        }
        
        docx_path = out_dir / "proposal.docx"
        
        # Find mmdc path
        mmdc_path = self._find_mmdc()
        
        try:
            # Create the document
            doc = Document()
            
            # Helper function for mermaid that uses the correct path
            def render_mermaid(mermaid_code: str, output_filename: str) -> Path:
                """Render mermaid diagram and return the image path."""
                if not mmdc_path:
                    raise RuntimeError("Mermaid CLI (mmdc) not found. Install with: npm install -g @mermaid-js/mermaid-cli")
                
                mmd_file = out_dir / f"{output_filename}.mmd"
                mmd_file.write_text(mermaid_code, encoding='utf-8')
                png_path = out_dir / f"{output_filename}.png"
                
                result = subprocess.run(
                    [mmdc_path, '-i', str(mmd_file), '-o', str(png_path), '-b', 'white'],
                    capture_output=True,
                    text=True
                )
                if result.returncode != 0:
                    raise RuntimeError(f"Mermaid failed: {result.stderr}")
                return png_path
            
            # Create execution environment with everything the LLM needs
            exec_globals = {
                # Builtins
                "__builtins__": __builtins__,
                # Document
                "doc": doc,
                "output_dir": out_dir,
                # python-docx
                "Inches": Inches,
                "Pt": Pt,
                "Cm": Cm,
                "WD_ALIGN_PARAGRAPH": WD_ALIGN_PARAGRAPH,
                "WD_TABLE_ALIGNMENT": WD_TABLE_ALIGNMENT,
                # Data science
                "plt": plt,
                "sns": sns,
                "np": np,
                "pd": pd,
                # System
                "subprocess": subprocess,
                "tempfile": tempfile,
                "os": os,
                "Path": Path,
                # Mermaid helper
                "render_mermaid": render_mermaid,
                "mmdc_path": mmdc_path,
            }
            
            # Execute the document code
            exec(response.document_code, exec_globals)
            
            # Close any open matplotlib figures
            plt.close('all')
            
            # Save the document
            doc.save(docx_path)
            stats["document_success"] = True
            logger.info(f"Document saved to {docx_path}")
            
        except Exception as e:
            error_msg = f"Document generation failed: {str(e)}"
            stats["errors"].append(error_msg)
            logger.error(error_msg, exc_info=True)
            
            # Create a fallback document with error info
            doc = Document()
            doc.add_heading("Document Generation Error", level=0)
            doc.add_paragraph(f"Error: {str(e)}")
            doc.add_heading("Generated Code", level=1)
            doc.add_paragraph(response.document_code[:5000])  # First 5000 chars
            doc.save(docx_path)
        
        logger.info(f"Code interpreter complete: {'success' if stats['document_success'] else 'failed'}")
        
        return docx_path, stats
